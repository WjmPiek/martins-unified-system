print('Loading app.py...', flush=True)
import os
import uuid
import re
import json
import io
import zipfile
import smtplib
import traceback
import shutil
import html
import time
import threading
from urllib.parse import quote_plus, urlencode, quote, unquote_plus
import urllib.request
from datetime import datetime, timedelta
from email.message import EmailMessage
from functools import wraps

import pandas as pd
try:
    from sqlalchemy import create_engine, text
except Exception:
    create_engine = None
    text = None
from openpyxl import load_workbook
from flask import Flask, render_template, request, redirect, url_for, send_file, flash, session, g, jsonify, make_response
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, 'uploads')
EXPORT_DIR = os.path.join(BASE_DIR, 'exports')
BACKUP_DIR = os.path.join(BASE_DIR, 'backups')
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)
os.makedirs(BACKUP_DIR, exist_ok=True)
DEFAULT_IMPORT_FILE = os.path.join(BASE_DIR, 'exports', 'policy_transaction_import_sample.xlsx')
CONFIG_FILE = os.path.join(BASE_DIR, 'franchise_config.json')
REPORT_LOGO = os.path.join(BASE_DIR, 'static', 'martins_logo.png')

app = Flask(__name__)

# Production/security configuration. Keep all secrets in environment variables
# on Render or in a local .env file. Do not hard-code production secrets.
APP_ENV = os.getenv('APP_ENV', 'local').strip().lower()
IS_PRODUCTION = APP_ENV in {'production', 'prod', 'live'}

app.secret_key = os.getenv('SECRET_KEY', 'change-this-secret-key-before-live-deployment')
app.permanent_session_lifetime = timedelta(minutes=int(os.getenv('SESSION_TIMEOUT_MINUTES', '30')))
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE=os.getenv('SESSION_COOKIE_SAMESITE', 'Lax'),
    SESSION_COOKIE_SECURE=os.getenv('SESSION_COOKIE_SECURE', '1' if IS_PRODUCTION else '0') == '1',
    REMEMBER_COOKIE_HTTPONLY=True,
    REMEMBER_COOKIE_SECURE=os.getenv('SESSION_COOKIE_SECURE', '1' if IS_PRODUCTION else '0') == '1',
)

APP_BASE_URL = os.getenv('APP_BASE_URL', '').strip().rstrip('/')
# Google Maps browser key used for the client heatmap and mini map.
# Production-safe: no hard-coded Google Maps key. Set GOOGLE_MAPS_API_KEY in Render.
def get_google_maps_api_key():
    return os.getenv('GOOGLE_MAPS_API_KEY', '').strip()
APP_VERSION = os.getenv('APP_VERSION', 'phase5-production-launch')
ALERT_EMAIL = os.getenv('ALERT_EMAIL', os.getenv('SMTP_ALERT_TO', '')).strip()
CRON_SECRET = os.getenv('CRON_SECRET', '').strip()
MAINTENANCE_MODE = os.getenv('MAINTENANCE_MODE', '0').strip() == '1'


DEFAULT_RATES = {
    'BrightRock': 10.00,
    'Inkulu': 2.50,  # Mkhulu / Inkulu business commission
    'MFF': 2.50,
}
DEFAULT_BOOK_VALUE = {
    'MFF Book Value': 2.50,
    'Franchise Book Value': 2.50,
}


def safe_float(value, default=0.0):
    """Convert form/query values to float; blank or invalid values use the default."""
    try:
        if value is None:
            return float(default)
        text = str(value).strip()
        if text == '':
            return float(default)
        return float(text)
    except Exception:
        return float(default)

LAST_CLAIMS_IMPORT_SUMMARY = {}
LAST_POLICY_IMPORT_SUMMARY = {}
LAST_POLICY_DETAIL_DF = pd.DataFrame()
LAST_CLAIMS_DF = pd.DataFrame()

LAST_RESULT = {
    'raw': pd.DataFrame(),
    'monthly': pd.DataFrame(),
    'periods': pd.DataFrame(),
    'portfolio': {},
    'rates': DEFAULT_RATES.copy(),
    'book_rates': DEFAULT_BOOK_VALUE.copy(),
}


DATABASE_URL = os.getenv('DATABASE_URL', '').strip()
DB_ENGINE = None
DB_STATUS = {'enabled': False, 'message': 'PostgreSQL not configured'}


def get_db_engine():
    """Return a SQLAlchemy engine when DATABASE_URL points to PostgreSQL."""
    global DB_ENGINE, DB_STATUS
    if DB_ENGINE is not None:
        return DB_ENGINE
    if not DATABASE_URL:
        DB_STATUS = {'enabled': False, 'message': 'DATABASE_URL is not set. Using in-memory mode.'}
        return None
    if create_engine is None or text is None:
        DB_STATUS = {'enabled': False, 'message': 'SQLAlchemy is not installed. Using in-memory mode.'}
        return None
    try:
        db_url = DATABASE_URL
        if db_url.startswith('postgresql://') and '+psycopg2' not in db_url:
            db_url = db_url.replace('postgresql://', 'postgresql+psycopg2://', 1)
        # Keep DB connection lightweight. Do NOT run schema migrations/index creation
        # during app startup or first page load; on large PolicyData tables that can make
        # python app.py look frozen for many minutes. Admin rebuild/repair pages perform
        # explicit setup when needed.
        connect_args = {}
        if 'postgres' in db_url:
            connect_args['connect_timeout'] = int(os.getenv('DB_CONNECT_TIMEOUT', '5'))
        DB_ENGINE = create_engine(db_url, pool_pre_ping=True, future=True, connect_args=connect_args)
        with DB_ENGINE.begin() as conn:
            conn.exec_driver_sql('SELECT 1')
        DB_STATUS = {'enabled': True, 'message': 'Connected to PostgreSQL'}
        print(DB_STATUS['message'], flush=True)
        return DB_ENGINE
    except Exception as exc:
        DB_ENGINE = None
        DB_STATUS = {'enabled': False, 'message': f'PostgreSQL connection failed: {exc}'}
        print(DB_STATUS['message'], flush=True)
        return None



# -----------------------------------------------------------------------------
# PostgreSQL self-healing schema support
# -----------------------------------------------------------------------------
REQUIRED_DB_SCHEMA = {
    'policydata_detail_raw': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'source_file': 'TEXT', 'import_month': 'DATE', 'row_number': 'INTEGER',
        'source_row_key': 'TEXT', 'client_address_o': 'TEXT', 'id_number_f': 'TEXT',
        'franchise_name': 'TEXT', 'relation': 'TEXT', 'is_mem': 'BOOLEAN DEFAULT false',
        'retail_premium': 'NUMERIC(18,2) DEFAULT 0', 'risk_premium': 'NUMERIC(18,2) DEFAULT 0',
        'original_risk_premium': 'NUMERIC(18,2) DEFAULT 0', 'mpia': 'NUMERIC(18,2) DEFAULT 0',
        'single_premium': 'NUMERIC(18,2) DEFAULT 0', 'single_monthly_premium': 'NUMERIC(18,2) DEFAULT 0',
        'r1_policy_fee': 'NUMERIC(18,2) DEFAULT 0', 'adv_fund_2_1_fee': 'NUMERIC(18,2) DEFAULT 0',
        'risk_after_r1': 'NUMERIC(18,2) DEFAULT 0', 'new_risk_premium': 'NUMERIC(18,2) DEFAULT 0',
        'raw_data': 'JSONB', 'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'policy_monthly_raw': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'franchise_name': 'TEXT', 'import_month': 'DATE',
        'retail_premium': 'NUMERIC(18,2) DEFAULT 0', 'risk_premium': 'NUMERIC(18,2) DEFAULT 0',
        'claims': 'NUMERIC(18,2) DEFAULT 0', 'claim_count': 'INTEGER DEFAULT 0',
        'claim_paid_franchise': 'NUMERIC(18,2) DEFAULT 0', 'claim_paid_client': 'NUMERIC(18,2) DEFAULT 0',
        'repudiated_pending': 'NUMERIC(18,2) DEFAULT 0', 'grand_total_claims': 'NUMERIC(18,2) DEFAULT 0',
        'policy_qty': 'NUMERIC(18,2) DEFAULT 0', 'original_risk_premium': 'NUMERIC(18,2) DEFAULT 0',
        'r1_policy_fee': 'NUMERIC(18,2) DEFAULT 0', 'underwriter_2_1_fee': 'NUMERIC(18,2) DEFAULT 0',
        'risk_after_r1': 'NUMERIC(18,2) DEFAULT 0', 'single_monthly_premium_total': 'NUMERIC(18,2) DEFAULT 0',
        'current_scenario': "TEXT DEFAULT '100% Claim Ratio'", 'source_file': 'TEXT',
        'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'claims_detail_raw': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'source_file': 'TEXT', 'claim_month': 'DATE', 'franchise_name': 'TEXT',
        'claim_amount': 'NUMERIC(18,2) DEFAULT 0', 'raw_data': 'JSONB',
        'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'claims_monthly_raw': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'claim_key': 'TEXT', 'claims_franchise_name': 'TEXT', 'claim_month': 'DATE',
        'claims_amount': 'NUMERIC(18,2) DEFAULT 0', 'claim_count': 'INTEGER DEFAULT 0',
        'claim_paid_franchise': 'NUMERIC(18,2) DEFAULT 0', 'claim_paid_client': 'NUMERIC(18,2) DEFAULT 0',
        'repudiated_pending': 'NUMERIC(18,2) DEFAULT 0', 'grand_total_claims': 'NUMERIC(18,2) DEFAULT 0',
        'source_file': 'TEXT', 'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'franchise_mapping_pg': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'source_name': 'TEXT', 'mapped_name': 'TEXT', 'approved': 'BOOLEAN DEFAULT true',
        'created_at': 'TIMESTAMP DEFAULT NOW()', 'updated_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'franchise_monthly_summary': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'franchise_name': 'TEXT', 'report_month': 'DATE',
        'retail_premium': 'NUMERIC(18,2) DEFAULT 0', 'risk_premium': 'NUMERIC(18,2) DEFAULT 0',
        'claims': 'NUMERIC(18,2) DEFAULT 0', 'claim_count': 'INTEGER DEFAULT 0',
        'claim_ratio': 'NUMERIC(10,4) DEFAULT 0',
        'brightrock_commission': 'NUMERIC(18,2) DEFAULT 0', 'mkhulu_commission': 'NUMERIC(18,2) DEFAULT 0',
        'mff_commission': 'NUMERIC(18,2) DEFAULT 0', 'r1_fee': 'NUMERIC(18,2) DEFAULT 0',
        'adv_fund_fee': 'NUMERIC(18,2) DEFAULT 0', 'net_risk_premium': 'NUMERIC(18,2) DEFAULT 0',
        'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'import_history': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'import_type': 'TEXT', 'source_file': 'TEXT', 'imported_months': 'TEXT[]',
        'row_count': 'INTEGER DEFAULT 0', 'status': 'TEXT', 'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_users': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'name': 'TEXT', 'email': 'TEXT', 'password_hash': 'TEXT',
        'role': "TEXT DEFAULT 'user'", 'is_active': 'BOOLEAN DEFAULT true',
        'is_super_admin': 'BOOLEAN DEFAULT false',
        'last_login': 'TIMESTAMP', 'failed_login_count': 'INTEGER DEFAULT 0',
        'last_failed_login': 'TIMESTAMP', 'last_activity': 'TIMESTAMP',
        'created_at': 'TIMESTAMP DEFAULT NOW()', 'updated_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_audit_log': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'user_id': 'INTEGER', 'user_email': 'TEXT', 'action': 'TEXT',
        'details': 'TEXT', 'ip_address': 'TEXT', 'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_user_franchise_access': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'user_id': 'INTEGER', 'franchise_name': 'TEXT',
        'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_password_resets': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'user_id': 'INTEGER', 'token': 'TEXT', 'used': 'BOOLEAN DEFAULT false',
        'expires_at': 'TIMESTAMP', 'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_backup_history': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'filename': 'TEXT', 'file_size': 'BIGINT DEFAULT 0',
        'table_count': 'INTEGER DEFAULT 0', 'row_count': 'INTEGER DEFAULT 0',
        'status': 'TEXT', 'created_by': 'TEXT', 'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_error_log': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'user_id': 'INTEGER', 'user_email': 'TEXT', 'route': 'TEXT', 'method': 'TEXT',
        'error_type': 'TEXT', 'error_message': 'TEXT', 'traceback_text': 'TEXT',
        'ip_address': 'TEXT', 'user_agent': 'TEXT', 'resolved': 'BOOLEAN DEFAULT false',
        'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_claim_cases': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'claim_ref': 'TEXT', 'franchise_name': 'TEXT', 'claimant_name': 'TEXT',
        'policy_number': 'TEXT', 'claim_date': 'DATE', 'claim_amount': 'NUMERIC(18,2) DEFAULT 0',
        'status': "TEXT DEFAULT 'New'", 'priority': "TEXT DEFAULT 'Normal'", 'description': 'TEXT',
        'created_by_id': 'INTEGER', 'created_by_email': 'TEXT', 'assigned_to_email': 'TEXT',
        'archived': 'BOOLEAN DEFAULT false', 'closed_at': 'TIMESTAMP',
        'created_at': 'TIMESTAMP DEFAULT NOW()', 'updated_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_claim_notes': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'claim_id': 'INTEGER', 'user_id': 'INTEGER', 'user_email': 'TEXT',
        'note_text': 'TEXT', 'old_status': 'TEXT', 'new_status': 'TEXT',
        'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_claim_attachments': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'claim_id': 'INTEGER', 'filename': 'TEXT', 'stored_filename': 'TEXT', 'file_path': 'TEXT',
        'content_type': 'TEXT', 'file_size': 'BIGINT DEFAULT 0',
        'document_type': 'TEXT', 'verification_status': "TEXT DEFAULT 'Pending'",
        'verification_details': 'TEXT', 'extracted_text': 'TEXT',
        'has_id_number': 'BOOLEAN DEFAULT false', 'is_certified': 'BOOLEAN DEFAULT false',
        'uploaded_by_id': 'INTEGER', 'uploaded_by_email': 'TEXT',
        'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_claim_document_types': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'document_type': 'TEXT', 'claim_type': "TEXT DEFAULT 'Funeral Claim'",
        'aliases': 'TEXT', 'required': 'BOOLEAN DEFAULT true',
        'is_active': 'BOOLEAN DEFAULT true', 'sort_order': 'INTEGER DEFAULT 100',
        'created_at': 'TIMESTAMP DEFAULT NOW()', 'updated_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_claim_document_rules': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'document_type_id': 'INTEGER', 'rule_key': 'TEXT', 'rule_label': 'TEXT',
        'is_enabled': 'BOOLEAN DEFAULT true', 'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_client_map_points': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'franchise_name': 'TEXT', 'franchise_key': 'TEXT', 'address_key': 'TEXT', 'display_address': 'TEXT',
        'client_name': 'TEXT', 'policy_number': 'TEXT', 'client_count': 'INTEGER DEFAULT 1',
        'lat': 'NUMERIC(12,8)', 'lng': 'NUMERIC(12,8)', 'status': "TEXT DEFAULT 'OK'",
        'updated_at': 'TIMESTAMP DEFAULT NOW()', 'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_cron_log': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'job_name': 'TEXT', 'status': 'TEXT', 'details': 'TEXT',
        'ip_address': 'TEXT', 'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_policy_relation_age_notifications': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'source_row_id': 'BIGINT', 'notification_key': 'TEXT', 'franchise_name': 'TEXT', 'relation': 'TEXT',
        'member_name': 'TEXT', 'id_number': 'TEXT', 'policy_number': 'TEXT',
        'age_years': 'INTEGER', 'age_months': 'INTEGER', 'age_band': 'TEXT',
        'notification_type': 'TEXT', 'message': 'TEXT', 'due_date': 'DATE',
        'status': "TEXT DEFAULT 'active'", 'resolved': 'BOOLEAN DEFAULT false',
        'created_at': 'TIMESTAMP DEFAULT NOW()', 'updated_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_policy_relation_summary': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'franchise_name': 'TEXT', 'relation': 'TEXT', 'member_count': 'INTEGER DEFAULT 0',
        'updated_at': 'TIMESTAMP DEFAULT NOW()',
    },
    'app_policydata_current_members': {
        'id': 'BIGSERIAL PRIMARY KEY',
        'member_key': 'TEXT', 'franchise_name': 'TEXT', 'relation': 'TEXT',
        'id_number': 'TEXT', 'policy_number': 'TEXT', 'member_name': 'TEXT',
        'age': 'INTEGER', 'age_band': 'TEXT', 'client_address_o': 'TEXT',
        'import_month': 'DATE', 'source_row_id': 'BIGINT', 'active': 'BOOLEAN DEFAULT true',
        'raw_data': 'JSONB', 'updated_at': 'TIMESTAMP DEFAULT NOW()', 'created_at': 'TIMESTAMP DEFAULT NOW()',
    },
}


_SCHEMA_CHECKED_ON_STARTUP = False

def _fast_dashboard_reads_enabled():
    return str(os.getenv('FAST_DASHBOARD_READS', '1')).lower() not in {'0', 'false', 'no'}


def _safe_ident(name):
    return '"' + str(name).replace('"', '""') + '"'


def ensure_database_schema(conn=None, force=False):
    """Create/repair all required PostgreSQL tables, columns and indexes.

    Safe to run repeatedly, but expensive on very large PostgreSQL databases.
    Dashboard reads must not re-check every index/column on every page load, so
    this function caches the startup schema check unless force=True.
    """
    global _SCHEMA_CHECKED_ON_STARTUP
    if conn is None and _SCHEMA_CHECKED_ON_STARTUP and not force:
        return {'ok': True, 'message': 'Database schema already checked', 'actions': [], 'errors': []}
    engine = None if conn is not None else get_db_engine()
    if conn is None and engine is None:
        return {'ok': False, 'message': 'PostgreSQL not connected', 'actions': [], 'errors': []}
    actions, errors = [], []

    def run(active_conn, sql, label):
        try:
            active_conn.exec_driver_sql(str(sql).replace('%', '%%'))
            actions.append(label)
        except Exception as exc:
            errors.append(f'{label}: {exc}')

    def work(active_conn):
        for table, columns in REQUIRED_DB_SCHEMA.items():
            id_def = columns.get('id', 'BIGSERIAL PRIMARY KEY')
            run(active_conn, f'CREATE TABLE IF NOT EXISTS {_safe_ident(table)} ("id" {id_def})', f'checked table {table}')
            for column, datatype in columns.items():
                if column == 'id':
                    continue
                run(active_conn, f'ALTER TABLE {_safe_ident(table)} ADD COLUMN IF NOT EXISTS {_safe_ident(column)} {datatype}', f'checked {table}.{column}')
        index_sql = [
            ('idx_policy_monthly_raw_month_franchise', 'CREATE INDEX IF NOT EXISTS idx_policy_monthly_raw_month_franchise ON policy_monthly_raw (import_month, franchise_name)'),
            ('idx_policydata_detail_raw_source_row_key', 'CREATE UNIQUE INDEX IF NOT EXISTS idx_policydata_detail_raw_source_row_key ON policydata_detail_raw (source_row_key) WHERE source_row_key IS NOT NULL'),
            ('idx_policydata_detail_raw_franchise_mem', 'CREATE INDEX IF NOT EXISTS idx_policydata_detail_raw_franchise_mem ON policydata_detail_raw (franchise_name, is_mem)'),
            ('idx_policydata_detail_raw_client_address_o', 'CREATE INDEX IF NOT EXISTS idx_policydata_detail_raw_client_address_o ON policydata_detail_raw (client_address_o) WHERE client_address_o IS NOT NULL'),
            ('idx_policydata_detail_raw_relation', 'CREATE INDEX IF NOT EXISTS idx_policydata_detail_raw_relation ON policydata_detail_raw (UPPER(TRIM(relation)))'),
            ('idx_policydata_detail_raw_id_number_f', 'CREATE INDEX IF NOT EXISTS idx_policydata_detail_raw_id_number_f ON policydata_detail_raw (id_number_f) WHERE id_number_f IS NOT NULL'),
            ('idx_claims_monthly_raw_month_key', 'CREATE INDEX IF NOT EXISTS idx_claims_monthly_raw_month_key ON claims_monthly_raw (claim_month, claim_key)'),
            ('idx_franchise_summary_month_franchise', 'CREATE INDEX IF NOT EXISTS idx_franchise_summary_month_franchise ON franchise_monthly_summary (report_month, franchise_name)'),
            ('uq_franchise_monthly_summary_month_franchise', 'CREATE UNIQUE INDEX IF NOT EXISTS uq_franchise_monthly_summary_month_franchise ON franchise_monthly_summary (franchise_name, report_month)'),
            ('uq_franchise_mapping_pg_source', 'CREATE UNIQUE INDEX IF NOT EXISTS uq_franchise_mapping_pg_source ON franchise_mapping_pg (LOWER(TRIM(source_name)))'),
            ('uq_app_users_email', 'CREATE UNIQUE INDEX IF NOT EXISTS uq_app_users_email ON app_users (LOWER(TRIM(email)))'),
            ('idx_user_franchise_access_user', 'CREATE INDEX IF NOT EXISTS idx_user_franchise_access_user ON app_user_franchise_access (user_id)'),
            ('uq_user_franchise_access', 'CREATE UNIQUE INDEX IF NOT EXISTS uq_user_franchise_access ON app_user_franchise_access (user_id, LOWER(TRIM(franchise_name)))'),
            ('idx_app_error_log_created_at', 'CREATE INDEX IF NOT EXISTS idx_app_error_log_created_at ON app_error_log (created_at DESC)'),
            ('idx_app_users_last_activity', 'CREATE INDEX IF NOT EXISTS idx_app_users_last_activity ON app_users (last_activity)'),
            ('uq_app_password_resets_token', 'CREATE UNIQUE INDEX IF NOT EXISTS uq_app_password_resets_token ON app_password_resets (token)'),
            ('idx_app_password_resets_user', 'CREATE INDEX IF NOT EXISTS idx_app_password_resets_user ON app_password_resets (user_id)'),
            ('idx_app_backup_history_created_at', 'CREATE INDEX IF NOT EXISTS idx_app_backup_history_created_at ON app_backup_history (created_at DESC)'),
            ('idx_app_cron_log_created_at', 'CREATE INDEX IF NOT EXISTS idx_app_cron_log_created_at ON app_cron_log (created_at DESC)'),
            ('idx_app_client_map_points_franchise', 'CREATE INDEX IF NOT EXISTS idx_app_client_map_points_franchise ON app_client_map_points (LOWER(TRIM(franchise_name)))'),
            ('idx_app_client_map_points_franchise_key', 'CREATE INDEX IF NOT EXISTS idx_app_client_map_points_franchise_key ON app_client_map_points (franchise_key)'),
            ('idx_app_client_map_points_latlng', 'CREATE INDEX IF NOT EXISTS idx_app_client_map_points_latlng ON app_client_map_points (lat, lng)'),
            ('uq_app_client_map_points_address', 'CREATE UNIQUE INDEX IF NOT EXISTS uq_app_client_map_points_address ON app_client_map_points (address_key)'),

            ('idx_app_claim_cases_status', 'CREATE INDEX IF NOT EXISTS idx_app_claim_cases_status ON app_claim_cases (status)'),
            ('idx_app_claim_cases_franchise', 'CREATE INDEX IF NOT EXISTS idx_app_claim_cases_franchise ON app_claim_cases (LOWER(TRIM(franchise_name)))'),
            ('idx_app_claim_cases_created_at', 'CREATE INDEX IF NOT EXISTS idx_app_claim_cases_created_at ON app_claim_cases (created_at DESC)'),
            ('idx_app_claim_notes_claim', 'CREATE INDEX IF NOT EXISTS idx_app_claim_notes_claim ON app_claim_notes (claim_id, created_at DESC)'),
            ('idx_app_claim_cases_archived', 'CREATE INDEX IF NOT EXISTS idx_app_claim_cases_archived ON app_claim_cases (archived)'),
            ('idx_app_claim_attachments_claim', 'CREATE INDEX IF NOT EXISTS idx_app_claim_attachments_claim ON app_claim_attachments (claim_id, created_at DESC)'),
            ('idx_claim_document_types_active', 'CREATE INDEX IF NOT EXISTS idx_claim_document_types_active ON app_claim_document_types (is_active, sort_order, document_type)'),
            ('idx_claim_document_rules_type', 'CREATE INDEX IF NOT EXISTS idx_claim_document_rules_type ON app_claim_document_rules (document_type_id, rule_key)'),
            ('idx_policy_age_notifications_active', 'CREATE INDEX IF NOT EXISTS idx_policy_age_notifications_active ON app_policy_relation_age_notifications (status, due_date, franchise_name)'),
            ('uq_policy_age_notifications_source_type', 'CREATE UNIQUE INDEX IF NOT EXISTS uq_policy_age_notifications_source_type ON app_policy_relation_age_notifications (source_row_id, notification_type) WHERE source_row_id IS NOT NULL'),
            ('uq_policy_age_notifications_key', 'CREATE UNIQUE INDEX IF NOT EXISTS uq_policy_age_notifications_key ON app_policy_relation_age_notifications (notification_key) WHERE notification_key IS NOT NULL'),
            ('idx_policy_relation_summary_franchise', 'CREATE INDEX IF NOT EXISTS idx_policy_relation_summary_franchise ON app_policy_relation_summary (franchise_name, relation)'),
            ('uq_policy_relation_summary_franchise_relation', 'CREATE UNIQUE INDEX IF NOT EXISTS uq_policy_relation_summary_franchise_relation ON app_policy_relation_summary (LOWER(TRIM(franchise_name)), UPPER(TRIM(relation)))'),
            ('idx_current_members_relation', 'CREATE INDEX IF NOT EXISTS idx_current_members_relation ON app_policydata_current_members (active, relation, franchise_name)'),
            ('idx_current_members_key', 'CREATE UNIQUE INDEX IF NOT EXISTS idx_current_members_key ON app_policydata_current_members (member_key)'),
            ('idx_policydata_raw_relation_franchise', 'CREATE INDEX IF NOT EXISTS idx_policydata_raw_relation_franchise ON policydata_detail_raw (relation, franchise_name) WHERE relation IS NOT NULL'),
        ]
        for label, sql in index_sql:
            run(active_conn, sql, f'checked index {label}')

    if conn is not None:
        work(conn)
    else:
        with engine.begin() as active_conn:
            work(active_conn)
    if conn is None and not errors:
        _SCHEMA_CHECKED_ON_STARTUP = True
    return {'ok': not errors, 'message': 'Database schema checked', 'actions': actions, 'errors': errors}


def database_health_report():
    engine = get_db_engine()
    if engine is None:
        return {'connected': False, 'message': DB_STATUS.get('message', 'PostgreSQL not connected'), 'tables': [], 'missing': []}
    ensure_database_schema()
    tables, missing = [], []
    with engine.begin() as conn:
        for table, columns in REQUIRED_DB_SCHEMA.items():
            exists = conn.execute(text("""
                SELECT EXISTS (
                    SELECT 1 FROM information_schema.tables
                    WHERE table_schema = 'public' AND table_name = :table
                )
            """), {'table': table}).scalar()
            row_count = 0
            existing_cols = set()
            if exists:
                try:
                    row_count = conn.execute(text(f'SELECT COUNT(*) FROM {_safe_ident(table)}')).scalar() or 0
                except Exception:
                    row_count = 0
                existing_cols = set(r[0] for r in conn.execute(text("""
                    SELECT column_name FROM information_schema.columns
                    WHERE table_schema = 'public' AND table_name = :table
                """), {'table': table}).fetchall())
            tables.append({'table': table, 'exists': bool(exists), 'rows': int(row_count)})
            for column, datatype in columns.items():
                if column not in existing_cols:
                    missing.append({'table': table, 'column': column, 'type': datatype})
    return {'connected': True, 'message': DB_STATUS.get('message', 'Connected to PostgreSQL'), 'tables': tables, 'missing': missing}

def _db_date(v):
    try:
        return pd.to_datetime(v).to_pydatetime().date()
    except Exception:
        return None


def _safe_num(v):
    try:
        return float(clean_money(v))
    except Exception:
        return 0.0



def display_source_filename(path_or_name):
    """Return the user-facing uploaded filename without the UUID prefix used on disk."""
    name = os.path.basename(str(path_or_name or ''))
    return re.sub(r'^[0-9a-fA-F-]{36}_', '', name)


def save_policy_detail_to_postgres(df, source_file=''):
    """Persist row-level PolicyData details to PostgreSQL for DBeaver audit/reconciliation.

    Monthly dashboard totals are stored in policy_monthly_raw. This table keeps the
    underlying PolicyData rows allocated to Column A franchise names so every import
    can be audited later in DBeaver.
    """
    engine = get_db_engine()
    if engine is None or df is None or df.empty:
        return False
    work = df.copy()
    work['import_month'] = pd.to_datetime(work['import_month'], errors='coerce').dt.to_period('M').dt.to_timestamp()
    work = work.dropna(subset=['import_month'])
    if work.empty:
        return False
    src = source_file or str(work.get('source_file', pd.Series([''])).iloc[0] or '')
    months = sorted({_db_date(x) for x in work['import_month'].dropna().unique() if _db_date(x) is not None})
    rows = []
    for _, r in work.iterrows():
        rows.append({
            'source_file': src or str(r.get('source_file','')),
            'import_month': _db_date(r.get('import_month')),
            'row_number': int(_safe_num(r.get('row_number',0))),
            'source_row_key': str(r.get('source_row_key') or f"{src or str(r.get('source_file',''))}|{_db_date(r.get('import_month'))}|{int(_safe_num(r.get('row_number',0)))}"),
            'client_address_o': _clean_map_value(r.get('client_address_o','')),
            'id_number_f': _clean_id_number(r.get('id_number_f') or _first_raw_value(r.get('raw_data', {}), ['id number', 'id no', 'identity number', 'sa id', 'id']) or _raw_value_by_excel_column(r.get('raw_data', {}), 'F')),
            'franchise_name': str(r.get('franchise','')).strip(),
            'relation': str(r.get('relation','')).strip(),
            'is_mem': bool(r.get('is_mem', False)),
            'retail_premium': _safe_num(r.get('retail_premium',0)),
            'original_risk_premium': _safe_num(r.get('original_risk_premium',0)),
            'mpia': _safe_num(r.get('mpia',0)),
            'single_premium': _safe_num(r.get('single_premium',0)),
            'r1_policy_fee': _safe_num(r.get('r1_policy_fee',0)),
            'adv_fund_2_1_fee': _safe_num(r.get('adv_fund_2_1_fee',0)),
            'risk_after_r1': _safe_num(r.get('risk_after_r1',0)),
            'new_risk_premium': _safe_num(r.get('new_risk_premium',0)),
            'raw_data': json.dumps(r.get('raw_data', {}), default=str),
        })
    rows = [r for r in rows if r['franchise_name'] and r['import_month']]
    if not rows:
        return False
    with engine.begin() as conn:
        for m in months:
            conn.execute(text('DELETE FROM policydata_detail_raw WHERE import_month = :m AND source_file = :src'), {'m': m, 'src': src})
        insert_sql = text("""
            INSERT INTO policydata_detail_raw (
                source_file, import_month, row_number, source_row_key, client_address_o, id_number_f, franchise_name, relation, is_mem,
                retail_premium, original_risk_premium, mpia, single_premium,
                r1_policy_fee, adv_fund_2_1_fee, risk_after_r1, new_risk_premium, raw_data
            ) VALUES (
                :source_file, :import_month, :row_number, :source_row_key, :client_address_o, :id_number_f, :franchise_name, :relation, :is_mem,
                :retail_premium, :original_risk_premium, :mpia, :single_premium,
                :r1_policy_fee, :adv_fund_2_1_fee, :risk_after_r1, :new_risk_premium, CAST(:raw_data AS jsonb)
            )
            ON CONFLICT (source_row_key) DO UPDATE SET
                source_file = EXCLUDED.source_file,
                import_month = EXCLUDED.import_month,
                row_number = EXCLUDED.row_number,
                client_address_o = EXCLUDED.client_address_o,
                id_number_f = EXCLUDED.id_number_f,
                franchise_name = EXCLUDED.franchise_name,
                relation = EXCLUDED.relation,
                is_mem = EXCLUDED.is_mem,
                retail_premium = EXCLUDED.retail_premium,
                original_risk_premium = EXCLUDED.original_risk_premium,
                mpia = EXCLUDED.mpia,
                single_premium = EXCLUDED.single_premium,
                r1_policy_fee = EXCLUDED.r1_policy_fee,
                adv_fund_2_1_fee = EXCLUDED.adv_fund_2_1_fee,
                risk_after_r1 = EXCLUDED.risk_after_r1,
                new_risk_premium = EXCLUDED.new_risk_premium,
                raw_data = EXCLUDED.raw_data
        """)
        for i in range(0, len(rows), 2000):
            conn.execute(insert_sql, rows[i:i+2000])
    return True


def save_policy_raw_to_postgres(df, source_file=''):
    """Persist monthly policy summaries to PostgreSQL. Same month re-import replaces that month."""
    engine = get_db_engine()
    if engine is None or df is None or df.empty:
        return False
    work = df.copy()
    work['month'] = pd.to_datetime(work['month'], errors='coerce').dt.to_period('M').dt.to_timestamp()
    work = work.dropna(subset=['month'])
    if work.empty:
        return False
    months = sorted({_db_date(x) for x in work['month'].dropna().unique() if _db_date(x) is not None})
    rows = []
    for _, r in work.iterrows():
        rows.append({
            'franchise_name': str(r.get('franchise','')).strip(),
            'import_month': _db_date(r.get('month')),
            'retail_premium': _safe_num(r.get('retail_premium',0)),
            'risk_premium': _safe_num(r.get('risk_premium',0)),
            'claims': _safe_num(r.get('claims',0)),
            'claim_count': _safe_num(r.get('claim_count',0)),
            'claim_paid_franchise': _safe_num(r.get('claim_paid_franchise',0)),
            'claim_paid_client': _safe_num(r.get('claim_paid_client',0)),
            'repudiated_pending': _safe_num(r.get('repudiated_pending',0)),
            'grand_total_claims': _safe_num(r.get('grand_total_claims',0)),
            'policy_qty': _safe_num(r.get('policy_qty',0)),
            'original_risk_premium': _safe_num(r.get('original_risk_premium',0)),
            'r1_policy_fee': _safe_num(r.get('r1_policy_fee_imported', r.get('r1_policy_fee',0))),
            'underwriter_2_1_fee': _safe_num(r.get('underwriter_2_1_fee',0)),
            'risk_after_r1': _safe_num(r.get('risk_after_r1',0)),
            'single_monthly_premium_total': _safe_num(r.get('single_monthly_premium_total',0)),
            'current_scenario': str(r.get('current_scenario','100% Claim Ratio') or '100% Claim Ratio'),
            'source_file': source_file or str(r.get('source_file','')),
        })
    rows = [r for r in rows if r['franchise_name'] and r['import_month']]
    if not rows:
        return False
    with engine.begin() as conn:
        for m in months:
            conn.execute(text('DELETE FROM policy_monthly_raw WHERE import_month = :m'), {'m': m})
        conn.execute(text("""
            INSERT INTO policy_monthly_raw (
                franchise_name, import_month, retail_premium, risk_premium, claims, claim_count,
                claim_paid_franchise, claim_paid_client, repudiated_pending, grand_total_claims,
                policy_qty, original_risk_premium, r1_policy_fee, underwriter_2_1_fee,
                risk_after_r1, single_monthly_premium_total, current_scenario, source_file
            ) VALUES (
                :franchise_name, :import_month, :retail_premium, :risk_premium, :claims, :claim_count,
                :claim_paid_franchise, :claim_paid_client, :repudiated_pending, :grand_total_claims,
                :policy_qty, :original_risk_premium, :r1_policy_fee, :underwriter_2_1_fee,
                :risk_after_r1, :single_monthly_premium_total, :current_scenario, :source_file
            )
        """), rows)
        conn.execute(text("""
            INSERT INTO import_history (import_type, source_file, imported_months, row_count, status)
            VALUES (:import_type, :source_file, :imported_months, :row_count, :status)
        """), {
            'import_type': 'policy',
            'source_file': source_file,
            'imported_months': [m.isoformat() for m in months],
            'row_count': len(rows),
            'status': 'success'
        })
    return True


def save_claims_raw_to_postgres(df, source_file=''):
    engine = get_db_engine()
    if engine is None or df is None or df.empty:
        return False
    work = df.copy()
    work['month'] = pd.to_datetime(work['month'], errors='coerce').dt.to_period('M').dt.to_timestamp()
    work = work.dropna(subset=['month'])
    if 'claim_key' not in work.columns:
        work['claim_key'] = work['franchise'].map(franchise_match_key).map(apply_franchise_alias)
    else:
        work['claim_key'] = work['claim_key'].map(apply_franchise_alias)
    months = sorted({_db_date(x) for x in work['month'].dropna().unique() if _db_date(x) is not None})
    rows = []
    for _, r in work.iterrows():
        rows.append({
            'claim_key': str(r.get('claim_key','')).strip(),
            'claims_franchise_name': str(r.get('franchise','')).strip(),
            'claim_month': _db_date(r.get('month')),
            'claims_amount': _safe_num(r.get('claims',0)),
            'claim_count': _safe_num(r.get('claim_count',0)),
            'claim_paid_franchise': _safe_num(r.get('claim_paid_franchise',0)),
            'claim_paid_client': _safe_num(r.get('claim_paid_client',0)),
            'repudiated_pending': _safe_num(r.get('repudiated_pending',0)),
            'grand_total_claims': _safe_num(r.get('grand_total_claims',0)),
            'source_file': source_file,
        })
    rows = [r for r in rows if r['claims_franchise_name'] and r['claim_month']]
    if not rows:
        return False
    with engine.begin() as conn:
        for m in months:
            conn.execute(text('DELETE FROM claims_monthly_raw WHERE claim_month = :m'), {'m': m})
        conn.execute(text("""
            INSERT INTO claims_monthly_raw (
                claim_key, claims_franchise_name, claim_month, claims_amount, claim_count,
                claim_paid_franchise, claim_paid_client, repudiated_pending, grand_total_claims, source_file
            ) VALUES (
                :claim_key, :claims_franchise_name, :claim_month, :claims_amount, :claim_count,
                :claim_paid_franchise, :claim_paid_client, :repudiated_pending, :grand_total_claims, :source_file
            )
        """), rows)
        conn.execute(text("""
            INSERT INTO import_history (import_type, source_file, imported_months, row_count, status)
            VALUES (:import_type, :source_file, :imported_months, :row_count, :status)
        """), {
            'import_type': 'claims',
            'source_file': source_file,
            'imported_months': [m.isoformat() for m in months],
            'row_count': len(rows),
            'status': 'success'
        })
    return True


def load_raw_from_postgres():
    """Load persisted policy and claims summaries from PostgreSQL for dashboard startup/reload."""
    global LAST_CLAIMS_DF, LAST_POLICY_IMPORT_SUMMARY
    engine = get_db_engine()
    if engine is None:
        return pd.DataFrame()
    try:
        policy = pd.read_sql("""
            SELECT franchise_name AS franchise, import_month AS month, retail_premium, risk_premium,
                   claims, claim_count, claim_paid_franchise, claim_paid_client, repudiated_pending,
                   grand_total_claims, policy_qty, original_risk_premium,
                   r1_policy_fee AS r1_policy_fee_imported, underwriter_2_1_fee, risk_after_r1,
                   single_monthly_premium_total, current_scenario
            FROM policy_monthly_raw
            ORDER BY franchise_name, import_month
        """, engine)
        if policy.empty:
            return pd.DataFrame()
        policy['month'] = pd.to_datetime(policy['month'], errors='coerce')
        claims = pd.read_sql("""
            SELECT claims_franchise_name AS franchise, claim_key, claim_month AS month,
                   claims_amount AS claims, claim_count, claim_paid_franchise, claim_paid_client,
                   repudiated_pending, grand_total_claims
            FROM claims_monthly_raw
            ORDER BY claims_franchise_name, claim_month
        """, engine)
        if not claims.empty:
            claims['month'] = pd.to_datetime(claims['month'], errors='coerce')
            LAST_CLAIMS_DF = claims.copy()
            raw = merge_claims_into_raw(policy, claims)
        else:
            raw = policy
        LAST_POLICY_IMPORT_SUMMARY = get_import_history_summary()
        return raw
    except Exception as exc:
        print(f'Could not load PostgreSQL data: {exc}')
        return pd.DataFrame()


def get_import_history_summary():
    engine = get_db_engine()
    if engine is None:
        return {}
    try:
        with engine.begin() as conn:
            row = conn.execute(text("""
                SELECT COUNT(*) AS imports,
                       COALESCE(SUM(row_count),0) AS rows_written,
                       MAX(created_at) AS last_import
                FROM import_history
            """)).mappings().first()
        return dict(row) if row else {}
    except Exception:
        return {}


def reload_dashboard_from_postgres():
    global LAST_RESULT
    raw = load_raw_from_postgres()
    if raw is not None and not raw.empty:
        monthly, periods, portfolio = analyse(raw, LAST_RESULT.get('rates', DEFAULT_RATES.copy()), LAST_RESULT.get('book_rates', DEFAULT_BOOK_VALUE.copy()))
        LAST_RESULT = {'raw': raw, 'monthly': monthly, 'periods': periods, 'portfolio': portfolio, 'rates': LAST_RESULT.get('rates', DEFAULT_RATES.copy()), 'book_rates': LAST_RESULT.get('book_rates', DEFAULT_BOOK_VALUE.copy())}
        return True
    return False


def load_franchise_config():
    """Read saved franchise inclusion/grouping rules."""
    default = {'excluded': [], 'groups': {}, 'use_groups': True, 'claims_aliases': {}}
    try:
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, 'r', encoding='utf-8') as fh:
                data = json.load(fh)
            if isinstance(data, dict):
                default.update({
                    'excluded': list(data.get('excluded', [])),
                    'groups': dict(data.get('groups', {})),
                    'use_groups': bool(data.get('use_groups', True)),
                    'claims_aliases': dict(data.get('claims_aliases', {})),
                })
    except Exception:
        pass
    return default


def save_franchise_config(config):
    with open(CONFIG_FILE, 'w', encoding='utf-8') as fh:
        json.dump(config, fh, indent=2, ensure_ascii=False)


def average_claim_ratio_for(df):
    """Claim ratio for a selected period: total claims divided by total risk premium."""
    if df is None or df.empty:
        return 0.0
    risk = pd.to_numeric(df.get('Risk Premium', pd.Series(dtype=float)), errors='coerce').fillna(0).sum()
    claims = pd.to_numeric(df.get('Claims', pd.Series(dtype=float)), errors='coerce').fillna(0).sum()
    return float((claims / risk * 100) if risk else 0.0)


def _dashboard_month_filter(monthly, period_view='six_months'):
    """Return only the latest policy month, latest 6 policy months, or latest 12 policy months.

    Claims files can contain months that are newer than the available PolicyData files.
    The dashboard must not switch to a claims-only month, otherwise premiums become R0
    and the claim ratio becomes 0%.  We therefore choose the period from months that
    have policy premium data, then include the matching claims for those same months.
    """
    if monthly is None or monthly.empty:
        return pd.DataFrame(), ''
    base = monthly.copy()
    if 'Month' not in base.columns:
        return base, 'All imported data'
    base['_month_dt'] = pd.to_datetime(base['Month'], errors='coerce')
    base = base.dropna(subset=['_month_dt'])
    if base.empty:
        return pd.DataFrame(), ''

    premium_cols = [c for c in ['Retail Premium', 'Risk Premium', 'Original Risk Premium'] if c in base.columns]
    if premium_cols:
        premium_total = sum(pd.to_numeric(base[c], errors='coerce').fillna(0) for c in premium_cols)
        policy_months = sorted(base.loc[premium_total > 0, '_month_dt'].dt.to_period('M').unique())
    else:
        policy_months = []
    month_values = policy_months or sorted(base['_month_dt'].dt.to_period('M').unique())

    if period_view == 'month':
        selected_months = month_values[-1:]
    elif period_view == 'year':
        selected_months = month_values[-12:]
    else:
        selected_months = month_values[-6:]
    filtered = base[base['_month_dt'].dt.to_period('M').isin(selected_months)].copy()
    filtered.drop(columns=['_month_dt'], inplace=True, errors='ignore')
    if not selected_months:
        label = ''
    elif len(selected_months) == 1:
        label = selected_months[0].strftime('%b %Y')
    else:
        label = f"{selected_months[0].strftime('%b %Y')} - {selected_months[-1].strftime('%b %Y')}"
    return filtered, label


def build_portfolio_for_period(monthly, periods, period_view='six_months'):
    """Dashboard KPI totals based on the selected dashboard period. All KPI cards use the same filtered data."""
    if monthly is None or monthly.empty:
        return {'total_franchises': 0, 'average_claim_ratio': 0.0, 'overall_claim_ratio': 0.0, 'dashboard_period_label': ''}

    base, period_label = _dashboard_month_filter(monthly, period_view)
    if base is None or base.empty:
        return {'total_franchises': 0, 'average_claim_ratio': 0.0, 'overall_claim_ratio': 0.0, 'dashboard_period_view': period_view, 'dashboard_period_label': period_label}

    def colsum(name):
        return float(pd.to_numeric(base.get(name, pd.Series(dtype=float)), errors='coerce').fillna(0).sum())

    total_claims = colsum('Claims')
    total_risk = colsum('Risk Premium')
    claim_ratio = float((total_claims / total_risk * 100) if total_risk else 0.0)

    p = {
        'total_franchises': int(base['Franchise'].nunique()) if 'Franchise' in base.columns else 0,
        'total_retail': colsum('Retail Premium'),
        'total_risk': total_risk,
        'total_original_risk': colsum('Original Risk Premium'),
        'total_underwriter_2_1_fee': colsum('Underwriter 2.1% Fee'),
        'total_claims': total_claims,
        'total_claim_count': colsum('Claim Count'),
        'average_claim_ratio': claim_ratio,
        'overall_claim_ratio': claim_ratio,
        'total_commission': colsum('Total Commission'),
        'total_paid_commissions': colsum('Total Paid Commissions'),
        'total_brightrock_commission': colsum('BrightRock Amount'),
        'total_mkhulu_commission': colsum('Inkulu Amount'),
        'total_inkulu_commission': colsum('Inkulu Amount'),
        'total_mff_commission': colsum('MFF Amount'),
        'total_r1_policy_fee': colsum('R1 Policy Fee'),
        'total_policy_qty': colsum('Policy Qty'),
        'total_brightrock_month_total': colsum('BrightRock Month Total') or colsum('Franchise Money'),
        'total_book_value': colsum('Total Book Value'),
        'total_mff_book_value': colsum('MFF Book Value 2.5%'),
        'total_franchise_book_value': colsum('Franchise Book Value 2.5%'),
        'dashboard_period_view': period_view,
        'dashboard_period_label': period_label,
    }

    # Recommendation counts are calculated per franchise over the same selected period.
    if 'Franchise' in base.columns:
        grouped = base.groupby('Franchise', as_index=False).agg({'Claims': 'sum', 'Risk Premium': 'sum'})
        grouped['Claim Ratio'] = grouped.apply(lambda r: (r['Claims'] / r['Risk Premium'] * 100) if r['Risk Premium'] else 0, axis=1)
        rec = grouped['Claim Ratio'].apply(lambda r: 'Move to BrightRock' if r < 50 else ('Can move to BrightRock' if r < 70 else 'High Risk - stay on 100%'))
    else:
        rec = pd.Series(dtype=str)
    p['move_count'] = int((rec == 'Move to BrightRock').sum()) if len(rec) else 0
    p['can_move_count'] = int((rec == 'Can move to BrightRock').sum()) if len(rec) else 0
    p['high_risk_count'] = int(rec.astype(str).str.contains('High Risk').sum()) if len(rec) else 0
    return p



def yearly_view_from_monthly(monthly):
    """Aggregate monthly franchise records into franchise/year rows for reporting."""
    if monthly is None or monthly.empty:
        return pd.DataFrame()
    base = monthly.copy()
    base['Year'] = pd.to_datetime(base['Month'], errors='coerce').dt.year.astype(str)
    numeric_cols = [c for c in base.columns if pd.api.types.is_numeric_dtype(base[c])]
    if not numeric_cols:
        return pd.DataFrame()
    out = base.groupby(['Franchise', 'Year'], as_index=False)[numeric_cols].sum()
    out['Period'] = out['Year']
    if 'Risk Premium' in out.columns and 'Claims' in out.columns:
        out['Claim Ratio'] = out.apply(lambda r: (r.get('Claims', 0) / r.get('Risk Premium', 0) * 100) if r.get('Risk Premium', 0) else 0, axis=1)
        out['Average Claim Ratio'] = out['Claim Ratio']
        out['Claim Ratio Label'] = out['Claim Ratio'].apply(status_label) if 'status_label' in globals() else ''
        out['Recommendation'] = out['Claim Ratio'].apply(lambda r: 'Move to BrightRock' if r < 50 else ('Can move to BrightRock' if r < 70 else 'High Risk - stay on 100%'))
    return out


def rebuild_periods_and_portfolio(monthly):
    """Rebuild period and portfolio tables after exclusions/grouping are applied.

    The old display used generic Period 1/2/3 labels.  Use real month ranges instead
    so every analysis/export shows the actual months being calculated.
    """
    if monthly is None or monthly.empty:
        return pd.DataFrame(), {'total_franchises': 0}
    work = monthly.copy().sort_values(['Franchise', 'Month'])
    work['Month'] = pd.to_datetime(work['Month'], errors='coerce')
    work = work.dropna(subset=['Month'])
    work['Period Index'] = work.groupby('Franchise').cumcount() // 6 + 1
    work['Period Start'] = work.groupby(['Franchise','Period Index'])['Month'].transform('min')
    work['Period End'] = work.groupby(['Franchise','Period Index'])['Month'].transform('max')
    work['Period'] = work.apply(lambda r: r['Period Start'].strftime('%b %Y') if r['Period Start'].to_period('M') == r['Period End'].to_period('M') else f"{r['Period Start'].strftime('%b %Y')} - {r['Period End'].strftime('%b %Y')}", axis=1)
    work['Claim Ratio'] = work.apply(lambda r: (r['Claims'] / r['Risk Premium'] * 100) if r.get('Risk Premium', 0) else 0, axis=1)
    work['Claim Ratio Status'] = work['Claim Ratio'].apply(status_colour)
    work['Claim Ratio Label'] = work['Claim Ratio'].apply(status_label)
    work['Recommendation'] = work['Claim Ratio'].apply(lambda r: 'Move to BrightRock' if r < 50 else ('Can move to BrightRock' if r < 70 else 'High Risk - stay on 100%'))
    work['BrightRock Running Balance'] = work.groupby('Franchise')['BrightRock Month Total'].cumsum() if 'BrightRock Month Total' in work.columns else 0

    numeric_cols = [c for c in work.columns if c not in {'Franchise', 'Month', 'Period', 'Period Start', 'Period End', 'Claim Ratio Status', 'Claim Ratio Label', 'Recommendation', 'Current Scenario', '100% Move Status', '100% Move Traffic', '100% Traffic Light', '100% Traffic Result', '100% Traffic Reason', 'BrightRock Traffic Light', 'BrightRock Money Result'} and pd.api.types.is_numeric_dtype(work[c])]
    agg = {c: 'sum' for c in numeric_cols}
    if 'BrightRock Running Balance' in agg:
        agg['BrightRock Running Balance'] = 'last'
    periods = work.groupby(['Franchise', 'Period'], as_index=False).agg(agg)
    periods['Average Claim'] = periods.apply(lambda r: (r['Claims'] / r['Claim Count']) if r.get('Claim Count', 0) else 0, axis=1)
    periods['Weighted Claim Ratio'] = periods.apply(lambda r: (r['Claims'] / r['Risk Premium'] * 100) if r.get('Risk Premium', 0) else 0, axis=1)
    periods['Claim Ratio Status'] = periods['Weighted Claim Ratio'].apply(status_colour)
    periods['Claim Ratio Label'] = periods['Weighted Claim Ratio'].apply(status_label)
    periods['Recommendation'] = periods.apply(lambda r: 'Move to BrightRock' if bool(r.get('BrightRock Eligible', False)) else ('Can move to BrightRock - needs 5 months below 75%' if float(r.get('Weighted Claim Ratio', 0) or 0) < 75 else 'High Risk - stay on 100%'), axis=1)
    periods['100% Traffic Light'] = periods.apply(lambda r: scenario_100_traffic(r.get('Claims',0), r.get('Risk Premium',0), r.get('Total Commission',0)), axis=1)
    periods['BrightRock Traffic Light'] = periods.get('BrightRock Month Total', pd.Series([0]*len(periods))).apply(brightrock_money_traffic)
    for metric in ['Retail Premium', 'Claims', 'Total Book Value', 'BrightRock Month Total', 'Scenario 1 Value', 'Scenario 2 Value']:
        if metric in periods.columns:
            total_col = periods.groupby('Period')[metric].transform('sum')
            periods[metric + ' Contribution %'] = periods[metric] / total_col * 100
            periods.loc[total_col == 0, metric + ' Contribution %'] = 0
    portfolio = {
        'total_franchises': int(work['Franchise'].nunique()),
        'total_retail': float(work.get('Retail Premium', pd.Series(dtype=float)).sum()),
        'total_risk': float(work.get('Risk Premium', pd.Series(dtype=float)).sum()),
        'total_original_risk': float(work.get('Original Risk Premium', pd.Series(dtype=float)).sum()),
        'total_underwriter_2_1_fee': float(work.get('Underwriter 2.1% Fee', pd.Series(dtype=float)).sum()),
        'total_claims': float(work.get('Claims', pd.Series(dtype=float)).sum()),
        'overall_claim_ratio': average_claim_ratio_for(periods),
        'average_claim_ratio': average_claim_ratio_for(periods),
        'total_commission': float(work.get('Total Commission', pd.Series(dtype=float)).sum()),
        'total_paid_commissions': float(work.get('Total Paid Commissions', pd.Series(dtype=float)).sum()),
        'total_brightrock_commission': float(work.get('BrightRock Amount', pd.Series(dtype=float)).sum()),
        'total_mkhulu_commission': float(work.get('Inkulu Amount', pd.Series(dtype=float)).sum()),
        'total_inkulu_commission': float(work.get('Inkulu Amount', pd.Series(dtype=float)).sum()),
        'total_mff_commission': float(work.get('MFF Amount', pd.Series(dtype=float)).sum()),
        'total_r1_policy_fee': float(work.get('R1 Policy Fee', pd.Series(dtype=float)).sum()),
        'total_policy_qty': float(work.get('Policy Qty', pd.Series(dtype=float)).sum()),
        'total_brightrock_month_total': float(work.get('BrightRock Month Total', pd.Series(dtype=float)).sum()),
        'total_book_value': float(work.get('Total Book Value', pd.Series(dtype=float)).sum()),
        'move_count': int((periods.get('Recommendation', pd.Series(dtype=str)) == 'Move to BrightRock').sum()) if not periods.empty else 0,
        'can_move_count': int((periods.get('Recommendation', pd.Series(dtype=str)) == 'Can move to BrightRock').sum()) if not periods.empty else 0,
        'high_risk_count': int((periods.get('Recommendation', pd.Series(dtype=str)).astype(str).str.contains('High Risk')).sum()) if not periods.empty else 0,
    }
    return periods, portfolio


def apply_franchise_config(monthly, config=None):
    """Apply include/exclude and grouping rules to the monthly dataset."""
    if config is None:
        config = load_franchise_config()
    if monthly is None or monthly.empty:
        return pd.DataFrame(), pd.DataFrame(), {'total_franchises': 0}
    work = monthly.copy()
    excluded = set(config.get('excluded', []))
    if excluded:
        work = work[~work['Franchise'].isin(excluded)].copy()
    if config.get('use_groups', True):
        groups = {k: v for k, v in config.get('groups', {}).items() if str(v).strip()}
        if groups:
            work['Franchise'] = work['Franchise'].map(lambda x: groups.get(x, x))
            numeric_cols = [c for c in work.columns if c not in {'Franchise', 'Month', 'Period', 'Claim Ratio Status', 'Claim Ratio Label', 'Recommendation', 'Current Scenario', '100% Move Status', '100% Move Traffic', '100% Traffic Light', '100% Traffic Result', '100% Traffic Reason', 'BrightRock Traffic Light', 'BrightRock Money Result'} and pd.api.types.is_numeric_dtype(work[c])]
            work = work.groupby(['Franchise', 'Month'], as_index=False)[numeric_cols].sum()
    periods, portfolio = rebuild_periods_and_portfolio(work)
    return work, periods, portfolio


COLUMN_ALIASES = {
    'franchise': ['franchise', 'franchise name', 'name', 'branch'],
    'month': ['month', 'date', 'period', 'report month'],
    'retail_premium': ['retail premium', 'retailpremium', 'retail_premium'],
        'risk_premium': ['risk premium', 'riskpremium', 'risk_premium'],
    'claims': ['claims', 'claim', 'total claims'],
    'policy_qty': ['policy qty', 'policy quantity', 'policies', 'qty policies', 'number of policies', 'policy count', 'policies sold'],
    'current_scenario': ['current scenario', 'scenario', 'current_scenario'],
}


def clean_money(value):
    if pd.isna(value):
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).replace('R', '').replace(',', '').replace(' ', '').strip()
    if text in ['', '-', 'nan', 'None']:
        return 0.0
    try:
        return float(text)
    except ValueError:
        return 0.0


def franchise_match_key(name):
    """Return a stable matching key for franchise/branch names from imports.

    Claims workbooks and policy exports often use slightly different wording, for example
    'Martins Funeral Agency Northcliff', 'Martins Funerals Northcliff CC', or 'Northcliff'.
    This helper removes common Martins wording, punctuation and legal suffixes so claims,
    policy data and grouped franchise settings can match reliably.
    """
    if name is None or (isinstance(name, float) and pd.isna(name)):
        return ''
    text = str(name).lower().strip()
    if text in {'', 'nan', 'none'}:
        return ''
    replacements = [
        'martins funerals agency', 'martins funeral agency', 'martins funerals',
        'martins funeral', 'martins direct', 'martins funeral brokers', 'martins brokers',
        'franchising', 'funeral services', 'funeral service', 'funerals', 'funeral',
        'agency', 'branch', 'undertakers', 'martins',
    ]
    for rep in replacements:
        text = text.replace(rep, ' ')
    # Remove common legal / descriptive suffixes that differ between reports.
    suffixes = [' pty ltd', ' (pty) ltd', ' cc', ' inc', ' ltd']
    for suffix in suffixes:
        text = text.replace(suffix, ' ')
    text = text.replace('&', 'and')
    text = re.sub(r'[^a-z0-9]+', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text




def resolve_claim_key_to_existing(claim_key, existing_keys):
    """Map a claims workbook branch key to the closest imported franchise key.

    The claims workbook often uses longer/shorter wording than the policy import. Exact
    matches are used first. If not found, use conservative token containment so names like
    "centurion pretoria east" can match "pretoria east" while avoiding unrelated branches.
    """
    if not claim_key:
        return claim_key
    existing = [k for k in existing_keys if k]
    if claim_key in existing:
        return claim_key
    claim_tokens = set(str(claim_key).split())
    if not claim_tokens:
        return claim_key
    best_key = None
    best_score = 0.0
    for key in existing:
        key_tokens = set(str(key).split())
        if not key_tokens:
            continue
        intersection = claim_tokens & key_tokens
        if not intersection:
            continue
        # Exact token containment or strong overlap only.
        containment = len(intersection) / max(1, min(len(claim_tokens), len(key_tokens)))
        union_score = len(intersection) / max(1, len(claim_tokens | key_tokens))
        score = containment * 0.7 + union_score * 0.3
        if (claim_tokens <= key_tokens or key_tokens <= claim_tokens or containment >= 0.80) and score > best_score:
            best_score = score
            best_key = key
    return best_key or claim_key



# Optional branch aliases for known wording differences. Keys and values are normalized.
FRANCHISE_ALIASES = {
    franchise_match_key('Martins Lydenburg - Belfast'): franchise_match_key('Martins Funerals Belfast'),
    franchise_match_key('Martins Funeral Centurion - Pretoria East'): franchise_match_key('Centurion Pretoria East'),
}


def apply_franchise_alias(key):
    if not key:
        return key
    key = FRANCHISE_ALIASES.get(key, key)
    try:
        user_aliases = load_franchise_config().get('claims_aliases', {})
        key = user_aliases.get(key, key)
    except Exception:
        pass
    return key


def build_claims_import_summary(claims_before_resolution, claims_after_resolution, existing_keys, written_claims=None, written_claim_count=None):
    """Create a diagnostic summary for the claims import screen.

    Matched claims are claims whose normalized branch key was resolved to an existing
    policy franchise key. Written claims are the matched claims actually applied to
    the analytics dataset. They should reconcile exactly.
    """
    try:
        before = claims_before_resolution.copy()
        after = claims_after_resolution.copy()
        existing = set([k for k in existing_keys if k])
        total_claims = float(before.get('claims', pd.Series(dtype=float)).sum()) if not before.empty else 0.0
        total_count = float(before.get('claim_count', pd.Series(dtype=float)).sum()) if not before.empty else 0.0
        matched = after[after['claim_key'].isin(existing)] if existing and 'claim_key' in after.columns else pd.DataFrame()
        matched_claims = float(matched.get('claims', pd.Series(dtype=float)).sum()) if not matched.empty else 0.0
        matched_count = float(matched.get('claim_count', pd.Series(dtype=float)).sum()) if not matched.empty else 0.0
        unmatched = after[~after['claim_key'].isin(existing)] if existing and 'claim_key' in after.columns else after
        unmatched_claims = float(unmatched.get('claims', pd.Series(dtype=float)).sum()) if not unmatched.empty else 0.0
        rows = []
        if not after.empty:
            sample = after.groupby(['franchise','claim_key'], as_index=False).agg({'claims':'sum','claim_count':'sum'}).sort_values('claims', ascending=False).head(30)
            for _, r in sample.iterrows():
                rows.append({
                    'claims_franchise': r.get('franchise',''),
                    'matched_key': r.get('claim_key',''),
                    'status': 'Matched' if r.get('claim_key','') in existing else 'Claims-only / unmatched',
                    'claims': float(r.get('claims',0) or 0),
                    'claim_count': float(r.get('claim_count',0) or 0),
                })
        if written_claims is None:
            written_claims = matched_claims
        if written_claim_count is None:
            written_claim_count = matched_count
        return {
            'total_claims_imported': total_claims,
            'total_claim_count_imported': total_count,
            'matched_claims': matched_claims,
            'matched_claim_count': matched_count,
            'written_claims': float(written_claims or 0),
            'written_claim_count': float(written_claim_count or 0),
            'unmatched_claims': unmatched_claims,
            'difference_claims': float(matched_claims - float(written_claims or 0)),
            'rows': rows,
        }
    except Exception:
        return {}

def month_from_filename(path):
    """Return first YYYYMMDD/YYYMM date found in filename as month start."""
    name = os.path.basename(path)
    m = re.search(r'(20\d{2})(\d{2})(\d{2})', name)
    if m:
        return pd.Timestamp(year=int(m.group(1)), month=int(m.group(2)), day=1)
    m = re.search(r'(20\d{2})[-_ ]?(\d{2})', name)
    if m:
        return pd.Timestamp(year=int(m.group(1)), month=int(m.group(2)), day=1)
    return pd.Timestamp(datetime.today().replace(day=1))


def parse_policy_transaction_sheet(df, source_path=''):
    """Parse detailed policy receipt export.

    Expected columns include Franchise, Relation (MEM for main members), Risk, Retail and MPIA.
    Commission/risk deductions are calculated on MEM rows only:
      R1 Policy Fee = MPIA x R1
      2.1% Fee = (Risk - R1 Policy Fee) x 2.1%
      Net Risk Premium = Risk - R1 Policy Fee - 2.1% Fee
    The R1 fee also represents the paid policy-month quantity.
    """
    if df.empty:
        return pd.DataFrame()
    lookup = {str(c).strip().lower(): c for c in df.columns}
    def find(*names):
        for n in names:
            key = n.strip().lower()
            if key in lookup:
                return lookup[key]
        return None
    franchise_col = find('franchise', 'branch', 'branch name', 'participating group', 'participating_group', 'franchise name', 'office')
    relation_col = find('relation', 'member relation', 'rel', 'member type')
    risk_col = find('risk', 'risk value', 'risk premium', 'risk amount')
    retail_col = find('retail', 'retail premium', 'retail amount')
    mpia_col = find('mpia')

    # Fallback to the fixed PolicyData column positions supplied by Martins Direct:
    # I = relation/MEM, M = risk amount, R = MPIA, W = retail.
    # Python indexes are zero-based: I=8, M=12, R=17, W=22.
    cols = list(df.columns)
    def by_index(idx):
        return cols[idx] if len(cols) > idx else None
    if relation_col is None:
        relation_col = by_index(8)
    if risk_col is None:
        risk_col = by_index(12)
    if mpia_col is None:
        mpia_col = by_index(17)
    if retail_col is None:
        retail_col = by_index(22)
    if franchise_col is None:
        # Most exports keep the branch/franchise name near the first few columns.
        # Try B first, then A as a final fallback.
        franchise_col = by_index(1) or by_index(0)

    if not all([franchise_col, relation_col, risk_col, retail_col, mpia_col]):
        return pd.DataFrame()

    work = df.copy()
    work['__franchise'] = work[franchise_col].astype(str).str.strip()
    work['__relation'] = work[relation_col].astype(str).str.strip().str.upper()
    work['__risk'] = work[risk_col].apply(clean_money)
    work['__retail'] = work[retail_col].apply(clean_money)
    work['__mpia'] = work[mpia_col].apply(clean_money).replace(0, 1)
    work.loc[work['__mpia'] < 1, '__mpia'] = 1
    mem = work[work['__relation'] == 'MEM'].copy()
    if mem.empty:
        return pd.DataFrame()
    mem['__single_monthly_premium'] = mem['__risk'] / mem['__mpia'].replace(0, 1)
    # Martins Direct R1/ADV rule, calculated per paid month and multiplied back by MPIA.
    mem['__r1_policy_fee'] = mem['__mpia'] * 1.0
    mem['__risk_after_r1_per_month'] = mem['__single_monthly_premium'] - 1.0
    mem.loc[mem['__risk_after_r1_per_month'] < 0, '__risk_after_r1_per_month'] = 0
    mem['__underwriter_2_1_fee'] = mem['__risk_after_r1_per_month'] * 0.021 * mem['__mpia']
    mem['__risk_after_r1'] = mem['__risk_after_r1_per_month'] * mem['__mpia']
    # This is the net premium payable after R1 and 2.1% are deducted.
    mem['__net_risk'] = mem['__risk_after_r1'] - mem['__underwriter_2_1_fee']
    # Retail is usually only populated on MEM rows in this export, but use MEM to keep commission base consistent.
    month = month_from_filename(source_path)
    grouped = mem.groupby('__franchise', as_index=False).agg({
        '__retail':'sum',
        '__risk':'sum',
        '__net_risk':'sum',
        '__r1_policy_fee':'sum',
        '__underwriter_2_1_fee':'sum',
        '__mpia':'sum',
        '__single_monthly_premium':'sum',
        '__risk_after_r1':'sum',
    })
    grouped = grouped[grouped['__franchise'].astype(str).str.strip() != '']
    out = pd.DataFrame({
        'franchise': grouped['__franchise'],
        'month': month,
        'retail_premium': grouped['__retail'],
        # Use net underwriter premium as the risk premium used by claim-ratio and scenarios.
        'risk_premium': grouped['__net_risk'],
        'claims': 0.0,
        'policy_qty': grouped['__mpia'],
        'original_risk_premium': grouped['__risk'],
        'r1_policy_fee_imported': grouped['__r1_policy_fee'],
        'underwriter_2_1_fee': grouped['__underwriter_2_1_fee'],
        'risk_after_r1': grouped['__risk_after_r1'],
        'single_monthly_premium_total': grouped['__single_monthly_premium'],
        'current_scenario': '100% Claim Ratio',
    })
    return out



def excel_serial_to_datetime(value):
    """Convert Excel serial dates or text month labels to pandas Timestamp."""
    if pd.isna(value):
        return pd.NaT
    if isinstance(value, datetime):
        return pd.Timestamp(value)
    if isinstance(value, (int, float)):
        # Excel Windows date system. Month headings in the supplied workbook are serial dates.
        return pd.to_datetime(value, unit='D', origin='1899-12-30', errors='coerce')
    return pd.to_datetime(value, errors='coerce')






def parse_claims_tab1_paid_detail_streaming(path):
    """Memory-safe parser for the first claims transaction tab.

    Uses only the business columns needed for claims import instead of loading
    the whole workbook/sheet into pandas:
      - E: franchise name
      - Y: Date Of Claim Paid
      - Z: Amount of Claim Paid Before Net Off
    """
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        ws = wb[wb.sheetnames[0]]
        agg = {}
        for row in ws.iter_rows(min_row=2, values_only=True):
            if len(row) < 26:
                continue
            franchise = row[4]
            paid_date = row[24]
            amount = row[25]
            if pd.isna(franchise) or str(franchise).strip() == '':
                continue
            paid_dt = excel_serial_to_datetime(paid_date)
            if pd.isna(paid_dt):
                continue
            claim_amount = clean_money(amount)
            if claim_amount == 0:
                continue
            month = pd.Timestamp(paid_dt).to_period('M').to_timestamp()
            key = (str(franchise).strip(), month)
            rec = agg.setdefault(key, {'claims': 0.0, 'claim_count': 0})
            rec['claims'] += claim_amount
            rec['claim_count'] += 1
        if not agg:
            return pd.DataFrame(columns=['franchise','month','claims','claim_count','claim_paid_franchise','claim_paid_client','repudiated_pending','grand_total_claims'])
        rows = []
        for (franchise, month), rec in agg.items():
            rows.append({
                'franchise': franchise,
                'month': month,
                'claims': rec['claims'],
                'claim_count': rec['claim_count'],
                'claim_paid_franchise': 0.0,
                'claim_paid_client': 0.0,
                'repudiated_pending': 0.0,
                'grand_total_claims': 0.0,
            })
        return pd.DataFrame(rows)
    finally:
        wb.close()


def worksheet_to_dataframe_limited(ws, max_rows=None, max_cols=None):
    """Convert a read-only worksheet to a small DataFrame for pivot fallbacks."""
    rows = []
    for r_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
        if max_rows is not None and r_idx > max_rows:
            break
        if max_cols is not None:
            row = row[:max_cols]
        rows.append(list(row))
    return pd.DataFrame(rows)

def parse_claims_tab1_paid_detail(df):
    """Parse claims from Tab 1 using fixed business columns.

    Business rule:
      - Column E (Excel E, zero index 4): franchise name
      - Column Y (Excel Y, zero index 24): Date Of Claim Paid
      - Column Z (Excel Z, zero index 25): Amount of Claim Paid Before Net Off

    Each non-zero paid claim row is allocated to the paid month and contributes
    one claim count. This source is preferred above the old pivot tabs because it
    is transaction-level and follows the actual paid date.
    """
    if df is None or df.empty or df.shape[1] < 26:
        return pd.DataFrame()
    rows = []
    for idx in range(1, len(df)):  # start at Excel row 2
        franchise = df.iat[idx, 4] if df.shape[1] > 4 else None
        paid_date = df.iat[idx, 24] if df.shape[1] > 24 else None
        amount = df.iat[idx, 25] if df.shape[1] > 25 else None
        if pd.isna(franchise) or str(franchise).strip() == '':
            continue
        paid_dt = excel_serial_to_datetime(paid_date)
        if pd.isna(paid_dt):
            continue
        claim_amount = clean_money(amount)
        if claim_amount == 0:
            continue
        rows.append({
            'franchise': str(franchise).strip(),
            'month': pd.Timestamp(paid_dt).to_period('M').to_timestamp(),
            'claims': claim_amount,
            'claim_count': 1,
            'claim_paid_franchise': 0.0,
            'claim_paid_client': 0.0,
            'repudiated_pending': 0.0,
            'grand_total_claims': 0.0,
        })
    if not rows:
        return pd.DataFrame(columns=['franchise','month','claims','claim_count','claim_paid_franchise','claim_paid_client','repudiated_pending','grand_total_claims'])
    out = pd.DataFrame(rows)
    return out.groupby(['franchise', 'month'], as_index=False).agg({
        'claims': 'sum',
        'claim_count': 'sum',
        'claim_paid_franchise': 'sum',
        'claim_paid_client': 'sum',
        'repudiated_pending': 'sum',
        'grand_total_claims': 'sum',
    })


def parse_claims_detail_sheet(df):
    """Parse detailed claims workbook and return Franchise/Month/Claims rows.

    Expected detailed columns include Participating Group, Date Of Claim Paid and Final Amount.
    Claims are allocated by Participating Group and claim paid month. This remains as a fallback
    for claims workbooks that do not include the pivot sheets.
    """
    if df is None or df.empty:
        return pd.DataFrame()
    lookup = {str(c).strip().lower(): c for c in df.columns}
    def find(*names):
        for n in names:
            key = n.strip().lower()
            if key in lookup:
                return lookup[key]
        return None
    franchise_col = find('participating group', 'franchise', 'branch', 'intermediary', 'beneficiary')
    date_col = find('date of claim paid', 'claim paid date', 'paid date', 'statusdate', 'approvaldate', 'claim reported date')
    amount_col = find('final amount', 'net amount paid via bank', 'amount of claim paid before net off', 'amount of claim approved', 'amount of claim reported')
    if not all([franchise_col, date_col, amount_col]):
        return pd.DataFrame()
    work = df.copy()
    work['__franchise'] = work[franchise_col].astype(str).str.strip()
    work['__date'] = pd.to_datetime(work[date_col], errors='coerce')
    work['__claims'] = work[amount_col].apply(clean_money)
    work = work[(work['__franchise'] != '') & work['__date'].notna()]
    work = work[work['__claims'] != 0]
    if work.empty:
        return pd.DataFrame()
    work['month'] = work['__date'].dt.to_period('M').dt.to_timestamp()
    grouped = work.groupby(['__franchise', 'month'], as_index=False)['__claims'].sum()
    grouped['claim_count'] = 0
    grouped['repudiated_pending'] = 0.0
    return grouped.rename(columns={'__franchise':'franchise', '__claims':'claims'})


def parse_claims_pivot_sheet(df):
    """Parse a generic Claims per Branch(sum) pivot-style sheet as a fallback."""
    if df is None or df.empty or df.shape[0] < 7:
        return pd.DataFrame()
    row_labels_idx = None
    row_labels_col = None
    for r in range(min(15, len(df))):
        for c in range(min(5, df.shape[1])):
            if str(df.iat[r, c]).strip().lower() == 'row labels':
                row_labels_idx = r
                row_labels_col = c
                break
        if row_labels_idx is not None:
            break
    if row_labels_idx is None:
        return pd.DataFrame()
    return parse_claims_named_pivot(df, value_name='claims', row_start=row_labels_idx + 1, row_end=None, row_labels_col=row_labels_col, year_row=row_labels_idx - 2, month_row=row_labels_idx - 1)


def parse_claims_named_pivot(df, value_name='claims', row_start=7, row_end=55, row_labels_col=1, year_row=4, month_row=5, first_value_col=2, last_value_col=None):
    """Parse the supplied Martins claims pivot tabs.

    Business rule for this workbook:
      - Tab 3: column B rows 8:56 contain the franchise names.
      - Tab 3: month columns start at C and continue until the column before Grand Total.
      - Tab 4: column B rows 8:56 contain the franchise names.
      - Tab 4: month columns start at C and continue until the column before Grand Total.

    This is intentionally dynamic because the claims workbook grows month by month
    (for example Apr 2025 through Apr 2026). The importer therefore does not stop at K;
    it scans the month header row and imports every month column it finds before Grand Total.
    Python indexes are zero-based, so B=1 and C=2. row_end is inclusive.
    """
    if df is None or df.empty:
        return pd.DataFrame()
    months = {'jan':1,'feb':2,'mar':3,'apr':4,'may':5,'jun':6,'jul':7,'aug':8,'sep':9,'oct':10,'nov':11,'dec':12}
    current_year = None
    col_months = {}
    max_cols = df.shape[1]
    if last_value_col is None:
        # Stop before a Grand Total column if one exists on the month header or year header rows.
        last_col = max_cols - 1
        for c in range(first_value_col, max_cols):
            header_values = []
            for rr in (year_row, month_row):
                if rr < len(df):
                    header_values.append(str(df.iat[rr, c]).strip().lower())
            if any('grand total' in hv or hv == 'total' for hv in header_values):
                last_col = c - 1
                break
    else:
        last_col = min(last_value_col, max_cols - 1)
    for c in range(first_value_col, last_col + 1):
        yv = df.iat[year_row, c] if year_row < len(df) else None
        if not pd.isna(yv):
            try:
                yf = float(str(yv).strip())
                if 1900 <= yf <= 2100:
                    current_year = int(yf)
            except Exception:
                ytxt = str(yv).strip()
                if ytxt.isdigit() and len(ytxt) == 4:
                    current_year = int(ytxt)
        mv = df.iat[month_row, c] if month_row < len(df) else None
        if pd.isna(mv) or current_year is None:
            continue
        mtxt = str(mv).strip()[:3].lower()
        if mtxt in months:
            col_months[c] = pd.Timestamp(year=current_year, month=months[mtxt], day=1)
    if not col_months:
        return pd.DataFrame()
    if row_end is None:
        row_end = len(df) - 1
    rows = []
    for r in range(row_start, min(row_end, len(df) - 1) + 1):
        franchise = df.iat[r, row_labels_col] if row_labels_col < df.shape[1] else None
        if pd.isna(franchise) or str(franchise).strip() == '':
            continue
        name = str(franchise).strip()
        if name.lower() in {'grand total', 'total'}:
            continue
        for c, month in col_months.items():
            val = clean_money(df.iat[r, c])
            if val:
                rows.append({'franchise': name, 'month': month, value_name: val})
    return pd.DataFrame(rows)


def parse_payable_to_sheet(df):
    """Parse Tab 5 Payable To using the fixed range B4:F52.

    Columns:
    - B: Franchise / branch name
    - C: Claim paid to franchise / entity
    - D: Claim paid to client / individual
    - E: Repudiated or pending
    - F: Grand total claims

    This sheet is franchise-level and does not contain month columns, so the values are
    attached once to the first imported claim month for that franchise during merge.
    """
    if df is None or df.empty:
        return pd.DataFrame()
    rows = []
    # Excel rows 4:52 are pandas indexes 3:51. Columns B:F are indexes 1:5.
    for r in range(3, min(52, len(df))):
        franchise = df.iat[r, 1] if df.shape[1] > 1 else None
        if pd.isna(franchise) or str(franchise).strip() == '':
            continue
        name = str(franchise).strip()
        if name.lower() in {'grand total', 'total', 'row labels'}:
            continue
        claim_paid_franchise = clean_money(df.iat[r, 2]) if df.shape[1] > 2 else 0
        claim_paid_client = clean_money(df.iat[r, 3]) if df.shape[1] > 3 else 0
        repudiated_pending = clean_money(df.iat[r, 4]) if df.shape[1] > 4 else 0
        grand_total_claims = clean_money(df.iat[r, 5]) if df.shape[1] > 5 else 0
        if claim_paid_franchise or claim_paid_client or repudiated_pending or grand_total_claims:
            rows.append({
                'franchise': name,
                'claim_paid_franchise': claim_paid_franchise,
                'claim_paid_client': claim_paid_client,
                'repudiated_pending': repudiated_pending,
                'grand_total_claims': grand_total_claims,
            })
    if not rows:
        return pd.DataFrame(columns=['franchise','claim_paid_franchise','claim_paid_client','repudiated_pending','grand_total_claims'])
    out = pd.DataFrame(rows)
    return out.groupby('franchise', as_index=False).agg({
        'claim_paid_franchise':'sum',
        'claim_paid_client':'sum',
        'repudiated_pending':'sum',
        'grand_total_claims':'sum',
    })


def read_claims_file(path):
    """Read a claims workbook and aggregate claims per franchise/month.

    This version avoids loading the full workbook into memory. The first claims
    transaction tab is streamed row-by-row, and only small pivot fallback ranges
    are converted to DataFrames.
    """
    claims_sum = pd.DataFrame()
    claims_count = pd.DataFrame()
    claims_status = pd.DataFrame()

    # Preferred source: stream Tab 1 instead of pd.read_excel(sheet_name=None).
    claims_sum = parse_claims_tab1_paid_detail_streaming(path)

    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        sheet_names = list(wb.sheetnames)
        sum_sheet = next((name for name in sheet_names if 'claims per branch' in name.lower() and 'sum' in name.lower()), None)
        count_sheet = next((name for name in sheet_names if 'claims per branch' in name.lower() and 'count' in name.lower()), None)
        payable_sheet = next((name for name in sheet_names if 'payable' in name.lower()), None)

        if claims_sum.empty and sum_sheet:
            df_sum = worksheet_to_dataframe_limited(wb[sum_sheet], max_rows=80)
            claims_sum = parse_claims_named_pivot(df_sum, value_name='claims', row_start=7, row_end=55, row_labels_col=1, year_row=4, month_row=5, first_value_col=2)
        if claims_sum.empty and count_sheet:
            df_count = worksheet_to_dataframe_limited(wb[count_sheet], max_rows=80)
            claims_count = parse_claims_named_pivot(df_count, value_name='claim_count', row_start=7, row_end=55, row_labels_col=1, year_row=4, month_row=5, first_value_col=2)
        elif count_sheet:
            df_count = worksheet_to_dataframe_limited(wb[count_sheet], max_rows=80)
            fallback_count = parse_claims_named_pivot(df_count, value_name='claim_count', row_start=7, row_end=55, row_labels_col=1, year_row=4, month_row=5, first_value_col=2)
            if claims_sum.get('claim_count') is None and not fallback_count.empty:
                claims_count = fallback_count
        if payable_sheet:
            df_payable = worksheet_to_dataframe_limited(wb[payable_sheet], max_rows=60, max_cols=8)
            claims_status = parse_payable_to_sheet(df_payable)

        if claims_sum.empty:
            frames = []
            # Last-resort fallback: inspect sheets one at a time and keep memory bounded.
            for name in sheet_names:
                df = worksheet_to_dataframe_limited(wb[name], max_rows=5000)
                parsed = parse_claims_pivot_sheet(df)
                if not parsed.empty:
                    frames.append(parsed)
            if frames:
                claims_sum = pd.concat(frames, ignore_index=True)
            else:
                raise ValueError('No usable claims data found. Expected Tab 1 paid-claim columns E/Y/Z, Claims per Branch pivot tabs, Payable To, or a supported detailed claims layout.')
    finally:
        wb.close()

    # Normalize tabs 3, 4 and 5 with the same matching key before joining them.
    claims = claims_sum.copy()
    claims['franchise'] = claims['franchise'].astype(str).str.strip()
    claims['match_key'] = claims['franchise'].map(franchise_match_key).map(apply_franchise_alias)
    claims['month'] = pd.to_datetime(claims['month'], errors='coerce').dt.to_period('M').dt.to_timestamp()
    claims = claims.dropna(subset=['month'])
    claims['claims'] = claims['claims'].apply(clean_money)
    claims = claims[claims['match_key'] != '']
    claims = claims.groupby(['match_key', 'month'], as_index=False).agg({'claims':'sum', 'franchise':'first'})

    existing_claim_count = None
    if 'claim_count' in claims.columns:
        existing_claim_count = claims[['match_key', 'month', 'claim_count']].copy()
        existing_claim_count['claim_count'] = existing_claim_count['claim_count'].apply(clean_money)
        claims = claims.drop(columns=['claim_count'])

    if not claims_count.empty:
        counts = claims_count.copy()
        counts['franchise'] = counts['franchise'].astype(str).str.strip()
        counts['match_key'] = counts['franchise'].map(franchise_match_key).map(apply_franchise_alias)
        counts['month'] = pd.to_datetime(counts['month'], errors='coerce').dt.to_period('M').dt.to_timestamp()
        counts = counts.dropna(subset=['month'])
        counts['claim_count'] = counts['claim_count'].apply(clean_money)
        counts = counts[counts['match_key'] != '']
        counts = counts.groupby(['match_key', 'month'], as_index=False)['claim_count'].sum()
        claims = claims.merge(counts, on=['match_key', 'month'], how='left')
    elif existing_claim_count is not None:
        existing_claim_count = existing_claim_count.groupby(['match_key', 'month'], as_index=False)['claim_count'].sum()
        claims = claims.merge(existing_claim_count, on=['match_key', 'month'], how='left')
    else:
        claims['claim_count'] = 0
    claims['claim_count'] = claims['claim_count'].fillna(0)
    for _c in ['claim_paid_franchise', 'claim_paid_client', 'repudiated_pending', 'grand_total_claims']:
        claims[_c] = 0.0

    if not claims_status.empty:
        status = claims_status.copy()
        status['franchise'] = status['franchise'].astype(str).str.strip()
        status['match_key'] = status['franchise'].map(franchise_match_key).map(apply_franchise_alias)
        for _c in ['claim_paid_franchise', 'claim_paid_client', 'repudiated_pending', 'grand_total_claims']:
            if _c not in status.columns:
                status[_c] = 0.0
            status[_c] = status[_c].apply(clean_money)
        status = status[status['match_key'] != '']
        status = status.groupby('match_key', as_index=False).agg({'claim_paid_franchise':'sum','claim_paid_client':'sum','repudiated_pending':'sum','grand_total_claims':'sum'})
        # Tab 5 is franchise-level, so attach the amounts once to the first month found for that franchise.
        claims = claims.sort_values(['match_key', 'month']).copy()
        first_idx = claims.groupby('match_key').head(1).index
        for _c in ['claim_paid_franchise', 'claim_paid_client', 'repudiated_pending', 'grand_total_claims']:
            status_map = dict(zip(status['match_key'], status[_c]))
            claims.loc[first_idx, _c] = claims.loc[first_idx, 'match_key'].map(status_map).fillna(0)

    claims = claims[(claims['claims'] != 0) | (claims['claim_count'] != 0) | (claims['repudiated_pending'] != 0) | (claims['claim_paid_franchise'] != 0) | (claims['claim_paid_client'] != 0) | (claims['grand_total_claims'] != 0)]
    claims = claims.rename(columns={'match_key':'claim_key'})
    return claims.groupby(['claim_key', 'franchise', 'month'], as_index=False).agg({'claims':'sum', 'claim_count':'sum', 'claim_paid_franchise':'sum', 'claim_paid_client':'sum', 'repudiated_pending':'sum', 'grand_total_claims':'sum'})



def merge_claims_into_raw(raw, claims_df):
    global LAST_CLAIMS_IMPORT_SUMMARY
    """Merge paid monthly claims into the current raw premium dataset.

    Claims are matched by one normalized franchise key used across PolicyData and
    claims tabs 3, 4 and 5. Claims that match an existing policy franchise are
    written back to that franchise (including months with no policy row yet).
    Claims that do not match an existing policy franchise stay in the import
    validation summary and are NOT added to the analytics dataset; this prevents
    duplicate/claims-only franchises from inflating Total Franchises.
    """
    if claims_df is None or claims_df.empty:
        return raw

    claims = claims_df.copy()
    claims['month'] = pd.to_datetime(claims['month'], errors='coerce').dt.to_period('M').dt.to_timestamp()
    claims = claims.dropna(subset=['month'])
    claims_before_resolution = claims.copy()
    if 'claim_key' not in claims.columns:
        claims['claim_key'] = claims['franchise'].map(franchise_match_key).map(apply_franchise_alias)
    else:
        claims['claim_key'] = claims['claim_key'].map(apply_franchise_alias)
    for col in ['claims', 'claim_count', 'claim_paid_franchise', 'claim_paid_client', 'repudiated_pending', 'grand_total_claims']:
        if col not in claims.columns:
            claims[col] = 0.0
        claims[col] = claims[col].apply(clean_money)
    claims = claims[claims['claim_key'].astype(str).str.strip() != '']

    if raw is None or raw.empty:
        LAST_CLAIMS_IMPORT_SUMMARY = build_claims_import_summary(claims_before_resolution, claims, [])
        return pd.DataFrame({
            'franchise': claims['franchise'],
            'month': claims['month'],
            'retail_premium': 0.0,
            'risk_premium': 0.0,
            'claims': claims['claims'],
            'claim_count': claims['claim_count'],
            'claim_paid_franchise': claims.get('claim_paid_franchise', 0),
            'claim_paid_client': claims.get('claim_paid_client', 0),
            'repudiated_pending': claims['repudiated_pending'],
            'grand_total_claims': claims.get('grand_total_claims', 0),
            'policy_qty': 0.0,
            'current_scenario': '100% Claim Ratio',
        })

    work = raw.copy()
    work['month'] = pd.to_datetime(work['month'], errors='coerce').dt.to_period('M').dt.to_timestamp()
    for col in ['claims', 'claim_count', 'claim_paid_franchise', 'claim_paid_client', 'repudiated_pending', 'grand_total_claims']:
        if col not in work.columns:
            work[col] = 0.0
        work[col] = work[col].apply(clean_money)
    work['claim_key'] = work['franchise'].map(franchise_match_key).map(apply_franchise_alias)

    # Build canonical policy franchise names from policy data only.
    canonical_franchise_by_key = {}
    for _key, _name in zip(work['claim_key'], work['franchise'].astype(str)):
        if _key and _key not in canonical_franchise_by_key:
            canonical_franchise_by_key[_key] = _name
    existing_keys = list(canonical_franchise_by_key.keys())

    # Resolve claim branch wording to the policy franchise keys and canonical names.
    claims['claim_key'] = claims['claim_key'].apply(lambda k: resolve_claim_key_to_existing(k, existing_keys))
    claims['franchise'] = claims.apply(lambda r: canonical_franchise_by_key.get(r['claim_key'], r.get('franchise', '')), axis=1)
    claims = claims.groupby(['claim_key', 'month'], as_index=False).agg({
        'claims':'sum',
        'claim_count':'sum',
        'claim_paid_franchise':'sum',
        'claim_paid_client':'sum',
        'repudiated_pending':'sum',
        'grand_total_claims':'sum',
        'franchise':'first'
    })

    # Only claims matched to active policy franchises are written into calculations.
    matched_claims = claims[claims['claim_key'].isin(existing_keys)].copy()
    written_claims = float(matched_claims['claims'].sum()) if not matched_claims.empty else 0.0
    written_claim_count = float(matched_claims['claim_count'].sum()) if not matched_claims.empty else 0.0
    LAST_CLAIMS_IMPORT_SUMMARY = build_claims_import_summary(
        claims_before_resolution,
        claims,
        existing_keys,
        written_claims=written_claims,
        written_claim_count=written_claim_count,
    )

    if matched_claims.empty:
        work = work.drop(columns=['claim_key'])
        return work.sort_values(['franchise', 'month']).reset_index(drop=True)

    supplied_months = set(matched_claims['month'].dropna().tolist())
    work.loc[work['month'].isin(supplied_months), ['claims', 'claim_count', 'claim_paid_franchise', 'claim_paid_client', 'repudiated_pending', 'grand_total_claims']] = 0.0

    merged = work.merge(
        matched_claims[['claim_key', 'month', 'claims', 'claim_count', 'claim_paid_franchise', 'claim_paid_client', 'repudiated_pending', 'grand_total_claims']],
        on=['claim_key', 'month'],
        how='left',
        suffixes=('', '_paid')
    )
    for col in ['claims', 'claim_count', 'claim_paid_franchise', 'claim_paid_client', 'repudiated_pending', 'grand_total_claims']:
        paid_col = f'{col}_paid'
        merged[col] = merged[paid_col].fillna(merged[col]).fillna(0)
        merged = merged.drop(columns=[paid_col])

    # Add matched claim months where the branch is active but no policy row exists yet.
    existing_pairs = set(zip(work['claim_key'], work['month']))
    missing_pairs = matched_claims[~matched_claims.apply(lambda r: (r['claim_key'], r['month']) in existing_pairs, axis=1)].copy()
    if not missing_pairs.empty:
        extras = pd.DataFrame({
            'franchise': missing_pairs['franchise'],
            'month': missing_pairs['month'],
            'retail_premium': 0.0,
            'risk_premium': 0.0,
            'claims': missing_pairs['claims'],
            'claim_count': missing_pairs['claim_count'],
            'claim_paid_franchise': missing_pairs.get('claim_paid_franchise', 0),
            'claim_paid_client': missing_pairs.get('claim_paid_client', 0),
            'repudiated_pending': missing_pairs['repudiated_pending'],
            'grand_total_claims': missing_pairs.get('grand_total_claims', 0),
            'policy_qty': 0.0,
            'current_scenario': '100% Claim Ratio',
        })
        # Preserve all existing raw calculation columns so analyse() can process extras safely.
        for col in merged.columns:
            if col not in extras.columns and col != 'claim_key':
                if col == 'current_scenario':
                    extras[col] = '100% Claim Ratio'
                elif col == 'franchise':
                    pass
                elif col == 'month':
                    pass
                else:
                    extras[col] = 0.0
        extras['claim_key'] = extras['franchise'].map(franchise_match_key).map(apply_franchise_alias)
        extras = extras[[c for c in merged.columns if c in extras.columns]]
        merged = pd.concat([merged, extras], ignore_index=True)

    merged = merged.drop(columns=['claim_key'], errors='ignore')
    return merged.sort_values(['franchise', 'month']).reset_index(drop=True)


def parse_wide_claim_ratio_sheet(df):
    """Parse Martins Direct wide claim-ratio sheets.

    Expected layout:
    Row 1: month/date headings every 4 columns
    Row 2: Franchise, Retail Premium, Risk Premium, Claims, Ratio repeated
    Rows below: one franchise per row
    """
    if df.shape[0] < 3:
        return pd.DataFrame()

    header_row_index = None
    for idx in range(min(10, len(df))):
        row_values = [str(v).strip().lower() for v in df.iloc[idx].tolist()]
        if 'franchise' in row_values and 'retail premium' in row_values and 'risk premium' in row_values and 'claims' in row_values:
            header_row_index = idx
            break
    if header_row_index is None:
        return pd.DataFrame()

    month_row_index = max(0, header_row_index - 1)
    header = [str(v).strip().lower() if not pd.isna(v) else '' for v in df.iloc[header_row_index].tolist()]
    month_row = df.iloc[month_row_index].tolist()

    # Locate repeated monthly groups. Ignore summary groups where the top heading is text like "12 Months".
    groups = []
    for col_idx, label in enumerate(header):
        if label == 'retail premium':
            month = excel_serial_to_datetime(month_row[col_idx])
            if pd.isna(month):
                continue
            if col_idx + 2 < len(header) and header[col_idx + 1] == 'risk premium' and header[col_idx + 2] == 'claims':
                groups.append((col_idx, month))

    if not groups:
        return pd.DataFrame()

    rows = []
    for r_idx in range(header_row_index + 1, len(df)):
        franchise = df.iat[r_idx, 0] if df.shape[1] else None
        if pd.isna(franchise) or str(franchise).strip() == '':
            continue
        franchise_name = str(franchise).strip()
        # Skip aggregate/summary rows if present.
        if franchise_name.lower() in {'franchise', 'total', 'grand total'}:
            continue
        for start_col, month in groups:
            retail = clean_money(df.iat[r_idx, start_col]) if start_col < df.shape[1] else 0
            risk = clean_money(df.iat[r_idx, start_col + 1]) if start_col + 1 < df.shape[1] else 0
            claims = clean_money(df.iat[r_idx, start_col + 2]) if start_col + 2 < df.shape[1] else 0
            # Keep rows with any value. Blank zero rows are not useful for analysis.
            if retail == 0 and risk == 0 and claims == 0:
                continue
            rows.append({
                'franchise': franchise_name,
                'month': month,
                'retail_premium': retail,
                'risk_premium': risk,
                'claims': claims,
                'policy_qty': 0,
                'current_scenario': '100% Claim Ratio',
            })
    return pd.DataFrame(rows)


def normalise_columns(df):
    lookup = {str(c).strip().lower(): c for c in df.columns}
    rename = {}
    for target, aliases in COLUMN_ALIASES.items():
        for alias in aliases:
            if alias in lookup:
                rename[lookup[alias]] = target
                break
    df = df.rename(columns=rename)
    required = ['franchise', 'month', 'retail_premium', 'risk_premium', 'claims']
    missing = [c for c in required if c not in df.columns]
    if missing:
        raise ValueError('Missing required columns: ' + ', '.join(missing))
    if 'policy_qty' not in df.columns:
        df['policy_qty'] = 0
    if 'current_scenario' not in df.columns:
        df['current_scenario'] = '100% Claim Ratio'
    return df




def _normalise_header(value):
    return re.sub(r'[^a-z0-9]+', ' ', str(value or '').strip().lower()).strip()



POLICY_RELATION_CODES = ['MEM', 'MEMBEN', 'EXT', 'EXTBEN', 'SPS', 'SPBEN', 'DEP', 'DEPBEN', 'DEPEXTRACHILD']
POLICY_RELATION_LABELS = {
    'MEM': 'MAIN MEMBERS', 'MEMBEN': 'MEMBER BENEFICIARY', 'EXT': 'EXTENDED FAMILY', 'EXTBEN': 'EXTENDED BENEFICIARY',
    'SPS': 'SPOUSE', 'SPBEN': 'SPOUSE BENEFICIARY', 'DEP': 'DEPENDANT', 'DEPBEN': 'DEPENDANT BENEFICIARY',
    'DEPEXTRACHILD': 'DEPENDANT EXTRA CHILD'
}
AGE_LIMIT_YEARS = [1, 5, 13, 21, 64, 70, 84, 85]


def _clean_id_number(value):
    s = re.sub(r'\D+', '', str(value or ''))
    return s[:13] if len(s) >= 6 else ''


def _dob_from_sa_id(id_number):
    s = _clean_id_number(id_number)
    if len(s) < 6:
        return None
    try:
        yy, mm, dd = int(s[:2]), int(s[2:4]), int(s[4:6])
        now = datetime.now().date()
        # South African IDs use YYMMDD. Use the century that gives a sensible age.
        year = 2000 + yy
        dob = datetime(year, mm, dd).date()
        if dob > now:
            dob = datetime(1900 + yy, mm, dd).date()
        if dob > now:
            return None
        return dob
    except Exception:
        return None


def _age_years_months(dob, today=None):
    if not dob:
        return None, None
    today = today or datetime.now().date()
    years = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
    months = (today.year - dob.year) * 12 + today.month - dob.month - (1 if today.day < dob.day else 0)
    return max(0, years), max(0, months)


def _date_at_age(dob, years):
    try:
        return dob.replace(year=dob.year + years)
    except ValueError:
        return dob.replace(month=2, day=28, year=dob.year + years)


def _policy_age_band(age_years, age_months):
    if age_months is None:
        return 'Unknown'
    if age_months <= 11:
        return '0-11 months'
    if age_years <= 5:
        return '1-5 years'
    if age_years <= 13:
        return '6-13 years'
    if age_years <= 21:
        return '14-21 years'
    if age_years <= 64:
        return '18-64 years'
    if age_years <= 84:
        return '65-84 years'
    return '85 years and older'


def _policy_age_notifications_for_member(relation, id_number, member_name='', franchise_name='', policy_number=''):
    dob = _dob_from_sa_id(id_number)
    if not dob:
        return []
    today = datetime.now().date()
    age_years, age_months = _age_years_months(dob, today)
    notes = []
    rel = str(relation or '').strip().upper()
    for limit in AGE_LIMIT_YEARS:
        due = _date_at_age(dob, limit)
        days = (due - today).days
        if 0 <= days <= 92:
            msg = f"{rel or 'Member'} is {days} days from age limit {limit}. Current age band: {_policy_age_band(age_years, age_months)}."
            notes.append({'notification_type': f'AGE_LIMIT_{limit}', 'message': msg, 'due_date': due, 'age_years': age_years, 'age_months': age_months})
    if rel in {'DEP', 'DEPBEN', 'DEPEXTRACHILD'}:
        due21 = _date_at_age(dob, 21)
        days21 = (due21 - today).days
        if 0 <= days21 <= 92:
            notes.append({
                'notification_type': 'CHILD_TURNING_21',
                'message': 'Child will need a new policy or pay extra within 3 months before turning 21.',
                'due_date': due21, 'age_years': age_years, 'age_months': age_months,
            })
    return notes


def _member_name_from_raw(raw_data):
    name = ' '.join([_first_raw_value(raw_data, ['full names', 'first names', 'member name', 'client name', 'policy holder', 'insured name']), _first_raw_value(raw_data, ['surname'])]).strip()
    return name[:180]


def _policy_number_from_raw(raw_data):
    return (_first_raw_value(raw_data, ['policy number', 'policy no', 'policy', 'certificate number']) or '')[:80]


def _ensure_policy_age_notification_table(conn):
    ensure_database_schema(conn)



def refresh_policy_age_notifications(limit=5000, selected='All'):
    """Refresh age-limit notifications from PolicyData rows without duplicates.

    The raw PolicyData table keeps every monthly import for audit, so the same
    person/policy can appear many times. Notifications are therefore de-duped by
    franchise + policy number + ID number + relation + warning type. Existing
    active notifications in the selected scope are marked inactive first, then
    the latest matching member row reactivates/updates one notification record.
    """
    engine = get_db_engine()
    if engine is None:
        return {'ok': False, 'message': 'Database not connected', 'processed': 0, 'created': 0}
    params = {'limit': int(limit), 'rels': POLICY_RELATION_CODES}
    relation_expr = "UPPER(TRIM(COALESCE(NULLIF(relation,''), raw_data->>'Relation', raw_data->>'RELATION', raw_data->>'I', '')))"
    where = [f"{relation_expr} = ANY(:rels)"]
    if selected and selected != 'All':
        where.append("LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))")
        params['franchise'] = selected
    sql = text(f"""
        SELECT DISTINCT ON (
            LOWER(TRIM(COALESCE(franchise_name,''))),
            COALESCE(NULLIF(TRIM(id_number_f),''), NULLIF(TRIM(raw_data->>'ID Number'),''), NULLIF(TRIM(raw_data->>'IDNUMBER'),''), NULLIF(TRIM(raw_data->>'F'),'')),
            COALESCE(NULLIF(TRIM(raw_data->>'Policy Number'),''), NULLIF(TRIM(raw_data->>'Policy'),''), NULLIF(TRIM(raw_data->>'Policy No'),''), id::text),
            {relation_expr}
        )
            id, franchise_name, {relation_expr} AS relation, id_number_f, raw_data
        FROM policydata_detail_raw
        WHERE {' AND '.join(where)}
        ORDER BY
            LOWER(TRIM(COALESCE(franchise_name,''))),
            COALESCE(NULLIF(TRIM(id_number_f),''), NULLIF(TRIM(raw_data->>'ID Number'),''), NULLIF(TRIM(raw_data->>'IDNUMBER'),''), NULLIF(TRIM(raw_data->>'F'),'')),
            COALESCE(NULLIF(TRIM(raw_data->>'Policy Number'),''), NULLIF(TRIM(raw_data->>'Policy'),''), NULLIF(TRIM(raw_data->>'Policy No'),''), id::text),
            {relation_expr},
            import_month DESC NULLS LAST,
            id DESC
        LIMIT :limit
    """)
    processed = created = 0
    seen_keys = set()
    with engine.begin() as conn:
        _ensure_policy_age_notification_table(conn)
        # Deactivate the selected scope first so old/stale duplicate rows disappear from the UI.
        if selected and selected != 'All':
            conn.execute(text("""
                UPDATE app_policy_relation_age_notifications
                SET status = 'inactive', resolved = true, updated_at = NOW()
                WHERE LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))
            """), {'franchise': selected})
        else:
            conn.execute(text("""
                UPDATE app_policy_relation_age_notifications
                SET status = 'inactive', resolved = true, updated_at = NOW()
            """))
        rows = [dict(r) for r in conn.execute(sql, params).mappings().all()]
        for r in rows:
            processed += 1
            raw = r.get('raw_data') or {}
            if isinstance(raw, str):
                try: raw = json.loads(raw)
                except Exception: raw = {}
            member_name = _member_name_from_raw(raw)
            policy_number = _policy_number_from_raw(raw)
            id_number = _clean_id_number(r.get('id_number_f') or _first_raw_value(raw, ['id number', 'id no', 'identity number', 'sa id', 'id']) or _raw_value_by_excel_column(raw, 'F'))
            if id_number and not r.get('id_number_f'):
                try:
                    conn.execute(text('UPDATE policydata_detail_raw SET id_number_f = :idn WHERE id = :id'), {'idn': id_number, 'id': r.get('id')})
                except Exception:
                    pass
            if not id_number:
                continue
            notes = _policy_age_notifications_for_member(r.get('relation'), id_number, member_name, r.get('franchise_name'), policy_number)
            for n in notes:
                notification_key = '|'.join([
                    str(r.get('franchise_name') or '').strip().upper(),
                    str(policy_number or '').strip().upper(),
                    str(id_number or '').strip().upper(),
                    str(r.get('relation') or '').strip().upper(),
                    str(n.get('notification_type') or '').strip().upper(),
                ])
                if notification_key in seen_keys:
                    continue
                seen_keys.add(notification_key)
                conn.execute(text("""
                    INSERT INTO app_policy_relation_age_notifications (
                        notification_key, source_row_id, franchise_name, relation, member_name, id_number, policy_number,
                        age_years, age_months, age_band, notification_type, message, due_date, status, resolved, updated_at
                    ) VALUES (
                        :notification_key, :source_row_id, :franchise_name, :relation, :member_name, :id_number, :policy_number,
                        :age_years, :age_months, :age_band, :notification_type, :message, :due_date, 'active', false, NOW()
                    )
                    ON CONFLICT (notification_key) WHERE notification_key IS NOT NULL DO UPDATE SET
                        source_row_id = EXCLUDED.source_row_id,
                        franchise_name = EXCLUDED.franchise_name,
                        relation = EXCLUDED.relation,
                        member_name = EXCLUDED.member_name,
                        id_number = EXCLUDED.id_number,
                        policy_number = EXCLUDED.policy_number,
                        age_years = EXCLUDED.age_years,
                        age_months = EXCLUDED.age_months,
                        age_band = EXCLUDED.age_band,
                        notification_type = EXCLUDED.notification_type,
                        message = EXCLUDED.message,
                        due_date = EXCLUDED.due_date,
                        status = 'active',
                        resolved = false,
                        updated_at = NOW()
                """), {
                    'notification_key': notification_key,
                    'source_row_id': r.get('id'), 'franchise_name': r.get('franchise_name'), 'relation': r.get('relation'),
                    'member_name': member_name, 'id_number': id_number, 'policy_number': policy_number,
                    'age_years': n.get('age_years'), 'age_months': n.get('age_months'), 'age_band': _policy_age_band(n.get('age_years'), n.get('age_months')), 'notification_type': n.get('notification_type'),
                    'message': n.get('message'), 'due_date': n.get('due_date')
                })
                created += 1
    return {'ok': True, 'processed': processed, 'created': created}



def get_policy_relation_counts(selected='All'):
    """Fast dashboard relation counts.

    Do not rebuild the relation summary table during dashboard rendering.  Rebuilding
    from policydata_detail_raw can scan millions of rows and makes the app look frozen.
    The summary/current tables are refreshed by imports, the daily job, or the manual
    Refresh Notifications action.
    """
    engine = get_db_engine()
    totals = {c: 0 for c in POLICY_RELATION_CODES}
    rows = []
    if engine is None:
        return {'totals': totals, 'rows': rows, 'total': 0}
    params = {'rels': POLICY_RELATION_CODES}
    selected_where = ''
    if selected and selected != 'All':
        selected_where = "WHERE LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))"
        params['franchise'] = selected
    try:
        with engine.begin() as conn:
            # Prefer the precomputed summary table. This is tiny and instant.
            summary_rows = [dict(r) for r in conn.execute(text(f"""
                SELECT UPPER(TRIM(relation)) AS relation, SUM(member_count)::integer AS count
                FROM app_policy_relation_summary
                {selected_where}
                GROUP BY UPPER(TRIM(relation))
            """), {k:v for k,v in params.items() if k != 'rels'}).mappings().all()]
            for r in summary_rows:
                rel = str(r.get('relation') or '').upper()
                if rel in totals:
                    totals[rel] = int(r.get('count') or 0)
            # If the new current-member table exists and the summary is still empty,
            # use it as a fast fallback. Do NOT fall back to policydata_detail_raw here.
            if sum(totals.values()) == 0:
                exists = conn.execute(text("""
                    SELECT EXISTS (
                        SELECT 1 FROM information_schema.tables
                        WHERE table_schema='public' AND table_name='app_policydata_current_members'
                    )
                """)).scalar()
                if exists:
                    where = ["active = true", "UPPER(TRIM(COALESCE(relation,''))) = ANY(:rels)"]
                    if selected and selected != 'All':
                        where.append("LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))")
                    for r in conn.execute(text(f"""
                        SELECT UPPER(TRIM(relation)) AS relation, COUNT(*)::integer AS count
                        FROM app_policydata_current_members
                        WHERE {' AND '.join(where)}
                        GROUP BY UPPER(TRIM(relation))
                    """), params).mappings().all():
                        rel = str(r.get('relation') or '').upper()
                        if rel in totals:
                            totals[rel] = int(r.get('count') or 0)
    except Exception as exc:
        print(f'Could not load relation counts quickly: {exc}')
    rows = [{'code': c, 'label': POLICY_RELATION_LABELS.get(c, c), 'count': totals.get(c, 0)} for c in POLICY_RELATION_CODES]
    return {'totals': totals, 'rows': rows, 'total': sum(totals.values())}


def refresh_policy_relation_summary():
    """Rebuild relation totals from the current-member snapshot only.

    This avoids scanning policydata_detail_raw during dashboard/admin page loads.
    If the current table is empty, run /admin/policydata_current_members/rebuild first.
    """
    engine = get_db_engine()
    if engine is None:
        return {'ok': False, 'rows': 0}
    with engine.begin() as conn:
        ensure_database_schema(conn)
        conn.execute(text("SET LOCAL lock_timeout = '3s'"))
        conn.execute(text("SET LOCAL statement_timeout = '60s'"))
        conn.execute(text("DELETE FROM app_policy_relation_summary"))
        result = conn.execute(text("""
            INSERT INTO app_policy_relation_summary (franchise_name, relation, member_count, updated_at)
            SELECT franchise_name, UPPER(TRIM(relation)) AS relation, COUNT(*)::integer AS member_count, NOW()
            FROM app_policydata_current_members
            WHERE active = true
              AND UPPER(TRIM(COALESCE(relation,''))) = ANY(:rels)
            GROUP BY franchise_name, UPPER(TRIM(relation))
        """), {'rels': POLICY_RELATION_CODES})
    return {'ok': True, 'rows': getattr(result, 'rowcount', 0) or 0}

_POLICY_CURRENT_REBUILD_STATUS = {'running': False, 'status': 'idle', 'processed': 0, 'error': '', 'updated_at': None}

def rebuild_policydata_current_members_job():
    """Fast rebuild of the current/latest PolicyData member snapshot.

    Important change: do NOT scan every historical PolicyData month. The current table
    should represent the latest imported PolicyData month only, then dedupe that month.
    This avoids freezing the app/database on 1.9M+ historical rows.
    """
    global _POLICY_CURRENT_REBUILD_STATUS
    engine = get_db_engine()
    _POLICY_CURRENT_REBUILD_STATUS.update({
        'running': True,
        'status': 'starting',
        'processed': 0,
        'error': '',
        'updated_at': datetime.now().isoformat()
    })
    print('Policy current member rebuild: starting', flush=True)
    try:
        if engine is None:
            raise RuntimeError('Database engine is not available')

        with engine.begin() as conn:
            ensure_database_schema(conn)
            conn.execute(text("SET LOCAL lock_timeout = '5s'"))
            conn.execute(text("SET LOCAL statement_timeout = '20min'"))

            latest_month = conn.execute(text("""
                SELECT MAX(import_month)
                FROM policydata_detail_raw
                WHERE import_month IS NOT NULL
            """)).scalar()

            if not latest_month:
                raise RuntimeError('No PolicyData rows with import_month found. Import PolicyData first.')

            _POLICY_CURRENT_REBUILD_STATUS.update({
                'status': f'clearing current table for latest month {latest_month}',
                'updated_at': datetime.now().isoformat()
            })
            print(f'Policy current member rebuild: latest import_month={latest_month}', flush=True)

            conn.execute(text('TRUNCATE TABLE app_policydata_current_members'))

            # Use only the latest month. Use explicit columns first and raw_data fallbacks second.
            # Column I = relation, Column F = ID number, Column O = address. Existing import code
            # stores these as relation, id_number_f, client_address_o when available.
            rel_expr = "UPPER(TRIM(COALESCE(NULLIF(relation,''), raw_data->>'Relation', raw_data->>'RELATION', raw_data->>'I', '')))"
            id_expr = "COALESCE(NULLIF(TRIM(id_number_f),''), NULLIF(TRIM(raw_data->>'ID Number'),''), NULLIF(TRIM(raw_data->>'IDNUMBER'),''), NULLIF(TRIM(raw_data->>'F'),''), '')"
            pol_expr = "COALESCE(NULLIF(TRIM(raw_data->>'Policy Number'),''), NULLIF(TRIM(raw_data->>'Policy'),''), NULLIF(TRIM(raw_data->>'Policy No'),''), id::text)"
            name_expr = "LEFT(TRIM(COALESCE(raw_data->>'Member Name', raw_data->>'Client Name', raw_data->>'Full Names', raw_data->>'Name', '')), 180)"
            addr_expr = "COALESCE(NULLIF(TRIM(client_address_o),''), NULLIF(TRIM(raw_data->>'Address'),''), NULLIF(TRIM(raw_data->>'ADDRESS'),''), NULLIF(TRIM(raw_data->>'O'),''), '')"

            _POLICY_CURRENT_REBUILD_STATUS.update({
                'status': 'building current members from latest PolicyData month',
                'updated_at': datetime.now().isoformat()
            })

            result = conn.execute(text(f"""
                INSERT INTO app_policydata_current_members (
                    member_key, franchise_name, relation, id_number, policy_number, member_name,
                    age, age_band, client_address_o, import_month, source_row_id, active, raw_data, updated_at, created_at
                )
                SELECT
                    md5(CONCAT_WS('|', LOWER(TRIM(COALESCE(franchise_name,''))), {id_expr}, {pol_expr}, {rel_expr})) AS member_key,
                    franchise_name,
                    {rel_expr} AS relation,
                    {id_expr} AS id_number,
                    {pol_expr} AS policy_number,
                    {name_expr} AS member_name,
                    age,
                    age_band,
                    {addr_expr} AS client_address_o,
                    import_month,
                    id AS source_row_id,
                    true,
                    raw_data,
                    NOW(),
                    NOW()
                FROM (
                    SELECT DISTINCT ON (
                        LOWER(TRIM(COALESCE(franchise_name,''))),
                        {id_expr},
                        {pol_expr},
                        {rel_expr}
                    ) *
                    FROM policydata_detail_raw
                    WHERE import_month = :latest_month
                      AND {rel_expr} = ANY(:rels)
                    ORDER BY
                        LOWER(TRIM(COALESCE(franchise_name,''))),
                        {id_expr},
                        {pol_expr},
                        {rel_expr},
                        id DESC
                ) latest
                WHERE COALESCE({id_expr}, {pol_expr}, '') <> ''
                ON CONFLICT (member_key) DO UPDATE SET
                    franchise_name = EXCLUDED.franchise_name,
                    relation = EXCLUDED.relation,
                    id_number = EXCLUDED.id_number,
                    policy_number = EXCLUDED.policy_number,
                    member_name = EXCLUDED.member_name,
                    age = EXCLUDED.age,
                    age_band = EXCLUDED.age_band,
                    client_address_o = EXCLUDED.client_address_o,
                    import_month = EXCLUDED.import_month,
                    source_row_id = EXCLUDED.source_row_id,
                    active = true,
                    raw_data = EXCLUDED.raw_data,
                    updated_at = NOW()
            """), {'rels': POLICY_RELATION_CODES, 'latest_month': latest_month})

            processed = getattr(result, 'rowcount', 0) or 0
            _POLICY_CURRENT_REBUILD_STATUS.update({
                'processed': processed,
                'status': 'refreshing relation summary',
                'updated_at': datetime.now().isoformat()
            })
            print(f'Policy current member rebuild: inserted/updated {processed}', flush=True)

            conn.execute(text('TRUNCATE TABLE app_policy_relation_summary'))
            summary_result = conn.execute(text("""
                INSERT INTO app_policy_relation_summary (franchise_name, relation, member_count, updated_at)
                SELECT franchise_name, UPPER(TRIM(relation)) AS relation, COUNT(*)::integer AS member_count, NOW()
                FROM app_policydata_current_members
                WHERE active = true
                  AND UPPER(TRIM(COALESCE(relation,''))) = ANY(:rels)
                GROUP BY franchise_name, UPPER(TRIM(relation))
            """), {'rels': POLICY_RELATION_CODES})

        _POLICY_CURRENT_REBUILD_STATUS.update({
            'running': False,
            'status': 'complete',
            'processed': processed,
            'error': '',
            'updated_at': datetime.now().isoformat()
        })
        print('Policy current member rebuild: complete', flush=True)
    except Exception as exc:
        _POLICY_CURRENT_REBUILD_STATUS.update({
            'running': False,
            'status': 'failed',
            'error': str(exc),
            'updated_at': datetime.now().isoformat()
        })
        print(f'Policy current member rebuild failed: {exc}', flush=True)

def run_daily_policy_age_check():
    """Scheduled daily age-limit notification refresh. Runs inside app context from APScheduler."""
    started = datetime.now()
    status = 'success'
    details = ''
    try:
        ensure_database_schema()
        summary = refresh_policy_relation_summary()
        limit = int(os.getenv('AGE_NOTIFICATION_DAILY_LIMIT', '600000'))
        result = refresh_policy_age_notifications(limit=limit, selected='All')
        details = f"relation summary rows={summary.get('rows', 0)}, processed={result.get('processed', 0)}, notifications={result.get('created', 0)}"
    except Exception as exc:
        status = 'failed'
        details = str(exc)
        print(f'Daily policy age check failed: {exc}')
    try:
        engine = get_db_engine()
        if engine is not None:
            with engine.begin() as conn:
                conn.execute(text("""
                    INSERT INTO app_cron_log (job_name, status, details, created_at)
                    VALUES ('daily_policy_age_check', :status, :details, NOW())
                """), {'status': status, 'details': details[:2000]})
    except Exception as log_exc:
        print(f'Could not write daily policy age check cron log: {log_exc}')
    return {'status': status, 'details': details, 'seconds': (datetime.now() - started).total_seconds()}


_POLICY_AGE_SCHEDULER = None

def start_policy_age_scheduler():
    """Start APScheduler for daily policy age checks. Safe if apscheduler is not installed."""
    global _POLICY_AGE_SCHEDULER
    if _POLICY_AGE_SCHEDULER is not None:
        return _POLICY_AGE_SCHEDULER
    if os.getenv('ENABLE_POLICY_AGE_SCHEDULER', '0') != '1':
        print('Policy age scheduler disabled by ENABLE_POLICY_AGE_SCHEDULER')
        return None
    if app.debug and os.environ.get('WERKZEUG_RUN_MAIN') != 'true':
        return None
    try:
        from apscheduler.schedulers.background import BackgroundScheduler
    except Exception as exc:
        print(f'APScheduler not installed; daily policy age check scheduler disabled: {exc}')
        return None
    hour = int(os.getenv('POLICY_AGE_CHECK_HOUR', '2'))
    minute = int(os.getenv('POLICY_AGE_CHECK_MINUTE', '0'))
    scheduler = BackgroundScheduler(daemon=True)
    def _job():
        with app.app_context():
            run_daily_policy_age_check()
    scheduler.add_job(_job, 'cron', hour=hour, minute=minute, id='daily_policy_age_check', replace_existing=True)
    scheduler.start()
    _POLICY_AGE_SCHEDULER = scheduler
    print(f'Daily policy age check scheduled for {hour:02d}:{minute:02d}')
    return scheduler


def get_policy_age_notifications(selected='All', limit=12):
    engine = get_db_engine()
    if engine is None:
        return []
    params = {'limit': int(limit)}
    where = ["status = 'active'"]
    if selected and selected != 'All':
        where.append("LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))")
        params['franchise'] = selected
    try:
        with engine.begin() as conn:
            rows = [dict(r) for r in conn.execute(text(f"""
                SELECT franchise_name, relation, member_name, policy_number, age_years, age_months, notification_type, message, due_date
                FROM app_policy_relation_age_notifications
                WHERE {' AND '.join(where)}
                ORDER BY due_date NULLS LAST, updated_at DESC
                LIMIT :limit
            """), params).mappings().all()]
            return rows
    except Exception as exc:
        print(f'Could not load age notifications: {exc}')
        return []




def _dedupe_age_notifications_cte():
    """SQL CTE that returns one active age notification per policy/person/relation/age band."""
    return """
        WITH deduped_age_notifications AS (
            SELECT *
            FROM (
                SELECT
                    n.*,
                    ROW_NUMBER() OVER (
                        PARTITION BY
                            LOWER(TRIM(COALESCE(n.franchise_name, ''))),
                            LOWER(TRIM(COALESCE(n.policy_number, ''))),
                            LOWER(TRIM(COALESCE(n.id_number, ''))),
                            LOWER(TRIM(COALESCE(n.member_name, ''))),
                            LOWER(TRIM(COALESCE(n.relation, ''))),
                            LOWER(TRIM(COALESCE(n.age_band, 'Unknown')))
                        ORDER BY n.updated_at DESC NULLS LAST, n.due_date ASC NULLS LAST, n.id DESC NULLS LAST
                    ) AS rn
                FROM app_policy_relation_age_notifications n
                WHERE n.status = 'active'
            ) x
            WHERE rn = 1
        )
    """


def get_policy_age_notification_summary(selected='All', limit=200):
    """Return franchise-level age-warning totals for the dashboard, deduped by policy/person/relation/age band."""
    engine = get_db_engine()
    if engine is None:
        return []
    params = {'limit': int(limit)}
    where = ["1=1"]
    if selected and selected != 'All':
        where.append("LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))")
        params['franchise'] = selected
    try:
        with engine.begin() as conn:
            rows = [dict(r) for r in conn.execute(text(f"""
                {_dedupe_age_notifications_cte()}
                SELECT
                    franchise_name,
                    COUNT(*)::integer AS notification_count,
                    COUNT(*) FILTER (WHERE notification_type = 'CHILD_TURNING_21')::integer AS child_21_count,
                    MIN(due_date) AS next_due_date,
                    MIN(message) AS sample_message
                FROM deduped_age_notifications
                WHERE {' AND '.join(where)}
                GROUP BY franchise_name
                ORDER BY notification_count DESC, franchise_name
                LIMIT :limit
            """), params).mappings().all()]
            return rows
    except Exception as exc:
        print(f'Could not load age notification summary: {exc}')
        return []


def get_policy_age_notification_members(franchise_name, limit=1000, age_band=None):
    """Detailed active age-limit notification list for one franchise, deduped by policy/person/relation/age band."""
    engine = get_db_engine()
    if engine is None:
        return []
    params = {'franchise': franchise_name, 'limit': int(limit)}
    extra = ''
    if age_band:
        extra = " AND COALESCE(NULLIF(TRIM(age_band), ''), 'Unknown') = :age_band"
        params['age_band'] = age_band
    try:
        with engine.begin() as conn:
            return [dict(r) for r in conn.execute(text(f"""
                {_dedupe_age_notifications_cte()}
                SELECT
                    franchise_name, relation, member_name, policy_number, id_number,
                    age_years, age_months, age_band, notification_type, message,
                    due_date, updated_at
                FROM deduped_age_notifications
                WHERE LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))
                  {extra}
                ORDER BY due_date NULLS LAST, relation, member_name, policy_number
                LIMIT :limit
            """), params).mappings().all()]
    except Exception as exc:
        print(f'Could not load age notification member list: {exc}')
        return []


def get_policy_age_notification_franchise_age_groups(franchise_name):
    """Return all age-limit notification bands for one franchise, including zero-count bands, deduped."""
    engine = get_db_engine()
    out = [{'age_band': b, 'count': 0, 'next_due_date': None} for b in AGE_NOTIFICATION_BANDS]
    if engine is None:
        return out
    try:
        with engine.begin() as conn:
            rows = [dict(r) for r in conn.execute(text(f"""
                {_dedupe_age_notifications_cte()}
                SELECT COALESCE(NULLIF(TRIM(age_band), ''), 'Unknown') AS age_band,
                       COUNT(*)::integer AS count,
                       MIN(due_date) AS next_due_date
                FROM deduped_age_notifications
                WHERE LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))
                GROUP BY COALESCE(NULLIF(TRIM(age_band), ''), 'Unknown')
            """), {'franchise': franchise_name}).mappings().all()]
        found = {str(r.get('age_band') or 'Unknown'): r for r in rows}
        return [{'age_band': b, 'count': int((found.get(b) or {}).get('count') or 0), 'next_due_date': (found.get(b) or {}).get('next_due_date')} for b in AGE_NOTIFICATION_BANDS]
    except Exception as exc:
        print(f'Could not load franchise age groups: {exc}')
        return out


def get_policy_age_notification_totals(selected='All'):
    engine = get_db_engine()
    if engine is None:
        return {'total': 0, 'franchises': 0, 'child_21': 0}
    params = {}
    where = ["1=1"]
    if selected and selected != 'All':
        where.append("LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))")
        params['franchise'] = selected
    try:
        with engine.begin() as conn:
            r = conn.execute(text(f"""
                {_dedupe_age_notifications_cte()}
                SELECT
                    COUNT(*)::integer AS total,
                    COUNT(DISTINCT franchise_name)::integer AS franchises,
                    COUNT(*) FILTER (WHERE notification_type = 'CHILD_TURNING_21')::integer AS child_21
                FROM deduped_age_notifications
                WHERE {' AND '.join(where)}
            """), params).mappings().first()
            return dict(r or {'total': 0, 'franchises': 0, 'child_21': 0})
    except Exception as exc:
        print(f'Could not load age notification totals: {exc}')
        return {'total': 0, 'franchises': 0, 'child_21': 0}


AGE_NOTIFICATION_BANDS = [
    '0-11 months', '1-5 years', '6-13 years', '14-21 years',
    '18-64 years', '65-84 years', '18-70 years', '71-85 years', '85 years and older', 'Unknown'
]

def get_policy_age_notification_age_bands(selected='All'):
    # Active age-limit notification counts grouped by age band. Always returns all bands.
    engine = get_db_engine()
    out = [{'age_band': b, 'count': 0} for b in AGE_NOTIFICATION_BANDS]
    if engine is None:
        return out
    params = {}
    where = ["status = 'active'"]
    if selected and selected != 'All':
        where.append("LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))")
        params['franchise'] = selected
    try:
        with engine.begin() as conn:
            rows = conn.execute(text(f"""
                SELECT COALESCE(NULLIF(TRIM(age_band), ''), 'Unknown') AS age_band,
                       COUNT(*)::integer AS count
                FROM app_policy_relation_age_notifications
                WHERE {' AND '.join(where)}
                GROUP BY COALESCE(NULLIF(TRIM(age_band), ''), 'Unknown')
            """), params).mappings().all()
        counts = {str(r.get('age_band') or 'Unknown'): int(r.get('count') or 0) for r in rows}
        return [{'age_band': b, 'count': counts.get(b, 0)} for b in AGE_NOTIFICATION_BANDS]
    except Exception as exc:
        print(f'Could not load age notification age-band totals: {exc}')
        return out

def read_policydata_streaming(path):
    """Fast streaming import for large PolicyData_YYYYMMDD_to_YYYYMMDD files.

    Allocates rows to franchise from the Franchise column, processes only MEM rows,
    and calculates R1 and the 2.1% underwriter fee using the MPIA months paid.
    """
    global LAST_POLICY_IMPORT_SUMMARY, LAST_POLICY_DETAIL_DF
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb[wb.sheetnames[0]]
    header = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
    header_map = {_normalise_header(v): i for i, v in enumerate(header)}

    def idx(*names, fallback=None):
        for name in names:
            key = _normalise_header(name)
            if key in header_map:
                return header_map[key]
        return fallback

    franchise_i = idx('Franchise', 'Branch', 'Participating Group', fallback=0)
    id_number_i = idx('ID Number', 'ID No', 'Identity Number', 'SA ID', fallback=5)  # F
    relation_i = idx('Relation', fallback=8)  # I
    risk_i = idx('AUL Risk', 'Risk', 'Risk Amount', 'Risk Premium', fallback=12)  # M
    retail_i = idx('Retail', 'Retail Premium', fallback=13)  # N in this PolicyData export
    mpia_i = idx('MPIA', fallback=17)  # R
    address_o_i = 14  # Excel column O, zero-based index 14

    agg = {}
    # Keep row-level detail lightweight. The dashboard only needs MEM rows for
    # client counts/map data, so we do not retain non-MEM rows in memory.
    detail_rows = []
    source_display = display_source_filename(path)
    month = month_from_filename(path)
    total_rows = 0
    mem_rows = 0
    skipped_rows = 0
    # Limit OpenPyXL to the columns we actually use. Some PolicyData exports have
    # many formatted/empty columns, and parsing every cell can exceed Render's
    # Gunicorn timeout before our code gets control again.
    max_needed_col = max(franchise_i, id_number_i, relation_i, risk_i, retail_i, mpia_i, address_o_i) + 1
    for row_number, row in enumerate(ws.iter_rows(min_row=2, max_col=max_needed_col, values_only=True), start=2):
        total_rows += 1
        try:
            relation = str(row[relation_i] or '').strip().upper()
        except Exception:
            skipped_rows += 1
            continue
        franchise = str(row[franchise_i] or '').strip() if franchise_i < len(row) else ''
        risk = clean_money(row[risk_i] if risk_i < len(row) else 0)
        retail = clean_money(row[retail_i] if retail_i < len(row) else 0)
        mpia = clean_money(row[mpia_i] if mpia_i < len(row) else 1)
        if mpia <= 0:
            mpia = 1
        single_premium = risk / mpia if mpia else risk
        is_mem = relation == 'MEM'
        r1_fee = mpia * 1.0 if is_mem else 0.0
        risk_after_r1_per_month = max(single_premium - 1.0, 0) if is_mem else 0.0
        underwriter_fee = risk_after_r1_per_month * 0.021 * mpia if is_mem else 0.0
        risk_after_r1 = risk_after_r1_per_month * mpia if is_mem else 0.0
        net_risk = risk_after_r1 - underwriter_fee if is_mem else 0.0
        # Avoid storing a full copy of every Excel row in memory during large imports.
        raw_data = {}
        client_address_o = _clean_map_value(row[address_o_i] if address_o_i < len(row) else '')
        id_number_f = _clean_id_number(row[id_number_i] if id_number_i < len(row) else '')
        if is_mem and franchise and franchise.lower() not in {'nan', 'none'}:
            detail_rows.append({
                'source_file': source_display,
                'import_month': month,
                'row_number': row_number,
                'source_row_key': f'{source_display}|{month}|{row_number}',
                'client_address_o': client_address_o,
                'id_number_f': id_number_f,
                'franchise': franchise,
                'relation': relation,
                'is_mem': is_mem,
                'retail_premium': retail,
                'original_risk_premium': risk,
                'mpia': mpia,
                'single_premium': single_premium,
                'r1_policy_fee': r1_fee,
                'adv_fund_2_1_fee': underwriter_fee,
                'risk_after_r1': risk_after_r1,
                'new_risk_premium': net_risk,
                'raw_data': raw_data,
            })
        if not is_mem:
            continue
        mem_rows += 1
        if not franchise or franchise.lower() in {'nan', 'none'}:
            skipped_rows += 1
            continue
        rec = agg.setdefault(franchise, {
            'retail_premium': 0.0,
            'risk_premium': 0.0,
            'claims': 0.0,
            'policy_qty': 0.0,
            'original_risk_premium': 0.0,
            'r1_policy_fee_imported': 0.0,
            'underwriter_2_1_fee': 0.0,
            'risk_after_r1': 0.0,
            'single_monthly_premium_total': 0.0,
        })
        rec['retail_premium'] += retail
        rec['risk_premium'] += net_risk
        rec['policy_qty'] += mpia
        rec['original_risk_premium'] += risk
        rec['r1_policy_fee_imported'] += r1_fee
        rec['underwriter_2_1_fee'] += underwriter_fee
        rec['risk_after_r1'] += risk_after_r1
        rec['single_monthly_premium_total'] += single_premium

    rows = []
    for franchise, rec in agg.items():
        rows.append({
            'franchise': franchise,
            'month': month,
            'retail_premium': rec['retail_premium'],
            'risk_premium': rec['risk_premium'],
            'claims': 0.0,
            'policy_qty': rec['policy_qty'],
            'original_risk_premium': rec['original_risk_premium'],
            'r1_policy_fee_imported': rec['r1_policy_fee_imported'],
            'underwriter_2_1_fee': rec['underwriter_2_1_fee'],
            'risk_after_r1': rec['risk_after_r1'],
            'single_monthly_premium_total': rec['single_monthly_premium_total'],
            'current_scenario': '100% Claim Ratio',
        })
    out = pd.DataFrame(rows)
    LAST_POLICY_DETAIL_DF = pd.DataFrame(detail_rows)
    try:
        wb.close()
    except Exception:
        pass
    LAST_POLICY_IMPORT_SUMMARY = {
        'file_name': display_source_filename(path),
        'month': month.strftime('%b %Y'),
        'total_rows': int(total_rows),
        'detail_rows_stored': int(len(detail_rows)),
        'mem_rows': int(mem_rows),
        'skipped_rows': int(skipped_rows),
        'franchises_found': int(out['franchise'].nunique()) if not out.empty else 0,
        'total_retail': float(out['retail_premium'].sum()) if not out.empty else 0.0,
        'total_risk': float(out['risk_premium'].sum()) if not out.empty else 0.0,
        'total_original_risk': float(out['original_risk_premium'].sum()) if not out.empty else 0.0,
        'total_policies': float(out['policy_qty'].sum()) if not out.empty else 0.0,
        'total_r1': float(out['r1_policy_fee_imported'].sum()) if not out.empty else 0.0,
        'total_underwriter_fee': float(out['underwriter_2_1_fee'].sum()) if not out.empty else 0.0,
    }
    return out


def looks_like_policydata_file(path):
    name = os.path.basename(path).lower()
    if name.startswith('policydata'):
        return True
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        ws = wb[wb.sheetnames[0]]
        header = next(ws.iter_rows(min_row=1, max_row=1, values_only=True))
        keys = {_normalise_header(v) for v in header}
        return {'franchise', 'relation', 'mpia'}.issubset(keys) and ('aul risk' in keys or 'risk' in keys) and 'retail' in keys
    except Exception:
        return False



def merge_policy_months(existing, incoming):
    """Append monthly PolicyData imports while replacing the same imported months.

    Users import one PolicyData file per month. If the same month is imported again,
    replace that month instead of duplicating it.
    """
    if incoming is None or incoming.empty:
        return existing if existing is not None else pd.DataFrame()
    if existing is None or existing.empty:
        return incoming.copy()
    old = existing.copy()
    new = incoming.copy()
    old['month'] = pd.to_datetime(old['month'], errors='coerce')
    new['month'] = pd.to_datetime(new['month'], errors='coerce')
    months = set(new['month'].dropna().dt.to_period('M').astype(str))
    if months:
        old = old[~old['month'].dt.to_period('M').astype(str).isin(months)].copy()
    combined = pd.concat([old, new], ignore_index=True)
    return combined.sort_values(['franchise', 'month']).reset_index(drop=True)

def read_excel_file(path):
    global LAST_POLICY_IMPORT_SUMMARY
    frames = []

    # Large Martins PolicyData exports are transaction-level files. Use a streaming parser
    # so big monthly files load reliably and allocate MEM rows to franchises.
    if looks_like_policydata_file(path):
        streamed = read_policydata_streaming(path)
        if not streamed.empty:
            df = streamed.copy()
            df['franchise'] = df['franchise'].astype(str).str.strip()
            df = df[df['franchise'] != '']
            df['month'] = pd.to_datetime(df['month'], errors='coerce')
            df = df.dropna(subset=['month'])
            for col in ['retail_premium', 'risk_premium', 'claims', 'claim_count', 'claim_paid_franchise', 'claim_paid_client', 'repudiated_pending', 'grand_total_claims', 'policy_qty', 'original_risk_premium', 'r1_policy_fee_imported', 'underwriter_2_1_fee', 'risk_after_r1', 'single_monthly_premium_total']:
                if col not in df.columns:
                    df[col] = 0
                df[col] = df[col].apply(clean_money)
            if 'current_scenario' not in df.columns:
                df['current_scenario'] = '100% Claim Ratio'
            df['current_scenario'] = df['current_scenario'].fillna('100% Claim Ratio').astype(str)
            return df.sort_values(['franchise', 'month']).reset_index(drop=True)

    # First support the normal long-format import template.
    all_sheets_standard = pd.read_excel(path, sheet_name=None)
    standard_errors = []
    for _, df in all_sheets_standard.items():
        if df.empty:
            continue
        try:
            normalised = normalise_columns(df.copy())
            frames.append(normalised)
        except Exception as exc:
            standard_errors.append(str(exc))

    # If no normal sheets worked, support detailed policy transaction receipt exports.
    if not frames:
        for _, df in all_sheets_standard.items():
            parsed = parse_policy_transaction_sheet(df, path)
            if not parsed.empty:
                frames.append(parsed)

    # If no transaction sheets worked, support the uploaded Martins wide claim-ratio workbook.
    if not frames:
        all_sheets_wide = pd.read_excel(path, sheet_name=None, header=None)
        for _, df in all_sheets_wide.items():
            parsed = parse_wide_claim_ratio_sheet(df)
            if not parsed.empty:
                frames.append(parsed)

    if not frames:
        raise ValueError('No usable data found. Required long-format columns: Franchise, Month, Retail Premium, Risk Premium, Claims; the detailed policy receipt export with Franchise/Relation/Risk/Retail/MPIA; or the Martins wide monthly claim-ratio layout.')

    df = pd.concat(frames, ignore_index=True)
    df['franchise'] = df['franchise'].astype(str).str.strip()
    df = df[df['franchise'] != '']
    df['month'] = pd.to_datetime(df['month'], errors='coerce')
    df = df.dropna(subset=['month'])
    for col in ['retail_premium', 'risk_premium', 'claims', 'claim_count', 'claim_paid_franchise', 'claim_paid_client', 'repudiated_pending', 'grand_total_claims', 'policy_qty', 'original_risk_premium', 'r1_policy_fee_imported', 'underwriter_2_1_fee', 'risk_after_r1', 'single_monthly_premium_total']:
        if col not in df.columns:
            df[col] = 0
        df[col] = df[col].apply(clean_money)
    df['current_scenario'] = df.get('current_scenario', '100% Claim Ratio')
    df['current_scenario'] = df['current_scenario'].fillna('100% Claim Ratio').astype(str)
    return df.sort_values(['franchise', 'month']).reset_index(drop=True)

def status_colour(ratio):
    if ratio < 50:
        return 'Blue'
    if ratio < 70:
        return 'Green'
    return 'Red'


def status_label(ratio):
    if ratio < 50:
        return 'Move to BrightRock'
    if ratio < 70:
        return 'Can move to BrightRock'
    return 'High Risk'


def can_move_to_brightrock(ratio):
    if ratio < 50:
        return 'Yes - move to BrightRock'
    if ratio < 70:
        return 'Can move to BrightRock'
    return 'No - High Risk'


def scenario_100_move_traffic(ratio):
    if ratio < 50:
        return 'Blue'
    if ratio < 70:
        return 'Green'
    return 'Red'


def scenario_100_traffic(claims, risk_premium, total_commission):
    threshold = risk_premium - total_commission
    if claims > threshold:
        return 'Red'
    return 'Green'


def scenario_100_result(claims, risk_premium, total_commission):
    threshold = risk_premium - total_commission
    if claims > threshold:
        return 'Claims over risk limit'
    return 'Claims within risk limit'


def brightrock_money_traffic(value):
    if value > 0:
        return 'Green'
    if value < 0:
        return 'Red'
    return 'Blue'


def brightrock_money_result(value):
    if value > 0:
        return 'Makes money'
    if value < 0:
        return 'Loses money'
    return 'Break even'


def recommendation(ratio, current):
    current_clean = str(current).strip().lower()
    if current_clean == 'brightrock':
        return 'Already on BrightRock'
    if ratio < 75:
        return 'Can move to BrightRock - requires 5 months below 75%'
    return 'High Risk - stay on 100%'


def apply_brightrock_five_month_rule(monthly):
    """A franchise may only move to BrightRock after 5 consecutive months below 75%."""
    if monthly is None or monthly.empty or 'Franchise' not in monthly.columns:
        return monthly
    monthly = monthly.sort_values(['Franchise', 'Month']).copy()
    ratios = pd.to_numeric(monthly.get('Claim Ratio', 0), errors='coerce').fillna(0)
    monthly['_below_75'] = ratios < 75
    streaks = []
    for _, grp in monthly.groupby('Franchise', sort=False):
        streak = 0
        for ok in grp['_below_75'].tolist():
            streak = streak + 1 if ok else 0
            streaks.append(streak)
    monthly['Months Below 75% Streak'] = streaks
    monthly['BrightRock Eligible'] = monthly['Months Below 75% Streak'] >= 5
    def _rec(row):
        current_clean = str(row.get('Current Scenario', '')).strip().lower()
        ratio = float(row.get('Claim Ratio', 0) or 0)
        if current_clean == 'brightrock':
            return 'Already on BrightRock'
        if row.get('BrightRock Eligible', False):
            return 'Move to BrightRock'
        if ratio < 75:
            return 'Can move to BrightRock - needs 5 months below 75%'
        return 'High Risk - stay on 100%'
    monthly['Recommendation'] = monthly.apply(_rec, axis=1)
    monthly['100% Move Status'] = monthly['Recommendation']
    monthly['100% Move Traffic'] = monthly.apply(lambda r: 'Blue' if r.get('BrightRock Eligible', False) else ('Green' if float(r.get('Claim Ratio',0) or 0) < 75 else 'Red'), axis=1)
    monthly.drop(columns=['_below_75'], inplace=True, errors='ignore')
    return monthly

def analyse(df, rates, book_rates):
    total_commission_rate = sum(rates.values())
    rows = []
    for franchise, group in df.groupby('franchise'):
        group = group.sort_values('month').copy().reset_index(drop=True)
        running_balance = 0.0
        prev_claims = 0.0
        for idx, row in group.iterrows():
            retail = row['retail_premium']
            risk = row['risk_premium']
            claims = row['claims']
            claim_count = row.get('claim_count', 0)
            claim_paid_franchise = row.get('claim_paid_franchise', 0)
            claim_paid_client = row.get('claim_paid_client', 0)
            repudiated_pending = row.get('repudiated_pending', 0)
            grand_total_claims = row.get('grand_total_claims', 0)
            policy_qty = row.get('policy_qty', 0)
            original_risk = row.get('original_risk_premium', risk)
            underwriter_2_1_fee = row.get('underwriter_2_1_fee', 0)
            risk_after_r1 = row.get('risk_after_r1', max(original_risk - row.get('r1_policy_fee_imported', 0), 0))
            single_monthly_premium_total = row.get('single_monthly_premium_total', 0)
            imported_r1_fee = row.get('r1_policy_fee_imported', 0)
            r1_policy_fee = imported_r1_fee if imported_r1_fee else policy_qty * 1.0
            commissions = {name: retail * pct / 100 for name, pct in rates.items()}
            percentage_commission = sum(commissions.values())
            total_commission = percentage_commission
            separate_paid_commissions = r1_policy_fee + underwriter_2_1_fee
            total_paid_commissions = total_commission + separate_paid_commissions
            admin_fee = retail - total_paid_commissions - risk
            claim_ratio = (claims / risk * 100) if risk else 0.0
            scenario_one_value = retail - risk - total_paid_commissions
            brightrock_month_total = retail - claims - total_paid_commissions
            scenario_two_value = brightrock_month_total
            running_balance += brightrock_month_total
            mff_book = retail * book_rates.get('MFF Book Value', 0) / 100
            franchise_book = retail * book_rates.get('Franchise Book Value', 0) / 100
            total_book_value = mff_book + franchise_book
            rows.append({
                'Franchise': franchise,
                'Month': row['month'],
                'Period Index': idx // 6 + 1,
                'Period': '',
                'Retail Premium': retail,
                'Original Risk Premium': original_risk,
                'Risk Premium': risk,
                'Risk After R1': risk_after_r1,
                'ADV Fee 2.1%': underwriter_2_1_fee,
                'Underwriter 2.1% Fee': underwriter_2_1_fee,
                'Single Monthly Premium Total': single_monthly_premium_total,
                'Claims': claims,
                'Claim Count': claim_count,
                'Average Claim': (claims / claim_count) if claim_count else 0.0,
                'Claim Paid to Franchise': claim_paid_franchise,
                'Claim Paid to Client': claim_paid_client,
                'Repudiated / Pending': repudiated_pending,
                'Grand Total Claims': grand_total_claims,
                'Policy Qty': policy_qty,
                'R1 Policy Fee': r1_policy_fee,
                'Percentage Commission': percentage_commission,
                'Claim Ratio': claim_ratio,
                'Claim Ratio Status': status_colour(claim_ratio),
                'Claim Ratio Label': status_label(claim_ratio),
                '100% Move Status': can_move_to_brightrock(claim_ratio),
                '100% Move Traffic': scenario_100_move_traffic(claim_ratio),
                '100% Traffic Light': scenario_100_traffic(claims, risk, total_commission),
                '100% Traffic Result': scenario_100_result(claims, risk, total_commission),
                '100% Traffic Reason': 'Claims are more than Risk Premium minus Commissions' if claims > (risk - total_commission) else 'Claims are within Risk Premium minus Commissions',
                'BrightRock Traffic Light': brightrock_money_traffic(brightrock_month_total),
                'BrightRock Money Result': brightrock_money_result(brightrock_month_total),
                'Current Scenario': row['current_scenario'],
                'Recommendation': recommendation(claim_ratio, row['current_scenario']),
                'Total Commission Rate': total_commission_rate,
                'Total Commission': total_commission,
                'Separate Paid Commissions': separate_paid_commissions,
                'Total Paid Commissions': total_paid_commissions,
                'Admin Fee': admin_fee,
                'Scenario 1 Value': scenario_one_value,
                'Scenario 2 Value': scenario_two_value,
                'MFF Book Value 2.5%': mff_book,
                'Franchise Book Value 2.5%': franchise_book,
                'Total Book Value': total_book_value,
                'Scenario 1 Book Value': total_book_value,
                'Scenario 2 Book Value': total_book_value,
                'Remaining Book Value': retail - mff_book - franchise_book,
                'Previous Month Claims': prev_claims,
                'BrightRock Month Total': brightrock_month_total,
                'BrightRock Running Balance': running_balance,
                **{f'{name} Rate': pct for name, pct in rates.items()},
                **{f'{name} Amount': amount for name, amount in commissions.items()},
            })
            prev_claims = claims
    monthly = pd.DataFrame(rows)
    if not monthly.empty and 'Month' in monthly.columns:
        monthly['Month'] = pd.to_datetime(monthly['Month'], errors='coerce')
        monthly = monthly.sort_values(['Franchise', 'Month']).reset_index(drop=True)
        monthly['Period Index'] = monthly.groupby('Franchise').cumcount() // 6 + 1
        monthly['_period_start'] = monthly.groupby(['Franchise', 'Period Index'])['Month'].transform('min')
        monthly['_period_end'] = monthly.groupby(['Franchise', 'Period Index'])['Month'].transform('max')
        monthly['Period'] = monthly.apply(
            lambda r: r['_period_start'].strftime('%b %Y')
            if pd.notna(r['_period_start']) and pd.notna(r['_period_end']) and r['_period_start'].to_period('M') == r['_period_end'].to_period('M')
            else (f"{r['_period_start'].strftime('%b %Y')} - {r['_period_end'].strftime('%b %Y')}" if pd.notna(r['_period_start']) and pd.notna(r['_period_end']) else ''),
            axis=1
        )
        monthly.drop(columns=['_period_start', '_period_end'], inplace=True, errors='ignore')
    monthly = apply_brightrock_five_month_rule(monthly)
    # Advertising Fund commission removed from the current model.
    for _c in ['Advertising Fund Rate', 'Advertising Fund Amount']:
        if _c in monthly.columns:
            monthly.drop(columns=[_c], inplace=True)

    agg = {
        'Retail Premium': 'sum',
        'Original Risk Premium': 'sum',
        'Risk Premium': 'sum',
        'Risk After R1': 'sum',
        'ADV Fee 2.1%': 'sum',
        'Underwriter 2.1% Fee': 'sum',
        'Single Monthly Premium Total': 'sum',
        'Claims': 'sum',
        'Claim Count': 'sum',
        'Average Claim': 'mean',
        'Claim Paid to Franchise': 'sum',
        'Claim Paid to Client': 'sum',
        'Repudiated / Pending': 'sum',
        'Grand Total Claims': 'sum',
        'Policy Qty': 'sum',
        'R1 Policy Fee': 'sum',
        'Percentage Commission': 'sum',
        'Claim Ratio': 'mean',
        'Total Commission': 'sum',
        'Separate Paid Commissions': 'sum',
        'Total Paid Commissions': 'sum',
        'Admin Fee': 'sum',
        'MFF Book Value 2.5%': 'sum',
        'Franchise Book Value 2.5%': 'sum',
        'Total Book Value': 'sum',
        'Remaining Book Value': 'sum',
        'Previous Month Claims': 'sum',
        'Scenario 1 Value': 'sum',
        'Scenario 2 Value': 'sum',
        'Scenario 1 Book Value': 'sum',
        'Scenario 2 Book Value': 'sum',
        'BrightRock Month Total': 'sum',
        'BrightRock Running Balance': 'last',
        'Months Below 75% Streak': 'last',
        'BrightRock Eligible': 'max',
    }
    for name in rates:
        agg[f'{name} Amount'] = 'sum'

    periods = monthly.groupby(['Franchise', 'Period'], as_index=False).agg(agg)
    periods['Average Claim'] = periods.apply(lambda r: (r['Claims'] / r['Claim Count']) if r.get('Claim Count', 0) else 0, axis=1)
    periods['Weighted Claim Ratio'] = periods.apply(lambda r: (r['Claims'] / r['Risk Premium'] * 100) if r['Risk Premium'] else 0, axis=1)
    periods['Claim Ratio Status'] = periods['Weighted Claim Ratio'].apply(status_colour)
    periods['Claim Ratio Label'] = periods['Weighted Claim Ratio'].apply(status_label)
    periods['100% Move Status'] = periods['Weighted Claim Ratio'].apply(can_move_to_brightrock)
    periods['100% Move Traffic'] = periods['Weighted Claim Ratio'].apply(scenario_100_move_traffic)
    periods['100% Traffic Light'] = periods.apply(lambda r: scenario_100_traffic(r['Claims'], r['Risk Premium'], r.get('Total Paid Commissions', r['Total Commission'])), axis=1)
    periods['100% Traffic Result'] = periods.apply(lambda r: scenario_100_result(r['Claims'], r['Risk Premium'], r.get('Total Paid Commissions', r['Total Commission'])), axis=1)
    periods['100% Traffic Reason'] = periods.apply(lambda r: 'Claims are more than Risk Premium minus Commissions' if r['Claims'] > (r['Risk Premium'] - r.get('Total Paid Commissions', r['Total Commission'])) else 'Claims are within Risk Premium minus Commissions', axis=1)
    periods['BrightRock Traffic Light'] = periods['BrightRock Month Total'].apply(brightrock_money_traffic)
    periods['BrightRock Money Result'] = periods['BrightRock Month Total'].apply(brightrock_money_result)
    periods['Recommendation'] = periods.apply(lambda r: 'Move to BrightRock' if bool(r.get('BrightRock Eligible', False)) else ('Can move to BrightRock - needs 5 months below 75%' if float(r.get('Weighted Claim Ratio', 0) or 0) < 75 else 'High Risk - stay on 100%'), axis=1)

    for metric in ['Retail Premium', 'Claims', 'Total Book Value', 'BrightRock Month Total', 'Scenario 1 Value', 'Scenario 2 Value']:
        total_col = periods.groupby('Period')[metric].transform('sum')
        contribution_col = metric.replace(' ', ' ') + ' Contribution %'
        periods[contribution_col] = periods[metric] / total_col * 100
        periods.loc[total_col == 0, contribution_col] = 0

    portfolio = {
        'total_franchises': int(monthly['Franchise'].nunique()) if not monthly.empty else 0,
        'total_retail': float(monthly['Retail Premium'].sum()) if not monthly.empty else 0,
        'total_original_risk': float(monthly['Original Risk Premium'].sum()) if not monthly.empty and 'Original Risk Premium' in monthly.columns else 0,
        'total_risk': float(monthly['Risk Premium'].sum()) if not monthly.empty else 0,
        'total_underwriter_2_1_fee': float(monthly['Underwriter 2.1% Fee'].sum()) if not monthly.empty and 'Underwriter 2.1% Fee' in monthly.columns else 0,
        'total_claims': float(monthly['Claims'].sum()) if not monthly.empty else 0,
        'total_claim_count': float(monthly['Claim Count'].sum()) if not monthly.empty and 'Claim Count' in monthly.columns else 0,
        'total_claim_paid_franchise': float(monthly['Claim Paid to Franchise'].sum()) if not monthly.empty and 'Claim Paid to Franchise' in monthly.columns else 0,
        'total_claim_paid_client': float(monthly['Claim Paid to Client'].sum()) if not monthly.empty and 'Claim Paid to Client' in monthly.columns else 0,
        'total_repudiated_pending': float(monthly['Repudiated / Pending'].sum()) if not monthly.empty and 'Repudiated / Pending' in monthly.columns else 0,
        'total_grand_total_claims': float(monthly['Grand Total Claims'].sum()) if not monthly.empty and 'Grand Total Claims' in monthly.columns else 0,
        'overall_claim_ratio': average_claim_ratio_for(periods),
        'average_claim_ratio': average_claim_ratio_for(periods),
        'total_commission': float(monthly['Total Commission'].sum()) if not monthly.empty else 0,
        'total_paid_commissions': float(monthly['Total Paid Commissions'].sum()) if not monthly.empty and 'Total Paid Commissions' in monthly.columns else 0,
        'total_brightrock_commission': float(monthly['BrightRock Amount'].sum()) if not monthly.empty and 'BrightRock Amount' in monthly.columns else 0,
        'total_mkhulu_commission': float(monthly['Inkulu Amount'].sum()) if not monthly.empty and 'Inkulu Amount' in monthly.columns else 0,
        'total_inkulu_commission': float(monthly['Inkulu Amount'].sum()) if not monthly.empty and 'Inkulu Amount' in monthly.columns else 0,
        'total_mff_commission': float(monthly['MFF Amount'].sum()) if not monthly.empty and 'MFF Amount' in monthly.columns else 0,
        'total_r1_policy_fee': float(monthly['R1 Policy Fee'].sum()) if not monthly.empty and 'R1 Policy Fee' in monthly.columns else 0,
        'total_policy_qty': float(monthly['Policy Qty'].sum()) if not monthly.empty and 'Policy Qty' in monthly.columns else 0,
        'total_admin_fee': float(monthly['Admin Fee'].sum()) if not monthly.empty else 0,
        'total_scenario_one_value': float(monthly['Scenario 1 Value'].sum()) if not monthly.empty else 0,
        'total_scenario_two_value': float(monthly['Scenario 2 Value'].sum()) if not monthly.empty else 0,
        'total_scenario_one_book_value': float(monthly['Scenario 1 Book Value'].sum()) if not monthly.empty else 0,
        'total_scenario_two_book_value': float(monthly['Scenario 2 Book Value'].sum()) if not monthly.empty else 0,
        'total_brightrock_accumulated': float(monthly.groupby('Franchise')['BrightRock Running Balance'].last().sum()) if not monthly.empty else 0,
        'total_brightrock_month_total': float(monthly['BrightRock Month Total'].sum()) if not monthly.empty else 0,
        'total_mff_book_value': float(monthly['MFF Book Value 2.5%'].sum()) if not monthly.empty else 0,
        'total_franchise_book_value': float(monthly['Franchise Book Value 2.5%'].sum()) if not monthly.empty else 0,
        'total_book_value': float(monthly['Total Book Value'].sum()) if not monthly.empty else 0,
        'move_count': int((periods['Recommendation'] == 'Move to BrightRock').sum()) if not periods.empty else 0,
        'can_move_count': int((periods['Recommendation'].astype(str).str.contains('Can move to BrightRock')).sum()) if not periods.empty else 0,
        'high_risk_count': int((periods['Recommendation'].astype(str).str.contains('High Risk')).sum()) if not periods.empty else 0,
        'blue_count': int((periods['Claim Ratio Status'] == 'Blue').sum()) if not periods.empty else 0,
        'green_count': int((periods['Claim Ratio Status'] == 'Green').sum()) if not periods.empty else 0,
        'red_count': int((periods['Claim Ratio Status'] == 'Red').sum()) if not periods.empty else 0,
        'scenario_100_red_count': int((periods['100% Traffic Light'] == 'Red').sum()) if not periods.empty else 0,
        'scenario_100_green_count': int((periods['100% Traffic Light'] == 'Green').sum()) if not periods.empty else 0,
        'brightrock_green_count': int((periods['BrightRock Traffic Light'] == 'Green').sum()) if not periods.empty else 0,
        'brightrock_red_count': int((periods['BrightRock Traffic Light'] == 'Red').sum()) if not periods.empty else 0,
    }
    return monthly, periods, portfolio




def build_quick_filter_view(monthly_view, periods_view, rates, book_rates, quick_filter, period_view, traffic_filter):
    """Build dashboard quick filter rows with chronological month sorting."""
    base = _period_base(monthly_view, periods_view, period_view, traffic_filter)
    if base is None or base.empty:
        return pd.DataFrame(), []
    base = base.copy()
    # Add auditable date columns for pages such as Move Recommendations.
    # These are display-only fields derived from the selected period rows; they do
    # not change imported data or the database.
    def _period_bounds_from_label(label):
        text_value = str(label or '').strip()
        if not text_value:
            return pd.NaT, pd.NaT
        parts = [p.strip() for p in text_value.split(' - ')]
        try:
            if len(parts) == 2:
                start_dt = pd.to_datetime(parts[0], format='%b %Y', errors='coerce')
                end_dt = pd.to_datetime(parts[1], format='%b %Y', errors='coerce')
            elif len(text_value) == 4 and text_value.isdigit():
                start_dt = pd.Timestamp(year=int(text_value), month=1, day=1)
                end_dt = pd.Timestamp(year=int(text_value), month=12, day=1)
            else:
                start_dt = pd.to_datetime(text_value, format='%b %Y', errors='coerce')
                end_dt = start_dt
        except Exception:
            start_dt, end_dt = pd.NaT, pd.NaT
        return start_dt, end_dt

    if 'Period View' in base.columns:
        bounds = base['Period View'].apply(_period_bounds_from_label)
        base['_analysis_start_dt'] = bounds.apply(lambda x: x[0])
        base['_analysis_end_dt'] = bounds.apply(lambda x: x[1])
    elif 'Month' in base.columns:
        base['_analysis_start_dt'] = pd.to_datetime(base['Month'], errors='coerce')
        base['_analysis_end_dt'] = base['_analysis_start_dt']
    else:
        base['_analysis_start_dt'] = pd.NaT
        base['_analysis_end_dt'] = pd.NaT
    base['Start Date'] = base['_analysis_start_dt'].apply(lambda d: d.strftime('%b %Y') if pd.notna(d) else '')
    base['End Date'] = base['_analysis_end_dt'].apply(lambda d: d.strftime('%b %Y') if pd.notna(d) else '')
    base['Months Analysed'] = base.apply(
        lambda r: int((r['_analysis_end_dt'].to_period('M') - r['_analysis_start_dt'].to_period('M')).n + 1)
        if pd.notna(r['_analysis_start_dt']) and pd.notna(r['_analysis_end_dt']) else 0,
        axis=1
    )
    base.drop(columns=['_analysis_start_dt','_analysis_end_dt'], inplace=True, errors='ignore')
    base['Retail Minus Risk Premium'] = base.get('Retail Premium', 0) - base.get('Risk Premium', 0)
    base['Accumulation Value'] = base.get('Retail Premium', 0) + base.get('Claims', 0) - base.get('Total Commission', 0)
    base['Accumulated Total'] = base.groupby('Franchise')['Accumulation Value'].cumsum() if 'Franchise' in base.columns else base['Accumulation Value'].cumsum()
    if 'Weighted Claim Ratio' not in base.columns:
        base['Weighted Claim Ratio'] = base.get('Claim Ratio', 0)

    if quick_filter == 'retail_minus_risk':
        cols = ['Franchise', 'Period View', 'Retail Premium', 'Risk Premium', 'Retail Minus Risk Premium', 'Claim Ratio Label']
    elif quick_filter == 'claim_ratio':
        base['Selected Claim Ratio'] = base.get('Claim Ratio', base.get('Weighted Claim Ratio', 0))
        cols = ['Franchise', 'Period View', 'Claims', 'Risk Premium', 'Selected Claim Ratio', 'Claim Ratio Label']
    elif quick_filter == 'commissions':
        cols = ['Franchise', 'Period View', 'Retail Premium', 'Original Risk Premium', 'Risk Premium', 'Policy Qty', 'Underwriter 2.1% Fee']
        for name in rates:
            cols.append(f'{name} Amount')
        cols += ['R1 Policy Fee', 'Total Commission', 'MFF Book Value 2.5%', 'Franchise Book Value 2.5%', 'Total Book Value']
    elif quick_filter == 'accumulation':
        cols = ['Franchise', 'Period View', 'Retail Premium', 'Claims', 'Total Commission', 'Accumulation Value', 'Accumulated Total', 'Claim Ratio Label']
    else:
        cols = ['Franchise', 'Period View', 'Retail Premium', 'Risk Premium', 'Claims', 'Claim Ratio Label']
    cols = [c for c in cols if c in base.columns]
    return base[cols], cols


def _float_arg(args, name, default):
    return safe_float(args.get(name), default)




def _fmt_month_label(value):
    try:
        return pd.to_datetime(value).strftime('%b %Y')
    except Exception:
        return str(value or '')




def _clean_map_value(value):
    """Return a Google Maps-friendly text value, or blank for PolicyData placeholders."""
    text_value = str(value or '').strip()
    text_value = text_value.replace('_x0000_', '').replace('\x00', '').strip()
    text_value = re.sub(r'\s+', ' ', text_value)
    bad_values = {'', 'nan', 'none', 'null', '0', '-', 'n/a', 'na', 'same'}
    if text_value.lower() in bad_values:
        return ''
    return text_value


def _first_raw_value(raw_data, names):
    """Find the first usable value in imported raw JSON data by fuzzy column name."""
    if not isinstance(raw_data, dict):
        try:
            raw_data = json.loads(raw_data or '{}')
        except Exception:
            raw_data = {}
    wanted = [re.sub(r'[^a-z0-9]+', '', str(n).lower()) for n in names]
    for key, value in raw_data.items():
        norm = re.sub(r'[^a-z0-9]+', '', str(key).lower())
        if any(w and (w == norm or w in norm or norm in w) for w in wanted):
            text_value = _clean_map_value(value)
            if text_value:
                return text_value
    return ''


def _raw_value_by_excel_column(raw_data, col_letter):
    """Return a value by Excel column letter from stored PolicyData raw_data.

    PolicyData client/member addresses are supplied in Excel column O.  The import
    stores raw_data as an ordered JSON object using the workbook headers.  This
    helper lets the map/cache system read the original column position even when
    the header name is not a predictable value such as Address or Physical Address.
    """
    if not isinstance(raw_data, dict):
        try:
            raw_data = json.loads(raw_data or '{}')
        except Exception:
            raw_data = {}
    col = str(col_letter or '').strip().upper()
    if not col:
        return ''
    index = 0
    for ch in col:
        if 'A' <= ch <= 'Z':
            index = index * 26 + (ord(ch) - ord('A') + 1)
    if index <= 0:
        return ''
    values = list(raw_data.values())
    if len(values) >= index:
        return _clean_map_value(values[index - 1])
    # Fallback for importers that stored generic names such as Column_15.
    for key, value in raw_data.items():
        key_norm = re.sub(r'[^a-z0-9]+', '', str(key).lower())
        if key_norm in {f'column{index}', f'col{index}', col.lower()}:
            text_value = _clean_map_value(value)
            if text_value:
                return text_value
    return ''


def _client_address_from_raw(raw_data, direct_address=''):
    """Extract a clean, Google Maps-friendly client address from PolicyData raw_data.

    The Martins PolicyData import has Address, Area, Suburb and Code columns. Many
    rows contain placeholders such as SAME or _x0000_. Those are ignored so the
    map never receives broken addresses like "_x0000_, South Africa". When a street
    address is not available, the system falls back to Suburb + Area + Code, which
    still creates a valid heat point for density reporting.
    """
    parts = []
    # Martins PolicyData client/member address is in Excel column O.  Use this as
    # the primary address source, then fall back to fuzzy header names for older
    # imports or manual templates.
    policydata_col_o_address = _clean_map_value(direct_address) or _raw_value_by_excel_column(raw_data, 'O')
    street = policydata_col_o_address or _first_raw_value(raw_data, [
        'street address', 'address line 1', 'address1', 'line 1', 'address',
        'client address', 'member address', 'residential address', 'physical address', 'postal address'
    ])
    suburb = _first_raw_value(raw_data, ['suburb', 'township'])
    area = _first_raw_value(raw_data, ['area', 'city', 'town'])
    province = _first_raw_value(raw_data, ['province', 'state'])
    code = _first_raw_value(raw_data, ['code', 'postal code', 'postcode', 'zip'])

    for value in [street, suburb, area, province, code]:
        value = _clean_map_value(value)
        if value and value.lower() not in {p.lower() for p in parts}:
            parts.append(value)
    if parts and not any('south africa' in p.lower() for p in parts):
        parts.append('South Africa')
    # Require at least a suburb/area/street plus country; country-only is not useful.
    return ', '.join(parts) if len(parts) >= 2 else ''


def _client_locality_address_from_raw(raw_data):
    """Return a geocodable locality-level address for heatmap density.

    Full client street addresses are too many to geocode all at once for the All
    view. This groups clients by suburb/town/province/postal code so the heatmap
    can show national density quickly while keeping the count weight accurate.
    """
    suburb = _clean_map_value(_first_raw_value(raw_data, ['suburb', 'township']))
    area = _clean_map_value(_first_raw_value(raw_data, ['area', 'city', 'town']))
    province = _clean_map_value(_first_raw_value(raw_data, ['province', 'state']))
    code = _clean_map_value(_first_raw_value(raw_data, ['code', 'postal code', 'postcode', 'zip']))
    parts = []
    for value in [suburb, area, province, code]:
        if value and value.lower() not in {p.lower() for p in parts}:
            parts.append(value)
    if parts:
        parts.append('South Africa')
    return ', '.join(parts) if len(parts) >= 2 else ''


def _map_franchise_key(value):
    """Normalize franchise names for client-map filtering. Handles extra spaces and punctuation."""
    return re.sub(r'[^a-z0-9]+', ' ', str(value or '').lower()).strip()


def _get_client_map_locations(selected='All', limit=250000, density_mode='auto'):
    """Return imported client/member addresses for the Google map panel.

    The system does not store latitude/longitude yet, so the browser geocodes these
    address strings with Google Maps when GOOGLE_MAPS_API_KEY is configured. Without
    an API key, the dashboard still shows a Google Maps search/embed and direct links.
    """
    engine = get_db_engine()
    if engine is None:
        return []
    rows_out = []
    try:
        params = {'limit': int(limit)}
        where = "franchise_name IS NOT NULL AND TRIM(franchise_name) <> ''"
        # For one franchise, filter in SQL with the same punctuation/spacing tolerant
        # key used in Python. This prevents the selected franchise from being missed
        # when old imports are beyond the general dashboard limit.
        selected_key = _map_franchise_key(selected) if selected and selected != 'All' else ''
        selected_terms = [t for t in selected_key.split() if t not in {'martins', 'funeral', 'funerals'}] if selected_key else []
        if selected_key:
            # Use a broad SQL pre-filter by the most specific token (usually the town),
            # then do the tolerant final match in Python.
            sql_token = selected_terms[-1] if selected_terms else selected_key.split()[-1]
            where += " AND TRIM(REGEXP_REPLACE(LOWER(franchise_name), '[^a-z0-9]+', ' ', 'g')) LIKE :selected_like"
            params['selected_like'] = '%' + sql_token + '%'
            params['limit'] = int(os.getenv('FRANCHISE_MAP_ADDRESS_LIMIT', '200000'))
        with engine.begin() as conn:
            rows = conn.execute(text(f"""
                SELECT franchise_name, client_address_o, raw_data, created_at
                FROM policydata_detail_raw
                WHERE {where}
                ORDER BY created_at DESC, id DESC
                LIMIT :limit
            """), params).mappings().all()
        grouped = {}
        for r in rows:
            franchise = str(r.get('franchise_name') or '').strip()
            if selected_key:
                franchise_key = _map_franchise_key(franchise)
                if franchise_key != selected_key:
                    terms = selected_terms or selected_key.split()
                    if terms and not all(t in franchise_key.split() for t in terms):
                        continue
            raw_data = r.get('raw_data') or {}
            direct_address = r.get('client_address_o') or ''
            relation = _first_raw_value(raw_data, ['relation', 'member relation', 'rel'])
            # Use main member rows for map density. Dependant rows usually carry blank
            # placeholders and would duplicate the same household location.
            if relation and relation.upper() != 'MEM':
                continue
            full_address = _client_address_from_raw(raw_data, direct_address=direct_address)
            locality_address = _client_locality_address_from_raw(raw_data)
            # All-franchise heatmaps must be fast and meaningful. Use locality-level
            # grouping for All so 171k policies become weighted heat points instead
            # of 171k geocoding calls. For a selected franchise, keep full addresses
            # where available so pins can represent individual client residences.
            use_locality = (str(selected or 'All') == 'All' and density_mode in {'auto', 'heat'})
            address = locality_address if use_locality else (full_address or locality_address)
            if not address:
                continue
            client_name = ' '.join([_first_raw_value(raw_data, ['full names', 'first names', 'member name', 'client name', 'policy holder', 'insured name']), _first_raw_value(raw_data, ['surname'])]).strip()
            policy_no = _first_raw_value(raw_data, ['policy number', 'policy no', 'policy', 'policynumber', 'system ref'])
            key = (franchise.lower() if not use_locality else 'all', address.lower())
            item = grouped.setdefault(key, {
                'franchise': franchise if not use_locality else 'All franchises',
                'client': client_name if not use_locality else 'Client density area',
                'policy_number': policy_no if not use_locality else '',
                'address': address,
                'full_address': full_address,
                'locality_address': locality_address,
                'count': 0,
                'maps_url': 'https://www.google.com/maps/search/?api=1&query=' + quote_plus(address),
            })
            item['count'] += 1
            if client_name and not item.get('client'):
                item['client'] = client_name
            if policy_no and not item.get('policy_number'):
                item['policy_number'] = policy_no
        rows_out = sorted(grouped.values(), key=lambda x: (-int(x.get('count') or 0), x.get('franchise',''), x.get('address','')))
    except Exception as exc:
        print(f'Could not load client address map data: {exc}')
    return rows_out


def _get_client_map_source_batch(selected='All', offset=0, limit=250):
    """Return grouped, unprocessed PolicyData address work for map preparation.

    This is built for very large books (2.5M+ policy rows): it reads only a small
    unprocessed window, groups duplicate franchise+address rows in Python, and the
    caller geocodes each unique group once.  A processed-source ledger prevents
    double-counting if the admin page is restarted.
    """
    engine = get_db_engine()
    if engine is None:
        return []
    selected = str(selected or 'All').strip() or 'All'
    limit = max(25, min(int(limit or 250), 2000))
    params = {'limit': limit}
    selected_key = _map_franchise_key(selected) if selected != 'All' else ''
    selected_terms = [t for t in selected_key.split() if t not in {'martins', 'funeral', 'funerals'}] if selected_key else []
    # MEM-only map preparation. Use the indexed/imported relation column first so
    # All-mode does not waste cycles on dependant/non-member rows. Also require a
    # usable PolicyData Column O/Address value before this row enters the work queue.
    # This prevents the All background job from stopping on a window of rows that
    # cannot produce map points while later franchises still have valid MEM addresses.
    where = """
        p.franchise_name IS NOT NULL
        AND TRIM(p.franchise_name) <> ''
        AND ps.source_id IS NULL
        AND UPPER(TRIM(COALESCE(p.relation,''))) = 'MEM'
        AND COALESCE(
            NULLIF(TRIM(COALESCE(p.client_address_o,'')), ''),
            NULLIF(TRIM(COALESCE(p.raw_data->>'Address','')), ''),
            NULLIF(TRIM(COALESCE(p.raw_data->>'ADDRESS','')), ''),
            NULLIF(TRIM(COALESCE(p.raw_data->>'address','')), '')
        ) IS NOT NULL
    """
    if selected_key:
        token = selected_terms[-1] if selected_terms else selected_key.split()[-1]
        where += " AND TRIM(REGEXP_REPLACE(LOWER(p.franchise_name), '[^a-z0-9]+', ' ', 'g')) LIKE :selected_like"
        params['selected_like'] = '%' + token + '%'
    try:
        with engine.begin() as conn:
            _ensure_client_map_points(conn)
            rows = conn.execute(text(f"""
                SELECT p.id, p.franchise_name, p.relation, p.client_address_o, p.raw_data, p.created_at
                FROM policydata_detail_raw p
                LEFT JOIN app_client_map_processed_source ps ON ps.source_id = p.id
                WHERE {where}
                ORDER BY p.id ASC
                LIMIT :limit
            """), params).mappings().all()
    except Exception as exc:
        print(f'Could not read client map source batch: {exc}')
        return []
    grouped = {}
    for r in rows:
        franchise = str(r.get('franchise_name') or '').strip()
        if selected_key:
            franchise_key = _map_franchise_key(franchise)
            terms = selected_terms or selected_key.split()
            if franchise_key != selected_key and not (terms and all(t in franchise_key.split() for t in terms)):
                continue
        raw_data = r.get('raw_data') or {}
        direct_address = r.get('client_address_o') or ''
        relation = str(r.get('relation') or _first_raw_value(raw_data, ['relation', 'member relation', 'rel']) or '').strip()
        if relation and relation.upper() != 'MEM':
            continue
        full_address = _client_address_from_raw(raw_data, direct_address=direct_address)
        locality_address = _client_locality_address_from_raw(raw_data)
        address = full_address or locality_address
        if not address:
            continue
        point_key = _client_map_point_key(franchise, address)
        if not point_key:
            continue
        client_name = ' '.join([_first_raw_value(raw_data, ['full names', 'first names', 'member name', 'client name', 'policy holder', 'insured name']), _first_raw_value(raw_data, ['surname'])]).strip()
        policy_no = _first_raw_value(raw_data, ['policy number', 'policy no', 'policy', 'policynumber', 'system ref'])
        item = grouped.setdefault(point_key, {
            'franchise': franchise,
            'client': client_name,
            'policy_number': policy_no,
            'address': address,
            'full_address': full_address,
            'locality_address': locality_address,
            'count': 0,
            'source_ids': [],
            'source_id': int(r.get('id') or 0),
        })
        item['count'] += 1
        item['source_ids'].append(int(r.get('id') or 0))
        item['source_id'] = max(int(item.get('source_id') or 0), int(r.get('id') or 0))
        if client_name and not item.get('client'):
            item['client'] = client_name
        if policy_no and not item.get('policy_number'):
            item['policy_number'] = policy_no
    return sorted(grouped.values(), key=lambda x: (x.get('source_id') or 0))


# -----------------------------------------------------------------------------
# Server-side Google geocoding cache for client heatmap
# -----------------------------------------------------------------------------
def _ensure_client_geocode_cache(conn):
    try:
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS app_client_geocode_cache (
                address_key TEXT PRIMARY KEY,
                address TEXT,
                lat NUMERIC(12,8),
                lng NUMERIC(12,8),
                status TEXT,
                updated_at TIMESTAMP DEFAULT NOW()
            )
        """))
    except Exception as exc:
        print(f'Could not ensure geocode cache: {exc}')


def _address_cache_key(address):
    return re.sub(r'[^a-z0-9]+', ' ', str(address or '').lower()).strip()[:500]


def _client_map_point_key(franchise, address):
    # Address cache keys are per-franchise so the same town/address can represent
    # different franchise books without overwriting each other.
    fkey = _map_franchise_key(franchise) or 'unknown'
    akey = _address_cache_key(address)
    return (fkey + '|' + akey)[:700] if akey else ''


def _ensure_client_map_points(conn):
    """Persistent, query-fast map points table used by dashboard/client heatmaps.

    Scales to millions of policy rows by storing only unique franchise+address
    coordinate groups and a processed-source ledger so repeated preparation runs do
    not double-count policy rows.
    """
    try:
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS app_client_map_points (
                id BIGSERIAL PRIMARY KEY,
                franchise_name TEXT,
                franchise_key TEXT,
                address_key TEXT UNIQUE,
                display_address TEXT,
                client_name TEXT,
                policy_number TEXT,
                client_count INTEGER DEFAULT 1,
                lat NUMERIC(12,8),
                lng NUMERIC(12,8),
                status TEXT DEFAULT 'OK',
                updated_at TIMESTAMP DEFAULT NOW(),
                created_at TIMESTAMP DEFAULT NOW()
            )
        """))
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS app_client_map_processed_source (
                source_id BIGINT PRIMARY KEY,
                franchise_key TEXT,
                address_key TEXT,
                processed_at TIMESTAMP DEFAULT NOW()
            )
        """))
        conn.execute(text("ALTER TABLE app_client_map_points ADD COLUMN IF NOT EXISTS franchise_key TEXT"))
        conn.execute(text("ALTER TABLE app_client_map_points ADD COLUMN IF NOT EXISTS client_count INTEGER DEFAULT 1"))
        conn.execute(text("ALTER TABLE app_client_map_processed_source ADD COLUMN IF NOT EXISTS franchise_key TEXT"))
        conn.execute(text("ALTER TABLE app_client_map_processed_source ADD COLUMN IF NOT EXISTS address_key TEXT"))
        conn.execute(text("UPDATE app_client_map_points SET franchise_key = LOWER(TRIM(REGEXP_REPLACE(COALESCE(franchise_name,''), '[^a-z0-9]+', ' ', 'g'))) WHERE franchise_key IS NULL OR franchise_key = ''"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_app_client_map_points_franchise ON app_client_map_points (LOWER(TRIM(franchise_name)))"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_app_client_map_points_franchise_key ON app_client_map_points (franchise_key)"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_app_client_map_points_latlng ON app_client_map_points (lat, lng)"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_app_client_map_processed_franchise_key ON app_client_map_processed_source (franchise_key)"))
        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_app_client_map_processed_address_key ON app_client_map_processed_source (address_key)"))
    except Exception as exc:
        print(f'Could not ensure client map points: {exc}')


def _upsert_client_map_point(item, address, loc):
    """Save or increment prepared client/member coordinates.

    One row represents one unique franchise+address group.  client_count is the
    number of policy/client rows represented by that location, so 2.5M policy rows
    can be rendered as a manageable weighted heatmap.
    """
    if not loc or loc.get('lat') is None or loc.get('lng') is None:
        return False
    franchise = str(item.get('franchise') or item.get('franchise_name') or '').strip()
    key = _client_map_point_key(franchise, address)
    if not key:
        return False
    engine = get_db_engine()
    if engine is None:
        return False
    try:
        with engine.begin() as conn:
            _ensure_client_map_points(conn)
            conn.execute(text("""
                INSERT INTO app_client_map_points (
                    franchise_name, franchise_key, address_key, display_address, client_name, policy_number,
                    client_count, lat, lng, status, updated_at
                ) VALUES (
                    :franchise_name, :franchise_key, :address_key, :display_address, :client_name, :policy_number,
                    :client_count, :lat, :lng, :status, NOW()
                )
                ON CONFLICT (address_key) DO UPDATE SET
                    franchise_name = EXCLUDED.franchise_name,
                    franchise_key = EXCLUDED.franchise_key,
                    display_address = EXCLUDED.display_address,
                    client_name = COALESCE(NULLIF(app_client_map_points.client_name,''), EXCLUDED.client_name),
                    policy_number = COALESCE(NULLIF(app_client_map_points.policy_number,''), EXCLUDED.policy_number),
                    client_count = COALESCE(app_client_map_points.client_count,0) + COALESCE(EXCLUDED.client_count,0),
                    lat = EXCLUDED.lat,
                    lng = EXCLUDED.lng,
                    status = EXCLUDED.status,
                    updated_at = NOW()
            """), {
                'franchise_name': franchise,
                'franchise_key': _map_franchise_key(franchise),
                'address_key': key,
                'display_address': str(address or '').strip(),
                'client_name': str(item.get('client') or '')[:500],
                'policy_number': str(item.get('policy_number') or '')[:200],
                'client_count': int(float(item.get('count') or 1)),
                'lat': float(loc['lat']),
                'lng': float(loc['lng']),
                'status': str(loc.get('status') or 'OK')[:100],
            })
        return True
    except Exception as exc:
        print(f'Could not save client map point: {exc}')
        return False


def _mark_client_map_sources_processed(source_ids, franchise, address):
    if not source_ids:
        return False
    engine = get_db_engine()
    if engine is None:
        return False
    franchise_key = _map_franchise_key(franchise)
    address_key = _client_map_point_key(franchise, address)
    rows = [{'source_id': int(x), 'franchise_key': franchise_key, 'address_key': address_key} for x in source_ids if x]
    if not rows:
        return False
    try:
        with engine.begin() as conn:
            _ensure_client_map_points(conn)
            conn.execute(text("""
                INSERT INTO app_client_map_processed_source (source_id, franchise_key, address_key, processed_at)
                VALUES (:source_id, :franchise_key, :address_key, NOW())
                ON CONFLICT (source_id) DO NOTHING
            """), rows)
        return True
    except Exception as exc:
        print(f'Could not mark client map source rows processed: {exc}')
        return False


def _query_client_map_points(selected='All', offset=0, limit=500, mode='auto'):
    """Fast map API source: read pre-prepared coordinates only.

    For selected franchises this uses the indexed franchise_key column. This prevents
    freezes caused by LOWER/TRIM comparisons over many rows and avoids raw PolicyData
    scans completely.
    """
    engine = get_db_engine()
    if engine is None:
        return {'total_groups': 0, 'clients_total': 0, 'locations': []}
    selected = str(selected or 'All').strip() or 'All'
    limit = max(50, min(int(limit or 500), 3000))
    offset = max(0, int(offset or 0))
    params = {'limit': limit, 'offset': offset}
    where = "WHERE lat IS NOT NULL AND lng IS NOT NULL"
    selected_key = ''
    if selected != 'All':
        selected_key = _map_franchise_key(selected)
        where += " AND franchise_key = :selected_key"
        params['selected_key'] = selected_key
    try:
        with engine.begin() as conn:
            _ensure_client_map_points(conn)
            # In case rows were created before franchise_key existed, repair only the
            # selected/all cached table, never the raw PolicyData table.
            if selected_key:
                conn.execute(text("""
                    UPDATE app_client_map_points
                    SET franchise_key = LOWER(TRIM(REGEXP_REPLACE(COALESCE(franchise_name,''), '[^a-z0-9]+', ' ', 'g')))
                    WHERE (franchise_key IS NULL OR franchise_key = '')
                      AND LOWER(TRIM(REGEXP_REPLACE(COALESCE(franchise_name,''), '[^a-z0-9]+', ' ', 'g'))) = :selected_key
                """), {'selected_key': selected_key})
            meta = conn.execute(text(f"""
                SELECT COUNT(*) AS total_groups, COALESCE(SUM(client_count),0) AS clients_total
                FROM app_client_map_points
                {where}
            """), params).mappings().first() or {}
            rows = [dict(r) for r in conn.execute(text(f"""
                SELECT franchise_name AS franchise, display_address AS address, display_address AS locality_address,
                       client_name AS client, policy_number, client_count AS count,
                       lat, lng, status AS geocode_status
                FROM app_client_map_points
                {where}
                ORDER BY client_count DESC, display_address
                LIMIT :limit OFFSET :offset
            """), params).mappings().all()]
        return {'total_groups': int(meta.get('total_groups') or 0), 'clients_total': int(meta.get('clients_total') or 0), 'locations': rows}
    except Exception as exc:
        print(f'Could not query client map points: {exc}')
        return {'total_groups': 0, 'clients_total': 0, 'locations': []}


def _client_map_points_summary(selected='All'):
    result = _query_client_map_points(selected=selected, offset=0, limit=1)
    return {'selected': selected, 'points_total': result['total_groups'], 'clients_total': result['clients_total']}


def _has_client_map_point(address, franchise=None):
    key = _client_map_point_key(franchise or '', address) if franchise else _address_cache_key(address)
    if not key:
        return False
    engine = get_db_engine()
    if engine is None:
        return False
    try:
        with engine.begin() as conn:
            _ensure_client_map_points(conn)
            if franchise:
                return bool(conn.execute(text("SELECT 1 FROM app_client_map_points WHERE address_key=:k AND lat IS NOT NULL AND lng IS NOT NULL"), {'k': key}).first())
            # Backward-compatible fallback for older callers: match any per-franchise key ending with |address.
            return bool(conn.execute(text("SELECT 1 FROM app_client_map_points WHERE (address_key=:k OR address_key LIKE :suffix) AND lat IS NOT NULL AND lng IS NOT NULL"), {'k': key, 'suffix': '%|' + key}).first())
    except Exception:
        return False


def _server_geocode_address(address):
    """Geocode one address with Google Geocoding API and cache the result.

    This avoids relying on thousands of browser-side geocoder calls, which was why
    the heatmap page opened but did not draw visible client points. The first run
    may take a little longer while addresses are cached; after that it is fast.
    """
    address = str(address or '').strip()
    key = _address_cache_key(address)
    if not address or not key:
        return None
    engine = get_db_engine()
    if engine is None:
        return None
    try:
        with engine.begin() as conn:
            _ensure_client_geocode_cache(conn)
            cached = conn.execute(text("""
                SELECT lat, lng, status FROM app_client_geocode_cache
                WHERE address_key = :address_key
            """), {'address_key': key}).mappings().first()
            if cached and cached.get('lat') is not None and cached.get('lng') is not None:
                return {'lat': float(cached['lat']), 'lng': float(cached['lng']), 'status': cached.get('status') or 'OK', 'cached': True}
            if cached and str(cached.get('status') or '').upper() in {'ZERO_RESULTS', 'REQUEST_DENIED', 'INVALID_REQUEST'}:
                return None
    except Exception:
        pass

    api_key = get_google_maps_api_key()
    if not api_key:
        return None
    try:
        url = 'https://maps.googleapis.com/maps/api/geocode/json?' + urlencode({
            'address': address,
            'components': 'country:ZA',
            'region': 'za',
            'key': api_key,
        })
        with urllib.request.urlopen(url, timeout=8) as resp:
            data = json.loads(resp.read().decode('utf-8'))
        status = str(data.get('status') or '')
        loc = None
        if status == 'OK' and data.get('results'):
            loc = data['results'][0]['geometry']['location']
        with engine.begin() as conn:
            _ensure_client_geocode_cache(conn)
            conn.execute(text("""
                INSERT INTO app_client_geocode_cache (address_key, address, lat, lng, status, updated_at)
                VALUES (:address_key, :address, :lat, :lng, :status, NOW())
                ON CONFLICT (address_key) DO UPDATE SET
                    address = EXCLUDED.address,
                    lat = EXCLUDED.lat,
                    lng = EXCLUDED.lng,
                    status = EXCLUDED.status,
                    updated_at = NOW()
            """), {
                'address_key': key,
                'address': address,
                'lat': loc.get('lat') if loc else None,
                'lng': loc.get('lng') if loc else None,
                'status': status,
            })
        if loc:
            return {'lat': float(loc['lat']), 'lng': float(loc['lng']), 'status': status, 'cached': False}
    except Exception as exc:
        print(f'Google geocode failed for {address}: {exc}')
    return None



def _get_cached_client_geocode(address):
    """Return cached coordinates for an address without calling Google."""
    address = str(address or '').strip()
    key = _address_cache_key(address)
    if not address or not key:
        return None
    engine = get_db_engine()
    if engine is None:
        return None
    try:
        with engine.begin() as conn:
            _ensure_client_geocode_cache(conn)
            row = conn.execute(text("""
                SELECT lat, lng, status FROM app_client_geocode_cache
                WHERE address_key = :address_key
            """), {'address_key': key}).mappings().first()
        if row and row.get('lat') is not None and row.get('lng') is not None:
            return {'lat': float(row['lat']), 'lng': float(row['lng']), 'status': row.get('status') or 'OK'}
    except Exception as exc:
        print(f'Could not read cached geocode for {address}: {exc}')
    return None


def _attach_cached_geocodes_to_locations(locations, selected='All', max_points=None):
    """Attach only cached coordinates. Never geocode during map page/API load."""
    out = []
    if not locations:
        return out
    if max_points is None:
        max_points = int(os.getenv('CLIENT_MAP_MAX_POINTS', '8000'))
    for item in locations:
        if len(out) >= max_points:
            break
        row = dict(item)
        candidates = []
        for field in ['address', 'full_address', 'locality_address']:
            val = str(item.get(field) or '').strip()
            if val and val.lower() not in {x.lower() for x in candidates}:
                candidates.append(val)
        loc = None
        used_address = None
        for address in candidates:
            loc = _get_cached_client_geocode(address)
            if loc:
                used_address = address
                break
        if not loc:
            continue
        row['address'] = used_address or row.get('address')
        row['lat'] = loc['lat']
        row['lng'] = loc['lng']
        row['geocode_status'] = loc.get('status') or 'OK_CACHED'
        row['needs_geocode'] = False
        out.append(row)
    return out




def _franchise_map_address(name):
    """Build a stable Google-geocodable franchise address/search string."""
    name = str(name or '').strip()
    if not name or name == 'All':
        return ''
    # Franchise names generally identify the branch/town well enough for Google.
    # Do not use client/member addresses here; dashboard maps are franchise-only.
    return name + ', South Africa'


def _get_franchise_map_locations(selected='All', max_items=500):
    """Return cached franchise map coordinates only. Never geocode during dashboard load."""
    try:
        names = [str(x).strip() for x in _available_franchises_for_user() if str(x).strip()]
    except Exception:
        names = []
    selected = str(selected or 'All').strip() or 'All'
    if selected != 'All':
        selected_key = _map_franchise_key(selected)
        names = [n for n in names if _map_franchise_key(n) == selected_key or selected_key in _map_franchise_key(n) or _map_franchise_key(n) in selected_key]
        if not names:
            names = [selected]
    out = []
    for name in names[:max_items]:
        address = _franchise_map_address(name)
        loc = _get_cached_client_geocode(address)
        if loc:
            out.append({'franchise': name, 'address': address, 'lat': loc['lat'], 'lng': loc['lng'], 'status': loc.get('status') or 'OK_CACHED'})
    return out


def _franchise_map_cache_summary(selected='All'):
    try:
        names = [str(x).strip() for x in _available_franchises_for_user() if str(x).strip()]
    except Exception:
        names = []
    selected = str(selected or 'All').strip() or 'All'
    if selected != 'All':
        sk = _map_franchise_key(selected)
        names = [n for n in names if _map_franchise_key(n) == sk or sk in _map_franchise_key(n) or _map_franchise_key(n) in sk] or [selected]
    cached = 0
    for name in names:
        if _get_cached_client_geocode(_franchise_map_address(name)):
            cached += 1
    return {'franchise_groups_total': len(names), 'franchise_groups_cached': cached, 'franchise_groups_missing': max(0, len(names)-cached)}

def _client_map_cache_summary(selected='All'):
    """Fast cache summary.

    This must never scan the full PolicyData/raw address table because the heatmap
    and admin pages poll it often.  It only reports already prepared coordinates
    from app_client_map_points, so it returns instantly even with 170k+ members.
    The prepare page shows progress by increasing these prepared counts after each
    small batch.
    """
    try:
        result = _query_client_map_points(selected=selected, offset=0, limit=1)
        groups_cached = int(result.get('total_groups') or 0)
        clients_cached = int(result.get('clients_total') or 0)
        return {
            'selected': selected,
            'groups_total': groups_cached,
            'groups_cached': groups_cached,
            'groups_missing': 0,
            'clients_total': clients_cached,
            'clients_cached': clients_cached,
            'clients_missing': 0,
        }
    except Exception as exc:
        print(f'Could not read fast client map cache summary: {exc}')
        return {
            'selected': selected,
            'groups_total': 0,
            'groups_cached': 0,
            'groups_missing': 0,
            'clients_total': 0,
            'clients_cached': 0,
            'clients_missing': 0,
        }

def _attach_geocodes_to_locations(locations, selected='All'):
    """Attach stored/server coordinates and keep ungeocoded rows for browser geocoding.

    Browser-restricted Google Maps keys cannot always call the Geocoding REST API
    from Flask.  Therefore this function never drops valid address rows.  Rows
    without stored coordinates are returned with ``needs_geocode=True`` so the
    Google Maps page can geocode them in the browser, save the coordinates back
    to PostgreSQL, and then draw heat points/pins immediately.
    """
    out = []
    if not locations:
        return out
    max_geocode = int(os.getenv('CLIENT_MAP_GEOCODE_LIMIT', '600' if str(selected or 'All') == 'All' else '2500'))
    for item in locations[:max_geocode]:
        row = dict(item)
        candidates = []
        for field in ['address', 'full_address', 'locality_address']:
            val = str(item.get(field) or '').strip()
            if val and val.lower() not in {x.lower() for x in candidates}:
                candidates.append(val)
        # Prefer existing coordinates if the row already carries them.
        try:
            if item.get('lat') is not None and item.get('lng') is not None:
                row['lat'] = float(item.get('lat'))
                row['lng'] = float(item.get('lng'))
                row['needs_geocode'] = False
                out.append(row)
                continue
        except Exception:
            pass
        loc = None
        for address in candidates:
            loc = _server_geocode_address(address)
            if loc:
                row['address'] = address
                break
        if loc:
            row['lat'] = loc['lat']
            row['lng'] = loc['lng']
            row['geocode_status'] = loc.get('status')
            row['needs_geocode'] = False
        else:
            row['needs_geocode'] = True
            row['geocode_status'] = 'PENDING_BROWSER'
            row['address'] = candidates[0] if candidates else str(item.get('address') or '')
        out.append(row)
    return out




def ensure_parlour_fee_settings_table():
    """Create per-franchise premium/parlour fee settings table when the DB is available."""
    engine = get_db_engine()
    if engine is None or text is None:
        return
    try:
        with engine.begin() as conn:
            conn.execute(text("""
                CREATE TABLE IF NOT EXISTS app_franchise_parlour_fee_settings (
                    franchise_name TEXT PRIMARY KEY,
                    commission_to_parlour_pct NUMERIC(10,4) DEFAULT 10.0,
                    admin_to_parlour_pct NUMERIC(10,4) DEFAULT 0.0,
                    facility_fee_to_parlour_pct NUMERIC(10,4) DEFAULT 0.0,
                    notes TEXT,
                    updated_at TIMESTAMP DEFAULT NOW()
                )
            """))
    except Exception as exc:
        print(f'Could not ensure parlour fee settings table: {exc}', flush=True)


def get_parlour_fee_settings(selected='All'):
    """Return per-franchise fee percentages, with safe defaults and no dashboard-heavy work."""
    defaults = {'commission_to_parlour_pct': 10.0, 'admin_to_parlour_pct': 0.0, 'facility_fee_to_parlour_pct': 0.0}
    engine = get_db_engine()
    if engine is None or text is None:
        return defaults.copy()
    try:
        ensure_parlour_fee_settings_table()
        with engine.begin() as conn:
            if selected and selected != 'All':
                row = conn.execute(text("""
                    SELECT commission_to_parlour_pct, admin_to_parlour_pct, facility_fee_to_parlour_pct
                    FROM app_franchise_parlour_fee_settings
                    WHERE LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))
                """), {'franchise': selected}).mappings().first()
                if row:
                    return {
                        'commission_to_parlour_pct': float(row.get('commission_to_parlour_pct') or 0),
                        'admin_to_parlour_pct': float(row.get('admin_to_parlour_pct') or 0),
                        'facility_fee_to_parlour_pct': float(row.get('facility_fee_to_parlour_pct') or 0),
                    }
            # All-franchise view uses the average of configured franchise rates; defaults if nothing configured.
            row = conn.execute(text("""
                SELECT AVG(commission_to_parlour_pct) AS commission_to_parlour_pct,
                       AVG(admin_to_parlour_pct) AS admin_to_parlour_pct,
                       AVG(facility_fee_to_parlour_pct) AS facility_fee_to_parlour_pct
                FROM app_franchise_parlour_fee_settings
            """)).mappings().first()
            if row and row.get('commission_to_parlour_pct') is not None:
                return {
                    'commission_to_parlour_pct': float(row.get('commission_to_parlour_pct') or defaults['commission_to_parlour_pct']),
                    'admin_to_parlour_pct': float(row.get('admin_to_parlour_pct') or 0),
                    'facility_fee_to_parlour_pct': float(row.get('facility_fee_to_parlour_pct') or 0),
                }
    except Exception as exc:
        print(f'Could not load parlour fee settings: {exc}', flush=True)
    return defaults.copy()


def _sum_col(df, col):
    if df is None or df.empty or col not in df.columns:
        return 0.0
    return float(pd.to_numeric(df[col], errors='coerce').fillna(0).sum())


def _period_month_window(filtered_df):
    if filtered_df is None or filtered_df.empty or 'Month' not in filtered_df.columns:
        return None, None
    months = pd.to_datetime(filtered_df['Month'], errors='coerce').dropna()
    if months.empty:
        return None, None
    return months.min(), months.max()


def _last_year_same_window(source_df, start_month, end_month):
    if source_df is None or source_df.empty or start_month is None or end_month is None or 'Month' not in source_df.columns:
        return source_df.iloc[0:0].copy() if source_df is not None else pd.DataFrame()
    tmp = source_df.copy()
    tmp['Month'] = pd.to_datetime(tmp['Month'], errors='coerce')
    ly_start = start_month - pd.DateOffset(years=1)
    ly_end = end_month - pd.DateOffset(years=1)
    return tmp[(tmp['Month'] >= ly_start) & (tmp['Month'] <= ly_end)]


def build_premium_parlour_analysis(source_df, current_df, selected='All'):
    """Build dashboard rows for gross/risk/last-year/parlour fee analysis."""
    settings = get_parlour_fee_settings(selected)
    retail = _sum_col(current_df, 'Retail Premium')
    risk = _sum_col(current_df, 'Risk Premium')
    current_risk_pct = (risk / retail * 100.0) if retail else 0.0
    start_month, end_month = _period_month_window(current_df)
    last_df = _last_year_same_window(source_df, start_month, end_month)
    last_retail = _sum_col(last_df, 'Retail Premium')
    last_risk = _sum_col(last_df, 'Risk Premium')
    last_risk_pct = (last_risk / last_retail * 100.0) if last_retail else 0.0
    diff_pct = current_risk_pct - last_risk_pct
    diff_amount = risk - last_risk
    commission_pct = float(settings.get('commission_to_parlour_pct') or 0)
    admin_pct = float(settings.get('admin_to_parlour_pct') or 0)
    facility_pct = float(settings.get('facility_fee_to_parlour_pct') or 0)
    total_pct = commission_pct + admin_pct + facility_pct
    total_amount = retail * total_pct / 100.0 if retail else 0.0
    return {
        'settings': settings,
        'current_period': '',
        'last_year_period': '',
        'rows': [
            {'label': 'Average Gross Premium p.m.', 'type': 'money', 'value': retail},
            {'label': 'Current Risk Premium %', 'type': 'pct', 'value': current_risk_pct},
            {'label': 'Current Risk Premium R p.m.', 'type': 'money', 'value': risk},
            {'label': 'Last year Risk Premium %', 'type': 'pct', 'value': last_risk_pct},
            {'label': 'Last year Risk Premium R p.m.', 'type': 'money', 'value': last_risk},
            {'label': 'Last year vs Current %', 'type': 'signed_pct_points', 'value': diff_pct},
            {'label': 'Last year vs Current R p.m.', 'type': 'signed_money', 'value': diff_amount},
            {'label': 'Commission to Parlour', 'type': 'pct', 'value': commission_pct},
            {'label': 'Admin to Parlour', 'type': 'pct', 'value': admin_pct},
            {'label': 'Facility Fee to Parlour', 'type': 'pct', 'value': facility_pct},
            {'label': 'Total to Parlour %', 'type': 'pct', 'value': total_pct},
            {'label': 'Total to Parlour R p.m.', 'type': 'money', 'value': total_amount},
        ]
    }


def build_executive_dashboard_context(monthly_view, selected='All', period_view='six_months', premium_analysis_month=''):
    """Build the executive dashboard data directly from imported system data.

    The calculations use the same analysed monthly dataframe as the rest of the
    system so KPI cards, charts and reports reconcile with exports.
    """
    empty = {
        'selected': selected or 'All', 'period_label': '', 'kpis': {}, 'status_cards': [],
        'claims_by_franchise': [], 'claims_by_status': [], 'claims_by_month': [],
        'loss_ratio_by_franchise': [], 'claim_trend': [], 'policy_trend': [],
        'heatmap_points': [], 'client_locations': [], 'map_query': '', 'top_franchise_rows': [], 'has_location_data': False,
        'status_breakdown': {'New': 0, 'In Review': 0, 'Pending': 0, 'Approved': 0, 'Rejected': 0, 'Closed': 0, 'Archived': 0},
        'relation_counts': {'rows': [], 'totals': {}, 'total': 0}, 'premium_parlour_analysis': {'rows': [], 'available_months': [], 'selected_month': ''}, 'age_notifications': [], 'age_notification_summary': [], 'age_notification_totals': {'total':0,'franchises':0,'child_21':0},
    }
    if monthly_view is None or monthly_view.empty:
        return empty
    source_df = monthly_view.copy()
    df = monthly_view.copy()
    try:
        df, _dashboard_label = _dashboard_month_filter(df, period_view)
    except Exception:
        _dashboard_label = ''
    if 'Month' in df.columns:
        df['Month'] = pd.to_datetime(df['Month'], errors='coerce')
        df = df.dropna(subset=['Month'])
    if df.empty:
        return empty
    for col in ['Retail Premium','Risk Premium','Claims','Claim Count','Policy Qty','Total Commission','Total Paid Commissions','BrightRock Month Total','Claim Ratio']:
        if col not in df.columns:
            df[col] = 0.0
        df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0.0)
    months = sorted(df['Month'].dt.to_period('M').unique()) if 'Month' in df.columns else []
    if _dashboard_label:
        period_label = _dashboard_label
    elif months:
        period_label = months[-1].strftime('%b %Y') if len(months)==1 else f"{months[0].strftime('%b %Y')} - {months[-1].strftime('%b %Y')}"
    else:
        period_label = ''
    total_claims = float(df['Claims'].sum())
    total_claim_count = float(df['Claim Count'].sum())
    total_risk = float(df['Risk Premium'].sum())
    total_retail = float(df['Retail Premium'].sum())
    total_policy_qty = float(df['Policy Qty'].sum())
    avg_claim_value = float(total_claims / total_claim_count) if total_claim_count else 0.0
    loss_ratio = float(total_claims / total_risk * 100) if total_risk else 0.0
    hit_ratio = float((df['Claims'] > 0).sum() / max(1, len(df)) * 100)
    kpis = {
        'premium_amount': total_retail,
        'claim_amount': total_claims,
        'claim_count': total_claim_count,
        'average_claim_value': avg_claim_value,
        'risk_premium': total_risk,
        'loss_ratio': loss_ratio,
        'policy_count': total_policy_qty,
        'hit_ratio': hit_ratio,
        'franchise_admin_fee': float(total_retail - total_risk),
        'franchise_money': float(df['BrightRock Month Total'].sum()),
        'franchises': int(df['Franchise'].nunique()) if 'Franchise' in df.columns else 0,
    }
    available_premium_months = []
    premium_current_df = df
    try:
        if 'Month' in source_df.columns:
            src_months = pd.to_datetime(source_df['Month'], errors='coerce').dropna().dt.to_period('M')
            available_premium_months = [str(x) for x in sorted(src_months.unique(), reverse=True)]
        selected_premium_month = str(premium_analysis_month or '').strip()
        if selected_premium_month and selected_premium_month != 'dashboard_period' and 'Month' in source_df.columns:
            src_copy = source_df.copy()
            src_copy['Month'] = pd.to_datetime(src_copy['Month'], errors='coerce')
            premium_current_df = src_copy[src_copy['Month'].dt.to_period('M').astype(str) == selected_premium_month].copy()
            if premium_current_df.empty:
                premium_current_df = df
                selected_premium_month = ''
        else:
            selected_premium_month = ''
    except Exception:
        selected_premium_month = ''
        premium_current_df = df
    premium_parlour_analysis = build_premium_parlour_analysis(source_df, premium_current_df, selected)
    premium_parlour_analysis['available_months'] = available_premium_months
    premium_parlour_analysis['selected_month'] = selected_premium_month
    premium_parlour_analysis['using_dashboard_period'] = not bool(selected_premium_month)
    # imported monthly claims by franchise
    by_fr = df.groupby('Franchise', as_index=False).agg({'Claims':'sum','Claim Count':'sum','Risk Premium':'sum','Retail Premium':'sum','Policy Qty':'sum'}).sort_values('Claims', ascending=False)
    by_fr['Claim Ratio'] = by_fr.apply(lambda r: (r['Claims']/r['Risk Premium']*100) if r['Risk Premium'] else 0, axis=1)
    claims_by_franchise = [{'label': str(r['Franchise']), 'value': float(r['Claims']), 'count': float(r['Claim Count']), 'ratio': float(r['Claim Ratio'])} for _,r in by_fr.head(12).iterrows()]
    loss_ratio_by_franchise = [{'label': str(r['Franchise']), 'value': float(r['Claim Ratio'])} for _,r in by_fr.sort_values('Claim Ratio', ascending=False).head(8).iterrows()]
    month_df = df.groupby(df['Month'].dt.to_period('M')).agg({'Claims':'sum','Claim Count':'sum','Retail Premium':'sum','Risk Premium':'sum','Policy Qty':'sum'}).reset_index()
    claims_by_month = [{'label': m.strftime('%b %Y'), 'value': float(v)} for m,v in zip(month_df['Month'], month_df['Claims'])]
    claim_trend = [{'label': m.strftime('%b'), 'claims': float(c), 'risk': float(r), 'ratio': float((c/r*100) if r else 0)} for m,c,r in zip(month_df['Month'], month_df['Claims'], month_df['Risk Premium'])]
    policy_trend = [{'label': m.strftime('%b'), 'value': float(v)} for m,v in zip(month_df['Month'], month_df['Policy Qty'])]
    # status from workflow database, falls back to financial claim buckets where workflow is empty.
    status_breakdown = dict(empty['status_breakdown'])
    workflow_claims = []
    overdue_count = 0
    ageing_buckets = {'0-7 days': 0, '8-14 days': 0, '15-30 days': 0, '31+ days': 0}
    try:
        engine = get_db_engine()
        if engine is not None:
            params = {}
            where = ['COALESCE(archived,false) = false']
            if selected and selected != 'All':
                where.append('LOWER(TRIM(franchise_name)) = LOWER(TRIM(:franchise))')
                params['franchise'] = selected
            with engine.begin() as conn:
                rows = [dict(r) for r in conn.execute(text(f"""
                    SELECT franchise_name, status, claim_amount, claim_date, created_at, closed_at
                    FROM app_claim_cases
                    WHERE {' AND '.join(where)}
                """), params).mappings().all()]
            today = datetime.today().date()
            for r in rows:
                st = str(r.get('status') or 'New')
                status_breakdown[st] = status_breakdown.get(st, 0) + 1
                amt = _safe_num(r.get('claim_amount'))
                workflow_claims.append({'status': st, 'amount': amt})
                if st not in {'Closed','Approved','Rejected'}:
                    try:
                        base_date = r.get('claim_date') or r.get('created_at')
                        d = pd.to_datetime(base_date).date()
                        age = (today - d).days
                    except Exception:
                        age = 0
                    if age > 14:
                        overdue_count += 1
                    if age <= 7: ageing_buckets['0-7 days'] += 1
                    elif age <= 14: ageing_buckets['8-14 days'] += 1
                    elif age <= 30: ageing_buckets['15-30 days'] += 1
                    else: ageing_buckets['31+ days'] += 1
    except Exception:
        pass
    if sum(status_breakdown.values()) == 0:
        # Financial import statuses available from the claims paid workbook.
        paid_fr = float(df.get('Claim Paid to Franchise', pd.Series(dtype=float)).sum()) if 'Claim Paid to Franchise' in df.columns else 0
        paid_client = float(df.get('Claim Paid to Client', pd.Series(dtype=float)).sum()) if 'Claim Paid to Client' in df.columns else 0
        pending = float(df.get('Repudiated / Pending', pd.Series(dtype=float)).sum()) if 'Repudiated / Pending' in df.columns else 0
        paid_count = int(total_claim_count) if total_claim_count else int((df['Claims'] > 0).sum())
        status_breakdown.update({'Paid': paid_count if total_claims else 0, 'Pending': int(pending > 0), 'Approved': int((paid_fr + paid_client) > 0)})
    claims_by_status = [{'label': k, 'value': int(v)} for k,v in status_breakdown.items() if v]
    status_cards = [
        {'label':'New', 'value': status_breakdown.get('New',0)},
        {'label':'Pending', 'value': status_breakdown.get('Pending',0) + status_breakdown.get('In Review',0)},
        {'label':'Approved/Paid', 'value': status_breakdown.get('Approved',0) + status_breakdown.get('Paid',0)},
        {'label':'Rejected/Closed', 'value': status_breakdown.get('Rejected',0) + status_breakdown.get('Closed',0)},
        {'label':'Overdue', 'value': overdue_count},
        {'label':'31+ Days', 'value': ageing_buckets.get('31+ days',0)},
    ]
    # Heatmap: the current imports do not include coordinates. Create deterministic
    # franchise intensity points so all imported client/claim/policy volume is visible,
    # and make the context ready for real coordinates when imported later.
    heatmap_points = []
    seed_positions = [(18,68),(24,58),(30,72),(36,51),(43,63),(50,47),(57,70),(64,55),(71,67),(78,49),(84,60),(40,78),(62,82),(24,82),(74,78)]
    max_intensity = max(1.0, float((by_fr['Policy Qty'] + by_fr['Claim Count']).max()) if not by_fr.empty else 1.0)
    for i, (_, r) in enumerate(by_fr.head(30).iterrows()):
        x,y = seed_positions[i % len(seed_positions)]
        ring = i // len(seed_positions)
        x = min(90, max(10, x + (ring * 3) - 3))
        y = min(88, max(15, y - (ring * 2) + 2))
        intensity = float((r.get('Policy Qty',0) + r.get('Claim Count',0)) / max_intensity)
        heatmap_points.append({'label': str(r['Franchise']), 'x': x, 'y': y, 'intensity': intensity, 'claims': float(r['Claims']), 'clients': float(r.get('Policy Qty',0))})
    # Keep the dashboard instant: do not load/geocode thousands of client addresses during page render.
    # The Google map fetches them asynchronously from /api/client-map-locations after the page has loaded.
    client_locations = []
    map_query = (selected + ', South Africa') if selected and selected != 'All' else 'South Africa'
    top_rows = []
    for _,r in by_fr.head(10).iterrows():
        top_rows.append({'franchise': str(r['Franchise']), 'premium': float(r['Retail Premium']), 'claims': float(r['Claims']), 'claim_count': float(r['Claim Count']), 'ratio': float(r['Claim Ratio'])})

    franchise_detail = None
    if selected and selected != 'All' and not df.empty:
        detail_cols = ['Retail Premium','Risk Premium','Original Risk Premium','Claims','Claim Count','Average Claim','Claim Ratio','Total Commission','Total Paid Commissions','BrightRock Month Total','MFF Book Value 2.5%','Franchise Book Value 2.5%','Total Book Value','R1 Policy Fee','Underwriter 2.1% Fee','Policy Qty']
        totals = {}
        for c in detail_cols:
            if c in df.columns:
                totals[c] = float(pd.to_numeric(df[c], errors='coerce').fillna(0).sum())
        if totals.get('Claim Count'):
            totals['Average Claim'] = totals.get('Claims',0) / totals.get('Claim Count',0)
        if totals.get('Risk Premium'):
            totals['Claim Ratio'] = totals.get('Claims',0) / totals.get('Risk Premium',0) * 100
        monthly_rows = []
        if 'Month' in df.columns:
            monthly_detail = df.sort_values('Month').copy()
            for _, rr in monthly_detail.iterrows():
                monthly_rows.append({
                    'month': pd.to_datetime(rr.get('Month')).strftime('%b %Y') if pd.notna(rr.get('Month')) else '',
                    'retail': float(rr.get('Retail Premium',0) or 0),
                    'risk': float(rr.get('Risk Premium',0) or 0),
                    'claims': float(rr.get('Claims',0) or 0),
                    'claim_count': float(rr.get('Claim Count',0) or 0),
                    'ratio': float(rr.get('Claim Ratio',0) or 0),
                    'commission': float(rr.get('Total Commission',0) or 0),
                    'paid_commission': float(rr.get('Total Paid Commissions',0) or 0),
                    'franchise_money': float(rr.get('BrightRock Month Total',0) or 0),
                    'book_value': float(rr.get('Total Book Value',0) or 0),
                })
        comparison = []
        if not by_fr.empty:
            all_totals = {
                'premium': float(by_fr['Retail Premium'].sum()),
                'risk': float(by_fr['Risk Premium'].sum()),
                'claims': float(by_fr['Claims'].sum()),
                'claim_count': float(by_fr['Claim Count'].sum()),
                'clients': float(by_fr['Policy Qty'].sum()),
            }
            selected_row = by_fr.iloc[0]
            comparison = [
                {'label':'Share of Premium', 'value': float(selected_row.get('Retail Premium',0) / all_totals['premium'] * 100) if all_totals['premium'] else 0},
                {'label':'Share of Risk Premium', 'value': float(selected_row.get('Risk Premium',0) / all_totals['risk'] * 100) if all_totals['risk'] else 0},
                {'label':'Share of Claims', 'value': float(selected_row.get('Claims',0) / all_totals['claims'] * 100) if all_totals['claims'] else 0},
                {'label':'Share of Claim Count', 'value': float(selected_row.get('Claim Count',0) / all_totals['claim_count'] * 100) if all_totals['claim_count'] else 0},
                {'label':'Share of Clients/Policies', 'value': float(selected_row.get('Policy Qty',0) / all_totals['clients'] * 100) if all_totals['clients'] else 0},
            ]
        franchise_detail = {'name': selected, 'totals': totals, 'monthly_rows': monthly_rows, 'comparisons': comparison, 'address_count': len(client_locations)}
    return {
        'selected': selected or 'All', 'period_label': period_label, 'kpis': kpis,
        'status_cards': status_cards, 'claims_by_franchise': claims_by_franchise,
        'claims_by_status': claims_by_status, 'claims_by_month': claims_by_month,
        'loss_ratio_by_franchise': loss_ratio_by_franchise, 'claim_trend': claim_trend,
        'policy_trend': policy_trend, 'heatmap_points': heatmap_points,
        'client_locations': client_locations, 'map_query': map_query,
        'top_franchise_rows': top_rows, 'has_location_data': bool(client_locations),
        'status_breakdown': status_breakdown,
        'franchise_detail': franchise_detail,
        'relation_counts': get_policy_relation_counts(selected),
        'premium_parlour_analysis': premium_parlour_analysis,
        'age_notifications': get_policy_age_notifications(selected, 12),
        'age_notification_summary': get_policy_age_notification_summary(selected, 100),
        'age_notification_totals': get_policy_age_notification_totals(selected),
        'age_notification_age_bands': get_policy_age_notification_age_bands(selected),
    }
def build_scenario_comparison_view(monthly_view, args):
    """Build the Excel-style Scenario Comparison commission calculator.

    This page intentionally does not write to PostgreSQL or change imported data.
    It only reads the already-calculated monthly view and applies the fixed
    commission structures supplied in comm.xlsx.
    """
    if monthly_view is None or monthly_view.empty:
        return pd.DataFrame(), []

    period_view = args.get('compare_period_view') or args.get('compare_period') or args.get('period_view') or 'six_months'
    if period_view not in {'month', 'six_months', 'year'}:
        period_view = 'six_months'

    def _rate_arg(name, default):
        try:
            return float(str(args.get(name, default)).replace('%', '').strip() or default)
        except Exception:
            return float(default)

    # Current structure is fixed from the Excel calculator. Scenario A and B are
    # editable from the page querystring only; nothing is written to the database.
    current_rates = {
        'Hollard': 10.00,
        'Mkhulu': 2.50,
        'MFF': 2.50,
        'R1': 0.66,
        'ADV': 1.39,
    }
    scenario_a_rates = {
        'Hollard': _rate_arg('a_hollard', 6.00),
        'Mkhulu': _rate_arg('a_mkhulu', 2.00),
        'MFF': _rate_arg('a_mff', 4.00),
        'R1': _rate_arg('a_r1', 3.00),
        'ADV': _rate_arg('a_adv', 5.00),
    }
    scenario_b_rates = {
        'Hollard': _rate_arg('b_hollard', 7.50),
        'Mkhulu': _rate_arg('b_mkhulu', 2.50),
        'MFF': _rate_arg('b_mff', 5.00),
        'R1': _rate_arg('b_r1', 2.00),
        'ADV': _rate_arg('b_adv', 3.00),
    }
    book_value_rate = _rate_arg('book_value_rate', 2.50)

    base = monthly_view.copy().sort_values(['Franchise', 'Month'])
    base['Month'] = pd.to_datetime(base['Month'], errors='coerce')
    base = base.dropna(subset=['Month'])
    if base.empty:
        return pd.DataFrame(), []

    if 'Retail Premium' not in base.columns:
        base['Retail Premium'] = 0.0
    if 'Claims' not in base.columns:
        base['Claims'] = 0.0
    if 'Risk Premium' not in base.columns:
        base['Risk Premium'] = 0.0

    # The calculator must use the actual selected reporting window, not all historic
    # period groups.  Choose the latest policy month, latest 6 policy months, or latest
    # 12 policy months from the imported monthly data, then calculate every percentage
    # from that selected-period Retail Premium total.
    premium_check_cols = [c for c in ['Retail Premium', 'Risk Premium', 'Original Risk Premium'] if c in base.columns]
    if premium_check_cols:
        premium_total = sum(pd.to_numeric(base[c], errors='coerce').fillna(0) for c in premium_check_cols)
        available_months = sorted(base.loc[premium_total > 0, 'Month'].dt.to_period('M').unique())
    else:
        available_months = sorted(base['Month'].dt.to_period('M').unique())
    if not available_months:
        return pd.DataFrame(), []
    if period_view == 'month':
        selected_months = available_months[-1:]
    elif period_view == 'year':
        selected_months = available_months[-12:]
    else:
        selected_months = available_months[-6:]
    base = base[base['Month'].dt.to_period('M').isin(selected_months)].copy()
    if len(selected_months) == 1:
        selected_label = selected_months[0].strftime('%b %Y')
    else:
        selected_label = f"{selected_months[0].strftime('%b %Y')} - {selected_months[-1].strftime('%b %Y')}"
    base['Period View'] = selected_label

    numeric_keep = [c for c in ['Retail Premium', 'Risk Premium', 'Claims', 'Claim Count', 'Policy Qty'] if c in base.columns]
    out = base.groupby(['Franchise', 'Period View'], as_index=False)[numeric_keep].sum()

    def apply_structure(prefix, rates, include_book_value=False):
        total_col = f'{prefix} Total Commission'
        out[total_col] = 0.0
        for component, rate in rates.items():
            out[f'{prefix} {component} %'] = float(rate)
            out[f'{prefix} {component} Amount'] = out['Retail Premium'] * float(rate) / 100.0
            out[total_col] += out[f'{prefix} {component} Amount']
        if include_book_value:
            out[f'{prefix} MFF Book Value %'] = book_value_rate
            out[f'{prefix} MFF Book Value'] = out['Retail Premium'] * book_value_rate / 100.0
        else:
            out[f'{prefix} MFF Book Value %'] = 0.0
            out[f'{prefix} MFF Book Value'] = 0.0
        out[f'{prefix} Total Benefit'] = out[total_col] + out[f'{prefix} MFF Book Value']

    apply_structure('Current', current_rates, include_book_value=False)
    apply_structure('Scenario A', scenario_a_rates, include_book_value=True)
    apply_structure('Scenario B', scenario_b_rates, include_book_value=True)

    out['Scenario A vs Current'] = out['Scenario A Total Benefit'] - out['Current Total Benefit']
    out['Scenario B vs Current'] = out['Scenario B Total Benefit'] - out['Current Total Benefit']
    out['Scenario B vs A'] = out['Scenario B Total Benefit'] - out['Scenario A Total Benefit']
    out['Best Scenario'] = out.apply(lambda r: 'Scenario A' if r['Scenario A Total Benefit'] > r['Scenario B Total Benefit'] else ('Scenario B' if r['Scenario B Total Benefit'] > r['Scenario A Total Benefit'] else 'Same'), axis=1)
    out['Best Additional Benefit'] = out[['Scenario A vs Current', 'Scenario B vs Current']].max(axis=1)

    if 'Risk Premium' in out.columns and 'Claims' in out.columns:
        out['Claim Ratio'] = out.apply(lambda r: (r['Claims'] / r['Risk Premium'] * 100) if r.get('Risk Premium', 0) else 0, axis=1)

    cols = [
        'Franchise', 'Period View', 'Retail Premium', 'Risk Premium', 'Claims', 'Claim Count', 'Policy Qty', 'Claim Ratio',
        'Current Hollard %', 'Current Hollard Amount', 'Current Mkhulu %', 'Current Mkhulu Amount', 'Current MFF %', 'Current MFF Amount', 'Current R1 %', 'Current R1 Amount', 'Current ADV %', 'Current ADV Amount', 'Current Total Commission', 'Current Total Benefit',
        'Scenario A Hollard %', 'Scenario A Hollard Amount', 'Scenario A Mkhulu %', 'Scenario A Mkhulu Amount', 'Scenario A MFF %', 'Scenario A MFF Amount', 'Scenario A R1 %', 'Scenario A R1 Amount', 'Scenario A ADV %', 'Scenario A ADV Amount', 'Scenario A Total Commission', 'Scenario A MFF Book Value %', 'Scenario A MFF Book Value', 'Scenario A Total Benefit',
        'Scenario B Hollard %', 'Scenario B Hollard Amount', 'Scenario B Mkhulu %', 'Scenario B Mkhulu Amount', 'Scenario B MFF %', 'Scenario B MFF Amount', 'Scenario B R1 %', 'Scenario B R1 Amount', 'Scenario B ADV %', 'Scenario B ADV Amount', 'Scenario B Total Commission', 'Scenario B MFF Book Value %', 'Scenario B MFF Book Value', 'Scenario B Total Benefit',
        'Scenario A vs Current', 'Scenario B vs Current', 'Scenario B vs A', 'Best Scenario', 'Best Additional Benefit'
    ]
    cols = [c for c in cols if c in out.columns]
    return out[cols], cols

@app.template_filter('money')
def money(value):
    try:
        return 'R{:,.2f}'.format(float(value or 0))
    except Exception:
        return 'R0.00'


@app.template_filter('pct')
def pct(value):
    try:
        return '{:,.2f}%'.format(float(value))
    except Exception:
        return value

# -----------------------------------------------------------------------------
# Authentication and user management
# -----------------------------------------------------------------------------
PUBLIC_ENDPOINTS = {'login', 'register', 'logout', 'forgot_password', 'reset_password', 'service_worker', 'static', 'healthz', 'cron_daily_backup'}
ADMIN_ENDPOINTS = {'database_health', 'repair_database', 'admin_users', 'admin_update_user', 'admin_delete_user', 'admin_audit_log', 'admin_backup_database', 'admin_backups', 'admin_create_backup', 'admin_download_backup', 'admin_system_health', 'admin_errors', 'admin_deployment_check', 'admin_launch_center', 'admin_cron_log', 'admin_prepare_client_map_locations', 'admin_prepare_client_map_batch', 'admin_claims_rules', 'admin_claims_rules_save', 'admin_claims_rules_seed', 'admin_map_cache_manager', 'admin_refresh_policy_age_notifications', 'admin_parlour_fee_settings'}


def _user_from_row(row):
    return dict(row) if row else None


def get_user_by_id(user_id):
    engine = get_db_engine()
    if engine is None or not user_id:
        return None
    try:
        with engine.begin() as conn:
            row = conn.execute(text("""
                SELECT id, name, email, role, is_active, COALESCE(is_super_admin, false) AS is_super_admin,
                       last_login, COALESCE(failed_login_count, 0) AS failed_login_count, last_failed_login, last_activity, created_at
                FROM app_users WHERE id = :id
            """), {'id': int(user_id)}).mappings().first()
        return _user_from_row(row)
    except Exception:
        return None


def get_user_by_email(email):
    engine = get_db_engine()
    if engine is None or not email:
        return None
    try:
        with engine.begin() as conn:
            row = conn.execute(text("""
                SELECT id, name, email, password_hash, role, is_active, COALESCE(is_super_admin, false) AS is_super_admin,
                       last_login, COALESCE(failed_login_count, 0) AS failed_login_count, last_failed_login, last_activity, created_at
                FROM app_users WHERE LOWER(TRIM(email)) = LOWER(TRIM(:email))
            """), {'email': email}).mappings().first()
        return _user_from_row(row)
    except Exception:
        return None


def app_user_count():
    engine = get_db_engine()
    if engine is None:
        return 0
    try:
        with engine.begin() as conn:
            return int(conn.execute(text('SELECT COUNT(*) FROM app_users')).scalar() or 0)
    except Exception:
        return 0


def record_login_success(user_id):
    engine = get_db_engine()
    if engine is None or not user_id:
        return False
    try:
        with engine.begin() as conn:
            conn.execute(text("""
                UPDATE app_users
                SET last_login = NOW(), failed_login_count = 0, updated_at = NOW()
                WHERE id = :id
            """), {'id': int(user_id)})
        return True
    except Exception as exc:
        print(f'Could not record login success: {exc}')
        return False


def record_login_failure(email):
    engine = get_db_engine()
    if engine is None or not email:
        return False
    try:
        with engine.begin() as conn:
            conn.execute(text("""
                UPDATE app_users
                SET failed_login_count = COALESCE(failed_login_count, 0) + 1,
                    last_failed_login = NOW(),
                    updated_at = NOW()
                WHERE LOWER(TRIM(email)) = LOWER(TRIM(:email))
            """), {'email': str(email).strip().lower()})
        return True
    except Exception as exc:
        print(f'Could not record login failure: {exc}')
        return False


def get_user_franchise_access(user_id):
    engine = get_db_engine()
    if engine is None or not user_id:
        return []
    try:
        with engine.begin() as conn:
            rows = conn.execute(text("""
                SELECT franchise_name
                FROM app_user_franchise_access
                WHERE user_id = :user_id
                ORDER BY franchise_name
            """), {'user_id': int(user_id)}).fetchall()
        return [str(r[0]).strip() for r in rows if str(r[0] or '').strip()]
    except Exception:
        return []


def apply_user_franchise_scope(df):
    user = getattr(g, 'user', None) or {}
    if df is None or df.empty or 'Franchise' not in df.columns:
        return df
    if user.get('role') == 'admin' or user.get('is_super_admin'):
        return df
    allowed = get_user_franchise_access(user.get('id'))
    if not allowed:
        return df.iloc[0:0].copy()
    allowed_keys = {str(x).strip().lower() for x in allowed}
    return df[df['Franchise'].astype(str).str.strip().str.lower().isin(allowed_keys)].copy()


def log_audit(action, details=''):
    """Write a lightweight audit entry when PostgreSQL is available."""
    engine = get_db_engine()
    if engine is None:
        return False
    try:
        user = getattr(g, 'user', None) or {}
        with engine.begin() as conn:
            conn.execute(text("""
                INSERT INTO app_audit_log (user_id, user_email, action, details, ip_address)
                VALUES (:user_id, :user_email, :action, :details, :ip_address)
            """), {
                'user_id': user.get('id'),
                'user_email': user.get('email'),
                'action': str(action or '')[:200],
                'details': str(details or '')[:2000],
                'ip_address': request.headers.get('X-Forwarded-For', request.remote_addr or '')[:200],
            })
        return True
    except Exception as exc:
        print(f'Audit log failed: {exc}')
        return False


def create_app_user(name, email, password, role='user', is_active=True):
    engine = get_db_engine()
    if engine is None:
        raise RuntimeError('Database is not connected. User accounts require PostgreSQL.')
    email = (email or '').strip().lower()
    name = (name or email).strip()
    if not email or not password:
        raise ValueError('Email and password are required.')
    password_hash = generate_password_hash(password)
    with engine.begin() as conn:
        row = conn.execute(text("""
            INSERT INTO app_users (name, email, password_hash, role, is_active, is_super_admin)
            VALUES (:name, :email, :password_hash, :role, :is_active, :is_super_admin)
            RETURNING id, name, email, role, is_active, COALESCE(is_super_admin, false) AS is_super_admin, created_at
        """), {
            'name': name, 'email': email, 'password_hash': password_hash,
            'role': role, 'is_active': bool(is_active), 'is_super_admin': bool(role == 'admin' and app_user_count() == 0)
        }).mappings().first()
    return _user_from_row(row)


def record_user_activity(user_id):
    """Track last user activity with a light throttle so every request does not write."""
    engine = get_db_engine()
    if engine is None or not user_id:
        return False
    now = datetime.now()
    last_update = session.get('last_activity_update')
    try:
        if last_update:
            last_dt = datetime.fromisoformat(last_update)
            if (now - last_dt).total_seconds() < 60:
                return True
    except Exception:
        pass
    try:
        with engine.begin() as conn:
            conn.execute(text("""
                UPDATE app_users
                SET last_activity = NOW()
                WHERE id = :id
            """), {'id': int(user_id)})
        session['last_activity_update'] = now.isoformat()
        return True
    except Exception as exc:
        print(f'Could not record user activity: {exc}')
        return False


def send_system_email(subject, body, to_email=None):
    """Send an operational email when SMTP is configured."""
    if not smtp_is_configured():
        return False
    target = (to_email or ALERT_EMAIL or os.getenv('SMTP_FROM', '')).strip()
    if not target:
        return False
    host = os.getenv('SMTP_HOST', '').strip()
    port = int(os.getenv('SMTP_PORT', '587'))
    username = os.getenv('SMTP_USERNAME', '').strip()
    password = os.getenv('SMTP_PASSWORD', '')
    use_tls = os.getenv('SMTP_USE_TLS', '1') == '1'
    from_email = os.getenv('SMTP_FROM', '').strip()
    from_name = os.getenv('SMTP_FROM_NAME', 'Martins Direct Analytics').strip()
    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = f'{from_name} <{from_email}>' if from_name else from_email
    msg['To'] = target
    msg.set_content(body)
    with smtplib.SMTP(host, port, timeout=20) as smtp:
        if use_tls:
            smtp.starttls()
        if username:
            smtp.login(username, password)
        smtp.send_message(msg)
    return True


def log_error_to_db(exc, route=None):
    """Persist application errors so admins can review production problems."""
    engine = get_db_engine()
    if engine is None:
        return False
    try:
        user = getattr(g, 'user', None) or {}
        tb = ''.join(traceback.format_exception(type(exc), exc, exc.__traceback__))
        with engine.begin() as conn:
            conn.execute(text("""
                INSERT INTO app_error_log (
                    user_id, user_email, route, method, error_type, error_message,
                    traceback_text, ip_address, user_agent
                ) VALUES (
                    :user_id, :user_email, :route, :method, :error_type, :error_message,
                    :traceback_text, :ip_address, :user_agent
                )
            """), {
                'user_id': user.get('id'),
                'user_email': user.get('email'),
                'route': str(route or request.path or '')[:500],
                'method': str(request.method or '')[:20],
                'error_type': type(exc).__name__[:200],
                'error_message': str(exc)[:2000],
                'traceback_text': tb[:12000],
                'ip_address': request.headers.get('X-Forwarded-For', request.remote_addr or '')[:200],
                'user_agent': str(request.headers.get('User-Agent', ''))[:500],
            })
        return True
    except Exception as log_exc:
        print(f'Error logging failed: {log_exc}')
        return False


@app.before_request
def load_logged_in_user():
    g.user = get_user_by_id(session.get('user_id')) if session.get('user_id') else None
    endpoint = request.endpoint or ''
    if MAINTENANCE_MODE and endpoint not in PUBLIC_ENDPOINTS and not endpoint.startswith('static'):
        if g.user and (g.user.get('role') == 'admin' or g.user.get('is_super_admin')):
            pass
        else:
            return '<h1>Maintenance Mode</h1><p>The Martins Direct analytics system is temporarily unavailable while maintenance is being performed.</p>', 503
    if endpoint in PUBLIC_ENDPOINTS or endpoint.startswith('static'):
        return None
    if g.user is None:
        return redirect(url_for('login', next=request.path))
    if not g.user.get('is_active'):
        session.clear()
        flash('Your account is waiting for admin approval or has been disabled. Contact the system administrator.', 'danger')
        return redirect(url_for('login'))
    if endpoint in ADMIN_ENDPOINTS and g.user.get('role') != 'admin':
        flash('Admin access required.', 'danger')
        return redirect(url_for('dashboard'))
    if g.user.get('role') == 'viewer' and request.method == 'POST':
        flash('View-only users cannot make changes.', 'danger')
        return redirect(url_for('dashboard'))
    record_user_activity(g.user.get('id'))
    return None


@app.context_processor
def inject_logged_in_user():
    return {'current_user': getattr(g, 'user', None)}


@app.after_request
def add_security_headers(response):
    response.headers.setdefault('X-Content-Type-Options', 'nosniff')
    response.headers.setdefault('X-Frame-Options', 'SAMEORIGIN')
    response.headers.setdefault('Referrer-Policy', 'strict-origin-when-cross-origin')
    if IS_PRODUCTION:
        response.headers.setdefault('Strict-Transport-Security', 'max-age=31536000; includeSubDomains')
    return response


@app.errorhandler(Exception)
def handle_unexpected_error(exc):
    code = getattr(exc, 'code', None)
    try:
        code = int(code)
    except Exception:
        code = 500

    if code < 500:
        return exc
    log_error_to_db(exc)
    try:
        if IS_PRODUCTION:
            send_system_email(
                'Martins Direct system error',
                f'An application error occurred.\n\nRoute: {request.path}\nError: {type(exc).__name__}: {exc}'
            )
    except Exception as notify_exc:
        print(f'Error notification failed: {notify_exc}')
    if IS_PRODUCTION:
        return render_auth_page('System Error', '<p class="muted">An unexpected error occurred. The system administrator has been notified.</p><p><a class="link" href="/dashboard">Back to dashboard</a></p>'), 500
    raise exc


AUTH_PAGE_CSS = """
<style>
:root{--mt-bg:#f7f8fb;--mt-primary:#6B4AA0;--mt-dark:#583b88;--mt-accent:#8f73bd;--mt-text:#18202a;--mt-muted:#64707d;--mt-border:#dce2e8;--mt-card:#fff;--mt-shadow:0 14px 36px rgba(24,32,42,.10);--mt-focus:0 0 0 3px rgba(107,74,160,.22)}*{box-sizing:border-box}html,body{min-height:100%}body{margin:0;background:var(--mt-bg);color:var(--mt-text);font-family:Arial,Helvetica,sans-serif;font-size:14px;font-weight:400;letter-spacing:0}h1,h2,h3,label,button,.btn,.link,.auth-links,.topnav a,th{font-weight:700}.auth-page{min-height:100vh;display:grid;place-items:center;padding:32px 16px}.auth-wrap{width:min(430px,100%);background:var(--mt-card);border:1px solid var(--mt-border);border-radius:12px;box-shadow:var(--mt-shadow);overflow:hidden}.auth-header{padding:0;background:var(--mt-card);color:var(--mt-text);text-align:center}.auth-header h1,.auth-header p{position:absolute;width:1px;height:1px;padding:0;margin:-1px;overflow:hidden;clip:rect(0,0,0,0);white-space:nowrap;border:0}.auth-body{width:100%;padding:34px 34px 30px;text-align:center}.auth-logo{display:block;width:210px;height:98px;max-width:100%;object-fit:contain;margin:0 auto 18px}.brand-name{margin:0;color:var(--mt-text);font-size:24px;line-height:1.18;font-weight:700}.module-name{margin:7px 0 24px;color:var(--mt-muted);font-size:14px;line-height:1.35;font-weight:400}.auth-form{display:grid;gap:14px;margin:0}.field{margin:0;text-align:left}.field label{display:block;margin:0 0 6px;color:var(--mt-text);font-size:13px}.field input,.field select,.field textarea,select,input,textarea{width:100%;min-height:42px;padding:10px 11px;border:1px solid var(--mt-border);border-radius:6px;background:#fff;color:var(--mt-text);font-family:Arial,Helvetica,sans-serif;font-size:14px;font-weight:400}.field input:focus,.field select:focus,.field textarea:focus,select:focus,input:focus,textarea:focus{outline:none;border-color:var(--mt-primary);box-shadow:var(--mt-focus)}.btn,button{display:inline-flex;align-items:center;justify-content:center;width:100%;min-height:43px;padding:10px 16px;border:0;border-radius:8px;background:var(--mt-primary);color:#fff;font-family:Arial,Helvetica,sans-serif;font-size:14px;font-weight:700;text-decoration:none;cursor:pointer}.btn:hover,button:hover{background:var(--mt-dark)}.btn-danger,.btn[style*="8b0000"]{background:var(--mt-primary)!important}.btn-muted{background:var(--mt-accent)}.auth-links{display:flex;align-items:center;justify-content:space-between;gap:12px;margin:17px 0 0;font-size:13px}.auth-links .muted{margin:0}.link{color:var(--mt-primary);text-decoration:none}.link:hover{color:var(--mt-dark);text-decoration:underline}.muted{color:var(--mt-muted);margin:11px 0;line-height:1.45}.flash{padding:10px 12px;border-radius:8px;margin:0 0 14px;background:#f3eff9;border:1px solid var(--mt-border);color:var(--mt-text);text-align:left}.flash.danger{background:#fff4f4}.flash.success{background:#f4fbf7}table{border-collapse:collapse;width:100%;background:#fff;font-size:12px}th,td{border-bottom:1px solid var(--mt-border);padding:7px;text-align:left;vertical-align:top}th{background:#f3eff9;color:var(--mt-text);font-size:11px}.admin-wrap{max-width:1180px;margin:24px auto;background:#fff;border:1px solid var(--mt-border);border-radius:12px;padding:20px;box-shadow:var(--mt-shadow)}.topnav{margin-bottom:14px}.small{font-size:11px;color:var(--mt-muted)}.inline-form{display:inline-block;margin:0 4px 4px 0}@media(max-width:520px){.auth-page{align-items:start;padding:18px 14px}.auth-body{padding:28px 22px 24px}.auth-logo{width:178px;height:84px}.brand-name{font-size:21px}.auth-links{display:grid;justify-content:center;text-align:center}}
</style>
"""


def get_flashed_messages_with_categories_safe():
    from flask import get_flashed_messages
    return get_flashed_messages(with_categories=True)


def render_auth_page(title, body_html):
    flashes = ''.join([f'<div class="flash {cat}">{msg}</div>' for cat, msg in get_flashed_messages_with_categories_safe()])
    module_name = 'Franchise Claims Analytics System'
    favicon = url_for('static', filename='martins_logo.png')
    login_css = url_for('static', filename='professional-login.css')
    return f'''<!doctype html><html><head><title>{title} - Martins System</title><meta name="viewport" content="width=device-width, initial-scale=1"><link rel="icon" href="{favicon}"><link rel="apple-touch-icon" href="{favicon}"><link rel="stylesheet" href="{login_css}">{AUTH_PAGE_CSS}</head><body><main class="auth-page"><div class="auth-wrap"><div class="auth-header"><h1>{title}</h1><p>{module_name}</p></div><div class="auth-body"><img class="auth-logo" src="{favicon}" alt="Martins logo"><h2 class="brand-name">Martins System</h2><p class="module-name">{module_name}</p>{flashes}{body_html}</div></div></main></body></html>'''


def build_external_url(endpoint, **values):
    """Build an absolute public URL when APP_BASE_URL is configured."""
    path = url_for(endpoint, **values)
    if APP_BASE_URL:
        return APP_BASE_URL + path
    return url_for(endpoint, _external=True, **values)


def smtp_is_configured():
    return bool(os.getenv('SMTP_HOST') and os.getenv('SMTP_FROM'))


def send_password_reset_email(to_email, reset_link):
    """Send the password reset email if SMTP is configured.

    Required env vars: SMTP_HOST, SMTP_FROM.
    Optional env vars: SMTP_PORT, SMTP_USERNAME, SMTP_PASSWORD, SMTP_USE_TLS, SMTP_FROM_NAME.
    """
    if not smtp_is_configured():
        return False

    host = os.getenv('SMTP_HOST', '').strip()
    port = int(os.getenv('SMTP_PORT', '587'))
    username = os.getenv('SMTP_USERNAME', '').strip()
    password = os.getenv('SMTP_PASSWORD', '')
    use_tls = os.getenv('SMTP_USE_TLS', '1') == '1'
    from_email = os.getenv('SMTP_FROM', '').strip()
    from_name = os.getenv('SMTP_FROM_NAME', 'Martins Direct Analytics').strip()

    msg = EmailMessage()
    msg['Subject'] = 'Reset your Martins Direct password'
    msg['From'] = f'{from_name} <{from_email}>' if from_name else from_email
    msg['To'] = to_email
    msg.set_content(f'''Hello,

A password reset was requested for your Martins Direct Analytics account.

Use this link to reset your password:
{reset_link}

This link expires in 1 hour and can only be used once.

If you did not request this, you can ignore this email.
''')
    msg.add_alternative(f'''
    <p>Hello,</p>
    <p>A password reset was requested for your Martins Direct Analytics account.</p>
    <p><a href="{reset_link}" style="display:inline-block;background:#0f2746;color:#fff;padding:12px 16px;border-radius:8px;text-decoration:none;font-weight:bold">Reset password</a></p>
    <p>This link expires in 1 hour and can only be used once.</p>
    <p>If you did not request this, you can ignore this email.</p>
    ''', subtype='html')

    with smtplib.SMTP(host, port, timeout=20) as smtp:
        if use_tls:
            smtp.starttls()
        if username:
            smtp.login(username, password)
        smtp.send_message(msg)
    return True


def create_password_reset_token(email):
    """Create a short-lived password reset token."""
    engine = get_db_engine()
    if engine is None:
        raise RuntimeError('Database is not connected.')
    user = get_user_by_email(email)
    if not user or not user.get('is_active'):
        log_audit('password_reset_requested_unknown', f"Password reset requested for: {email}")
        return None
    token_value = uuid.uuid4().hex + uuid.uuid4().hex
    expires_at = datetime.now() + timedelta(hours=1)
    with engine.begin() as conn:
        conn.execute(text("""
            UPDATE app_password_resets
            SET used = true
            WHERE user_id = :user_id AND used = false
        """), {'user_id': user.get('id')})
        conn.execute(text("""
            INSERT INTO app_password_resets (user_id, token, used, expires_at)
            VALUES (:user_id, :token, false, :expires_at)
        """), {'user_id': user.get('id'), 'token': token_value, 'expires_at': expires_at})
    log_audit('password_reset_requested', f"Password reset token created for: {email}")
    return token_value


def get_valid_password_reset(token_value):
    engine = get_db_engine()
    if engine is None or not token_value:
        return None
    with engine.begin() as conn:
        row = conn.execute(text("""
            SELECT pr.id, pr.user_id, pr.token, pr.used, pr.expires_at, u.email
            FROM app_password_resets pr
            JOIN app_users u ON u.id = pr.user_id
            WHERE pr.token = :token
              AND pr.used = false
              AND pr.expires_at > NOW()
              AND u.is_active = true
        """), {'token': token_value}).mappings().first()
    return dict(row) if row else None


@app.route('/login', methods=['GET', 'POST'])
def login():
    if g.user:
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        email = (request.form.get('email') or '').strip().lower()
        password = request.form.get('password') or ''
        user = get_user_by_email(email)
        if not user or not check_password_hash(user.get('password_hash') or '', password):
            record_login_failure(email)
            log_audit('login_failed', f"Failed login attempt for: {email}")
            try:
                failed_user = get_user_by_email(email)
                if failed_user and int(failed_user.get('failed_login_count') or 0) >= 5:
                    send_system_email('Martins Direct multiple failed logins', f"Account {email} has {failed_user.get('failed_login_count')} failed login attempts.")
            except Exception:
                pass
            flash('Incorrect email or password.', 'danger')
        elif not user.get('is_active'):
            log_audit('login_blocked', f"Inactive or pending account attempted login: {email}")
            flash('This account is waiting for admin approval or has been disabled.', 'danger')
        else:
            session.clear()
            session.permanent = True
            session['user_id'] = int(user['id'])
            session['user_role'] = user.get('role', 'user')
            record_login_success(user['id'])
            log_audit('login_success', f"User logged in: {user.get('email')}")
            flash('Logged in successfully.', 'success')
            return redirect(request.args.get('next') or url_for('dashboard'))
    body = """
    <form class="auth-form" method="post">
      <div class="field"><label>Username</label><input name="email" type="email" autocomplete="username" required autofocus></div>
      <div class="field"><label>Password</label><input name="password" type="password" required></div>
      <button class="btn" type="submit">Log in</button>
    </form>
    <div class="auth-links">
      <p class="muted"><a class="link" href="/forgot_password">Forgot password?</a></p>
      <p class="muted"><a class="link" href="/register">Register account</a></p>
    </div>
    """
    return render_auth_page('Martins Login', body)


@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    reset_token = None
    reset_link = None
    email_sent = False
    email_error = None

    if request.method == 'POST':
        email = (request.form.get('email') or '').strip().lower()
        if not email:
            flash('Enter your email address first.', 'danger')
        else:
            try:
                reset_token = create_password_reset_token(email)
                flash('If that email exists and is active, password reset instructions were created.', 'success')
                if reset_token:
                    reset_link = build_external_url('reset_password', token=reset_token)
                    try:
                        email_sent = send_password_reset_email(email, reset_link)
                        if email_sent:
                            log_audit('password_reset_email_sent', f"Password reset email sent to: {email}")
                    except Exception as mail_exc:
                        email_error = str(mail_exc)
                        log_audit('password_reset_email_failed', f"Password reset email failed for {email}: {email_error}")
            except Exception as exc:
                flash(f'Could not create password reset: {exc}', 'danger')

    token_html = ''
    if reset_token and reset_link:
        if email_sent:
            token_html = """
            <div class="flash success">
              <strong>Password reset email sent.</strong><br>
              <span class="small">Check your inbox. This link expires in 1 hour.</span>
            </div>
            """
        elif IS_PRODUCTION:
            token_html = """
            <div class="flash danger">
              <strong>Email is not configured.</strong><br>
              <span class="small">Ask the system administrator to configure SMTP_HOST and SMTP_FROM in the production environment.</span>
            </div>
            """
        else:
            error_note = f'<br><span class="small">Email error: {email_error}</span>' if email_error else ''
            token_html = f"""
            <div class="flash success">
              <strong>Local reset link:</strong><br>
              <a class="link" href="{reset_link}">{reset_link}</a><br>
              <span class="small">This local testing link expires in 1 hour.</span>{error_note}
            </div>
            """

    body = f"""
    <p class="muted">Enter your email address and the system will send a one-time password reset link.</p>
    <form method="post">
      <div class="field"><label>Email</label><input name="email" type="email" required autofocus></div>
      <button class="btn" type="submit">Send reset link</button>
    </form>
    {token_html}
    <p class="muted"><a class="link" href="/login">Back to login</a></p>
    """
    return render_auth_page('Forgot Password', body)


@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    reset_row = get_valid_password_reset(token)
    if not reset_row:
        body = """
        <p class="muted">This reset link is invalid, expired, or already used.</p>
        <p><a class="link" href="/forgot_password">Request a new reset link</a></p>
        <p><a class="link" href="/login">Back to login</a></p>
        """
        return render_auth_page('Reset Password', body)
    if request.method == 'POST':
        password = request.form.get('password') or ''
        confirm = request.form.get('confirm_password') or ''
        if password != confirm:
            flash('Passwords do not match.', 'danger')
        elif len(password) < 8:
            flash('Password must be at least 8 characters.', 'danger')
        else:
            engine = get_db_engine()
            with engine.begin() as conn:
                conn.execute(text("""
                    UPDATE app_users
                    SET password_hash = :password_hash, updated_at = NOW()
                    WHERE id = :user_id
                """), {'user_id': reset_row.get('user_id'), 'password_hash': generate_password_hash(password)})
                conn.execute(text("""
                    UPDATE app_password_resets
                    SET used = true
                    WHERE id = :reset_id
                """), {'reset_id': reset_row.get('id')})
            log_audit('password_reset_completed', f"Password reset completed for: {reset_row.get('email')}")
            flash('Password reset successfully. You can now log in with your new password.', 'success')
            return redirect(url_for('login'))
    body = f"""
    <p class="muted">Set a new password for <strong>{reset_row.get('email')}</strong>.</p>
    <form method="post">
      <div class="field"><label>New Password</label><input name="password" type="password" minlength="8" required autofocus></div>
      <div class="field"><label>Confirm Password</label><input name="confirm_password" type="password" minlength="8" required></div>
      <button class="btn" type="submit">Reset password</button>
    </form>
    <p class="muted"><a class="link" href="/login">Back to login</a></p>
    """
    return render_auth_page('Reset Password', body)


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        name = (request.form.get('name') or '').strip()
        email = (request.form.get('email') or '').strip().lower()
        password = request.form.get('password') or ''
        confirm = request.form.get('confirm_password') or ''
        if password != confirm:
            flash('Passwords do not match.', 'danger')
        elif len(password) < 8:
            flash('Password must be at least 8 characters.', 'danger')
        elif get_user_by_email(email):
            flash('An account with that email already exists.', 'danger')
        else:
            first_user = app_user_count() == 0
            role = 'admin' if first_user else 'user'
            is_active = True if first_user else False
            try:
                user = create_app_user(name, email, password, role=role, is_active=is_active)
                if first_user:
                    session.clear()
                    session.permanent = True
                    session['user_id'] = int(user['id'])
                    session['user_role'] = role
                    log_audit('register_admin', f"First admin account created: {email}")
                    flash('Admin account created successfully. The first account is automatically approved as admin.', 'success')
                    return redirect(url_for('dashboard'))
                log_audit('register_pending', f"New account waiting for admin approval: {email}")
                flash('Account created. An admin must approve your account before you can log in.', 'success')
                return redirect(url_for('login'))
            except Exception as exc:
                flash(f'Could not create account: {exc}', 'danger')
    body = """
    <p class="muted">Create a user account. The first account becomes admin. All other accounts wait for admin approval.</p>
    <form method="post">
      <div class="field"><label>Name</label><input name="name" required></div>
      <div class="field"><label>Email</label><input name="email" type="email" required></div>
      <div class="field"><label>Password</label><input name="password" type="password" minlength="8" required></div>
      <div class="field"><label>Confirm Password</label><input name="confirm_password" type="password" minlength="8" required></div>
      <button class="btn" type="submit">Register</button>
    </form>
    <p class="muted">Already registered? <a class="link" href="/login">Log in</a></p>
    """
    return render_auth_page('Register', body)


@app.route('/logout')
def logout():
    if getattr(g, 'user', None):
        log_audit('logout', f"User logged out: {g.user.get('email')}")
    session.clear()
    flash('Logged out.', 'success')
    return redirect(url_for('login'))


@app.route('/admin/users')
def admin_users():
    engine = get_db_engine()
    users = []
    if engine is not None:
        with engine.begin() as conn:
            users = [dict(r) for r in conn.execute(text("""
                SELECT id, name, email, role, is_active, COALESCE(is_super_admin, false) AS is_super_admin,
                       last_login, COALESCE(failed_login_count, 0) AS failed_login_count, last_failed_login, last_activity, created_at
                FROM app_users ORDER BY created_at DESC, id DESC
            """)).mappings().all()]
            for _u in users:
                _access = conn.execute(text("""
                    SELECT franchise_name FROM app_user_franchise_access
                    WHERE user_id = :user_id
                    ORDER BY franchise_name
                """), {'user_id': _u.get('id')}).fetchall()
                _u['allowed_franchises'] = ', '.join([str(r[0]) for r in _access])
    rows = []
    for u in users:
        active_selected = 'selected' if u.get('is_active') else ''
        disabled_selected = '' if u.get('is_active') else 'selected'
        admin_selected = 'selected' if u.get('role') == 'admin' else ''
        user_selected = 'selected' if u.get('role') == 'user' else ''
        viewer_selected = 'selected' if u.get('role') == 'viewer' else ''
        is_super_admin = bool(u.get('is_super_admin'))
        protected_badge = '<span class="small" style="font-weight:700;color:#007a2f">Protected Super Admin</span>' if is_super_admin else ''
        lock_attrs = 'disabled' if is_super_admin else ''
        hidden_locked_fields = ''
        if is_super_admin:
            hidden_locked_fields = '<input type="hidden" name="role" value="admin"><input type="hidden" name="is_active" value="1">'
        delete_form = ''
        if not is_super_admin:
            delete_form = f'''
            <form method="post" action="/admin/users/{u.get('id')}/delete" onsubmit="return confirm('Delete this user?')" style="margin-top:6px">
                <button class="btn" style="background:#8b0000" type="submit">Delete User</button>
            </form>
            '''
        rows.append(f'''
        <tr>
          <td>{u.get('name','')}<br>{protected_badge}</td>
          <td>{u.get('email','')}</td>
          <td>{u.get('created_at','')}</td>
          <td>{u.get('last_login') or '-'}</td>
          <td>{u.get('failed_login_count', 0)}</td>
          <td colspan="2">
            <form method="post" action="/admin/users/{u.get('id')}/update" style="display:flex;gap:8px;align-items:center;flex-wrap:wrap">
              {hidden_locked_fields}
              <select name="role" {lock_attrs}><option value="admin" {admin_selected}>Admin</option><option value="user" {user_selected}>User</option><option value="viewer" {viewer_selected}>View only</option></select>
              <select name="is_active" {lock_attrs}><option value="1" {active_selected}>Approved / Active</option><option value="0" {disabled_selected}>Pending / Disabled</option></select>
              <input name="allowed_franchises" value="{u.get('allowed_franchises','')}" placeholder="Allowed franchises, comma separated" style="width:320px" {'disabled' if is_super_admin else ''}>
              <input name="new_password" type="password" placeholder="New password (optional)" style="width:210px">
              <button class="btn" type="submit">Save</button>
            </form>
            {delete_form}
          </td>
        </tr>
        ''')
    body = f'''<!doctype html><html><head><title>User Management</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/dashboard">Back to dashboard</a> &nbsp; | &nbsp; <a class="link" href="/admin/system_health">System Health</a> &nbsp; | &nbsp; <a class="link" href="/admin/claims_rules">Claims Rules</a> &nbsp; | &nbsp; <a class="link" href="/admin/audit_log">Audit Log</a> &nbsp; | &nbsp; <a class="link" href="/admin/backups">Backups</a> &nbsp; | &nbsp; <a class="link" href="/logout">Logout</a></div>
    <h1>User Management</h1><p class="muted">Admin users can approve pending accounts, change roles, assign franchise access, disable accounts, delete users, or reset passwords.</p>
    <table><tr><th>Name</th><th>Email</th><th>Created</th><th>Last Login</th><th>Failed Logins</th><th colspan="2">Manage</th></tr>{''.join(rows) or '<tr><td colspan="7">No users found.</td></tr>'}</table>
    <p class="small">Roles: Admin can manage users and database health. User can use assigned franchises. View only can view assigned franchises but cannot submit changes. Leave franchise access blank only for admin accounts.</p>
    </div></body></html>'''
    return body


@app.route('/admin/users/<int:user_id>/update', methods=['POST'])
def admin_update_user(user_id):
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('admin_users'))

    with engine.begin() as conn:
        user = conn.execute(text("""
            SELECT id, name, email, role, is_active, COALESCE(is_super_admin, false) AS is_super_admin
            FROM app_users
            WHERE id = :id
        """), {'id': user_id}).mappings().first()

        if not user:
            flash('User not found.', 'danger')
            return redirect(url_for('admin_users'))

        user = dict(user)
        current_user_id = (g.user or {}).get('id')
        role = request.form.get('role') if request.form.get('role') in {'admin', 'user', 'viewer'} else 'user'
        is_active = request.form.get('is_active') == '1'
        new_password = request.form.get('new_password') or ''
        allowed_franchises_text = request.form.get('allowed_franchises') or ''
        allowed_franchises = [x.strip() for x in allowed_franchises_text.split(',') if x.strip()]

        if user.get('is_super_admin'):
            role = 'admin'
            is_active = True
            if user.get('id') != current_user_id and new_password.strip():
                flash('Protected Super Admin password can only be changed by the Super Admin account itself.', 'danger')
                return redirect(url_for('admin_users'))

        if new_password.strip():
            conn.execute(text("""
                UPDATE app_users
                SET role=:role, is_active=:is_active, password_hash=:password_hash, updated_at=NOW()
                WHERE id=:id
            """), {'id': user_id, 'role': role, 'is_active': is_active, 'password_hash': generate_password_hash(new_password)})
        else:
            conn.execute(text("""
                UPDATE app_users
                SET role=:role, is_active=:is_active, updated_at=NOW()
                WHERE id=:id
            """), {'id': user_id, 'role': role, 'is_active': is_active})

        if not user.get('is_super_admin'):
            conn.execute(text('DELETE FROM app_user_franchise_access WHERE user_id = :user_id'), {'user_id': user_id})
            seen_franchises = set()
            for franchise_name in allowed_franchises:
                key = franchise_name.lower()
                if key in seen_franchises:
                    continue
                seen_franchises.add(key)
                conn.execute(text("""
                    INSERT INTO app_user_franchise_access (user_id, franchise_name)
                    VALUES (:user_id, :franchise_name)
                """), {'user_id': user_id, 'franchise_name': franchise_name})

    log_audit('admin_update_user', f"Updated user id {user_id}: role={role}, active={is_active}, franchises={len(allowed_franchises)}, password_reset={'yes' if new_password.strip() else 'no'}")
    flash('User updated.', 'success')
    return redirect(url_for('admin_users'))


@app.route('/admin/users/<int:user_id>/delete', methods=['POST'])
def admin_delete_user(user_id):
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('admin_users'))
    current_user_id = (g.user or {}).get('id')
    if int(user_id) == int(current_user_id or 0):
        flash('You cannot delete your own account while logged in.', 'danger')
        return redirect(url_for('admin_users'))
    with engine.begin() as conn:
        user = conn.execute(text("""
            SELECT id, email, COALESCE(is_super_admin, false) AS is_super_admin
            FROM app_users
            WHERE id = :id
        """), {'id': user_id}).mappings().first()
        if not user:
            flash('User not found.', 'danger')
            return redirect(url_for('admin_users'))
        user = dict(user)
        if user.get('is_super_admin'):
            flash('Protected Super Admin account cannot be deleted.', 'danger')
            return redirect(url_for('admin_users'))
        conn.execute(text('DELETE FROM app_user_franchise_access WHERE user_id = :id'), {'id': user_id})
        conn.execute(text('DELETE FROM app_users WHERE id = :id'), {'id': user_id})
    log_audit('admin_delete_user', f"Deleted user id {user_id}: {user.get('email')}")
    flash('User deleted.', 'success')
    return redirect(url_for('admin_users'))


@app.route('/profile/change_password', methods=['GET', 'POST'])
def change_password():
    if not getattr(g, 'user', None):
        return redirect(url_for('login'))
    if request.method == 'POST':
        current_password = request.form.get('current_password') or ''
        new_password = request.form.get('new_password') or ''
        confirm_password = request.form.get('confirm_password') or ''
        user = get_user_by_email(g.user.get('email'))
        if not user or not check_password_hash(user.get('password_hash') or '', current_password):
            flash('Current password is incorrect.', 'danger')
        elif len(new_password) < 8:
            flash('New password must be at least 8 characters.', 'danger')
        elif new_password != confirm_password:
            flash('New passwords do not match.', 'danger')
        else:
            engine = get_db_engine()
            if engine is None:
                flash('Database not connected.', 'danger')
                return redirect(url_for('change_password'))
            with engine.begin() as conn:
                conn.execute(text("""
                    UPDATE app_users
                    SET password_hash = :password_hash, updated_at = NOW()
                    WHERE id = :id
                """), {'id': g.user.get('id'), 'password_hash': generate_password_hash(new_password)})
            log_audit('change_own_password', 'User changed own password')
            flash('Password changed successfully.', 'success')
            return redirect(url_for('dashboard'))
    body = """
    <p class="muted">Change your own password.</p>
    <form method="post">
      <div class="field"><label>Current Password</label><input name="current_password" type="password" required></div>
      <div class="field"><label>New Password</label><input name="new_password" type="password" minlength="8" required></div>
      <div class="field"><label>Confirm New Password</label><input name="confirm_password" type="password" minlength="8" required></div>
      <button class="btn" type="submit">Change Password</button>
    </form>
    <p class="muted"><a class="link" href="/dashboard">Back to dashboard</a></p>
    """
    return render_auth_page('Change Password', body)


def _build_database_backup_zip():
    """Create a complete CSV-based database backup and return (filename, bytes, table_count, row_count)."""
    engine = get_db_engine()
    if engine is None:
        raise RuntimeError('Database not connected.')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f'martins_database_backup_{timestamp}.zip'
    buffer = io.BytesIO()
    table_count = 0
    row_count = 0
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('backup_info.txt', f'Martins Direct database backup created {timestamp}\n')
        zf.writestr('restore_note.txt', 'Restore only from trusted backups created by this system. Always create a fresh backup before restoring.\n')
        with engine.begin() as conn:
            for table in REQUIRED_DB_SCHEMA.keys():
                try:
                    df = pd.read_sql(f'SELECT * FROM {_safe_ident(table)}', conn)
                    zf.writestr(f'{table}.csv', df.to_csv(index=False))
                    table_count += 1
                    row_count += int(len(df))
                except Exception as exc:
                    zf.writestr(f'{table}_ERROR.txt', str(exc))
    buffer.seek(0)
    return filename, buffer.getvalue(), table_count, row_count


def _record_backup_history(filename, file_size, table_count, row_count, status='created'):
    engine = get_db_engine()
    if engine is None:
        return False
    try:
        with engine.begin() as conn:
            conn.execute(text("""
                INSERT INTO app_backup_history (filename, file_size, table_count, row_count, status, created_by)
                VALUES (:filename, :file_size, :table_count, :row_count, :status, :created_by)
            """), {
                'filename': filename,
                'file_size': int(file_size or 0),
                'table_count': int(table_count or 0),
                'row_count': int(row_count or 0),
                'status': status,
                'created_by': (getattr(g, 'user', None) or {}).get('email', ''),
            })
        return True
    except Exception as exc:
        print(f'Backup history failed: {exc}')
        return False


def _fmt_status(ok):
    return '<span style="font-weight:700;color:#047857">OK</span>' if ok else '<span style="font-weight:700;color:#b91c1c">Needs attention</span>'


def _latest_backup_summary():
    engine = get_db_engine()
    if engine is None:
        return None
    try:
        with engine.begin() as conn:
            row = conn.execute(text("""
                SELECT filename, file_size, row_count, status, created_at
                FROM app_backup_history
                ORDER BY created_at DESC, id DESC
                LIMIT 1
            """)).mappings().first()
        return dict(row) if row else None
    except Exception:
        return None



def _record_cron_log(job_name, status, details=''):
    engine = get_db_engine()
    if engine is None:
        return False
    try:
        with engine.begin() as conn:
            conn.execute(text("""
                INSERT INTO app_cron_log (job_name, status, details, ip_address)
                VALUES (:job_name, :status, :details, :ip_address)
            """), {
                'job_name': str(job_name or '')[:200],
                'status': str(status or '')[:80],
                'details': str(details or '')[:2000],
                'ip_address': request.headers.get('X-Forwarded-For', request.remote_addr or '')[:200],
            })
        return True
    except Exception as exc:
        print(f'Cron log failed: {exc}')
        return False


def _cron_authorized():
    supplied = request.headers.get('X-Cron-Secret') or request.args.get('token') or ''
    return bool(CRON_SECRET) and supplied == CRON_SECRET


@app.route('/healthz')
def healthz():
    """Lightweight Render health endpoint. Does not require login."""
    db_ok = bool(get_db_engine() is not None and DB_STATUS.get('enabled'))
    return ({'ok': db_ok, 'app': 'franchise_claims_system', 'version': APP_VERSION, 'database': DB_STATUS.get('message', '')}, 200 if db_ok else 503)



@app.route('/cron/daily_policy_age_check', methods=['GET', 'POST'])
def cron_daily_policy_age_check():
    """Protected endpoint for Render Cron Job to refresh age notifications and relation summaries."""
    if not _cron_authorized():
        _record_cron_log('daily_policy_age_check', 'unauthorized', 'Unauthorized cron attempt')
        return {'ok': False, 'error': 'unauthorized'}, 401
    result = run_daily_policy_age_check()
    status_code = 200 if result.get('status') == 'success' else 500
    return {'ok': result.get('status') == 'success', **result}, status_code


@app.route('/cron/daily_backup', methods=['GET', 'POST'])
def cron_daily_backup():
    """Protected endpoint for Render Cron Job or external scheduler daily backups."""
    if not _cron_authorized():
        _record_cron_log('daily_backup', 'unauthorized', 'Unauthorized cron attempt')
        return {'ok': False, 'error': 'unauthorized'}, 401
    try:
        filename, data, table_count, row_count = _build_database_backup_zip()
        path = os.path.join(BACKUP_DIR, os.path.basename(filename))
        with open(path, 'wb') as fh:
            fh.write(data)
        _record_backup_history(filename, len(data), table_count, row_count, status='cron_created')
        _record_cron_log('daily_backup', 'success', f'Created {filename}; tables={table_count}; rows={row_count}')
        return {'ok': True, 'filename': filename, 'tables': table_count, 'rows': row_count}
    except Exception as exc:
        _record_cron_log('daily_backup', 'failed', str(exc))
        try:
            send_system_email('Martins Direct scheduled backup failed', f'Scheduled backup failed: {exc}')
        except Exception:
            pass
        return {'ok': False, 'error': str(exc)}, 500


@app.route('/admin/cron_log')
def admin_cron_log():
    ensure_database_schema()
    logs = []
    engine = get_db_engine()
    if engine is not None:
        with engine.begin() as conn:
            logs = [dict(r) for r in conn.execute(text("""
                SELECT created_at, job_name, status, details, ip_address
                FROM app_cron_log
                ORDER BY created_at DESC, id DESC
                LIMIT 100
            """)).mappings().all()]
    rows = []
    for row in logs:
        status = str(row.get('status',''))
        color = '#047857' if status == 'success' else ('#b91c1c' if status in {'failed','unauthorized'} else '#92400e')
        rows.append(f"""
        <tr><td>{row.get('created_at','')}</td><td>{row.get('job_name','')}</td><td style="font-weight:700;color:{color}">{status}</td><td>{row.get('details','')}</td><td>{row.get('ip_address','')}</td></tr>
        """)
    body = f'''<!doctype html><html><head><title>Cron Log</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/admin/launch_center">Launch Center</a> &nbsp; | &nbsp; <a class="link" href="/admin/system_health">System Health</a> &nbsp; | &nbsp; <a class="link" href="/admin/claims_rules">Claims Rules</a> &nbsp; | &nbsp; <a class="link" href="/dashboard">Back to dashboard</a></div>
    <h1>Cron Log</h1><p class="muted">Scheduled backup and maintenance job activity.</p>
    <table><tr><th>Date</th><th>Job</th><th>Status</th><th>Details</th><th>IP</th></tr>{''.join(rows) or '<tr><td colspan="5">No cron activity yet.</td></tr>'}</table>
    </div></body></html>'''
    return body


@app.route('/admin/launch_center')
def admin_launch_center():
    base_url = APP_BASE_URL or request.host_url.rstrip('/')
    cron_url = f"{base_url}/cron/daily_backup"
    checks = []
    checks.append(('Production environment', IS_PRODUCTION))
    checks.append(('Database connected', bool(get_db_engine() is not None and DB_STATUS.get('enabled'))))
    checks.append(('SMTP configured', smtp_is_configured()))
    checks.append(('Secure cookies', bool(app.config.get('SESSION_COOKIE_SECURE')) if IS_PRODUCTION else True))
    checks.append(('App base URL set', bool(APP_BASE_URL) if IS_PRODUCTION else True))
    checks.append(('Cron secret set', bool(CRON_SECRET)))
    checks.append(('Recent backup exists', bool(_latest_backup_summary())))
    ready = all(ok for _, ok in checks)
    check_rows = ''.join([f'<tr><td>{name}</td><td>{_fmt_status(ok)}</td></tr>' for name, ok in checks])
    summary = '<div class="flash success">Ready for Render production launch.</div>' if ready else '<div class="flash danger">Complete the items below before inviting live users.</div>'
    body = f'''<!doctype html><html><head><title>Production Launch Center</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/dashboard">Back to dashboard</a> &nbsp; | &nbsp; <a class="link" href="/admin/launch_center">Launch Center</a> &nbsp; | &nbsp; <a class="link" href="/admin/deployment_check">Deployment Check</a> &nbsp; | &nbsp; <a class="link" href="/admin/system_health">System Health</a> &nbsp; | &nbsp; <a class="link" href="/admin/claims_rules">Claims Rules</a> &nbsp; | &nbsp; <a class="link" href="/admin/cron_log">Cron Log</a></div>
    <h1>Production Launch Center</h1>{summary}
    <h2>Launch readiness</h2><table><tr><th>Item</th><th>Status</th></tr>{check_rows}</table>
    <h2>Render Cron Job</h2>
    <p class="muted">Create a Render Cron Job that calls the protected backup endpoint once per day.</p>
    <table><tr><th>Setting</th><th>Value</th></tr>
      <tr><td>Endpoint</td><td><code>{cron_url}</code></td></tr>
      <tr><td>Header</td><td><code>X-Cron-Secret: your CRON_SECRET value</code></td></tr>
      <tr><td>Recommended schedule</td><td><code>0 2 * * *</code></td></tr>
    </table>
    <h2>Go-live sequence</h2>
    <ol>
      <li>Create a fresh backup from <a class="link" href="/admin/backups">Backups</a>.</li>
      <li>Confirm <a class="link" href="/admin/deployment_check">Deployment Check</a> has no red items.</li>
      <li>Set Render environment variables from <code>.env.example</code>.</li>
      <li>Deploy from GitHub to Render.</li>
      <li>Open <code>/healthz</code> on the live URL.</li>
      <li>Test login, forgot password email, and a normal franchise user.</li>
      <li>Create the Render Cron Job for daily backup.</li>
    </ol>
    </div></body></html>'''
    return body


@app.route('/admin/system_health')
def admin_system_health():
    ensure_database_schema()
    engine = get_db_engine()
    db_ok = engine is not None and DB_STATUS.get('enabled')
    stats = {
        'total_users': 0, 'active_users': 0, 'pending_users': 0,
        'active_30m': 0, 'failed_login_users': 0,
        'recent_errors': 0, 'recent_audit': 0,
    }
    if engine is not None:
        try:
            with engine.begin() as conn:
                row = conn.execute(text("""
                    SELECT
                        COUNT(*) AS total_users,
                        SUM(CASE WHEN is_active THEN 1 ELSE 0 END) AS active_users,
                        SUM(CASE WHEN NOT is_active THEN 1 ELSE 0 END) AS pending_users,
                        SUM(CASE WHEN last_activity >= NOW() - INTERVAL '30 minutes' THEN 1 ELSE 0 END) AS active_30m,
                        SUM(CASE WHEN COALESCE(failed_login_count,0) > 0 THEN 1 ELSE 0 END) AS failed_login_users
                    FROM app_users
                """)).mappings().first()
                if row:
                    stats.update(dict(row))
                stats['recent_errors'] = int(conn.execute(text("SELECT COUNT(*) FROM app_error_log WHERE created_at >= NOW() - INTERVAL '24 hours'")).scalar() or 0)
                stats['recent_audit'] = int(conn.execute(text("SELECT COUNT(*) FROM app_audit_log WHERE created_at >= NOW() - INTERVAL '24 hours'")).scalar() or 0)
        except Exception as exc:
            print(f'Health stats failed: {exc}')
    latest_backup = _latest_backup_summary()
    backup_ok = bool(latest_backup)
    try:
        total, used, free = shutil.disk_usage(BASE_DIR)
        disk_html = f'{free / 1024 / 1024 / 1024:,.2f} GB free of {total / 1024 / 1024 / 1024:,.2f} GB'
    except Exception:
        disk_html = 'Unavailable'
    rows = [
        ('Application version', APP_VERSION),
        ('Environment', APP_ENV),
        ('Database', f"{_fmt_status(db_ok)} - {DB_STATUS.get('message','')}"),
        ('SMTP email', _fmt_status(smtp_is_configured())),
        ('Latest backup', f"{_fmt_status(backup_ok)} - {(latest_backup or {}).get('created_at','No backup yet')}"),
        ('Users', f"{int(stats.get('active_users') or 0)} active / {int(stats.get('total_users') or 0)} total"),
        ('Active last 30 min', int(stats.get('active_30m') or 0)),
        ('Users with failed logins', int(stats.get('failed_login_users') or 0)),
        ('Errors last 24h', f"{_fmt_status(int(stats.get('recent_errors') or 0) == 0)} - {int(stats.get('recent_errors') or 0)}"),
        ('Audit events last 24h', int(stats.get('recent_audit') or 0)),
        ('Disk usage', disk_html),
    ]
    row_html = ''.join([f'<tr><td>{name}</td><td>{value}</td></tr>' for name, value in rows])
    body = f"""<!doctype html><html><head><title>System Health</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/dashboard">Back to dashboard</a> &nbsp; | &nbsp; <a class="link" href="/admin/launch_center">Launch Center</a> &nbsp; | &nbsp; <a class="link" href="/admin/deployment_check">Deployment Check</a> &nbsp; | &nbsp; <a class="link" href="/admin/errors">Errors</a> &nbsp; | &nbsp; <a class="link" href="/admin/prepare_client_map_locations">Prepare Client Map</a> &nbsp; | &nbsp; <a class="link" href="/admin/backups">Backups</a> &nbsp; | &nbsp; <a class="link" href="/admin/users">User Management</a></div>
    <h1>System Health</h1><p class="muted">Operational monitoring for the live Martins Direct analytics system.</p>
    <table><tr><th>Check</th><th>Status</th></tr>{row_html}</table>
    </div></body></html>"""
    return body


@app.route('/admin/deployment_check')
def admin_deployment_check():
    checks = []
    checks.append(('SECRET_KEY configured', bool(os.getenv('SECRET_KEY')) and app.secret_key != 'change-this-secret-key-before-live-deployment'))
    checks.append(('DATABASE_URL configured', bool(DATABASE_URL)))
    checks.append(('Database connected', bool(get_db_engine() is not None and DB_STATUS.get('enabled'))))
    checks.append(('APP_ENV production for live server', IS_PRODUCTION))
    checks.append(('Secure session cookies in production', bool(app.config.get('SESSION_COOKIE_SECURE')) if IS_PRODUCTION else True))
    checks.append(('SMTP configured for password reset emails', smtp_is_configured()))
    checks.append(('APP_BASE_URL configured', bool(APP_BASE_URL) if IS_PRODUCTION else True))
    checks.append(('CRON_SECRET configured for scheduled backups', bool(CRON_SECRET)))
    checks.append(('At least one backup created', bool(_latest_backup_summary())))
    checks.append(('Render Procfile present', os.path.exists(os.path.join(BASE_DIR, 'Procfile'))))
    checks.append(('Requirements file present', os.path.exists(os.path.join(BASE_DIR, 'requirements.txt'))))
    rows = ''.join([f'<tr><td>{name}</td><td>{_fmt_status(ok)}</td></tr>' for name, ok in checks])
    ready = all(ok for _, ok in checks)
    summary = '<div class="flash success">Deployment checklist passed.</div>' if ready else '<div class="flash danger">Some deployment checks still need attention before going live.</div>'
    body = f"""<!doctype html><html><head><title>Deployment Check</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/admin/system_health">System Health</a> &nbsp; | &nbsp; <a class="link" href="/admin/claims_rules">Claims Rules</a> &nbsp; | &nbsp; <a class="link" href="/dashboard">Back to dashboard</a></div>
    <h1>Deployment Check</h1>{summary}
    <table><tr><th>Requirement</th><th>Status</th></tr>{rows}</table>
    <p class="small">For local testing, APP_ENV may remain local. For Render/live, set APP_ENV=production and configure all environment variables.</p>
    </div></body></html>"""
    return body


@app.route('/admin/errors')
def admin_errors():
    ensure_database_schema()
    errors = []
    engine = get_db_engine()
    if engine is not None:
        with engine.begin() as conn:
            errors = [dict(r) for r in conn.execute(text("""
                SELECT created_at, user_email, route, method, error_type, error_message, traceback_text, ip_address
                FROM app_error_log
                ORDER BY created_at DESC, id DESC
                LIMIT 100
            """)).mappings().all()]
    rows = []
    for e in errors:
        tb = str(e.get('traceback_text') or '')
        if len(tb) > 2500:
            tb = tb[:2500] + '...'
        rows.append(f"""
        <tr>
          <td>{e.get('created_at','')}</td><td>{e.get('user_email') or '-'}</td><td>{e.get('method','')} {e.get('route','')}</td>
          <td><strong>{e.get('error_type','')}</strong><br>{e.get('error_message','')}<details><summary>Traceback</summary><pre style="white-space:pre-wrap;max-width:900px">{tb}</pre></details></td>
          <td>{e.get('ip_address') or '-'}</td>
        </tr>
        """)
    body = f"""<!doctype html><html><head><title>Error Log</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/admin/system_health">System Health</a> &nbsp; | &nbsp; <a class="link" href="/admin/claims_rules">Claims Rules</a> &nbsp; | &nbsp; <a class="link" href="/dashboard">Back to dashboard</a></div>
    <h1>Error Log</h1><p class="muted">Latest unexpected application errors. 404/favicon errors are not logged here.</p>
    <table><tr><th>Date</th><th>User</th><th>Route</th><th>Error</th><th>IP</th></tr>{''.join(rows) or '<tr><td colspan="5">No application errors logged.</td></tr>'}</table>
    </div></body></html>"""
    return body


@app.route('/admin/backups')
def admin_backups():
    ensure_database_schema()
    backups = []
    engine = get_db_engine()
    if engine is not None:
        with engine.begin() as conn:
            backups = [dict(r) for r in conn.execute(text("""
                SELECT filename, file_size, table_count, row_count, status, created_by, created_at
                FROM app_backup_history
                ORDER BY created_at DESC, id DESC
                LIMIT 100
            """)).mappings().all()]
    rows = []
    for b in backups:
        filename = str(b.get('filename') or '')
        size_mb = (float(b.get('file_size') or 0) / 1024 / 1024)
        rows.append(f'''
        <tr>
          <td>{b.get('created_at','')}</td>
          <td>{filename}</td>
          <td>{size_mb:,.2f} MB</td>
          <td>{b.get('table_count',0)}</td>
          <td>{b.get('row_count',0)}</td>
          <td>{b.get('created_by') or '-'}</td>
          <td>{b.get('status') or '-'}</td>
          <td><a class="link" href="/admin/backups/download/{filename}">Download</a></td>
        </tr>
        ''')
    body = f'''<!doctype html><html><head><title>Database Backups</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/dashboard">Back to dashboard</a> &nbsp; | &nbsp; <a class="link" href="/admin/system_health">System Health</a> &nbsp; | &nbsp; <a class="link" href="/admin/claims_rules">Claims Rules</a> &nbsp; | &nbsp; <a class="link" href="/admin/users">User Management</a> &nbsp; | &nbsp; <a class="link" href="/admin/audit_log">Audit Log</a> &nbsp; | &nbsp; <a class="link" href="/logout">Logout</a></div>
    <h1>Database Backups</h1>
    <p class="muted">Create and download full database backups before imports, repairs, or production changes.</p>
    <form method="post" action="/admin/backups/create"><button class="btn" type="submit">Create Backup Now</button></form>
    <p class="small">Backups are ZIP files containing CSV exports for every managed table. Store downloaded backups somewhere safe.</p>
    <table><tr><th>Date</th><th>Filename</th><th>Size</th><th>Tables</th><th>Rows</th><th>Created By</th><th>Status</th><th>Download</th></tr>{''.join(rows) or '<tr><td colspan="8">No backups created yet.</td></tr>'}</table>
    </div></body></html>'''
    return body


@app.route('/admin/backups/create', methods=['POST'])
def admin_create_backup():
    try:
        filename, data, table_count, row_count = _build_database_backup_zip()
        path = os.path.join(BACKUP_DIR, os.path.basename(filename))
        with open(path, 'wb') as fh:
            fh.write(data)
        _record_backup_history(filename, len(data), table_count, row_count, status='created')
        log_audit('database_backup_created', f'Created database backup {filename} tables={table_count} rows={row_count}')
        flash('Database backup created successfully.', 'success')
    except Exception as exc:
        log_audit('database_backup_failed', str(exc))
        try:
            send_system_email('Martins Direct backup failed', f'Database backup failed: {exc}')
        except Exception:
            pass
        flash(f'Could not create backup: {exc}', 'danger')
    return redirect(url_for('admin_backups'))


@app.route('/admin/backups/download/<path:filename>')
def admin_download_backup(filename):
    safe_name = os.path.basename(filename)
    path = os.path.join(BACKUP_DIR, safe_name)
    if not os.path.exists(path):
        flash('Backup file was not found on disk. Create a new backup.', 'danger')
        return redirect(url_for('admin_backups'))
    log_audit('database_backup_downloaded', f'Downloaded database backup {safe_name}')
    return send_file(path, as_attachment=True, download_name=safe_name, mimetype='application/zip')


@app.route('/admin/backup_database')
def admin_backup_database():
    """Backward-compatible old backup link: create and download a fresh backup immediately."""
    try:
        filename, data, table_count, row_count = _build_database_backup_zip()
        path = os.path.join(BACKUP_DIR, os.path.basename(filename))
        with open(path, 'wb') as fh:
            fh.write(data)
        _record_backup_history(filename, len(data), table_count, row_count, status='created_downloaded')
        log_audit('database_backup', f'Downloaded database backup {filename}')
        return send_file(io.BytesIO(data), as_attachment=True, download_name=filename, mimetype='application/zip')
    except Exception as exc:
        flash(f'Could not create backup: {exc}', 'danger')
        return redirect(url_for('admin_backups'))


@app.route('/admin/audit_log')
def admin_audit_log():
    engine = get_db_engine()
    logs = []
    if engine is not None:
        with engine.begin() as conn:
            logs = [dict(r) for r in conn.execute(text("""
                SELECT created_at, user_email, action, details, ip_address
                FROM app_audit_log ORDER BY created_at DESC, id DESC LIMIT 300
            """)).mappings().all()]
    rows = []
    for row in logs:
        rows.append(f"""
        <tr><td>{row.get('created_at','')}</td><td>{row.get('user_email') or '-'}</td><td>{row.get('action','')}</td><td>{row.get('details','')}</td><td>{row.get('ip_address','')}</td></tr>
        """)
    body = f'''<!doctype html><html><head><title>Audit Log</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/dashboard">Back to dashboard</a> &nbsp; | &nbsp; <a class="link" href="/admin/users">User Management</a> &nbsp; | &nbsp; <a class="link" href="/logout">Logout</a></div>
    <h1>Audit Log</h1><p class="muted">Latest user login, registration, and admin-management activity.</p>
    <table><tr><th>Date</th><th>User</th><th>Action</th><th>Details</th><th>IP Address</th></tr>{''.join(rows) or '<tr><td colspan="5">No audit activity yet.</td></tr>'}</table>
    </div></body></html>'''
    return body


@app.route('/', methods=['GET', 'POST'])
def index():
    return redirect(url_for('dashboard'))


@app.route('/dashboard', methods=['GET', 'POST'])
def dashboard():
    global LAST_RESULT, LAST_CLAIMS_DF, LAST_POLICY_IMPORT_SUMMARY
    if request.method == 'POST':
        rates = {
            'BrightRock': safe_float(request.form.get('brightrock_rate'), DEFAULT_RATES['BrightRock']),
            'Inkulu': safe_float(request.form.get('inkulu_rate'), DEFAULT_RATES['Inkulu']),
            'MFF': safe_float(request.form.get('mff_rate'), DEFAULT_RATES['MFF']),
        }
        book_rates = {
            'MFF Book Value': safe_float(request.form.get('mff_book_rate'), DEFAULT_BOOK_VALUE['MFF Book Value']),
            'Franchise Book Value': safe_float(request.form.get('franchise_book_rate'), DEFAULT_BOOK_VALUE['Franchise Book Value']),
        }
        imported_any = False
        raw = LAST_RESULT['raw']
        try:
            policy_files = []
            for field_name in ('excel_file', 'excel_files'):
                policy_files.extend([f for f in request.files.getlist(field_name) if f and f.filename])

            claims_files = []
            for field_name in ('claims_file', 'claims_files'):
                claims_files.extend([f for f in request.files.getlist(field_name) if f and f.filename])

            policy_summaries = []
            incoming_policy_frames = []
            for file in policy_files:
                filename = f'{uuid.uuid4()}_{file.filename}'
                path = os.path.join(UPLOAD_DIR, filename)
                file.save(path)
                imported_policy = read_excel_file(path)
                if LAST_POLICY_DETAIL_DF is not None and not LAST_POLICY_DETAIL_DF.empty:
                    save_policy_detail_to_postgres(LAST_POLICY_DETAIL_DF, source_file=display_source_filename(file.filename))
                if imported_policy is not None and not imported_policy.empty:
                    incoming_policy_frames.append(imported_policy)
                policy_summaries.append(dict(LAST_POLICY_IMPORT_SUMMARY))
                imported_any = True

            if incoming_policy_frames:
                imported_policy_all = pd.concat(incoming_policy_frames, ignore_index=True)
                save_policy_raw_to_postgres(imported_policy_all, source_file=', '.join([x.get('file_name','') for x in policy_summaries if x.get('file_name')]))
                try:
                    recalculate_claims_summary_from_postgres()
                except Exception as exc:
                    print(f'Could not rebuild PostgreSQL summary after policy import: {exc}')
                db_raw = load_raw_from_postgres()
                raw = db_raw if db_raw is not None and not db_raw.empty else merge_policy_months(raw, imported_policy_all)
                # If claims were already imported, re-apply them after all policy months are added/replaced.
                if LAST_CLAIMS_DF is not None and not LAST_CLAIMS_DF.empty:
                    raw = merge_claims_into_raw(raw, LAST_CLAIMS_DF)
                total_mem = sum(int(x.get('mem_rows', 0) or 0) for x in policy_summaries)
                total_policies = sum(float(x.get('total_policies', 0) or 0) for x in policy_summaries)
                months = ', '.join([str(x.get('month', '')) for x in policy_summaries if x.get('month')])
                franchises_allocated = int(imported_policy_all['franchise'].nunique()) if 'franchise' in imported_policy_all.columns else 0
                detail_rows = sum(int(x.get('detail_rows_stored', 0) or 0) for x in policy_summaries)
                flash('Policy/premium import complete. Files: {}. Months: {}. Total rows stored: {}. MEM rows calculated: {}. Franchises allocated: {}. Policies: {}.'.format(len(policy_summaries), months, detail_rows, total_mem, franchises_allocated, int(total_policies)), 'success')

            claims_frames = []
            claims_summaries = []
            for claims_file in claims_files:
                claims_filename = f'{uuid.uuid4()}_{claims_file.filename}'
                claims_path = os.path.join(UPLOAD_DIR, claims_filename)
                claims_file.save(claims_path)
                claims_df = read_claims_file(claims_path)
                if claims_df is not None and not claims_df.empty:
                    claims_frames.append(claims_df)
                claims_summaries.append(dict(LAST_CLAIMS_IMPORT_SUMMARY))
                imported_any = True

            if claims_frames:
                combined_claims = pd.concat(claims_frames, ignore_index=True)
                LAST_CLAIMS_DF = combined_claims.copy()
                save_claims_raw_to_postgres(combined_claims, source_file=', '.join([x.get('file_name','') for x in claims_summaries if x.get('file_name')]))
                try:
                    recalculate_claims_summary_from_postgres()
                except Exception as exc:
                    print(f'Could not rebuild PostgreSQL summary after claims import: {exc}')
                db_raw = load_raw_from_postgres()
                raw = db_raw if db_raw is not None and not db_raw.empty else merge_claims_into_raw(raw, combined_claims)
                total_matched = sum(float(x.get('matched_claims', 0) or 0) for x in claims_summaries)
                total_unmatched = sum(float(x.get('unmatched_claims', 0) or 0) for x in claims_summaries)
                flash('Claims import complete. Files: {}. Claim rows read: {}. Matched claims: {}. Unmatched/claims-only: {}.'.format(len(claims_summaries), len(combined_claims), money(total_matched), money(total_unmatched)), 'success')

            if raw is not None and not raw.empty:
                monthly, periods, portfolio = analyse(raw, rates, book_rates)
                LAST_RESULT = {'raw': raw, 'monthly': monthly, 'periods': periods, 'portfolio': portfolio, 'rates': rates, 'book_rates': book_rates}
                if not imported_any:
                    flash('Scenario rates updated.', 'success')
            elif not imported_any:
                flash('No data available yet. Import a policy/premium file or claims file.', 'warning')
        except Exception as exc:
            flash(str(exc), 'danger')
        # Post/Redirect/Get prevents the browser from asking to resubmit the large import on refresh/back.
        return redirect(url_for('dashboard', dashboard_period_view=request.args.get('dashboard_period_view', 'six_months')))

    if request.method == 'GET' and (LAST_RESULT.get('raw') is None or LAST_RESULT['raw'].empty):
        reload_dashboard_from_postgres()

    selected = request.args.get('franchise', 'All')
    selected_scenario = request.args.get('scenario', 'scenario_one')
    quick_filter = request.args.get('quick_filter', 'retail_minus_risk')
    period_view = request.args.get('period_view', 'month')
    dashboard_period_view = request.args.get('dashboard_period_view', 'six_months')
    premium_analysis_month = request.args.get('premium_analysis_month', '').strip()
    traffic_filter = request.args.get('traffic_filter', 'All')
    if selected_scenario not in {'scenario_one', 'scenario_two'}:
        selected_scenario = 'scenario_one'
    if quick_filter not in {'retail_minus_risk', 'claim_ratio', 'commissions', 'accumulation'}:
        quick_filter = 'retail_minus_risk'
    if period_view not in {'month', 'six_months', 'year'}:
        period_view = 'month'
    if dashboard_period_view not in {'month', 'six_months', 'year'}:
        dashboard_period_view = 'six_months'
    if traffic_filter not in {'All', 'Blue', 'Green', 'Red'}:
        traffic_filter = 'All'
    search_text = request.args.get('q', '').strip()
    # A typed search should not override an explicit franchise selection.
    if selected and selected != 'All':
        search_text = ''
    monthly = apply_user_franchise_scope(LAST_RESULT['monthly'])
    periods = apply_user_franchise_scope(LAST_RESULT['periods'])
    franchises = sorted(monthly['Franchise'].unique()) if not monthly.empty else []

    if search_text and not monthly.empty:
        exact_matches = [f for f in franchises if f.lower() == search_text.lower()]
        if exact_matches:
            selected = exact_matches[0]
        else:
            partial_matches = [f for f in franchises if search_text.lower() in f.lower()]
            if partial_matches:
                selected = partial_matches[0]
            else:
                flash(f'No franchise found matching "{search_text}".', 'warning')
                selected = 'All'

    if selected != 'All' and not monthly.empty:
        monthly_view = monthly[monthly['Franchise'] == selected]
        periods_view = periods[periods['Franchise'] == selected]
    else:
        monthly_view = monthly
        periods_view = periods

    value_col = 'Scenario 1 Value' if selected_scenario == 'scenario_one' else 'Scenario 2 Value'
    book_col = 'Scenario 1 Book Value' if selected_scenario == 'scenario_one' else 'Scenario 2 Book Value'
    selected_summary = {
        'scenario_name': 'Scenario 1 - 100%' if selected_scenario == 'scenario_one' else 'Scenario 2 - BrightRock',
        'value_label': 'Scenario 1 Value' if selected_scenario == 'scenario_one' else 'Franchise Money per Month',
        'total_value': float(monthly_view[value_col].sum()) if not monthly_view.empty else 0,
        'total_book_value': float(monthly_view[book_col].sum()) if not monthly_view.empty else 0,
        'mff_book_value': float(monthly_view['MFF Book Value 2.5%'].sum()) if not monthly_view.empty else 0,
        'franchise_book_value': float(monthly_view['Franchise Book Value 2.5%'].sum()) if not monthly_view.empty else 0,
        'retail': float(monthly_view['Retail Premium'].sum()) if not monthly_view.empty else 0,
        'risk': float(monthly_view['Risk Premium'].sum()) if not monthly_view.empty else 0,
        'claims': float(monthly_view['Claims'].sum()) if not monthly_view.empty else 0,
        'commission': float(monthly_view['Total Commission'].sum()) if not monthly_view.empty else 0,
    }
    quick_view, quick_columns = build_quick_filter_view(monthly_view, periods_view, LAST_RESULT['rates'], LAST_RESULT['book_rates'], quick_filter, period_view, traffic_filter)
    compare_view, compare_columns = build_scenario_comparison_view(monthly_view, request.args)
    configured_monthly, configured_periods, _configured_portfolio_all = apply_franchise_config(apply_user_franchise_scope(LAST_RESULT['monthly']), load_franchise_config())
    configured_portfolio = build_portfolio_for_period(configured_monthly, configured_periods, dashboard_period_view)
    executive_dashboard = build_executive_dashboard_context(monthly_view, selected=selected, period_view=dashboard_period_view, premium_analysis_month=premium_analysis_month)
    return render_template('dashboard.html', workspace='home', portfolio=configured_portfolio, rates=LAST_RESULT['rates'], book_rates=LAST_RESULT['book_rates'], franchises=franchises, selected=selected, search_text=search_text, selected_scenario=selected_scenario, selected_summary=selected_summary, period_view=period_view, traffic_filter=traffic_filter, claims_import_summary=LAST_CLAIMS_IMPORT_SUMMARY, policy_import_summary=LAST_POLICY_IMPORT_SUMMARY, dashboard_period_view=dashboard_period_view, premium_analysis_month=premium_analysis_month, executive_dashboard=executive_dashboard, google_maps_api_key=get_google_maps_api_key())



def _map_cache_manager_rows(selected='All'):
    """Return per-franchise map cache coverage without scanning/geocoding addresses.

    This is intentionally lightweight: it counts imported PolicyData rows per franchise
    and prepared coordinate rows from app_client_map_points. It does not parse every
    raw address or call Google, so the admin screen opens quickly even with large data.
    """
    engine = get_db_engine()
    if engine is None:
        return [], {'franchises': 0, 'imported_clients': 0, 'cached_clients': 0, 'cached_groups': 0, 'missing_clients': 0}
    selected = str(selected or 'All').strip() or 'All'
    selected_key = _map_franchise_key(selected) if selected != 'All' else ''
    rows = []
    try:
        with engine.begin() as conn:
            _ensure_client_map_points(conn)
            source_sql = """
                SELECT franchise_name, COUNT(*) AS imported_clients
                FROM policydata_detail_raw
                WHERE franchise_name IS NOT NULL AND TRIM(franchise_name) <> ''
            """
            params = {}
            if selected_key:
                # Use broad token prefilter to keep the query quick, then exact-normalize in Python.
                terms = [t for t in selected_key.split() if t not in {'martins', 'funeral', 'funerals'}]
                token = terms[-1] if terms else selected_key.split()[-1]
                source_sql += " AND TRIM(REGEXP_REPLACE(LOWER(franchise_name), '[^a-z0-9]+', ' ', 'g')) LIKE :selected_like"
                params['selected_like'] = '%' + token + '%'
            source_sql += " GROUP BY franchise_name ORDER BY franchise_name"
            source = [dict(r) for r in conn.execute(text(source_sql), params).mappings().all()]
            cache = [dict(r) for r in conn.execute(text("""
                SELECT franchise_key,
                       MIN(franchise_name) AS franchise_name,
                       COUNT(*) AS cached_groups,
                       COALESCE(SUM(client_count),0) AS cached_clients
                FROM app_client_map_points
                WHERE lat IS NOT NULL AND lng IS NOT NULL
                GROUP BY franchise_key
            """)).mappings().all()]
    except Exception as exc:
        print(f'Could not build map cache manager rows: {exc}')
        return [], {'franchises': 0, 'imported_clients': 0, 'cached_clients': 0, 'cached_groups': 0, 'missing_clients': 0}

    cache_by_key = {str(c.get('franchise_key') or '').strip(): c for c in cache}
    merged = {}
    for r in source:
        fname = str(r.get('franchise_name') or '').strip()
        if not fname:
            continue
        key = _map_franchise_key(fname)
        if selected_key and key != selected_key:
            terms = [t for t in selected_key.split() if t not in {'martins', 'funeral', 'funerals'}]
            if terms and not all(t in key.split() for t in terms):
                continue
        rec = merged.setdefault(key, {
            'franchise': fname,
            'franchise_key': key,
            'imported_clients': 0,
            'cached_groups': 0,
            'cached_clients': 0,
            'missing_clients': 0,
            'coverage': 0.0,
        })
        rec['imported_clients'] += int(r.get('imported_clients') or 0)
    # Include cached-only rows as well, useful after older imports/cache rebuilds.
    for key, c in cache_by_key.items():
        if selected_key and key != selected_key:
            continue
        rec = merged.setdefault(key, {
            'franchise': c.get('franchise_name') or key or 'Unknown',
            'franchise_key': key,
            'imported_clients': 0,
            'cached_groups': 0,
            'cached_clients': 0,
            'missing_clients': 0,
            'coverage': 0.0,
        })
        rec['cached_groups'] = int(c.get('cached_groups') or 0)
        rec['cached_clients'] = int(c.get('cached_clients') or 0)
    for key, rec in merged.items():
        c = cache_by_key.get(key, {})
        rec['cached_groups'] = int(c.get('cached_groups') or rec.get('cached_groups') or 0)
        rec['cached_clients'] = int(c.get('cached_clients') or rec.get('cached_clients') or 0)
        rec['missing_clients'] = max(0, int(rec.get('imported_clients') or 0) - int(rec.get('cached_clients') or 0))
        rec['coverage'] = (float(rec['cached_clients']) / float(rec['imported_clients']) * 100.0) if rec.get('imported_clients') else (100.0 if rec.get('cached_clients') else 0.0)
    rows = sorted(merged.values(), key=lambda x: (x['coverage'], -x['imported_clients'], x['franchise']))
    total_imported = sum(int(r.get('imported_clients') or 0) for r in rows)
    total_cached_clients = sum(int(r.get('cached_clients') or 0) for r in rows)
    total_cached_groups = sum(int(r.get('cached_groups') or 0) for r in rows)
    totals = {
        'franchises': len(rows),
        'imported_clients': total_imported,
        'cached_clients': total_cached_clients,
        'cached_groups': total_cached_groups,
        'missing_clients': max(0, total_imported - total_cached_clients),
        'coverage': (total_cached_clients / total_imported * 100.0) if total_imported else 0.0,
    }
    return rows, totals


@app.route('/admin/map_cache_manager', methods=['GET'])
def admin_map_cache_manager():
    """Admin overview of franchise/client map-cache coverage.

    The page is read-only and fast. Actual geocoding still happens through the
    existing small-batch preparation screen so admins can start/continue batches
    without browser/server timeouts.
    """
    ensure_database_schema()
    selected = (request.args.get('franchise') or 'All').strip() or 'All'
    rows, totals = _map_cache_manager_rows(selected)
    try:
        franchises = ['All'] + [str(x).strip() for x in _available_franchises_for_user() if str(x).strip()]
    except Exception:
        franchises = ['All']
    options = ''.join([f'<option value="{_esc(f)}" {"selected" if f == selected else ""}>{_esc(f)}</option>' for f in franchises])
    row_html = []
    for r in rows:
        coverage = float(r.get('coverage') or 0)
        status = 'Ready' if coverage >= 99.9 and int(r.get('missing_clients') or 0) == 0 else ('Partial' if int(r.get('cached_clients') or 0) else 'Not prepared')
        prepare_url = url_for('admin_prepare_client_map_locations', franchise=r.get('franchise') or 'All')
        heatmap_url = url_for('client_heatmap_page', franchise=r.get('franchise') or 'All')
        row_html.append(f"""
        <tr>
          <td><b>{_esc(r.get('franchise'))}</b></td>
          <td>{int(r.get('imported_clients') or 0):,}</td>
          <td>{int(r.get('cached_clients') or 0):,}</td>
          <td>{int(r.get('cached_groups') or 0):,}</td>
          <td>{int(r.get('missing_clients') or 0):,}</td>
          <td>{coverage:,.1f}%</td>
          <td>{_esc(status)}</td>
          <td><a class='btn btn-muted' href='{prepare_url}'>Prepare</a> <a class='btn btn-muted' href='{heatmap_url}'>Heatmap</a></td>
        </tr>
        """)
    rows_html = ''.join(row_html) or '<tr><td colspan="8">No imported franchise/client data found yet.</td></tr>'
    body = f"""<!doctype html><html><head><title>Map Cache Manager</title>{AUTH_PAGE_CSS}</head><body><div class='admin-wrap'>
    <div class='topnav'><a class='link' href='/dashboard'>Back to dashboard</a> &nbsp; | &nbsp; <a class='link' href='/admin/prepare_client_map_locations'>Prepare Map Locations</a> &nbsp; | &nbsp; <a class='link' href='/admin/claims_rules'>Claims Rules</a> &nbsp; | &nbsp; <a class='link' href='/admin/system_health'>System Health</a></div>
    <h1>Map Cache Manager</h1>
    <p class='muted'>Admin overview of prepared franchise and client map coordinates. Counts are policy rows versus unique mapped address groups, so this is safe for 2.5M+ policy rows.</p>
    <form method='get' style='display:flex;gap:10px;align-items:end;flex-wrap:wrap;margin:12px 0 18px'>
      <label>Franchise<br><select name='franchise' style='min-width:340px'>{options}</select></label>
      <button class='btn' type='submit'>View</button>
      <a class='btn btn-muted' href='/admin/prepare_client_map_locations?franchise={quote_plus(selected)}'>Start / Continue Preparation</a>
    </form>
    <div class='kpi-grid' style='grid-template-columns:repeat(5,1fr);margin-bottom:16px'>
      <div class='kpi'><b>{totals.get('franchises',0):,}</b><span>Franchises</span></div>
      <div class='kpi'><b>{totals.get('imported_clients',0):,}</b><span>Imported policy/client rows</span></div>
      <div class='kpi'><b>{totals.get('cached_clients',0):,}</b><span>Client count cached</span></div>
      <div class='kpi'><b>{totals.get('cached_groups',0):,}</b><span>Unique mapped addresses cached</span></div>
      <div class='kpi'><b>{totals.get('coverage',0):,.1f}%</b><span>Coverage</span></div>
    </div>
    <table>
      <tr><th>Franchise</th><th>Policy rows</th><th>Mapped clients</th><th>Unique mapped addresses</th><th>Missing</th><th>Coverage</th><th>Status</th><th>Actions</th></tr>
      {rows_html}
    </table>
    <p class='small'>Tip: if a franchise has imported rows but 0 cached clients, click Prepare. The preparation screen runs small batches and saves latitude/longitude so maps load instantly afterwards.</p>
    </div></body></html>"""
    return body



# -----------------------------------------------------------------------------
# Background high-speed map preparation job
# -----------------------------------------------------------------------------
MAP_PREP_JOBS = {}
MAP_PREP_LOCK = threading.Lock()


def _ensure_map_prep_job_table(conn):
    conn.execute(text("""
        CREATE TABLE IF NOT EXISTS app_map_prep_jobs (
            id TEXT PRIMARY KEY,
            selected_franchise TEXT,
            status TEXT,
            started_at TIMESTAMP DEFAULT NOW(),
            updated_at TIMESTAMP DEFAULT NOW(),
            finished_at TIMESTAMP,
            processed_rows BIGINT DEFAULT 0,
            prepared_points BIGINT DEFAULT 0,
            cached_hits BIGINT DEFAULT 0,
            failed_points BIGINT DEFAULT 0,
            message TEXT
        )
    """))


def _clean_geocode_part(value):
    value = _clean_map_value(value)
    value = re.sub(r'\s+', ' ', str(value or '')).strip(' ,')
    return value


def _address_parts_from_display(address):
    text_value = _clean_map_value(address)
    if not text_value:
        return []
    parts = []
    for p in re.split(r'[,;|]+', text_value):
        p = _clean_geocode_part(p)
        if p and p.lower() not in {x.lower() for x in parts}:
            parts.append(p)
    return parts


def _geocode_candidates_for_address(address):
    """Return high-success, low-volume geocode candidates.

    The original Column O address is preserved for display.  For speed we geocode
    progressively broader location strings so thousands of members in the same
    suburb/town reuse one coordinate.  This avoids days of one-by-one full-address
    geocoding while still showing street/suburb/town on pin click.
    """
    parts = _address_parts_from_display(address)
    if not parts:
        return []
    candidates = []
    def add(parts_list):
        vals = [_clean_geocode_part(x) for x in parts_list if _clean_geocode_part(x)]
        if not vals:
            return
        if not any('south africa' in x.lower() for x in vals):
            vals.append('South Africa')
        cand = ', '.join(vals)
        if cand.lower() not in {x.lower() for x in candidates}:
            candidates.append(cand)
    # Fast locality-level first.  If Column O is "Street, Suburb, Town, Code",
    # this turns it into "Suburb, Town, Code, South Africa".
    if len(parts) >= 4:
        add(parts[-3:])
    if len(parts) >= 3:
        add(parts[-2:])
        add(parts[1:])
    if len(parts) >= 2:
        add(parts[-1:])
        add(parts)
    else:
        add(parts)
    return candidates[:5]


def _server_geocode_address_fast(address):
    """Geocode using cached locality/street candidates and return coordinates.

    Full display addresses remain in app_client_map_points.display_address; the
    geocoded candidate may be suburb/town level to make large books feasible.
    """
    for level, candidate in enumerate(_geocode_candidates_for_address(address), start=1):
        loc = _get_cached_client_geocode(candidate)
        if not loc:
            loc = _server_geocode_address(candidate)
        if loc:
            loc = dict(loc)
            loc['status'] = f"{loc.get('status') or 'OK'}_L{level}"
            loc['geocoded_address'] = candidate
            return loc
    return None


def _write_map_prep_job(job_id, **updates):
    engine = get_db_engine()
    if engine is None:
        return
    allowed = {'selected_franchise','status','processed_rows','prepared_points','cached_hits','failed_points','message'}
    updates = {k:v for k,v in updates.items() if k in allowed}
    if not updates:
        return
    try:
        with engine.begin() as conn:
            _ensure_map_prep_job_table(conn)
            conn.execute(text("""
                INSERT INTO app_map_prep_jobs (id, selected_franchise, status, message, updated_at)
                VALUES (:id, :selected_franchise, :status, :message, NOW())
                ON CONFLICT (id) DO NOTHING
            """), {
                'id': job_id,
                'selected_franchise': str(updates.get('selected_franchise') or ''),
                'status': str(updates.get('status') or 'running'),
                'message': str(updates.get('message') or '')[:2000],
            })
            set_parts = ["updated_at = NOW()"]
            params = {'id': job_id}
            for k,v in updates.items():
                set_parts.append(f"{k} = :{k}")
                params[k] = v
            if updates.get('status') in {'complete','stopped','error'}:
                set_parts.append("finished_at = NOW()")
            conn.execute(text("UPDATE app_map_prep_jobs SET " + ', '.join(set_parts) + " WHERE id = :id"), params)
    except Exception as exc:
        print('Could not write map prep job:', exc)


def _read_map_prep_job(job_id):
    engine = get_db_engine()
    if engine is None or not job_id:
        return None
    try:
        with engine.begin() as conn:
            _ensure_map_prep_job_table(conn)
            row = conn.execute(text("SELECT * FROM app_map_prep_jobs WHERE id=:id"), {'id': job_id}).mappings().first()
            return dict(row) if row else None
    except Exception:
        return None


def _background_prepare_client_map_locations(job_id, selected='All', scan_size=2500, point_limit=500):
    """Run map preparation on the server side until complete or stopped.

    This avoids browser session expiry and processes unique MEM Column-O address
    groups.  It still respects Google quotas by geocoding at most point_limit new
    coordinates per cycle, while cached duplicates are processed quickly.
    """
    selected = str(selected or 'All').strip() or 'All'
    processed = prepared = cached_hits = failed = 0
    _write_map_prep_job(job_id, selected_franchise=selected, status='running', message='Starting background map preparation')
    try:
        # Prepare franchise address coordinates first.
        f_names = [str(x).strip() for x in _available_franchises_for_user() if str(x).strip()]
        if selected != 'All':
            sk = _map_franchise_key(selected)
            f_names = [n for n in f_names if _map_franchise_key(n) == sk or sk in _map_franchise_key(n) or _map_franchise_key(n) in sk] or [selected]
        for fname in f_names:
            with MAP_PREP_LOCK:
                state = MAP_PREP_JOBS.get(job_id, {})
                if state.get('stop'):
                    _write_map_prep_job(job_id, status='stopped', processed_rows=processed, prepared_points=prepared, cached_hits=cached_hits, failed_points=failed, message='Stopped by user')
                    return
            faddr = _franchise_map_address(fname)
            if faddr and not _get_cached_client_geocode(faddr):
                if _server_geocode_address_fast(faddr):
                    prepared += 1
                else:
                    failed += 1

        cycles_without_work = 0
        while True:
            with MAP_PREP_LOCK:
                state = MAP_PREP_JOBS.get(job_id, {})
                if state.get('stop'):
                    _write_map_prep_job(job_id, status='stopped', processed_rows=processed, prepared_points=prepared, cached_hits=cached_hits, failed_points=failed, message='Stopped by user')
                    return
            raw = _get_client_map_source_batch(selected=selected, limit=scan_size)
            if not raw:
                _write_map_prep_job(job_id, status='complete', processed_rows=processed, prepared_points=prepared, cached_hits=cached_hits, failed_points=failed, message='Complete: no more unprocessed MEM Column O address groups for selected scope')
                return
            cycle_prepared = cycle_processed = 0
            for item in raw:
                source_ids = item.get('source_ids') or []
                count = int(item.get('count') or len(source_ids) or 1)
                address = str(item.get('full_address') or item.get('address') or item.get('locality_address') or '').strip()
                franchise = str(item.get('franchise') or '').strip()
                if not address or not franchise:
                    _mark_client_map_sources_processed(source_ids, franchise, address)
                    processed += count
                    continue
                loc = None
                if _has_client_map_point(address, franchise):
                    loc = _get_cached_client_geocode(address) or _get_cached_client_geocode((_geocode_candidates_for_address(address) or [''])[0])
                    cached_hits += 1
                if not loc:
                    # Check candidate cache before calling Google.
                    for cand in _geocode_candidates_for_address(address):
                        loc = _get_cached_client_geocode(cand)
                        if loc:
                            cached_hits += 1
                            break
                if not loc:
                    if cycle_prepared >= point_limit:
                        break
                    loc = _server_geocode_address_fast(address)
                    if loc:
                        prepared += 1
                        cycle_prepared += 1
                    else:
                        failed += 1
                if loc:
                    _upsert_client_map_point(item, address, loc)
                # Mark source rows processed even if failed so bad addresses do not loop forever.
                _mark_client_map_sources_processed(source_ids, franchise, address)
                processed += count
                cycle_processed += count
                if processed % 500 == 0:
                    _write_map_prep_job(job_id, status='running', processed_rows=processed, prepared_points=prepared, cached_hits=cached_hits, failed_points=failed, message=f'Processing {selected}: {processed:,} source rows represented')
            if cycle_processed == 0 and cycle_prepared == 0:
                cycles_without_work += 1
            else:
                cycles_without_work = 0
            _write_map_prep_job(job_id, status='running', processed_rows=processed, prepared_points=prepared, cached_hits=cached_hits, failed_points=failed, message=f'Running: processed {processed:,}; geocoded {prepared:,}; cache hits {cached_hits:,}; failed {failed:,}')
            if cycles_without_work >= 3:
                _write_map_prep_job(job_id, status='complete', processed_rows=processed, prepared_points=prepared, cached_hits=cached_hits, failed_points=failed, message='Complete or no more geocodable source rows')
                return
            time.sleep(float(os.getenv('MAP_PREP_CYCLE_DELAY_SECONDS','0.2')))
    except Exception as exc:
        _write_map_prep_job(job_id, status='error', processed_rows=processed, prepared_points=prepared, cached_hits=cached_hits, failed_points=failed, message=str(exc))
        print('Background map preparation failed:', exc)


def _background_prepare_client_map_locations_with_context(job_id, selected='All', scan_size=2500, point_limit=500):
    """Thread wrapper: Flask DB helpers need an app context outside request threads."""
    with app.app_context():
        return _background_prepare_client_map_locations(job_id, selected, scan_size, point_limit)


@app.route('/admin/prepare_client_map_locations/background_start', methods=['POST'])
def admin_prepare_client_map_locations_background_start():
    ensure_database_schema()
    selected = (request.form.get('franchise') or (request.get_json(silent=True) or {}).get('franchise') or request.values.get('franchise') or 'All').strip() or 'All'
    scan_size = int(safe_float(request.form.get('scan_size') or (request.get_json(silent=True) or {}).get('scan_size') or 2500, 2500))
    point_limit = int(safe_float(request.form.get('point_limit') or (request.get_json(silent=True) or {}).get('point_limit') or 500, 500))
    scan_size = max(250, min(scan_size, 10000))
    point_limit = max(25, min(point_limit, 2000))
    job_id = uuid.uuid4().hex
    with MAP_PREP_LOCK:
        MAP_PREP_JOBS[job_id] = {'stop': False, 'selected': selected}
    _write_map_prep_job(job_id, selected_franchise=selected, status='running', message='Queued')
    t = threading.Thread(target=_background_prepare_client_map_locations_with_context, args=(job_id, selected, scan_size, point_limit), daemon=True)
    t.start()
    return jsonify({'ok': True, 'job_id': job_id, 'selected': selected, 'scan_size': scan_size, 'point_limit': point_limit})


@app.route('/admin/prepare_client_map_locations/background_stop', methods=['POST'])
def admin_prepare_client_map_locations_background_stop():
    payload = request.get_json(silent=True) or request.form or request.values
    job_id = str(payload.get('job_id') or '').strip()
    with MAP_PREP_LOCK:
        if job_id in MAP_PREP_JOBS:
            MAP_PREP_JOBS[job_id]['stop'] = True
    _write_map_prep_job(job_id, status='stopped', message='Stop requested')
    return jsonify({'ok': True, 'job_id': job_id})


@app.route('/admin/prepare_client_map_locations/background_status')
def admin_prepare_client_map_locations_background_status():
    job_id = (request.args.get('job_id') or '').strip()
    job = _read_map_prep_job(job_id) if job_id else None
    selected = (job or {}).get('selected_franchise') or request.args.get('franchise') or 'All'
    summary = _client_map_points_summary(selected)
    summary.update(_franchise_map_cache_summary(selected))
    return jsonify({'ok': True, 'job': job or {}, 'summary': summary})

@app.route('/admin/prepare_client_map_locations', methods=['GET'])
def admin_prepare_client_map_locations():
    """Admin page for preparing client map coordinates in small non-blocking batches."""
    ensure_database_schema()
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('admin_system_health'))
    selected = (request.values.get('franchise') or 'All').strip() or 'All'
    batch_size = int(safe_float(request.values.get('batch_size'), 10))
    batch_size = max(1, min(batch_size, 25))
    try:
        franchises = ['All'] + [x for x in _available_franchises_for_user() if str(x).strip()]
    except Exception:
        franchises = ['All']
    # Fast summary only: do not scan all imported addresses on page load.
    summary = _client_map_points_summary(selected)
    summary.update(_franchise_map_cache_summary(selected))
    summary.setdefault('groups_total', summary.get('points_total', 0))
    summary.setdefault('groups_cached', summary.get('points_total', 0))
    summary.setdefault('groups_missing', 0)
    summary.setdefault('clients_cached', summary.get('clients_total', 0))
    options = ''.join([f'<option value="{_esc(f)}" {"selected" if f == selected else ""}>{_esc(f)}</option>' for f in franchises])
    body = f"""<!doctype html><html><head><title>Prepare Client Map Locations</title>{AUTH_PAGE_CSS}</head><body><div class='admin-wrap'>
    <div class='topnav'><a class='link' href='/admin/system_health'>System Health</a> &nbsp; | &nbsp; <a class='link' href='/admin/claims_rules'>Claims Rules</a> &nbsp; | &nbsp; <a class='link' href='/dashboard'>Back to dashboard</a></div>
    <h1>Prepare Client Map Locations</h1>
    <p class='muted'>This page prepares member/client map coordinates in small background batches. PolicyData client addresses are read from Excel column O. Keep it open until completed. The map pages will then load instantly because they only read saved latitude/longitude.</p>
    <div class='notice'><b>Correct local URL:</b> http://127.0.0.1:5000/admin/prepare_client_map_locations</div>
    <div style='display:flex;gap:10px;align-items:end;flex-wrap:wrap;margin:16px 0'>
      <div><label>Franchise</label><select id='prepFranchise' style='min-width:360px'>{options}</select></div>
      <div><label>Geocode limit per cycle</label><input id='prepBatch' type='number' min='25' max='2000' value='500' style='width:120px'></div>
      <button class='btn' id='startPrep' type='button'>Start / Continue</button>
      <button class='btn btn-muted' id='stopPrep' type='button'>Pause</button>
    </div>
    <div class='card' style='padding:14px;margin:14px 0'>
      <b>Status:</b> <span id='prepStatus'>Ready</span><br>
      <b>Last batch:</b> <span id='lastBatch'>None yet</span>
      <div style='height:14px;background:#f3eff9;border-radius:8px;overflow:hidden;margin-top:10px'><div id='prepBar' style='height:14px;background:#6B4AA0;width:0%'></div></div>
    </div>
    <table><tr><th>Metric</th><th>Value</th></tr>
      <tr><td>Selected</td><td id='mSelected'>{_esc(summary.get('selected'))}</td></tr>
      <tr><td>Franchise addresses found</td><td id='mFrTotal'>{int(summary.get('franchise_groups_total') or 0):,}</td></tr>
      <tr><td>Franchise addresses cached</td><td id='mFrCached'>{int(summary.get('franchise_groups_cached') or 0):,}</td></tr>
      <tr><td>Franchise addresses missing coordinates</td><td id='mFrMissing'>{int(summary.get('franchise_groups_missing') or 0):,}</td></tr>
      <tr><td>Client address groups found</td><td id='mTotal'>{int(summary.get('groups_total') or 0):,}</td></tr>
      <tr><td>Address groups cached</td><td id='mCached'>{int(summary.get('groups_cached') or 0):,}</td></tr>
      <tr><td>Address groups still missing coordinates</td><td id='mMissing'>{int(summary.get('groups_missing') or 0):,}</td></tr>
      <tr><td>Client/policy count represented</td><td id='mClients'>{int(summary.get('clients_total') or 0):,}</td></tr>
      <tr><td>Client/policy count already mapped</td><td id='mClientCached'>{int(summary.get('clients_cached') or 0):,}</td></tr>
    </table>
    <p class='small'>Live use: this works on Render/PostgreSQL as long as GOOGLE_MAPS_API_KEY has Geocoding API enabled. Batches avoid browser/server timeouts and Google rate spikes.</p>
    <script>
    let prepRunning = false;
    let prepOffset = 0;
    function fmt(n) {{ return Number(n || 0).toLocaleString(); }}
    function updateMetrics(summary) {{
      document.getElementById('mSelected').textContent = summary.selected || '';
      document.getElementById('mTotal').textContent = fmt(summary.groups_total);
      document.getElementById('mCached').textContent = fmt(summary.groups_cached);
      document.getElementById('mMissing').textContent = fmt(summary.groups_missing);
      if(document.getElementById('mFrTotal')) document.getElementById('mFrTotal').textContent = fmt(summary.franchise_groups_total);
      if(document.getElementById('mFrCached')) document.getElementById('mFrCached').textContent = fmt(summary.franchise_groups_cached);
      if(document.getElementById('mFrMissing')) document.getElementById('mFrMissing').textContent = fmt(summary.franchise_groups_missing);
      document.getElementById('mClients').textContent = fmt(summary.clients_total);
      document.getElementById('mClientCached').textContent = fmt(summary.clients_cached);
      const total = Number(summary.groups_total || 0) + Number(summary.franchise_groups_total || 0), cached = Number(summary.groups_cached || 0) + Number(summary.franchise_groups_cached || 0);
      document.getElementById('prepBar').style.width = total ? Math.min(100, Math.round(cached / total * 100)) + '%' : '0%';
    }}
    async function runBatch() {{
      if (!prepRunning) return;
      const franchise = document.getElementById('prepFranchise').value || 'All';
      const batch = Math.max(25, Math.min(2000, Number(document.getElementById('prepBatch').value || 500))); // legacy browser batch, overridden by background job
      document.getElementById('prepStatus').textContent = 'Working...';
      try {{
        const res = await fetch('/admin/prepare_client_map_locations/batch', {{
          method: 'POST',
          headers: {{'Content-Type': 'application/json'}},
          body: JSON.stringify({{franchise: franchise, batch_size: batch, offset: prepOffset}})
        }});
        const data = await res.json();
        if (!data.ok) throw new Error(data.error || 'Batch failed');
        updateMetrics(data.summary || {{}});
        prepOffset = (data.next_offset === null || data.next_offset === undefined) ? prepOffset : Number(data.next_offset || 0);
        document.getElementById('lastBatch').textContent = 'prepared ' + fmt(data.prepared) + ', cached/skipped ' + fmt(data.skipped_cached) + ', failed ' + fmt(data.failed);
        const sm = data.summary || {{}};
        const missingTotal = Number(sm.groups_missing || 0) + Number(sm.franchise_groups_missing || 0);
        if (data.done) {{
          prepRunning = false;
          document.getElementById('prepStatus').textContent = 'Complete or no more geocodable addresses.';
          return;
        }}
        setTimeout(runBatch, 750);
      }} catch (err) {{
        prepRunning = false;
        document.getElementById('prepStatus').textContent = 'Stopped: ' + err.message;
      }}
    }}
    document.getElementById('startPrep').onclick = function() {{ prepRunning = true; runBatch(); }};
    document.getElementById('stopPrep').onclick = function() {{ prepRunning = false; document.getElementById('prepStatus').textContent = 'Paused'; }};
    document.getElementById('prepFranchise').onchange = function() {{ window.location = '/admin/prepare_client_map_locations?franchise=' + encodeURIComponent(this.value); }};
    // Phase 10.2: replace browser batch loop with server-side background job.
    let bgJobId = '';
    async function pollBackgroundJob() {{
      if (!bgJobId) return;
      try {{
        const res = await fetch('/admin/prepare_client_map_locations/background_status?job_id=' + encodeURIComponent(bgJobId));
        const data = await res.json();
        const job = data.job || {{}};
        updateMetrics(data.summary || {{}});
        document.getElementById('prepStatus').textContent = (job.status || 'running') + ': ' + (job.message || '');
        document.getElementById('lastBatch').textContent = 'processed ' + fmt(job.processed_rows) + ', geocoded ' + fmt(job.prepared_points) + ', cache hits ' + fmt(job.cached_hits) + ', failed ' + fmt(job.failed_points);
        if (['complete','stopped','error'].indexOf(job.status) === -1) setTimeout(pollBackgroundJob, 1500);
      }} catch (err) {{
        document.getElementById('prepStatus').textContent = 'Status error: ' + err.message;
        setTimeout(pollBackgroundJob, 3000);
      }}
    }}
    document.getElementById('startPrep').onclick = async function() {{
      prepRunning = false;
      const franchise = document.getElementById('prepFranchise').value || 'All';
      const pointLimit = Math.max(25, Math.min(2000, Number(document.getElementById('prepBatch').value || 500)));
      document.getElementById('prepStatus').textContent = 'Starting background job...';
      const res = await fetch('/admin/prepare_client_map_locations/background_start', {{
        method: 'POST', headers: {{'Content-Type': 'application/json'}},
        body: JSON.stringify({{franchise: franchise, point_limit: pointLimit, scan_size: 5000}})
      }});
      const data = await res.json();
      bgJobId = data.job_id || '';
      document.getElementById('prepStatus').textContent = 'Background job started. You can leave this page open to monitor, or come back later.';
      pollBackgroundJob();
    }};
    document.getElementById('stopPrep').onclick = async function() {{
      prepRunning = false;
      if (bgJobId) await fetch('/admin/prepare_client_map_locations/background_stop', {{method:'POST', headers:{{'Content-Type':'application/json'}}, body:JSON.stringify({{job_id:bgJobId}})}});
      document.getElementById('prepStatus').textContent = 'Stop requested';
    }};
    </script>
    </div></body></html>"""
    return body


@app.route('/admin/prepare_client_map_locations/batch', methods=['POST'])
def admin_prepare_client_map_batch():
    """Prepare one small geocode batch and return quickly so the browser never times out.

    Scalable flow:
    - Dashboard franchise addresses are prepared first.
    - Client source rows are read from an unprocessed ledger window.
    - Duplicate rows are grouped by franchise+address.
    - Each unique address is geocoded once and client_count is incremented.
    """
    ensure_database_schema()
    engine = get_db_engine()
    if engine is None:
        return {'ok': False, 'error': 'Database not connected.'}, 500
    payload = request.get_json(silent=True) or {}
    selected = (payload.get('franchise') or request.values.get('franchise') or 'All').strip() or 'All'
    batch_size = int(safe_float(payload.get('batch_size') or request.values.get('batch_size'), 25))
    batch_size = max(1, min(batch_size, 100))
    scan_size = int(os.getenv('CLIENT_MAP_PREPARE_SCAN_SIZE', '1000'))
    raw = _get_client_map_source_batch(selected=selected, limit=scan_size)
    prepared, skipped, failed, checked, counted_existing = 0, 0, 0, 0, 0

    # Prepare franchise address coordinates first. These are only 60-ish rows.
    try:
        f_names = [str(x).strip() for x in _available_franchises_for_user() if str(x).strip()]
        if selected != 'All':
            sk = _map_franchise_key(selected)
            f_names = [n for n in f_names if _map_franchise_key(n) == sk or sk in _map_franchise_key(n) or _map_franchise_key(n) in sk] or [selected]
        for fname in f_names:
            if prepared >= batch_size:
                break
            faddr = _franchise_map_address(fname)
            if not faddr:
                continue
            checked += 1
            if _get_cached_client_geocode(faddr):
                skipped += 1
                continue
            if _server_geocode_address(faddr):
                prepared += 1
            else:
                failed += 1
                if failed >= 5 and prepared == 0:
                    break
    except Exception as exc:
        print('Franchise map preparation failed:', exc)

    for item in raw:
        checked += int(item.get('count') or 1)
        candidates = []
        for field in ['address', 'full_address', 'locality_address']:
            val = str(item.get(field) or '').strip()
            if val and val.lower() not in {x.lower() for x in candidates}:
                candidates.append(val)
        if not candidates:
            skipped += int(item.get('count') or 1)
            _mark_client_map_sources_processed(item.get('source_ids') or [], item.get('franchise') or '', '')
            continue
        # Existing prepared point: increment count for these not-yet-processed source rows without geocoding.
        existing_loc = None
        existing_addr = None
        for addr in candidates:
            if _has_client_map_point(addr, item.get('franchise') or ''):
                existing_loc = _get_cached_client_geocode(addr)
                existing_addr = addr
                break
        if existing_loc:
            _upsert_client_map_point(item, existing_addr or candidates[0], existing_loc)
            _mark_client_map_sources_processed(item.get('source_ids') or [], item.get('franchise') or '', existing_addr or candidates[0])
            counted_existing += int(item.get('count') or 1)
            skipped += 1
            continue
        if prepared >= batch_size:
            # Leave these source rows unprocessed for the next small browser batch.
            break
        loc = None
        used_addr = None
        for addr in candidates:
            loc = _get_cached_client_geocode(addr)
            if loc:
                used_addr = addr
                break
        if not loc:
            for addr in candidates:
                loc = _server_geocode_address(addr)
                if loc:
                    used_addr = addr
                    break
        if loc:
            _upsert_client_map_point(item, used_addr or candidates[0], loc)
            _mark_client_map_sources_processed(item.get('source_ids') or [], item.get('franchise') or '', used_addr or candidates[0])
            prepared += 1
        else:
            # Mark failed address group as processed so the job can keep moving.
            _mark_client_map_sources_processed(item.get('source_ids') or [], item.get('franchise') or '', candidates[0])
            failed += 1
            if failed >= 10 and prepared == 0:
                break

    done = len(raw) == 0
    summary = _client_map_points_summary(selected)
    summary.update(_franchise_map_cache_summary(selected))
    summary.setdefault('groups_total', summary.get('points_total', 0))
    summary.setdefault('groups_cached', summary.get('points_total', 0))
    summary.setdefault('groups_missing', 0)
    summary.setdefault('clients_cached', summary.get('clients_total', 0))
    log_audit('prepare_client_map_locations_batch', f'{selected}: prepared {prepared}, counted_existing {counted_existing}, skipped {skipped}, failed {failed}')
    return {'ok': True, 'selected': selected, 'prepared': prepared, 'counted_existing': counted_existing, 'skipped_cached': skipped, 'failed': failed, 'checked': checked, 'offset': 0, 'next_offset': 0, 'done': done, 'summary': summary}


@app.route('/admin/claims_rules', methods=['GET'])
def admin_claims_rules():
    ensure_database_schema()
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('dashboard'))
    with engine.begin() as conn:
        _seed_claim_document_rule_defaults(conn)
        rows = [dict(r) for r in conn.execute(text("""
            SELECT id, document_type, claim_type, aliases, required, is_active, sort_order
            FROM app_claim_document_types
            ORDER BY sort_order, document_type
        """)).mappings().all()]
        rules = [dict(r) for r in conn.execute(text("""
            SELECT document_type_id, rule_key, rule_label, is_enabled
            FROM app_claim_document_rules
            ORDER BY document_type_id, id
        """)).mappings().all()]
    rules_by_type = {}
    for r in rules:
        rules_by_type.setdefault(int(r['document_type_id']), []).append(r)
    cards = []
    for t in rows:
        type_id = int(t['id'])
        chk_required = 'checked' if t.get('required') else ''
        chk_active = 'checked' if t.get('is_active') else ''
        rule_boxes = []
        existing_keys = {r['rule_key'] for r in rules_by_type.get(type_id, [])}
        for r in rules_by_type.get(type_id, []):
            checked = 'checked' if r.get('is_enabled') else ''
            rule_boxes.append(f"<label style='display:block;margin:4px 0'><input type='checkbox' name='rule_{type_id}_{_esc(r.get('rule_key'))}' value='1' {checked}> {_esc(r.get('rule_label'))}</label>")
        for key, label in CLAIM_RULE_LABELS.items():
            if key not in existing_keys:
                rule_boxes.append(f"<label style='display:block;margin:4px 0'><input type='checkbox' name='rule_{type_id}_{_esc(key)}' value='1'> {_esc(label)}</label>")
        cards.append(f"""
        <div class='card' style='padding:14px;margin:12px 0;border:1px solid #c9c9c9;border-radius:12px;background:#fff'>
          <input type='hidden' name='type_ids' value='{type_id}'>
          <div style='display:grid;grid-template-columns:1fr 1fr 100px;gap:10px'>
            <div><label>Document Type</label><input name='document_type_{type_id}' value='{_esc(t.get('document_type'))}' style='width:100%'></div>
            <div><label>Claim Type</label><input name='claim_type_{type_id}' value='{_esc(t.get('claim_type') or 'Funeral Claim')}' style='width:100%'></div>
            <div><label>Order</label><input name='sort_order_{type_id}' type='number' value='{int(t.get('sort_order') or 100)}' style='width:100%'></div>
          </div>
          <div class='field'><label>Aliases / Keywords, comma separated</label><input name='aliases_{type_id}' value='{_esc(t.get('aliases'))}' style='width:100%'></div>
          <label><input type='checkbox' name='required_{type_id}' value='1' {chk_required}> Required document</label>
          &nbsp;&nbsp;<label><input type='checkbox' name='active_{type_id}' value='1' {chk_active}> Active</label>
          <h4 style='margin:12px 0 4px'>Validation Checklist</h4>
          {''.join(rule_boxes)}
        </div>""")
    body = f"""<!doctype html><html><head><title>Claims Rules</title>{AUTH_PAGE_CSS}</head><body><div class='admin-wrap'>
    <div class='topnav'><a class='link' href='/dashboard'>Back to dashboard</a> &nbsp; | &nbsp; <a class='link' href='/admin/system_health'>System Health</a> &nbsp; | &nbsp; <a class='link' href='/admin/claims_rules'>Claims Rules</a> &nbsp; | &nbsp; <a class='link' href='/claims'>Claims</a></div>
    <h1>Phase 10 - Claims Rules & Document Validation</h1>
    <p class='muted'>Admin-only checklist rules. These rules decide which uploaded claim documents are required and what each document must contain before a claim becomes Approved.</p>
    <form method='post' action='/admin/claims_rules/save'>
      {''.join(cards)}
      <h2>Add New Document Rule</h2>
      <div class='card' style='padding:14px;margin:12px 0;background:#fff;border:1px solid #c9c9c9;border-radius:12px'>
        <div style='display:grid;grid-template-columns:1fr 1fr;gap:10px'>
          <div><label>New document type</label><input name='new_document_type' placeholder='e.g. Police Report' style='width:100%'></div>
          <div><label>Claim type</label><input name='new_claim_type' value='Funeral Claim' style='width:100%'></div>
        </div>
        <div class='field'><label>Aliases / Keywords</label><input name='new_aliases' placeholder='police report, accident report, case number' style='width:100%'></div>
        <label><input type='checkbox' name='new_required' value='1'> Required</label>
        <h4>Validation Checklist</h4>
        {''.join([f"<label style='display:block;margin:4px 0'><input type='checkbox' name='new_rule_{_esc(k)}' value='1'> {_esc(v)}</label>" for k,v in CLAIM_RULE_LABELS.items()])}
      </div>
      <button class='btn' type='submit'>Save Claims Rules</button>
    </form>
    <form method='post' action='/admin/claims_rules/seed' style='margin-top:12px' onsubmit="return confirm('Reset all claim document rules to defaults?')"><button class='btn btn-muted' type='submit'>Reset to Default Rules</button></form>
    </div></body></html>"""
    return body


@app.route('/admin/claims_rules/save', methods=['POST'])
def admin_claims_rules_save():
    ensure_database_schema()
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('dashboard'))
    with engine.begin() as conn:
        _seed_claim_document_rule_defaults(conn)
        type_ids = [int(x) for x in request.form.getlist('type_ids') if str(x).isdigit()]
        for type_id in type_ids:
            conn.execute(text("""
                UPDATE app_claim_document_types
                SET document_type=:document_type, claim_type=:claim_type, aliases=:aliases,
                    required=:required, is_active=:is_active, sort_order=:sort_order, updated_at=NOW()
                WHERE id=:id
            """), {
                'id': type_id,
                'document_type': (request.form.get(f'document_type_{type_id}') or '').strip() or 'Document',
                'claim_type': (request.form.get(f'claim_type_{type_id}') or '').strip() or 'Funeral Claim',
                'aliases': (request.form.get(f'aliases_{type_id}') or '').strip(),
                'required': bool(request.form.get(f'required_{type_id}')),
                'is_active': bool(request.form.get(f'active_{type_id}')),
                'sort_order': int(safe_float(request.form.get(f'sort_order_{type_id}'), 100)),
            })
            existing = {r[0] for r in conn.execute(text('SELECT rule_key FROM app_claim_document_rules WHERE document_type_id=:id'), {'id': type_id}).fetchall()}
            for key, label in CLAIM_RULE_LABELS.items():
                enabled = bool(request.form.get(f'rule_{type_id}_{key}'))
                if key in existing:
                    conn.execute(text('UPDATE app_claim_document_rules SET is_enabled=:enabled, rule_label=:label WHERE document_type_id=:id AND rule_key=:key'), {'enabled': enabled, 'label': label, 'id': type_id, 'key': key})
                else:
                    conn.execute(text('INSERT INTO app_claim_document_rules (document_type_id, rule_key, rule_label, is_enabled) VALUES (:id,:key,:label,:enabled)'), {'id': type_id, 'key': key, 'label': label, 'enabled': enabled})
        new_doc = (request.form.get('new_document_type') or '').strip()
        if new_doc:
            row = conn.execute(text("""
                INSERT INTO app_claim_document_types (document_type, claim_type, aliases, required, is_active, sort_order, updated_at)
                VALUES (:document_type, :claim_type, :aliases, :required, true, 999, NOW()) RETURNING id
            """), {
                'document_type': new_doc,
                'claim_type': (request.form.get('new_claim_type') or 'Funeral Claim').strip() or 'Funeral Claim',
                'aliases': (request.form.get('new_aliases') or '').strip(),
                'required': bool(request.form.get('new_required')),
            }).mappings().first()
            new_id = int(row['id'])
            for key, label in CLAIM_RULE_LABELS.items():
                conn.execute(text('INSERT INTO app_claim_document_rules (document_type_id, rule_key, rule_label, is_enabled) VALUES (:id,:key,:label,:enabled)'), {'id': new_id, 'key': key, 'label': label, 'enabled': bool(request.form.get(f'new_rule_{key}'))})
    log_audit('claims_rules_updated', 'Admin updated claim document validation rules')
    flash('Claim document validation rules saved.', 'success')
    return redirect(url_for('admin_claims_rules'))


@app.route('/admin/claims_rules/seed', methods=['POST'])
def admin_claims_rules_seed():
    ensure_database_schema()
    engine = get_db_engine()
    if engine is not None:
        with engine.begin() as conn:
            _seed_claim_document_rule_defaults(conn, force=True)
    flash('Default claim document validation rules restored.', 'success')
    return redirect(url_for('admin_claims_rules'))

@app.route('/client-heatmap')
def client_heatmap_page():
    """Full-page client heatmap opened from the Policies / Clients dashboard card."""
    selected = (request.args.get('franchise') or 'All').strip() or 'All'
    try:
        franchises = ['All'] + [x for x in _available_franchises_for_user() if str(x).strip()]
    except Exception:
        franchises = ['All']
    return render_template(
        'client_heatmap.html',
        selected=selected,
        franchises=franchises,
        google_maps_api_key=get_google_maps_api_key(),
    )




@app.route('/api/franchise-map-locations')
def api_franchise_map_locations():
    """Return cached franchise address coordinates for the dashboard mini map.

    This endpoint is intentionally franchise-only and never returns client/member
    addresses. It is instant because it only reads saved lat/lng from the cache.
    """
    selected = (request.args.get('franchise') or 'All').strip() or 'All'
    rows = _get_franchise_map_locations(selected=selected, max_items=int(os.getenv('DASHBOARD_FRANCHISE_MAP_MAX_POINTS', '500')))
    summary = _franchise_map_cache_summary(selected)
    return jsonify({'selected': selected, 'count': len(rows), 'locations': rows, 'summary': summary})



@app.route('/api/client-map-cache-summary')
def api_client_map_cache_summary():
    """Return live client-map coordinate preparation progress for the heatmap page."""
    selected = (request.args.get('franchise') or 'All').strip() or 'All'
    try:
        summary = _client_map_cache_summary(selected)
        summary.update(_franchise_map_cache_summary(selected))
        groups_total = int(summary.get('groups_total') or 0)
        groups_cached = int(summary.get('groups_cached') or 0)
        groups_missing = int(summary.get('groups_missing') or 0)
        clients_total = int(summary.get('clients_total') or 0)
        clients_cached = int(summary.get('clients_cached') or 0)
        pct = round((groups_cached / groups_total * 100), 2) if groups_total else 100
        return jsonify({
            'ok': True,
            'selected': selected,
            'groups_total': groups_total,
            'groups_cached': groups_cached,
            'groups_missing': groups_missing,
            'clients_total': clients_total,
            'clients_cached': clients_cached,
            'percent': pct,
            'summary': summary,
        })
    except Exception as exc:
        return jsonify({'ok': False, 'error': str(exc), 'selected': selected}), 500

@app.route('/api/client-map-locations-page')
def api_client_map_locations_page():
    """Return prepared client map coordinates in pages. No geocoding or address scanning here."""
    selected = (request.args.get('franchise') or 'All').strip() or 'All'
    density_mode = (request.args.get('mode') or 'auto').strip().lower()
    offset = max(0, int(safe_float(request.args.get('offset'), 0)))
    page_size = int(safe_float(request.args.get('limit'), os.getenv('CLIENT_MAP_PAGE_SIZE', '500')))
    page_size = max(50, min(page_size, 3000))
    data = _query_client_map_points(selected=selected, offset=offset, limit=page_size, mode=density_mode)
    total_groups = int(data.get('total_groups') or 0)
    locations = data.get('locations') or []
    next_offset = offset + page_size if offset + page_size < total_groups else None
    mapped_total = sum(int(x.get('count') or 0) for x in locations)
    return jsonify({
        'selected': selected,
        'offset': offset,
        'page_size': page_size,
        'next_offset': next_offset,
        'done': next_offset is None,
        'total_groups': total_groups,
        'raw_count': len(locations),
        'count': len(locations),
        'client_total': int(data.get('clients_total') or 0),
        'mapped_client_total': mapped_total,
        'locations': locations,
    })

@app.route('/api/client-map-points-summary')
def api_client_map_points_summary():
    selected = (request.args.get('franchise') or 'All').strip() or 'All'
    return jsonify({'ok': True, 'selected': selected, 'summary': _client_map_points_summary(selected)})

@app.route('/api/client-map-locations')
def api_client_map_locations():
    """Fast compatibility endpoint for prepared Google map points.

    This route must never scan raw PolicyData rows. It reads only the prepared
    app_client_map_points table so opening the URL or loading a map cannot freeze
    the browser/server. Use /api/client-map-locations-page for pagination on the
    full heatmap page.
    """
    selected = (request.args.get('franchise') or 'All').strip() or 'All'
    density_mode = (request.args.get('mode') or 'auto').strip().lower()
    dashboard_preview = str(request.args.get('dashboard_preview') or '').strip() == '1'
    try:
        requested_limit = int(safe_float(request.args.get('limit'), 0))
    except Exception:
        requested_limit = 0
    if requested_limit <= 0:
        requested_limit = int(os.getenv('CLIENT_MAP_COMPAT_LIMIT', '1000'))
    if dashboard_preview:
        requested_limit = min(requested_limit, int(os.getenv('DASHBOARD_MAP_PREVIEW_LIMIT', '350')))
    # Hard cap protects localhost and Render/PostgreSQL from massive single JSON payloads.
    requested_limit = max(50, min(requested_limit, int(os.getenv('CLIENT_MAP_COMPAT_HARD_LIMIT', '3000'))))
    data = _query_client_map_points(selected=selected, offset=0, limit=requested_limit, mode=density_mode)
    locations = data.get('locations') or []
    return jsonify({
        'ok': True,
        'selected': selected,
        'count': len(locations),
        'raw_count': len(locations),
        'total_groups': int(data.get('total_groups') or 0),
        'client_total': int(data.get('clients_total') or 0),
        'mapped_client_total': sum(int(x.get('count') or 0) for x in locations),
        'limit': requested_limit,
        'truncated': int(data.get('total_groups') or 0) > len(locations),
        'message': 'Fast endpoint uses prepared coordinates only. Use /api/client-map-locations-page for paged loading.',
        'locations': locations,
    })


@app.route('/api/client-map-geocode-cache', methods=['POST'])
def api_client_map_geocode_cache():
    """Save browser-geocoded client coordinates into the PostgreSQL cache.

    This supports browser-restricted Google Maps keys.  The page geocodes client
    addresses with the Maps JavaScript API, posts the coordinates here, and future
    loads use the saved lat/lng directly.
    """
    engine = get_db_engine()
    if engine is None:
        return jsonify({'ok': False, 'saved': 0, 'error': 'Database not connected'}), 200
    payload = request.get_json(silent=True) or {}
    rows = payload.get('rows') or []
    saved = 0
    try:
        with engine.begin() as conn:
            _ensure_client_geocode_cache(conn)
            for r in rows[:500]:
                address = str(r.get('address') or '').strip()
                if not address:
                    continue
                try:
                    lat = float(r.get('lat'))
                    lng = float(r.get('lng'))
                except Exception:
                    continue
                if not (-35.5 <= lat <= -21.0 and 16.0 <= lng <= 33.5):
                    # Keep the map focused on South Africa and avoid bad geocodes.
                    continue
                conn.execute(text("""
                    INSERT INTO app_client_geocode_cache (address_key, address, lat, lng, status, updated_at)
                    VALUES (:address_key, :address, :lat, :lng, 'OK_BROWSER', NOW())
                    ON CONFLICT (address_key) DO UPDATE SET
                        address = EXCLUDED.address,
                        lat = EXCLUDED.lat,
                        lng = EXCLUDED.lng,
                        status = EXCLUDED.status,
                        updated_at = NOW()
                """), {
                    'address_key': _address_cache_key(address),
                    'address': address,
                    'lat': lat,
                    'lng': lng,
                })
                saved += 1
    except Exception as exc:
        return jsonify({'ok': False, 'saved': saved, 'error': str(exc)}), 200
    return jsonify({'ok': True, 'saved': saved})


def _filter_by_franchise(monthly, periods, selected):
    if selected and selected != 'All' and not monthly.empty:
        return monthly[monthly['Franchise'] == selected], periods[periods['Franchise'] == selected]
    return monthly, periods


def _period_base(monthly_view, periods_view, period_view, traffic_filter='All'):
    """Return the selected analysis base with real month/year/6-month labels."""
    rates = LAST_RESULT.get('rates', DEFAULT_RATES.copy())
    if monthly_view is None or monthly_view.empty:
        return pd.DataFrame()

    def _existing_agg(frame):
        agg = {
            'Retail Premium':'sum','Original Risk Premium':'sum','Risk Premium':'sum',
            'Underwriter 2.1% Fee':'sum','Claims':'sum','Claim Count':'sum',
            'Claim Paid to Franchise':'sum','Claim Paid to Client':'sum',
            'Repudiated / Pending':'sum','Grand Total Claims':'sum','Policy Qty':'sum',
            'R1 Policy Fee':'sum','Total Commission':'sum','Total Paid Commissions':'sum',
            'Separate Paid Commissions':'sum','BrightRock Month Total':'sum',
            'BrightRock Running Balance':'last','MFF Book Value 2.5%':'sum',
            'Franchise Book Value 2.5%':'sum','Total Book Value':'sum',
            'Scenario 1 Value':'sum','Scenario 2 Value':'sum',
            'Months Below 75% Streak':'last','BrightRock Eligible':'max'
        }
        for name in rates:
            agg[f'{name} Amount'] = 'sum'
        return {k:v for k,v in agg.items() if k in frame.columns}

    if period_view == 'year':
        base = monthly_view.copy()
        base['Month'] = pd.to_datetime(base.get('Month'), errors='coerce')
        base = base.dropna(subset=['Month'])
        base['_sort_month'] = base['Month'].dt.to_period('Y').dt.to_timestamp()
        base['Period View'] = base['_sort_month'].dt.year.astype(str)
        base = base.groupby(['Franchise', 'Period View', '_sort_month'], as_index=False).agg(_existing_agg(base))
    elif period_view == 'six_months':
        if periods_view is not None and not periods_view.empty and 'Period' in periods_view.columns:
            base = periods_view.copy()
            base['Period View'] = base['Period'].astype(str)
            extracted = base['Period View'].str.extract(r'([A-Za-z]{3} \d{4})')[0]
            base['_sort_month'] = pd.to_datetime(extracted, format='%b %Y', errors='coerce')
            if base['_sort_month'].isna().all():
                temp = monthly_view.copy()
                temp['Month'] = pd.to_datetime(temp.get('Month'), errors='coerce')
                temp = temp.dropna(subset=['Month']).sort_values(['Franchise','Month'])
                temp['Period Index'] = temp.groupby('Franchise').cumcount() // 6 + 1
                sort_lookup = temp.groupby(['Franchise','Period Index'])['Month'].min().reset_index()
                if 'Period Index' in base.columns:
                    base = base.merge(sort_lookup.rename(columns={'Month':'_sort_lookup'}), on=['Franchise','Period Index'], how='left')
                    base['_sort_month'] = base['_sort_lookup']
                    base.drop(columns=['_sort_lookup'], inplace=True, errors='ignore')
        else:
            base = monthly_view.copy()
            base['Month'] = pd.to_datetime(base.get('Month'), errors='coerce')
            base = base.dropna(subset=['Month']).sort_values(['Franchise','Month'])
            base['Period Index'] = base.groupby('Franchise').cumcount() // 6 + 1
            base['_period_start'] = base.groupby(['Franchise','Period Index'])['Month'].transform('min')
            base['_period_end'] = base.groupby(['Franchise','Period Index'])['Month'].transform('max')
            base['Period View'] = base.apply(lambda r: r['_period_start'].strftime('%b %Y') if r['_period_start'].to_period('M') == r['_period_end'].to_period('M') else f"{r['_period_start'].strftime('%b %Y')} - {r['_period_end'].strftime('%b %Y')}", axis=1)
            base['_sort_month'] = base['_period_start']
            base = base.groupby(['Franchise','Period View','_sort_month'], as_index=False).agg(_existing_agg(base))
    else:
        base = monthly_view.copy()
        base['Month'] = pd.to_datetime(base.get('Month'), errors='coerce')
        base = base.dropna(subset=['Month'])
        base['_sort_month'] = base['Month']
        base['Period View'] = base['Month'].dt.strftime('%b %Y')

    if base.empty:
        return base
    base['Claim Ratio'] = base.apply(lambda r: (r.get('Claims', 0) / r.get('Risk Premium', 0) * 100) if r.get('Risk Premium', 0) else 0, axis=1)
    base['Average Claim'] = base.apply(lambda r: (r.get('Claims', 0) / r.get('Claim Count', 0)) if r.get('Claim Count', 0) else 0, axis=1)
    base['Claim Ratio Status'] = base['Claim Ratio'].apply(status_colour)
    base['Claim Ratio Label'] = base['Claim Ratio'].apply(status_label)
    if traffic_filter in {'Blue','Green','Red'}:
        base = base[base['Claim Ratio Status'] == traffic_filter]
    if '_sort_month' not in base.columns:
        base['_sort_month'] = pd.to_datetime(base.get('Month'), errors='coerce')
    sort_cols = [c for c in ['Franchise','_sort_month','Period View'] if c in base.columns]
    base = base.sort_values(sort_cols).drop(columns=['_sort_month'], errors='ignore')
    return base


def _money_view(base):
    if base.empty:
        return base
    out = base.copy()
    if 'Total Paid Commissions' not in out.columns:
        out['Total Paid Commissions'] = out.get('Total Commission', 0) + out.get('R1 Policy Fee', 0) + out.get('Underwriter 2.1% Fee', 0)
    out['Franchise Money'] = out['Retail Premium'] - out['Claims'] - out['Total Paid Commissions']
    out['Accumulated Balance'] = out.groupby('Franchise')['Franchise Money'].cumsum()
    out['Money Status'] = out['Franchise Money'].apply(lambda v: 'Losing money' if v < 0 else 'Making money')
    out['Money Traffic'] = out['Franchise Money'].apply(lambda v: 'Red' if v < 0 else 'Green')
    return out[[c for c in ['Franchise','Period View','Retail Premium','Risk Premium','Claims','Claim Count','Average Claim','Claim Paid to Franchise','Claim Paid to Client','Repudiated / Pending','Grand Total Claims','Claim Ratio','Total Commission','R1 Policy Fee','Underwriter 2.1% Fee','Total Paid Commissions','Franchise Money','Accumulated Balance','Money Status','Money Traffic'] if c in out.columns]]


def _commission_view(base):
    if base.empty:
        return base
    out = base.copy()
    if 'Policy Qty' not in out.columns:
        out['Policy Qty'] = 0
    if 'R1 Policy Fee' not in out.columns:
        out['R1 Policy Fee'] = out['Policy Qty'] * 1.0
    return out


def _book_view(base):
    if base.empty:
        return base
    out=base.copy()
    for col in ['MFF Book Value 2.5%','Franchise Book Value 2.5%','Total Book Value']:
        if col not in out.columns:
            out[col]=0.0
    return out[[c for c in ['Franchise','Period View','Retail Premium','Risk Premium','Claims','Claim Count','Claim Paid to Franchise','Claim Paid to Client','Repudiated / Pending','Grand Total Claims','Claim Ratio','MFF Book Value 2.5%','Franchise Book Value 2.5%','Total Book Value'] if c in out.columns]]


def _compare_full(monthly_view, args):
    base, _ = build_scenario_comparison_view(monthly_view, args)
    return base


def _advanced_view(periods_view):
    if periods_view.empty:
        return periods_view
    out = periods_view.copy().sort_values(['Franchise','Period'])
    out['Weighted Claim Ratio'] = out.apply(lambda r: (r.get('Claims', 0) / r.get('Risk Premium', 0) * 100) if r.get('Risk Premium', 0) else 0, axis=1)
    out['Claim Ratio'] = out['Weighted Claim Ratio']
    out['Period View'] = out.get('Period', '')
    out['Risk Score'] = (out['Weighted Claim Ratio'] / 10).round(1)
    out['Previous Ratio'] = out.groupby('Franchise')['Weighted Claim Ratio'].shift(1)
    out['Trend'] = out.apply(lambda r: 'Improving' if pd.notna(r['Previous Ratio']) and r['Weighted Claim Ratio'] < r['Previous Ratio'] else ('Worsening' if pd.notna(r['Previous Ratio']) and r['Weighted Claim Ratio'] > r['Previous Ratio'] else 'Stable'), axis=1)
    out['Opportunity'] = out['Scenario 2 Value'] - out['Scenario 1 Value']
    out['Franchise Money on 100%'] = out.get('Scenario 1 Value', 0)
    out['Franchise Money on BrightRock'] = out.get('Scenario 2 Value', 0)
    return out.sort_values('Weighted Claim Ratio', ascending=False)


def _ai_answer(question):
    q=(question or '').lower()
    periods=LAST_RESULT['periods']
    if not question:
        return None
    if periods.empty:
        return {'summary':'No imported data is available yet.','columns':[],'rows':[]}
    if 'move' in q or 'brightrock' in q:
        df=periods[['Franchise','Period','Weighted Claim Ratio','Claim Ratio Label']].copy().sort_values('Weighted Claim Ratio')
        df['Weighted Claim Ratio']=df['Weighted Claim Ratio'].map(lambda v: f'{v:.2f}%')
        return {'summary':'These are the franchises ranked by readiness to move to BrightRock.','columns':list(df.columns),'rows':df.head(20).to_dict('records')}
    if 'risk' in q or 'worst' in q:
        df=periods[['Franchise','Period','Weighted Claim Ratio','Claims']].copy().sort_values('Weighted Claim Ratio', ascending=False)
        df['Weighted Claim Ratio']=df['Weighted Claim Ratio'].map(lambda v: f'{v:.2f}%')
        df['Claims']=df['Claims'].map(money)
        return {'summary':'These are the highest risk franchises by claim ratio.','columns':list(df.columns),'rows':df.head(20).to_dict('records')}
    if 'book' in q:
        df=periods[['Franchise','Period','Retail Premium','Total Book Value']].copy().sort_values('Total Book Value', ascending=False)
        for c in ['Retail Premium','Total Book Value']: df[c]=df[c].map(money)
        return {'summary':'These franchises contribute the most book value.','columns':list(df.columns),'rows':df.head(20).to_dict('records')}
    df=periods[['Franchise','Period','Retail Premium','Risk Premium','Claims','Weighted Claim Ratio','BrightRock Month Total','Total Book Value']].head(20).copy()
    df['Weighted Claim Ratio']=df['Weighted Claim Ratio'].map(lambda v: f'{v:.2f}%')
    for c in ['Retail Premium','Risk Premium','Claims','BrightRock Month Total','Total Book Value']: df[c]=df[c].map(money)
    return {'summary':'Here is a general summary from the imported data. Ask about risk, move to BrightRock, book value, or profitability for more focused answers.','columns':list(df.columns),'rows':df.to_dict('records')}




def _commission_business_view(base):
    """Commission comparison by business for 100% and BrightRock workspaces.

    In this build both scenarios use the selected commission settings, but they are shown
    side-by-side so users can validate supplier totals by period.
    """
    if base.empty:
        return base
    out = base.copy()
    out['BR 100% Amount'] = out.get('BrightRock Amount', 0)
    out['BR BrightRock Amount'] = out.get('BrightRock Amount', 0)
    out['Inkulu 100% Amount'] = out.get('Inkulu Amount', 0)
    out['Inkulu BrightRock Amount'] = out.get('Inkulu Amount', 0)
    out['MFF 100% Amount'] = out.get('MFF Amount', 0)
    out['MFF BrightRock Amount'] = out.get('MFF Amount', 0)
    return out


def _underwriter_allocation_view(base):
    """Reconciliation view for PolicyData underwriter calculation.

    Shows the exact steps: Original Risk - R1 per paid policy month, then 2.1% ADV fee,
    leaving the net underwriter premium used as Risk Premium in the analytics.
    """
    if base.empty:
        return base
    out = base.copy()
    if 'Original Risk Premium' not in out.columns:
        out['Original Risk Premium'] = out.get('Risk Premium', 0)
    if 'R1 Policy Fee' not in out.columns:
        out['R1 Policy Fee'] = out.get('Policy Qty', 0) * 1.0
    if 'Risk After R1' not in out.columns:
        out['Risk After R1'] = out['Original Risk Premium'] - out['R1 Policy Fee']
    out.loc[out['Risk After R1'] < 0, 'Risk After R1'] = 0
    if 'ADV Fee 2.1%' not in out.columns:
        out['ADV Fee 2.1%'] = out.get('Underwriter 2.1% Fee', out['Risk After R1'] * 0.021)
    out['Calculated ADV Check'] = out['Risk After R1'] * 0.021
    out['Underwriter Premium Payable'] = out['Original Risk Premium'] - out['R1 Policy Fee'] - out['ADV Fee 2.1%']
    out['Retail Less Payover'] = out.get('Retail Premium', 0) - out['Underwriter Premium Payable']
    return out[[c for c in ['Franchise','Period View','Retail Premium','Original Risk Premium','Policy Qty','R1 Policy Fee','Risk After R1','ADV Fee 2.1%','Calculated ADV Check','Underwriter Premium Payable','Retail Less Payover','Risk Premium','Claims','Claim Ratio'] if c in out.columns]]



def policy_rows_only(raw):
    """Return only rows that came from PolicyData, removing old claims-only rows."""
    if raw is None or raw.empty:
        return raw
    work = raw.copy()
    for col in ['retail_premium', 'risk_premium', 'policy_qty']:
        if col not in work.columns:
            work[col] = 0
    retail = pd.to_numeric(work['retail_premium'], errors='coerce').fillna(0)
    risk = pd.to_numeric(work['risk_premium'], errors='coerce').fillna(0)
    policies = pd.to_numeric(work['policy_qty'], errors='coerce').fillna(0)
    return work[(retail != 0) | (risk != 0) | (policies != 0)].copy()


def save_claim_aliases_to_postgres(aliases):
    """Persist saved claim mappings to PostgreSQL so DBeaver/dashboard rebuilds use them."""
    engine = get_db_engine()
    if engine is None or not aliases:
        return False
    rows = []
    for src, dst in aliases.items():
        src_key = franchise_match_key(src)
        dst_key = franchise_match_key(dst)
        if src_key and dst_key:
            rows.append({'source_name': src_key, 'mapped_name': dst_key, 'approved': True})
    if not rows:
        return False
    with engine.begin() as conn:
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS franchise_mapping_pg (
                id BIGSERIAL PRIMARY KEY,
                source_name TEXT NOT NULL,
                mapped_name TEXT NOT NULL,
                approved BOOLEAN DEFAULT true,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """))
        conn.execute(text("""
            CREATE UNIQUE INDEX IF NOT EXISTS uq_franchise_mapping_pg_source_key
            ON franchise_mapping_pg (LOWER(TRIM(source_name)))
        """))
        for row in rows:
            existing = conn.execute(text("""
                SELECT id FROM franchise_mapping_pg
                WHERE LOWER(TRIM(source_name)) = LOWER(TRIM(:source_name))
                LIMIT 1
            """), row).mappings().first()
            if existing:
                conn.execute(text("""
                    UPDATE franchise_mapping_pg
                    SET mapped_name = :mapped_name, approved = true
                    WHERE id = :id
                """), {'mapped_name': row['mapped_name'], 'id': existing['id']})
            else:
                conn.execute(text("""
                    INSERT INTO franchise_mapping_pg (source_name, mapped_name, approved)
                    VALUES (:source_name, :mapped_name, :approved)
                """), row)
    return True


def recalculate_claims_summary_from_postgres():
    """Rebuild franchise_monthly_summary from policy + mapped claims raw tables.

    PolicyData is the base list for dashboard rows. Claims are then written in once,
    using claim_key -> franchise_mapping_pg.source_name -> mapped_name. If a mapped
    claims key has no matching policy row for that month, it is inserted as a
    claims-only reconciliation row with zero premiums so total claims still audit.
    """
    engine = get_db_engine()
    if engine is None:
        return False
    with engine.begin() as conn:
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS franchise_mapping_pg (
                id BIGSERIAL PRIMARY KEY,
                source_name TEXT NOT NULL,
                mapped_name TEXT NOT NULL,
                approved BOOLEAN DEFAULT true,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """))
        conn.execute(text("""
            CREATE TABLE IF NOT EXISTS franchise_monthly_summary (
                id BIGSERIAL PRIMARY KEY,
                franchise_name TEXT NOT NULL,
                report_month DATE NOT NULL,
                retail_premium NUMERIC(18,2) DEFAULT 0,
                risk_premium NUMERIC(18,2) DEFAULT 0,
                claims NUMERIC(18,2) DEFAULT 0,
                claim_count INTEGER DEFAULT 0,
                claim_ratio NUMERIC(10,4) DEFAULT 0,
                brightrock_commission NUMERIC(18,2) DEFAULT 0,
                mkhulu_commission NUMERIC(18,2) DEFAULT 0,
                mff_commission NUMERIC(18,2) DEFAULT 0,
                r1_fee NUMERIC(18,2) DEFAULT 0,
                adv_fund_fee NUMERIC(18,2) DEFAULT 0,
                net_risk_premium NUMERIC(18,2) DEFAULT 0,
                created_at TIMESTAMP DEFAULT NOW(),
                UNIQUE(franchise_name, report_month)
            )
        """))
        conn.execute(text("ALTER TABLE franchise_monthly_summary ADD COLUMN IF NOT EXISTS claim_count INTEGER DEFAULT 0"))
        conn.execute(text("DELETE FROM franchise_monthly_summary"))
        conn.execute(text("""
            INSERT INTO franchise_monthly_summary (
                franchise_name, report_month, retail_premium, risk_premium,
                claims, claim_count, claim_ratio,
                brightrock_commission, mkhulu_commission, mff_commission,
                r1_fee, adv_fund_fee, net_risk_premium
            )
            SELECT
                p.franchise_name,
                p.import_month,
                COALESCE(p.retail_premium, 0),
                COALESCE(p.risk_premium, 0),
                0,
                0,
                0,
                ROUND(COALESCE(p.retail_premium, 0) * 0.10, 2),
                ROUND(COALESCE(p.retail_premium, 0) * 0.025, 2),
                ROUND(COALESCE(p.retail_premium, 0) * 0.025, 2),
                COALESCE(p.r1_policy_fee, 0),
                COALESCE(p.underwriter_2_1_fee, 0),
                COALESCE(p.risk_after_r1, p.risk_premium, 0)
            FROM policy_monthly_raw p
            ON CONFLICT (franchise_name, report_month)
            DO UPDATE SET
                retail_premium = EXCLUDED.retail_premium,
                risk_premium = EXCLUDED.risk_premium,
                brightrock_commission = EXCLUDED.brightrock_commission,
                mkhulu_commission = EXCLUDED.mkhulu_commission,
                mff_commission = EXCLUDED.mff_commission,
                r1_fee = EXCLUDED.r1_fee,
                adv_fund_fee = EXCLUDED.adv_fund_fee,
                net_risk_premium = EXCLUDED.net_risk_premium,
                claims = 0,
                claim_count = 0,
                claim_ratio = 0
        """))
        # Push claims into existing policy rows using the mapped key contained in the full policy franchise name.
        conn.execute(text("""
            WITH mapped_claims AS (
                SELECT
                    LOWER(TRIM(COALESCE(m.mapped_name, c.claim_key))) AS mapped_key,
                    DATE_TRUNC('month', c.claim_month)::date AS claim_month,
                    SUM(c.claims_amount) AS claims,
                    SUM(c.claim_count) AS claim_count
                FROM claims_monthly_raw c
                LEFT JOIN franchise_mapping_pg m
                    ON LOWER(TRIM(c.claim_key)) = LOWER(TRIM(m.source_name))
                   AND m.approved = true
                GROUP BY
                    LOWER(TRIM(COALESCE(m.mapped_name, c.claim_key))),
                    DATE_TRUNC('month', c.claim_month)::date
            )
            UPDATE franchise_monthly_summary s
            SET
                claims = COALESCE(x.claims, 0),
                claim_count = COALESCE(x.claim_count, 0),
                claim_ratio = CASE
                    WHEN COALESCE(s.risk_premium, 0) = 0 THEN 0
                    ELSE ROUND((COALESCE(x.claims, 0) / s.risk_premium) * 100, 2)
                END
            FROM mapped_claims x
            WHERE DATE_TRUNC('month', s.report_month)::date = x.claim_month
              AND LOWER(TRIM(s.franchise_name)) LIKE '%' || x.mapped_key || '%'
        """))
        # Insert mapped claim rows that still did not find a policy row for that month.
        conn.execute(text("""
            WITH mapped_claims AS (
                SELECT
                    UPPER(TRIM(COALESCE(m.mapped_name, c.claim_key))) AS franchise_name,
                    LOWER(TRIM(COALESCE(m.mapped_name, c.claim_key))) AS mapped_key,
                    DATE_TRUNC('month', c.claim_month)::date AS report_month,
                    SUM(c.claims_amount) AS claims,
                    SUM(c.claim_count) AS claim_count
                FROM claims_monthly_raw c
                LEFT JOIN franchise_mapping_pg m
                    ON LOWER(TRIM(c.claim_key)) = LOWER(TRIM(m.source_name))
                   AND m.approved = true
                GROUP BY
                    UPPER(TRIM(COALESCE(m.mapped_name, c.claim_key))),
                    LOWER(TRIM(COALESCE(m.mapped_name, c.claim_key))),
                    DATE_TRUNC('month', c.claim_month)::date
            ), unmatched AS (
                SELECT x.*
                FROM mapped_claims x
                WHERE NOT EXISTS (
                    SELECT 1
                    FROM franchise_monthly_summary s
                    WHERE DATE_TRUNC('month', s.report_month)::date = x.report_month
                      AND LOWER(TRIM(s.franchise_name)) LIKE '%' || x.mapped_key || '%'
                )
            )
            INSERT INTO franchise_monthly_summary (
                franchise_name, report_month, retail_premium, risk_premium,
                claims, claim_count, claim_ratio
            )
            SELECT franchise_name, report_month, 0, 0, claims, claim_count, 0
            FROM unmatched
            ON CONFLICT (franchise_name, report_month)
            DO UPDATE SET
                claims = EXCLUDED.claims,
                claim_count = EXCLUDED.claim_count,
                claim_ratio = 0
        """))
    return True


def get_mapping_health_from_postgres():
    """Return mapping health totals, independent from claim-ratio performance.

    A claim row is only treated as healthy when its claim_key has an approved saved
    mapping AND the mapped key resolves to an existing policy franchise.  This fixes
    the case where the summary shows an unmatched Rand value but the mapping table is
    empty because the source key was saved, while the target key did not resolve to a
    policy franchise.
    """
    engine = get_db_engine()
    empty = {'imported_claims':0.0,'mapped_claims':0.0,'unmapped_claims':0.0,'unmapped_count':0,'status':'No claims imported'}
    if engine is None:
        return empty
    try:
        with engine.begin() as conn:
            row = conn.execute(text("""
                WITH policy_keys AS (
                    SELECT DISTINCT LOWER(TRIM(franchise_name)) AS franchise_key
                    FROM policy_monthly_raw
                    WHERE franchise_name IS NOT NULL AND TRIM(franchise_name) <> ''
                    UNION
                    SELECT DISTINCT LOWER(TRIM(franchise_name)) AS franchise_key
                    FROM franchise_monthly_summary
                    WHERE franchise_name IS NOT NULL AND TRIM(franchise_name) <> ''
                ), checked AS (
                    SELECT
                        c.claims_amount,
                        CASE
                            WHEN m.id IS NULL THEN false
                            WHEN EXISTS (
                                SELECT 1 FROM policy_keys p
                                WHERE p.franchise_key LIKE '%%' || LOWER(TRIM(COALESCE(m.mapped_name, c.claim_key))) || '%%'
                                   OR LOWER(TRIM(COALESCE(m.mapped_name, c.claim_key))) LIKE '%%' || p.franchise_key || '%%'
                            ) THEN true
                            ELSE false
                        END AS is_mapped
                    FROM claims_monthly_raw c
                    LEFT JOIN franchise_mapping_pg m
                        ON LOWER(TRIM(c.claim_key)) = LOWER(TRIM(m.source_name))
                       AND m.approved = true
                )
                SELECT
                    COALESCE(SUM(claims_amount), 0) AS imported_claims,
                    COALESCE(SUM(CASE WHEN is_mapped THEN claims_amount ELSE 0 END), 0) AS mapped_claims,
                    COALESCE(SUM(CASE WHEN NOT is_mapped THEN claims_amount ELSE 0 END), 0) AS unmapped_claims,
                    COUNT(*) FILTER (WHERE NOT is_mapped) AS unmapped_count
                FROM checked
            """)).mappings().first()
        data = dict(row) if row else empty.copy()
        data['status'] = 'Healthy' if float(data.get('unmapped_claims') or 0) == 0 else 'Needs Mapping'
        return data
    except Exception:
        return empty


def get_claims_mapping_rows_from_postgres(limit=500):
    """Rows that still need attention on Claims Mapping.

    Includes both truly unmapped claims keys and saved mappings whose target does not
    resolve to an existing policy franchise key.  This prevents a non-zero unmatched
    total with an empty mapping table.
    """
    engine = get_db_engine()
    if engine is None:
        return []
    try:
        with engine.begin() as conn:
            rows = conn.execute(text("""
                WITH policy_keys AS (
                    SELECT DISTINCT LOWER(TRIM(franchise_name)) AS franchise_key
                    FROM policy_monthly_raw
                    WHERE franchise_name IS NOT NULL AND TRIM(franchise_name) <> ''
                    UNION
                    SELECT DISTINCT LOWER(TRIM(franchise_name)) AS franchise_key
                    FROM franchise_monthly_summary
                    WHERE franchise_name IS NOT NULL AND TRIM(franchise_name) <> ''
                ), grouped AS (
                    SELECT
                        c.claim_key,
                        MIN(c.claims_franchise_name) AS claims_franchise,
                        SUM(c.claims_amount) AS claims,
                        SUM(c.claim_count) AS claim_count,
                        MIN(c.claim_month) AS first_month,
                        MAX(c.claim_month) AS last_month,
                        MAX(m.id) AS mapping_id,
                        MAX(m.mapped_name) AS mapped_name,
                        BOOL_OR(
                            EXISTS (
                                SELECT 1 FROM policy_keys p
                                WHERE p.franchise_key LIKE '%%' || LOWER(TRIM(COALESCE(m.mapped_name, c.claim_key))) || '%%'
                                   OR LOWER(TRIM(COALESCE(m.mapped_name, c.claim_key))) LIKE '%%' || p.franchise_key || '%%'
                            )
                        ) AS target_resolves
                    FROM claims_monthly_raw c
                    LEFT JOIN franchise_mapping_pg m
                        ON LOWER(TRIM(c.claim_key)) = LOWER(TRIM(m.source_name))
                       AND m.approved = true
                    GROUP BY c.claim_key
                )
                SELECT
                    claims_franchise,
                    claim_key AS matched_key,
                    CASE
                        WHEN mapping_id IS NULL THEN 'Claims-only / unmatched'
                        WHEN NOT COALESCE(target_resolves, false) THEN 'Mapped target not found in PolicyData'
                        ELSE 'Matched'
                    END AS status,
                    claims,
                    claim_count,
                    mapped_name,
                    first_month,
                    last_month
                FROM grouped
                WHERE mapping_id IS NULL OR NOT COALESCE(target_resolves, false)
                ORDER BY claims DESC, claims_franchise
                LIMIT :limit
            """), {'limit': int(limit)}).mappings().all()
        return [dict(r) for r in rows]
    except Exception:
        return []

def _claims_reconciliation_view():
    """Build a franchise-by-franchise audit of claims workbook values vs analytics records.

    The reconciliation proves that claims imported from the claims workbook were written to
    the same franchise/month records used by the dashboard and claim ratio calculations.
    """
    claims_df = LAST_CLAIMS_DF.copy() if LAST_CLAIMS_DF is not None else pd.DataFrame()
    raw_monthly = LAST_RESULT.get('monthly', pd.DataFrame()).copy()
    if claims_df.empty and raw_monthly.empty:
        return {
            'summary': {'workbook_total':0.0,'system_total':0.0,'matched_total':0.0,'unmatched_total':0.0,'difference':0.0,'status':'No claims imported'},
            'rows': pd.DataFrame(),
            'monthly': pd.DataFrame(),
            'unmatched': pd.DataFrame(),
        }

    if not raw_monthly.empty:
        raw_monthly['audit_key'] = raw_monthly['Franchise'].map(franchise_match_key).map(apply_franchise_alias)
        existing_keys = list(raw_monthly['audit_key'].dropna().astype(str).unique())
        canonical_by_key = {}
        for k, name in zip(raw_monthly['audit_key'], raw_monthly['Franchise'].astype(str)):
            if k and k not in canonical_by_key:
                canonical_by_key[k] = name
    else:
        existing_keys = []
        canonical_by_key = {}

    if not claims_df.empty:
        claims = claims_df.copy()
        claims['month'] = pd.to_datetime(claims['month'], errors='coerce').dt.to_period('M').dt.to_timestamp()
        claims = claims.dropna(subset=['month'])
        if 'claim_key' not in claims.columns:
            claims['claim_key'] = claims['franchise'].map(franchise_match_key).map(apply_franchise_alias)
        else:
            claims['claim_key'] = claims['claim_key'].map(apply_franchise_alias)
        for col in ['claims','claim_count','claim_paid_franchise','claim_paid_client','repudiated_pending','grand_total_claims']:
            if col not in claims.columns:
                claims[col] = 0.0
            claims[col] = claims[col].apply(clean_money)
        claims['resolved_key'] = claims['claim_key'].apply(lambda k: resolve_claim_key_to_existing(k, existing_keys)) if existing_keys else claims['claim_key']
        claims['system_franchise'] = claims['resolved_key'].map(canonical_by_key).fillna('')
    else:
        claims = pd.DataFrame(columns=['franchise','claim_key','resolved_key','system_franchise','month','claims','claim_count'])

    workbook_total = float(claims['claims'].sum()) if not claims.empty else 0.0
    matched_claims = claims[claims['resolved_key'].isin(existing_keys)].copy() if existing_keys and not claims.empty else pd.DataFrame(columns=claims.columns)
    unmatched_claims = claims[~claims['resolved_key'].isin(existing_keys)].copy() if existing_keys and not claims.empty else claims.copy()
    matched_total = float(matched_claims['claims'].sum()) if not matched_claims.empty else 0.0
    unmatched_total = float(unmatched_claims['claims'].sum()) if not unmatched_claims.empty else 0.0

    system = pd.DataFrame()
    if not raw_monthly.empty:
        system = raw_monthly.copy()
        for col in ['Claims','Claim Count','Claim Paid to Franchise','Claim Paid to Client','Repudiated / Pending','Grand Total Claims']:
            if col not in system.columns:
                system[col] = 0.0
        system['month'] = pd.to_datetime(system['Month'], errors='coerce').dt.to_period('M').dt.to_timestamp()
        system = system.groupby(['audit_key','month'], as_index=False).agg({'Claims':'sum','Claim Count':'sum','Franchise':'first'})
    system_total = float(system['Claims'].sum()) if not system.empty else 0.0

    workbook_by_key = pd.DataFrame()
    if not claims.empty:
        if 'resolved_key' not in claims.columns:
            claims['resolved_key'] = claims.get('claim_key', pd.Series(dtype=str)).map(apply_franchise_alias)
        workbook_by_key = claims.groupby('resolved_key', as_index=False).agg({'claims':'sum','claim_count':'sum','franchise':'first','system_franchise':'first'})
    if 'resolved_key' not in workbook_by_key.columns:
        workbook_by_key = pd.DataFrame(columns=['resolved_key','claims','claim_count','franchise','system_franchise'])
    system_by_key = system.groupby('audit_key', as_index=False).agg({'Claims':'sum','Claim Count':'sum','Franchise':'first'}) if not system.empty else pd.DataFrame(columns=['audit_key','Claims','Claim Count','Franchise'])
    rows = workbook_by_key.merge(system_by_key, left_on='resolved_key', right_on='audit_key', how='outer')
    if rows.empty:
        rows = pd.DataFrame(columns=['Claims Franchise','System Franchise','Workbook Claims','System Claims','Difference','Claim Count','Status'])
    else:
        rows['Claims Franchise'] = rows.get('franchise','').fillna('')
        rows['System Franchise'] = rows.get('Franchise','').fillna(rows.get('system_franchise','')).fillna('')
        rows['Workbook Claims'] = rows.get('claims',0).fillna(0).astype(float)
        rows['System Claims'] = rows.get('Claims',0).fillna(0).astype(float)
        rows['Claim Count'] = rows.get('claim_count',0).fillna(rows.get('Claim Count',0)).fillna(0).astype(float)
        rows['Difference'] = rows['Workbook Claims'] - rows['System Claims']
        # Claims Reconciliation is now a performance traffic-light view.
        # Mapping health is shown separately; a financial variance is not a mapping failure.
        rows['Risk Premium'] = 0.0
        if not raw_monthly.empty and 'Risk Premium' in raw_monthly.columns:
            risk_by_key = raw_monthly.groupby('audit_key', as_index=False).agg({'Risk Premium':'sum'})
            rows = rows.merge(risk_by_key, left_on='resolved_key', right_on='audit_key', how='left', suffixes=('', '_risk'))
            if 'Risk Premium_risk' in rows.columns:
                rows['Risk Premium'] = rows['Risk Premium_risk'].fillna(0).astype(float)
                rows = rows.drop(columns=['Risk Premium_risk'])
        rows['Claim Ratio'] = rows.apply(lambda r: (float(r.get('System Claims',0)) / float(r.get('Risk Premium',0)) * 100) if float(r.get('Risk Premium',0) or 0) else 0.0, axis=1)
        def _status(r):
            if str(r.get('System Franchise','')).strip() == '':
                return 'Unmapped'
            ratio = float(r.get('Claim Ratio', 0) or 0)
            if ratio <= 60:
                return 'Green'
            if ratio <= 80:
                return 'Amber'
            return 'Red'
        rows['Status'] = rows.apply(_status, axis=1)
        rows = rows[['Claims Franchise','System Franchise','Workbook Claims','System Claims','Difference','Claim Count','Risk Premium','Claim Ratio','Status']].sort_values(['Status','Claim Ratio','System Claims'], ascending=[True,False,False])

    monthly = pd.DataFrame()
    if not claims.empty:
        wb_month = claims.groupby(['resolved_key','month'], as_index=False).agg({'claims':'sum','claim_count':'sum','franchise':'first','system_franchise':'first'})
        sys_month = system.groupby(['audit_key','month'], as_index=False).agg({'Claims':'sum','Claim Count':'sum','Franchise':'first'}) if not system.empty else pd.DataFrame(columns=['audit_key','month','Claims','Claim Count','Franchise'])
        monthly = wb_month.merge(sys_month, left_on=['resolved_key','month'], right_on=['audit_key','month'], how='outer')
        if not monthly.empty:
            monthly['Claims Franchise'] = monthly.get('franchise','').fillna('')
            monthly['System Franchise'] = monthly.get('Franchise','').fillna(monthly.get('system_franchise','')).fillna('')
            monthly['Month'] = pd.to_datetime(monthly['month'], errors='coerce').dt.strftime('%b %Y')
            monthly['Workbook Claims'] = monthly.get('claims',0).fillna(0).astype(float)
            monthly['System Claims'] = monthly.get('Claims',0).fillna(0).astype(float)
            monthly['Claim Count'] = monthly.get('claim_count',0).fillna(monthly.get('Claim Count',0)).fillna(0).astype(float)
            monthly['Difference'] = monthly['Workbook Claims'] - monthly['System Claims']
            monthly = monthly[['Claims Franchise','System Franchise','Month','Workbook Claims','System Claims','Difference','Claim Count']].sort_values(['System Franchise','Claims Franchise','Month'])

    mapping_health = get_mapping_health_from_postgres()
    summary = {
        'workbook_total': workbook_total,
        'matched_total': mapping_health.get('mapped_claims', matched_total),
        'unmatched_total': mapping_health.get('unmapped_claims', unmatched_total),
        'system_total': system_total,
        'difference': workbook_total - system_total,
        'status': mapping_health.get('status', 'Healthy'),
        'unmatched_count': int(mapping_health.get('unmapped_count', 0) or 0),
        'green_count': int((rows['Status'] == 'Green').sum()) if not rows.empty and 'Status' in rows.columns else 0,
        'amber_count': int((rows['Status'] == 'Amber').sum()) if not rows.empty and 'Status' in rows.columns else 0,
        'red_count': int((rows['Status'] == 'Red').sum()) if not rows.empty and 'Status' in rows.columns else 0,
        'mapping_health': mapping_health,
    }
    return {'summary': summary, 'rows': rows, 'monthly': monthly, 'unmatched': unmatched_claims, 'mapping_health': mapping_health}


def _empty_mapping_health(status='Not available'):
    return {
        'imported_claims': 0.0,
        'mapped_claims': 0.0,
        'unmapped_claims': 0.0,
        'unmapped_count': 0,
        'status': status,
    }


def _safe_claims_reconciliation_view():
    try:
        data = _claims_reconciliation_view()
        if not isinstance(data, dict):
            data = {}
        data.setdefault('mapping_health', _empty_mapping_health())
        data.setdefault('summary', {})
        data['summary'].setdefault('mapping_health', data['mapping_health'])
        return data
    except Exception as exc:
        mapping_health = _empty_mapping_health('Error')
        return {
            'summary': {
                'error': str(exc),
                'workbook_total': 0.0,
                'matched_total': 0.0,
                'unmatched_total': 0.0,
                'system_total': 0.0,
                'difference': 0.0,
                'status': 'Error',
                'green_count': 0,
                'amber_count': 0,
                'red_count': 0,
                'mapping_health': mapping_health,
            },
            'rows': pd.DataFrame(columns=['Claims Franchise','System Franchise','Workbook Claims','System Claims','Difference','Claim Count','Risk Premium','Claim Ratio','Status']),
            'monthly': pd.DataFrame(columns=['Claims Franchise','System Franchise','Month','Workbook Claims','System Claims','Difference','Claim Count']),
            'unmatched': pd.DataFrame(),
            'mapping_health': mapping_health,
        }

def _workspace_context(workspace):
    config = load_franchise_config()
    original_monthly = apply_user_franchise_scope(LAST_RESULT.get('monthly', pd.DataFrame()))
    monthly, periods, configured_portfolio = apply_franchise_config(original_monthly, config)
    franchises = sorted(monthly['Franchise'].unique()) if not monthly.empty and 'Franchise' in monthly.columns else []
    original_franchises = sorted(original_monthly['Franchise'].unique()) if not original_monthly.empty and 'Franchise' in original_monthly.columns else []

    selected = request.args.get('franchise', 'All')
    period_view = request.args.get('period_view', 'six_months')
    traffic_filter = request.args.get('traffic_filter', 'All')
    commission_focus = request.args.get('commission_focus', 'All')
    scenario_filter = request.args.get('scenario_filter', 'Both')

    if period_view not in {'month', 'six_months', 'year'}:
        period_view = 'six_months'
    if traffic_filter not in {'All', 'Blue', 'Green', 'Red'}:
        traffic_filter = 'All'
    if commission_focus not in {'All', 'BrightRock', 'Inkulu', 'MFF', 'R1 Policy Fee', 'Adv Fee', 'MFF Book Value', 'Franchise Book Value', 'Total Commissions'}:
        commission_focus = 'All'
    if scenario_filter not in {'Both', '100%', 'BrightRock'}:
        scenario_filter = 'Both'

    monthly_view, periods_view = _filter_by_franchise(monthly, periods, selected)

    titles = {
      'claim_ratio': ('Claim Ratio Analysis', 'Retail premium, risk premium, claims and claim ratio for the selected period.'),
      'move_recommendations': ('Move Recommendations', 'Only franchise name, claim ratio and recommendation.'),
      'franchise_money': ('Franchise Money Accumulation', 'Retail premium minus claims minus total commissions. Book value is excluded.'),
      'commission_analysis': ('Commission Analysis', 'Commission percentages and amounts per franchise.'),
      'commission_business': ('Commission Comparison by Business', 'Compare each business commission under 100% and BrightRock by month, 6 months or 12 months.'),
      'underwriter_allocation': ('Underwriter Allocation', 'Reconcile R1 policy fee, ADV fee 2.1%, and net underwriter premium from PolicyData.'),
      'franchise_configuration': ('Franchise Configuration', 'Exclude franchises from calculations/extracts and combine franchises into reporting groups.'),
      'claims_mapping': ('Claims Mapping', 'Map claims-only branch names to the correct policy franchise and save the mapping for future imports.'),
      'claims_reconciliation': ('Claims Reconciliation', 'Audit workbook claims against the franchise analytics records before trusting claim ratios.'),
      'book_value': ('Book Value Analysis', 'Book value calculated on retail premium.'),
      'scenario_comparison': ('Scenario Comparison', 'Excel-style commission allocation calculator using Current, Scenario A and Scenario B percentages.'),
      'advanced_analytics': ('Advanced Analytics', 'Risk ranking, trend, contributions, opportunity and early warning views.'),
      'ai_assistant': ('AI Data Assistant', 'Ask questions about the imported franchise data.'),
    }
    title, subtitle = titles.get(workspace, ('Workspace', ''))

    # Safe defaults for all templates. These prevent UnboundLocalError on pages that
    # do not build scenario comparison, mapping, reconciliation, or AI data.
    base = pd.DataFrame()
    money_view = pd.DataFrame()
    commission_view = pd.DataFrame()
    book_view = pd.DataFrame()
    compare_full = pd.DataFrame()
    compare_totals = pd.DataFrame()
    compare_summary = {}
    compare_args = {
      'compare_period_view': request.args.get('compare_period_view') or request.args.get('period_view', 'six_months'),
      'a_hollard': request.args.get('a_hollard', '6.00'),
      'a_mkhulu': request.args.get('a_mkhulu', '2.00'),
      'a_mff': request.args.get('a_mff', '4.00'),
      'a_r1': request.args.get('a_r1', '3.00'),
      'a_adv': request.args.get('a_adv', '5.00'),
      'b_hollard': request.args.get('b_hollard', '7.50'),
      'b_mkhulu': request.args.get('b_mkhulu', '2.50'),
      'b_mff': request.args.get('b_mff', '5.00'),
      'b_r1': request.args.get('b_r1', '2.00'),
      'b_adv': request.args.get('b_adv', '3.00'),
      'book_value_rate': request.args.get('book_value_rate', '2.50'),
    }
    advanced_view = pd.DataFrame()
    commission_business_view = pd.DataFrame()
    underwriter_view = pd.DataFrame()
    claims_reconciliation = {'summary': {}, 'rows': pd.DataFrame(), 'monthly': pd.DataFrame(), 'unmatched': pd.DataFrame(), 'mapping_health': _empty_mapping_health()}
    mh = None
    claims_mapping_unmatched_rows = pd.DataFrame()
    ai_answer = None

    # Build only what the requested page needs. This keeps KPI menu clicks fast and
    # avoids expensive reconciliation/mapping/scenario calculations in the background.
    if workspace in {'claim_ratio', 'move_recommendations', 'franchise_money', 'commission_analysis', 'commission_business', 'underwriter_allocation', 'book_value'}:
        base = _period_base(monthly_view, periods_view, period_view, traffic_filter)

    if workspace == 'franchise_money':
        money_view = _money_view(base)
    elif workspace == 'commission_analysis':
        commission_view = _commission_view(base)
    elif workspace == 'book_value':
        book_view = _book_view(base)
    elif workspace == 'commission_business':
        commission_business_view = _commission_business_view(base)
    elif workspace == 'underwriter_allocation':
        underwriter_view = _underwriter_allocation_view(base)
    elif workspace == 'scenario_comparison':
        compare_full = _compare_full(monthly_view, request.args)
        if not compare_full.empty:
            total_cols = [c for c in compare_full.columns if pd.api.types.is_numeric_dtype(compare_full[c]) and c not in {'Claim Ratio'}]
            compare_totals = compare_full.groupby('Period View', as_index=False)[total_cols].sum() if total_cols else pd.DataFrame()
            current_total = float(compare_full.get('Current Total Benefit', pd.Series(dtype=float)).sum())
            scenario_a_total = float(compare_full.get('Scenario A Total Benefit', pd.Series(dtype=float)).sum())
            scenario_b_total = float(compare_full.get('Scenario B Total Benefit', pd.Series(dtype=float)).sum())
            scenario_a_commission = float(compare_full.get('Scenario A Total Commission', pd.Series(dtype=float)).sum())
            scenario_b_commission = float(compare_full.get('Scenario B Total Commission', pd.Series(dtype=float)).sum())
            current_commission = float(compare_full.get('Current Total Commission', pd.Series(dtype=float)).sum())
            best_scenario = 'Scenario A' if scenario_a_total > scenario_b_total else ('Scenario B' if scenario_b_total > scenario_a_total else 'Scenario A = Scenario B')
            best_total = max(scenario_a_total, scenario_b_total)
            compare_summary = {
                'retail_total': float(compare_full.get('Retail Premium', pd.Series(dtype=float)).sum()),
                'current_total': current_total,
                'scenario_a_total': scenario_a_total,
                'scenario_b_total': scenario_b_total,
                'current_commission': current_commission,
                'scenario_a_commission': scenario_a_commission,
                'scenario_b_commission': scenario_b_commission,
                'scenario_a_commission_diff': scenario_a_commission - current_commission,
                'scenario_b_commission_diff': scenario_b_commission - current_commission,
                'scenario_b_commission_vs_a': scenario_b_commission - scenario_a_commission,
                'scenario_a_vs_current': scenario_a_total - current_total,
                'scenario_b_vs_current': scenario_b_total - current_total,
                'scenario_b_vs_a': scenario_b_total - scenario_a_total,
                'a_wins': int((compare_full.get('Best Scenario', pd.Series(dtype=str)) == 'Scenario A').sum()),
                'b_wins': int((compare_full.get('Best Scenario', pd.Series(dtype=str)) == 'Scenario B').sum()),
                'same': int((compare_full.get('Best Scenario', pd.Series(dtype=str)) == 'Same').sum()),
                'best_scenario': best_scenario,
                'best_total': best_total,
                'best_additional_benefit': best_total - current_total,
            }
            for prefix_key, prefix_col in [('current', 'Current'), ('scenario_a', 'Scenario A'), ('scenario_b', 'Scenario B')]:
                for component in ['Hollard', 'Mkhulu', 'MFF', 'R1', 'ADV']:
                    compare_summary[f'{prefix_key}_{component.lower()}_amount'] = float(compare_full.get(f'{prefix_col} {component} Amount', pd.Series(dtype=float)).sum())
                compare_summary[f'{prefix_key}_book_value'] = float(compare_full.get(f'{prefix_col} MFF Book Value', pd.Series(dtype=float)).sum())
    elif workspace == 'advanced_analytics':
        advanced_view = _advanced_view(periods_view)
    elif workspace == 'claims_reconciliation':
        claims_reconciliation = _safe_claims_reconciliation_view()
    elif workspace == 'claims_mapping':
        mh = get_mapping_health_from_postgres()
        claims_mapping_unmatched_rows = get_claims_mapping_rows_from_postgres()
    elif workspace == 'ai_assistant':
        ai_answer = _ai_answer(request.args.get('ask', ''))

    claims_import_summary_effective = dict(LAST_CLAIMS_IMPORT_SUMMARY or {})
    if workspace == 'claims_mapping' and mh:
        claims_import_summary_effective.setdefault('total_claims_imported', mh.get('imported_claims', 0))
        claims_import_summary_effective['matched_claims'] = mh.get('mapped_claims', claims_import_summary_effective.get('matched_claims', 0))
        claims_import_summary_effective['written_claims'] = mh.get('mapped_claims', claims_import_summary_effective.get('written_claims', 0))
        claims_import_summary_effective['unmatched_claims'] = mh.get('unmapped_claims', claims_import_summary_effective.get('unmatched_claims', 0))

    return dict(
        workspace=workspace, workspace_title=title, workspace_subtitle=subtitle,
        current_timestamp=datetime.now().strftime('%d %B %Y %H:%M'),
        portfolio=configured_portfolio, rates=LAST_RESULT.get('rates', DEFAULT_RATES.copy()), book_rates=LAST_RESULT.get('book_rates', DEFAULT_BOOK_VALUE.copy()),
        franchises=franchises, selected=selected, period_view=period_view, traffic_filter=traffic_filter,
        commission_focus=commission_focus, scenario_filter=scenario_filter,
        quick_view=base, money_view=money_view, commission_view=commission_view, book_view=book_view,
        compare_full=compare_full, compare_totals=compare_totals, compare_args=compare_args, compare_summary=compare_summary,
        advanced_view=advanced_view, ai_question=request.args.get('ask', ''), ai_answer=ai_answer,
        franchise_config=config, original_franchises=original_franchises,
        commission_business_view=commission_business_view, underwriter_view=underwriter_view,
        claims_import_summary=claims_import_summary_effective, policy_import_summary=LAST_POLICY_IMPORT_SUMMARY,
        claims_reconciliation=claims_reconciliation, mh=mh, claims_mapping_unmatched_rows=claims_mapping_unmatched_rows,
        claims_mapping_options=[{'name': f, 'key': apply_franchise_alias(franchise_match_key(f))} for f in original_franchises]
    )


@app.route('/workspace/<workspace>', methods=['GET','POST'])
def workspace_page(workspace):
    allowed={'claim_ratio','move_recommendations','franchise_money','commission_analysis','commission_business','underwriter_allocation','book_value','scenario_comparison','advanced_analytics','ai_assistant','franchise_configuration','claims_mapping','claims_reconciliation'}
    if workspace not in allowed:
        return redirect(url_for('dashboard'))
    if request.method == 'POST' and workspace == 'claims_mapping':
        global LAST_RESULT, LAST_CLAIMS_DF, LAST_POLICY_IMPORT_SUMMARY
        config = load_franchise_config()
        aliases = dict(config.get('claims_aliases', {}))
        claim_keys = request.form.getlist('claim_key')
        for claim_key in claim_keys:
            target_key = request.form.get(f'target_{claim_key}', '').strip()
            if target_key:
                aliases[claim_key] = target_key
        # allow removal of a saved mapping
        for remove_key in request.form.getlist('remove_alias'):
            aliases.pop(remove_key, None)
        config['claims_aliases'] = aliases
        save_franchise_config(config)
        save_claim_aliases_to_postgres(aliases)
        recalculated_pg = recalculate_claims_summary_from_postgres()
        # Re-apply latest claims file immediately when available. Start from clean
        # PolicyData rows only so old claims-only rows cannot duplicate totals.
        if LAST_CLAIMS_DF is not None and not LAST_CLAIMS_DF.empty and LAST_RESULT.get('raw') is not None and not LAST_RESULT['raw'].empty:
            clean_policy_raw = policy_rows_only(LAST_RESULT['raw'])
            raw = merge_claims_into_raw(clean_policy_raw, LAST_CLAIMS_DF)
            monthly, periods, portfolio = analyse(raw, LAST_RESULT['rates'], LAST_RESULT['book_rates'])
            LAST_RESULT = {'raw': raw, 'monthly': monthly, 'periods': periods, 'portfolio': portfolio, 'rates': LAST_RESULT['rates'], 'book_rates': LAST_RESULT['book_rates']}
            flash('Claims mappings saved, PostgreSQL summary rebuilt, and claims were re-applied.', 'success')
        elif recalculated_pg:
            reload_dashboard_from_postgres()
            flash('Claims mappings saved and PostgreSQL summary rebuilt from existing imports.', 'success')
        else:
            flash('Claims mappings saved. Re-import the claims file to apply them.', 'success')
        return redirect(url_for('workspace_page', workspace='claims_mapping'))
    if request.method == 'POST' and workspace == 'franchise_configuration':
        franchises = request.form.getlist('franchise_name')
        excluded = []
        groups = {}
        for idx, franchise in enumerate(franchises):
            if request.form.get(f'include_{idx}') != 'on':
                excluded.append(franchise)
            group_name = request.form.get(f'group_{idx}', '').strip()
            if group_name:
                groups[franchise] = group_name
        save_franchise_config({'excluded': excluded, 'groups': groups, 'use_groups': request.form.get('use_groups') == 'on', 'claims_aliases': load_franchise_config().get('claims_aliases', {})})
        flash('Franchise configuration saved.', 'success')
        return redirect(url_for('workspace_page', workspace='franchise_configuration'))
    if request.method == 'POST':
        return dashboard()
    return render_template('dashboard.html', **_workspace_context(workspace))


@app.route('/franchise/<path:franchise>')
def franchise_page(franchise):
    # Opens claim ratio workspace filtered to one franchise; users can switch workspace with Back to Dashboard.
    return redirect(url_for('workspace_page', workspace='claim_ratio', franchise=franchise, period_view='month'))




def _friendly_export_df(df):
    """Prepare tables for export/print with current naming."""
    if df is None or df.empty:
        return pd.DataFrame()
    out = df.copy()
    rename = {
        'Weighted Claim Ratio': 'Average Claim Ratio',
        'BrightRock Month Total': 'Franchise Money',
        'BrightRock Running Balance': 'Franchise Money Running Balance',
    }
    out = out.rename(columns={k: v for k, v in rename.items() if k in out.columns})
    return out


def _format_excel_sheet(writer, sheet_name, df):
    """Autofit, freeze, filter and format exported sheets."""
    workbook = writer.book
    ws = writer.sheets[sheet_name]
    header_fmt = workbook.add_format({'bold': True, 'bg_color': '#6B4AA0', 'font_color': '#FFFFFF', 'border': 1})
    money_fmt = workbook.add_format({'num_format': 'R#,##0.00'})
    pct_fmt = workbook.add_format({'num_format': '0.00%'})
    int_fmt = workbook.add_format({'num_format': '#,##0'})
    text_fmt = workbook.add_format({})
    for col_num, col_name in enumerate(df.columns):
        ws.write(0, col_num, col_name, header_fmt)
        # Use positional selection so duplicate column names cannot return a DataFrame.
        series = df.iloc[:, col_num] if len(df.columns) else pd.Series(dtype=str)
        values = series.astype(str).head(500).tolist() if len(df) else []
        max_len = max([len(str(col_name))] + [len(v) for v in values])
        width = max(10, min(38, max_len + 2))
        col_lower = str(col_name).lower()
        fmt = text_fmt
        if any(key in col_lower for key in ['premium', 'claim', 'commission', 'fee', 'value', 'money', 'amount', 'payover', 'risk', 'retail', 'book']):
            if 'ratio' not in col_lower and 'count' not in col_lower and 'qty' not in col_lower and '%' not in col_lower:
                fmt = money_fmt
        if 'ratio' in col_lower or '%' in col_lower or 'contribution' in col_lower:
            fmt = pct_fmt
        if any(key in col_lower for key in ['qty', 'count', 'policies', 'franchises']):
            fmt = int_fmt
        ws.set_column(col_num, col_num, width, fmt)
    rows = max(len(df), 1)
    cols = max(len(df.columns), 1)
    ws.freeze_panes(1, 1)
    ws.autofilter(0, 0, rows, cols - 1)



PDF_HEADER_ABBREVIATIONS = {
    'Average Claim Ratio': 'Avg Claim %',
    'Weighted Claim Ratio': 'Avg Claim %',
    'Claim Ratio Label': 'Claim Label',
    'Recommendation': 'Reco',
    'Retail Premium': 'Retail',
    'Risk Premium': 'Risk',
    'Original Risk Premium': 'Orig Risk',
    'Total Commission': 'Tot Comm',
    'Total Paid Commissions': 'Paid Comm',
    'Franchise Money': 'Franchise $',
    'Franchise Money Running Balance': 'Run Balance',
    'BrightRock Running Balance': 'BR Total',
    'R1 Policy Fee': 'R1 Fee',
    'Policy Qty': 'Policies',
    'Underwriter 2.1% Fee': '2.1% Fee',
    'BrightRock Amount': 'BrightRock',
    'Inkulu Amount': 'Mkhulu',
    'MFF Amount': 'MFF',
    'MFF Book Value 2.5%': 'MFF Book',
    'Franchise Book Value 2.5%': 'Franchise Book',
    'Total Book Value': 'Book Value',
}


def _short_pdf_header(value):
    return PDF_HEADER_ABBREVIATIONS.get(str(value), str(value))


def _short_pdf_cell(value, col_name=''):
    text = str(value)
    if str(col_name).lower() == 'franchise' and len(text) > 28:
        return text[:25] + '...'
    if len(text) > 34:
        return text[:31] + '...'
    return text


def _pdf_table(data, page_width, font_size=5.4, first_col_weight=1.35):
    """Create a fitted reportlab table for wide data."""
    if not data:
        data = [['No data']]
    col_count = max(len(data[0]), 1)
    available = page_width - 1.35 * cm
    weights = [1.0] * col_count
    if col_count:
        weights[0] = first_col_weight
    total_weight = sum(weights)
    col_widths = [available * w / total_weight for w in weights]
    if col_count > 18:
        font_size = min(font_size, 4.2)
    elif col_count > 14:
        font_size = min(font_size, 4.6)
    elif col_count > 10:
        font_size = min(font_size, 5.0)
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#6B4AA0')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('GRID', (0,0), (-1,-1), 0.2, colors.lightgrey),
        ('FONTSIZE', (0,0), (-1,-1), font_size),
        ('PADDING', (0,0), (-1,-1), 1.1),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
    ]))
    return t


def _rows_for_pdf(df, cols, limit=70):
    """Convert dataframe to PDF display rows with currency/percent formatting."""
    rows = [[_short_pdf_header(c) for c in cols]]
    if df is None or df.empty:
        return rows + [['No data'] + [''] * (len(cols) - 1)]
    for _, r in df.head(limit).iterrows():
        row = []
        for c in cols:
            v = r.get(c, '')
            cl = str(c).lower()
            try:
                if any(k in cl for k in ['premium','claims','commission','fee','value','money','amount','payover','risk','retail','book']):
                    if 'ratio' not in cl and '%' not in cl and 'count' not in cl and 'qty' not in cl:
                        v = money(v)
                elif 'ratio' in cl or '%' in cl or 'contribution' in cl:
                    v = pct(v)
                elif 'qty' in cl or 'count' in cl or 'policies' in cl:
                    v = int(float(v))
            except Exception:
                pass
            row.append(_short_pdf_cell(v, c))
        rows.append(row)
    return rows


def _page_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont('Helvetica', 7)
    canvas.drawString(doc.leftMargin, 0.5 * cm, "Martin's Funerals South Africa | Confidential | " + datetime.now().strftime('Generated %d %b %Y %H:%M'))
    canvas.drawRightString(doc.pagesize[0] - doc.rightMargin, 0.5 * cm, f'Page {doc.page}')
    canvas.restoreState()



def _report_subtitle(name='Board Report'):
    return name or 'Report'


def _add_pdf_cover(story, styles, subtitle, period_text=''):
    story.append(Spacer(1, 1.2*cm))
    if os.path.exists(REPORT_LOGO):
        story.append(Image(REPORT_LOGO, width=12.5*cm, height=5.7*cm, kind='proportional'))
        story.append(Spacer(1, 0.6*cm))
    story.append(Paragraph("Martin's Funerals South Africa", styles['Title']))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(_report_subtitle(subtitle), styles['Heading1']))
    story.append(Spacer(1, 0.4*cm))
    generated = datetime.now().strftime('%d %B %Y %H:%M')
    info = f'Generated: {generated}' + (f'<br/>Period: {period_text}' if period_text else '')
    story.append(Paragraph(info, styles['Normal']))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph('Confidential', styles['Normal']))
    story.append(PageBreak())


def _add_excel_cover(writer, report_name, period_text=''):
    wb = writer.book
    cover = wb.add_worksheet('Cover')
    title_fmt = wb.add_format({'bold': True, 'font_size': 22, 'font_color': '#18202a'})
    sub_fmt = wb.add_format({'bold': True, 'font_size': 16, 'font_color': '#374151'})
    text_fmt = wb.add_format({'font_size': 11})
    cover.set_column('A:A', 24)
    cover.set_column('B:F', 18)
    try:
        if os.path.exists(REPORT_LOGO):
            cover.insert_image('B2', REPORT_LOGO, {'x_scale': 0.35, 'y_scale': 0.35})
    except Exception:
        pass
    cover.write('B12', "Martin's Funerals South Africa", title_fmt)
    cover.write('B14', report_name, sub_fmt)
    cover.write('B16', 'Generated', text_fmt)
    cover.write('C16', datetime.now().strftime('%d %B %Y %H:%M'), text_fmt)
    cover.write('B17', 'Period', text_fmt)
    cover.write('C17', period_text or 'All data', text_fmt)
    cover.write('B19', 'Confidential', text_fmt)

@app.route('/export_payover')
def export_payover():
    """Export monthly/period underwriter payover sheet for insurers.

    Source: PolicyData transaction import, MEM rows only.
    Payover Premium = Original Risk Premium - R1 Policy Fee - ADV Fee 2.1%.
    Retail Less Payover is included for reconciliation against all retail premiums paid.
    """
    if LAST_RESULT['monthly'].empty:
        flash('No data to export.', 'warning')
        return redirect(url_for('dashboard'))
    insurer = request.args.get('insurer', 'All') or 'All'
    config = load_franchise_config()
    monthly, periods, _ = apply_franchise_config(apply_user_franchise_scope(LAST_RESULT['monthly']), config)
    period_view = request.args.get('period_view', 'six_months')
    if period_view == 'month':
        base = monthly.copy()
        base['Period View'] = base['Month'].dt.strftime('%b %Y') if 'Month' in base.columns else ''
    elif period_view == 'year':
        agg = {'Retail Premium':'sum','Original Risk Premium':'sum','Risk Premium':'sum','Policy Qty':'sum','R1 Policy Fee':'sum','Underwriter 2.1% Fee':'sum','Claims':'sum'}
        base = monthly.copy()
        if 'Year' not in base.columns and 'Month' in base.columns:
            base['Year'] = base['Month'].dt.year.astype(str)
        base = base.groupby(['Franchise','Year'], as_index=False).agg({k:v for k,v in agg.items() if k in base.columns})
        base['Period View'] = base['Year']
    else:
        base = periods.copy()
        base['Period View'] = base.get('Period', '')
    selected = request.args.get('franchise', 'All') or 'All'
    if selected != 'All' and 'Franchise' in base.columns:
        base = base[base['Franchise'] == selected]
    view = _underwriter_allocation_view(base)
    if view.empty:
        flash('No payover data available for export.', 'warning')
        return redirect(url_for('workspace_page', workspace='underwriter_allocation'))
    export_df = view.rename(columns={
        'Original Risk Premium': 'Original Risk',
        'Policy Qty': 'Policies Paid (MPIA)',
        'ADV Fee 2.1%': 'ADV Fee 2.1%',
        'Underwriter Premium Payable': 'Payover Premium',
    }).copy()
    export_df.insert(0, 'Insurer', insurer)
    cols = [c for c in ['Insurer','Franchise','Period View','Retail Premium','Original Risk','Policies Paid (MPIA)','R1 Policy Fee','Risk After R1','ADV Fee 2.1%','Payover Premium','Retail Less Payover','Claims','Claim Ratio'] if c in export_df.columns]
    export_df = export_df[cols]
    totals = {}
    for c in export_df.columns:
        if pd.api.types.is_numeric_dtype(export_df[c]):
            totals[c] = export_df[c].sum()
        elif c == 'Franchise':
            totals[c] = 'TOTAL'
        elif c == 'Insurer':
            totals[c] = insurer
        else:
            totals[c] = ''
    export_df = pd.concat([export_df, pd.DataFrame([totals])], ignore_index=True)
    safe_insurer = re.sub(r'[^A-Za-z0-9]+', '_', insurer).strip('_') or 'All'
    path = os.path.join(EXPORT_DIR, f'{safe_insurer}_payover_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx')
    with pd.ExcelWriter(path, engine='xlsxwriter') as writer:
        export_df.to_excel(writer, index=False, sheet_name='Payover Export')
        workbook = writer.book
        ws = writer.sheets['Payover Export']
        money_fmt = workbook.add_format({'num_format': 'R#,##0.00'})
        pct_fmt = workbook.add_format({'num_format': '0.00%'})
        for idx, col in enumerate(export_df.columns):
            width = max(14, min(32, len(str(col)) + 4))
            ws.set_column(idx, idx, width)
            if col in ['Retail Premium','Original Risk','R1 Policy Fee','Risk After R1','ADV Fee 2.1%','Payover Premium','Retail Less Payover','Claims']:
                ws.set_column(idx, idx, width, money_fmt)
            if col == 'Claim Ratio':
                ws.set_column(idx, idx, width, pct_fmt)
    return send_file(path, as_attachment=True)

@app.route('/export')
def export():
    if LAST_RESULT['monthly'].empty:
        flash('No data to export.', 'warning')
        return redirect(url_for('dashboard'))
    path = os.path.join(EXPORT_DIR, f'claims_analytics_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx')
    monthly_scoped = apply_user_franchise_scope(LAST_RESULT['monthly'])
    periods_scoped = apply_user_franchise_scope(LAST_RESULT['periods'])
    monthly_export = _friendly_export_df(monthly_scoped)
    periods_export = _friendly_export_df(periods_scoped)
    portfolio_export = pd.DataFrame([build_portfolio_for_period(monthly_scoped, periods_scoped, 'six_months')]).rename(columns={'overall_claim_ratio': 'average_claim_ratio'})
    with pd.ExcelWriter(path, engine='xlsxwriter') as writer:
        _add_excel_cover(writer, 'Franchise Claims Analytics Export', 'All imported data')
        monthly_export.to_excel(writer, index=False, sheet_name='Monthly Detail')
        periods_export.to_excel(writer, index=False, sheet_name='Six Month Periods')
        portfolio_export.to_excel(writer, index=False, sheet_name='Portfolio Summary')
        _format_excel_sheet(writer, 'Monthly Detail', monthly_export)
        _format_excel_sheet(writer, 'Six Month Periods', periods_export)
        _format_excel_sheet(writer, 'Portfolio Summary', portfolio_export)
    return send_file(path, as_attachment=True)


@app.route('/board_report')
def board_report():
    if LAST_RESULT['periods'].empty:
        flash('No data to report.', 'warning')
        return redirect(url_for('dashboard'))
    period_choice = (request.args.get('dashboard_period_view') or request.args.get('period_view') or 'six_months')
    if period_choice not in {'month', 'six_months', 'year'}:
        period_choice = 'six_months'
    period_text_map = {'month': 'Monthly', 'six_months': 'Every 6 Months', 'year': '12 Months / Yearly'}
    config = load_franchise_config()
    configured_monthly, configured_periods, _ = apply_franchise_config(apply_user_franchise_scope(LAST_RESULT['monthly']), config)
    portfolio = build_portfolio_for_period(configured_monthly, configured_periods, period_choice)
    path = os.path.join(EXPORT_DIR, f'board_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf')
    page_size = landscape(A4)
    doc = SimpleDocTemplate(path, pagesize=page_size, rightMargin=0.45*cm, leftMargin=0.45*cm, topMargin=0.55*cm, bottomMargin=0.55*cm)
    styles = getSampleStyleSheet()
    story = []
    _add_pdf_cover(story, styles, 'Executive Board Report', period_text_map.get(period_choice, 'Every 6 Months'))
    summary_data = [
        ['Metric', 'Value'],
        ['Total Franchises', portfolio.get('total_franchises', 0)],
        ['Total Retail Premium', money(portfolio.get('total_retail', 0))],
        ['Total Claims', money(portfolio.get('total_claims', 0))],
        ['Average Claim Ratio', pct(portfolio.get('average_claim_ratio', portfolio.get('overall_claim_ratio', 0)))],
        ['BrightRock Commission', money(portfolio.get('total_brightrock_commission', 0))],
        ['Mkhulu Commission', money(portfolio.get('total_mkhulu_commission', portfolio.get('total_inkulu_commission', 0)))],
        ['MFF Commission', money(portfolio.get('total_mff_commission', 0))],
        ['R1 Policy Fee', money(portfolio.get('total_r1_policy_fee', 0))],
        ['2.1% Underwriter Fee', money(portfolio.get('total_underwriter_2_1_fee', 0))],
        ['Total Commission', money(portfolio.get('total_commission', 0))],
        ['Policies Sold', int(portfolio.get('total_policy_qty', 0))],
        ['Franchise Money Accumulated', money(portfolio.get('total_brightrock_month_total', portfolio.get('total_brightrock_accumulated', 0)))],
        ['Total Book Value', money(portfolio.get('total_book_value', 0))],
        ['MFF Book Value', money(portfolio.get('total_mff_book_value', 0))],
        ['Franchise Book Value', money(portfolio.get('total_franchise_book_value', 0))],
        ['Move to BrightRock', portfolio.get('move_count', 0)],
        ['Can Move to BrightRock', portfolio.get('can_move_count', 0)],
        ['High Risk', portfolio.get('high_risk_count', 0)],
    ]
    story.append(_pdf_table(summary_data, page_size[0], font_size=5.4, first_col_weight=1.25))
    story.append(PageBreak())

    detail_source = configured_periods if period_choice == 'six_months' else (configured_monthly.copy() if period_choice == 'month' else yearly_view_from_monthly(configured_monthly))
    periods = _friendly_export_df(detail_source)
    story.append(Paragraph(f'{period_text_map.get(period_choice, "Every 6 Months")} Scenario Recommendations', styles['Heading2']))
    cols = [c for c in ['Franchise','Period','Retail Premium','Risk Premium','Claims','Average Claim Ratio','Claim Ratio Label','Recommendation','Policy Qty','R1 Policy Fee','Total Commission','Franchise Money','Franchise Money Running Balance','Total Book Value'] if c in periods.columns]
    story.append(_pdf_table(_rows_for_pdf(periods, cols, limit=80), page_size[0], font_size=4.6, first_col_weight=1.45))
    story.append(PageBreak())

    # Commission summary
    story.append(Paragraph('Commission Summary', styles['Heading2']))
    commission_cols = [c for c in ['Franchise','Period','Retail Premium','BrightRock Amount','Inkulu Amount','MFF Amount','R1 Policy Fee','Underwriter 2.1% Fee','Total Commission','Total Paid Commissions'] if c in periods.columns]
    story.append(_pdf_table(_rows_for_pdf(periods, commission_cols, limit=80), page_size[0], font_size=4.6, first_col_weight=1.45))
    story.append(PageBreak())

    story.append(Paragraph('Book Value Summary', styles['Heading2']))
    book_cols = [c for c in ['Franchise','Period','Retail Premium','MFF Book Value 2.5%','Franchise Book Value 2.5%','Total Book Value'] if c in periods.columns]
    story.append(_pdf_table(_rows_for_pdf(periods, book_cols, limit=100), page_size[0], font_size=4.8, first_col_weight=1.45))
    doc.build(story, onFirstPage=_page_footer, onLaterPages=_page_footer)
    return send_file(path, as_attachment=True)



@app.route('/service-worker.js')
def service_worker():
    resp = make_response('', 204)
    resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    return resp

# -----------------------------------------------------------------------------
# Claims workflow engine
# -----------------------------------------------------------------------------
CLAIM_STATUSES = ['New', 'In Review', 'Pending Documents', 'Approved', 'Rejected', 'Closed']
CLAIM_PRIORITIES = ['Low', 'Normal', 'High', 'Urgent']


CLAIM_EMAIL_ENABLED = os.getenv('CLAIM_EMAIL_ENABLED', '1').strip() == '1'
CLAIM_OVERDUE_DAYS = int(os.getenv('CLAIM_OVERDUE_DAYS', '7'))
CLAIM_ATTACHMENT_DIR = os.path.join(UPLOAD_DIR, 'claim_attachments')
os.makedirs(CLAIM_ATTACHMENT_DIR, exist_ok=True)

DEFAULT_CLAIM_DOCUMENT_RULES = [
    {
        'type': 'ID Document',
        'claim_type': 'Funeral Claim',
        'aliases': ['id document', 'identity document', 'id copy', 'identification', 'smart id', 'sa id'],
        'required': True,
        'rules': {
            'detect_id_number': True,
            'certified_copy': True,
            'document_wording': True,
            'commissioner_stamp': False,
            'certification_date': False,
            'certification_not_older_than_3_months': False,
            'death_certificate_wording': False,
            'deceased_name': False,
            'date_of_death': False,
            'bank_details': False,
        },
    },
    {
        'type': 'Death Certificate',
        'claim_type': 'Funeral Claim',
        'aliases': ['death certificate', 'dha-1663', 'death notice', 'notice of death'],
        'required': True,
        'rules': {
            'detect_id_number': False,
            'certified_copy': False,
            'document_wording': False,
            'commissioner_stamp': False,
            'certification_date': False,
            'certification_not_older_than_3_months': False,
            'death_certificate_wording': True,
            'deceased_name': False,
            'date_of_death': False,
            'bank_details': False,
        },
    },
    {
        'type': 'Proof of Banking',
        'claim_type': 'Funeral Claim',
        'aliases': ['bank statement', 'proof of bank', 'banking details', 'account number', 'bank account'],
        'required': False,
        'rules': {
            'detect_id_number': False,
            'certified_copy': False,
            'document_wording': False,
            'commissioner_stamp': False,
            'certification_date': False,
            'certification_not_older_than_3_months': False,
            'death_certificate_wording': False,
            'deceased_name': False,
            'date_of_death': False,
            'bank_details': True,
        },
    },
]

CLAIM_RULE_LABELS = {
    'detect_id_number': 'Detect South African ID number',
    'certified_copy': 'Certified copy wording/stamp required',
    'document_wording': 'Document wording must match selected type',
    'commissioner_stamp': 'Commissioner/SAPS stamp required',
    'certification_date': 'Certification date required',
    'certification_not_older_than_3_months': 'Certification must not be older than 3 months',
    'death_certificate_wording': 'Death certificate / DHA-1663 wording required',
    'deceased_name': 'Deceased name required',
    'date_of_death': 'Date of death required',
    'bank_details': 'Bank details/account number required',
}


def _default_claim_document_rules():
    return [dict(x) for x in DEFAULT_CLAIM_DOCUMENT_RULES]


def _seed_claim_document_rule_defaults(conn=None, force=False):
    """Create default admin-editable claim document rules if none exist."""
    engine = None if conn is not None else get_db_engine()
    if conn is None and engine is None:
        return False
    def work(active_conn):
        existing = int(active_conn.execute(text('SELECT COUNT(*) FROM app_claim_document_types')).scalar() or 0)
        if existing and not force:
            return False
        if force:
            active_conn.execute(text('DELETE FROM app_claim_document_rules'))
            active_conn.execute(text('DELETE FROM app_claim_document_types'))
        for idx, cfg in enumerate(DEFAULT_CLAIM_DOCUMENT_RULES, start=1):
            row = active_conn.execute(text("""
                INSERT INTO app_claim_document_types (document_type, claim_type, aliases, required, is_active, sort_order, updated_at)
                VALUES (:document_type, :claim_type, :aliases, :required, true, :sort_order, NOW())
                RETURNING id
            """), {
                'document_type': cfg['type'],
                'claim_type': cfg.get('claim_type') or 'Funeral Claim',
                'aliases': ', '.join(cfg.get('aliases') or []),
                'required': bool(cfg.get('required', True)),
                'sort_order': idx * 10,
            }).mappings().first()
            type_id = int(row['id'])
            for key, label in CLAIM_RULE_LABELS.items():
                active_conn.execute(text("""
                    INSERT INTO app_claim_document_rules (document_type_id, rule_key, rule_label, is_enabled)
                    VALUES (:document_type_id, :rule_key, :rule_label, :is_enabled)
                """), {
                    'document_type_id': type_id,
                    'rule_key': key,
                    'rule_label': label,
                    'is_enabled': bool((cfg.get('rules') or {}).get(key, False)),
                })
        return True
    if conn is not None:
        return work(conn)
    with engine.begin() as active_conn:
        return work(active_conn)


def _claim_required_documents(claim_type='Funeral Claim'):
    """Load admin-configured claim document requirements; fall back to safe defaults."""
    engine = get_db_engine()
    if engine is None:
        return _default_claim_document_rules()
    try:
        with engine.begin() as conn:
            _seed_claim_document_rule_defaults(conn)
            types = [dict(r) for r in conn.execute(text("""
                SELECT id, document_type, claim_type, aliases, required, is_active, sort_order
                FROM app_claim_document_types
                WHERE COALESCE(is_active, true) = true
                  AND COALESCE(claim_type, 'Funeral Claim') = :claim_type
                ORDER BY sort_order, document_type
            """), {'claim_type': claim_type or 'Funeral Claim'}).mappings().all()]
            if not types:
                types = [dict(r) for r in conn.execute(text("""
                    SELECT id, document_type, claim_type, aliases, required, is_active, sort_order
                    FROM app_claim_document_types
                    WHERE COALESCE(is_active, true) = true
                    ORDER BY sort_order, document_type
                """)).mappings().all()]
            out = []
            for t in types:
                rules = [dict(r) for r in conn.execute(text("""
                    SELECT rule_key, rule_label, is_enabled
                    FROM app_claim_document_rules
                    WHERE document_type_id = :document_type_id
                    ORDER BY id
                """), {'document_type_id': int(t['id'])}).mappings().all()]
                rule_map = {r['rule_key']: bool(r.get('is_enabled')) for r in rules}
                aliases = [x.strip().lower() for x in str(t.get('aliases') or '').split(',') if x.strip()]
                if str(t.get('document_type') or '').strip().lower() not in aliases:
                    aliases.append(str(t.get('document_type') or '').strip().lower())
                out.append({
                    'id': int(t['id']),
                    'type': str(t.get('document_type') or '').strip(),
                    'claim_type': str(t.get('claim_type') or 'Funeral Claim'),
                    'aliases': aliases,
                    'required': bool(t.get('required')),
                    'rules': rule_map,
                    'needs_id_number': bool(rule_map.get('detect_id_number')),
                    'needs_certified': bool(rule_map.get('certified_copy')),
                })
            return out or _default_claim_document_rules()
    except Exception as exc:
        print('Could not load claim document rules:', exc)
        return _default_claim_document_rules()


def _normalise_doc_text(value):
    return re.sub(r'\s+', ' ', str(value or '').replace('\x00', ' ')).strip()


def _extract_text_from_attachment(path, filename=''):
    text_bits = [str(filename or '')]
    ext = os.path.splitext(str(filename or path or '').lower())[1]
    try:
        if ext == '.pdf':
            try:
                import PyPDF2
                with open(path, 'rb') as fh:
                    reader = PyPDF2.PdfReader(fh)
                    for page in reader.pages[:8]:
                        try:
                            text_bits.append(page.extract_text() or '')
                        except Exception:
                            pass
            except Exception:
                pass
        elif ext in {'.txt', '.csv'}:
            with open(path, 'r', encoding='utf-8', errors='ignore') as fh:
                text_bits.append(fh.read(50000))
        elif ext in {'.docx'}:
            try:
                from docx import Document
                doc = Document(path)
                text_bits.extend([p.text for p in doc.paragraphs[:300]])
            except Exception:
                pass
    except Exception:
        pass
    return _normalise_doc_text(' '.join([x for x in text_bits if x]))


def _detect_claim_document_type(filename, extracted_text):
    hay = _normalise_doc_text(f'{filename} {extracted_text}').lower()
    for rule in _claim_required_documents():
        if any(alias in hay for alias in rule['aliases']):
            return rule['type']
    if re.search(r'\b\d{13}\b', hay) and ('id' in hay or 'identity' in hay):
        return 'ID Document'
    return 'Other Document'


def _verify_claim_attachment(filename, path):
    extracted = _extract_text_from_attachment(path, filename)
    hay = extracted.lower()
    document_type = _detect_claim_document_type(filename, extracted)
    has_id_number = bool(re.findall(r'(?<!\d)(\d{13})(?!\d)', extracted))
    certified_terms = ['certified true copy', 'certified copy', 'certified', 'commissioner of oaths', 'commissioner of oath', 'saps', 'police station', 'stamp', 'date stamp', 'certify']
    is_certified = any(term in hay for term in certified_terms)
    rules_by_type = {r['type']: r for r in _claim_required_documents()}
    rule_cfg = rules_by_type.get(document_type, {})
    enabled = rule_cfg.get('rules') or {}
    issues = []
    if not rule_cfg and document_type == 'Other Document':
        issues.append('Document type could not be matched to a configured claim document rule.')
    if enabled.get('detect_id_number') and not has_id_number:
        issues.append('South African ID number not detected.')
    if enabled.get('certified_copy') and not is_certified:
        issues.append('Certified copy wording/stamp not detected.')
    if enabled.get('document_wording'):
        aliases = [str(x).lower() for x in rule_cfg.get('aliases') or []]
        if aliases and not any(a and a in hay for a in aliases):
            issues.append(f'{document_type} wording not detected.')
    if enabled.get('commissioner_stamp') and not any(term in hay for term in ['commissioner', 'saps', 'police station', 'oaths', 'stamp']):
        issues.append('Commissioner/SAPS stamp not detected.')
    if enabled.get('certification_date') and not re.search(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b20\d{2}[/-]\d{1,2}[/-]\d{1,2}\b', hay):
        issues.append('Certification/date stamp date not detected.')
    if enabled.get('death_certificate_wording') and not any(term in hay for term in ['death certificate', 'death notice', 'dha-1663', 'notice of death']):
        issues.append('Death certificate / DHA-1663 wording not detected.')
    if enabled.get('deceased_name') and not any(term in hay for term in ['deceased', 'surname', 'forenames', 'name of deceased']):
        issues.append('Deceased name wording not detected.')
    if enabled.get('date_of_death') and not any(term in hay for term in ['date of death', 'died on', 'death date']):
        issues.append('Date of death wording not detected.')
    if enabled.get('bank_details') and not (re.search(r'\b\d{6,16}\b', hay) and any(term in hay for term in ['bank', 'account', 'branch code', 'branch'])):
        issues.append('Bank account details not detected.')
    status = 'Verified' if not issues else 'Pending'
    details = '; '.join(issues) if issues else f'{document_type} verified against admin checklist.'
    return {'document_type': document_type, 'verification_status': status, 'verification_details': details, 'extracted_text': extracted[:50000], 'has_id_number': has_id_number, 'is_certified': is_certified}


def _claim_document_status(claim_id):
    engine = get_db_engine()
    result = {'overall_status': 'Pending', 'approved': False, 'items': [], 'missing': [], 'summary': 'No documents uploaded yet.'}
    if engine is None:
        return result
    with engine.begin() as conn:
        docs = [dict(r) for r in conn.execute(text("""
            SELECT id, filename, document_type, verification_status, verification_details,
                   has_id_number, is_certified, created_at
            FROM app_claim_attachments
            WHERE claim_id = :claim_id
            ORDER BY created_at DESC, id DESC
        """), {'claim_id': int(claim_id)}).mappings().all()]
    by_type = {}
    for d in docs:
        by_type.setdefault(d.get('document_type') or 'Other Document', []).append(d)
    items, missing = [], []
    for rule in _claim_required_documents():
        docs_for_type = by_type.get(rule['type'], [])
        if not docs_for_type:
            missing.append(f"{rule['type']} required.")
            items.append({'type': rule['type'], 'status': 'Missing', 'details': f"{rule['type']} required."})
            continue
        passed = False
        details = []
        for d in docs_for_type:
            ok = str(d.get('verification_status') or '').lower() == 'verified'
            if rule.get('needs_id_number') and not d.get('has_id_number'):
                ok = False
            if rule.get('needs_certified') and not d.get('is_certified'):
                ok = False
            details.append(f"{d.get('filename')}: {d.get('verification_details') or d.get('verification_status')}")
            if ok:
                passed = True
        if passed:
            items.append({'type': rule['type'], 'status': 'Verified', 'details': '; '.join(details)})
        else:
            missing.append(f"{rule['type']} uploaded but still requires review/fix.")
            items.append({'type': rule['type'], 'status': 'Pending', 'details': '; '.join(details)})
    known = {r['type'] for r in _claim_required_documents()}
    for d in docs:
        if (d.get('document_type') or 'Other Document') not in known:
            items.append({'type': d.get('document_type') or 'Other Document', 'status': d.get('verification_status') or 'Pending', 'details': f"{d.get('filename')}: {d.get('verification_details') or ''}"})
    result['items'] = items
    result['missing'] = missing
    result['approved'] = not missing and bool(docs)
    result['overall_status'] = 'Approved' if result['approved'] else 'Pending Documents'
    if result['approved']:
        result['summary'] = 'All required claim documents are verified.'
    elif docs:
        result['summary'] = 'Pending: ' + '; '.join(missing)
    return result


def _update_claim_document_status(claim_id, conn=None):
    status = _claim_document_status(claim_id)
    note_text = f"Document checklist status: {status['overall_status']}. {status['summary']}"
    engine = get_db_engine()
    if engine is None:
        return status
    def work(active_conn):
        active_conn.execute(text("""
            UPDATE app_claim_cases
            SET status = CASE
                    WHEN :approved THEN 'Approved'
                    WHEN status IN ('New', 'In Review', 'Approved') THEN 'Pending Documents'
                    ELSE status
                END,
                updated_at = NOW()
            WHERE id = :claim_id
        """), {'claim_id': int(claim_id), 'approved': bool(status['approved'])})
        active_conn.execute(text("""
            INSERT INTO app_claim_notes (claim_id, user_id, user_email, note_text)
            VALUES (:claim_id, :user_id, :user_email, :note_text)
        """), {'claim_id': int(claim_id), 'user_id': (getattr(g, 'user', None) or {}).get('id'), 'user_email': (getattr(g, 'user', None) or {}).get('email'), 'note_text': note_text})
    if conn is not None:
        work(conn)
    else:
        with engine.begin() as active_conn:
            work(active_conn)
    return status


def _split_emails(value):
    emails = []
    for part in re.split(r'[;,]', str(value or '')):
        email = part.strip()
        if email and '@' in email and email.lower() not in [e.lower() for e in emails]:
            emails.append(email)
    return emails


def _claim_link(claim_id):
    path = f'/claims/{int(claim_id)}'
    return (APP_BASE_URL + path) if APP_BASE_URL else path


def _claim_email_recipients(claim=None, extra_recipients=None):
    recipients = []
    recipients += _split_emails(os.getenv('CLAIMS_ALERT_EMAIL', ALERT_EMAIL))
    if claim:
        recipients += _split_emails(claim.get('assigned_to_email'))
        recipients += _split_emails(claim.get('created_by_email'))
    recipients += _split_emails(extra_recipients or '')
    unique = []
    for email in recipients:
        if email.lower() not in [x.lower() for x in unique]:
            unique.append(email)
    return unique


def _send_claim_notification(subject, body, claim=None, extra_recipients=None):
    if not CLAIM_EMAIL_ENABLED:
        return False
    sent_any = False
    for recipient in _claim_email_recipients(claim, extra_recipients):
        try:
            if send_system_email(subject, body, to_email=recipient):
                sent_any = True
        except Exception as exc:
            print(f'Claim email failed to {recipient}: {exc}')
    return sent_any


def _claim_email_body(title, claim, extra_lines=None):
    lines = [
        title,
        '',
        f"Claim ref: {claim.get('claim_ref') or claim.get('id')}",
        f"Franchise: {claim.get('franchise_name') or ''}",
        f"Claimant: {claim.get('claimant_name') or ''}",
        f"Policy number: {claim.get('policy_number') or ''}",
        f"Claim date: {claim.get('claim_date') or ''}",
        f"Claim amount: R{float(claim.get('claim_amount') or 0):,.2f}",
        f"Status: {claim.get('status') or ''}",
        f"Priority: {claim.get('priority') or ''}",
        f"Assigned to: {claim.get('assigned_to_email') or ''}",
    ]
    if extra_lines:
        lines += ['', *[x for x in extra_lines if x]]
    if claim.get('id'):
        lines += ['', f"Open claim: {_claim_link(claim.get('id'))}"]
    return '\n'.join(lines)

def _esc(value):
    return html.escape(str(value or ''))


def _can_access_franchise_name(franchise_name):
    user = getattr(g, 'user', None) or {}
    if user.get('role') == 'admin' or user.get('is_super_admin'):
        return True
    allowed = get_user_franchise_access(user.get('id'))
    allowed_keys = {str(x).strip().lower() for x in allowed}
    return str(franchise_name or '').strip().lower() in allowed_keys


def _available_franchises_for_user():
    user = getattr(g, 'user', None) or {}
    if user.get('role') == 'admin' or user.get('is_super_admin'):
        names = []
        try:
            monthly = LAST_RESULT.get('monthly', pd.DataFrame())
            if monthly is not None and not monthly.empty and 'Franchise' in monthly.columns:
                names = sorted({str(x).strip() for x in monthly['Franchise'].dropna().tolist() if str(x).strip()})
        except Exception:
            names = []
        if not names:
            engine = get_db_engine()
            if engine is not None:
                try:
                    with engine.begin() as conn:
                        rows = conn.execute(text("""
                            SELECT DISTINCT franchise_name FROM policy_monthly_raw
                            WHERE franchise_name IS NOT NULL AND TRIM(franchise_name) <> ''
                            ORDER BY franchise_name
                        """)).fetchall()
                    names = [str(r[0]) for r in rows]
                except Exception:
                    names = []
        return names
    return get_user_franchise_access(user.get('id'))


def _claim_row_to_dict(row):
    return dict(row) if row else None


def _get_claim_case(claim_id):
    engine = get_db_engine()
    if engine is None:
        return None
    with engine.begin() as conn:
        row = conn.execute(text("""
            SELECT * FROM app_claim_cases WHERE id = :id
        """), {'id': int(claim_id)}).mappings().first()
    claim = _claim_row_to_dict(row)
    if claim and not _can_access_franchise_name(claim.get('franchise_name')):
        return None
    return claim


def _status_options(selected):
    return ''.join([f'<option value="{_esc(x)}" {"selected" if x == selected else ""}>{_esc(x)}</option>' for x in CLAIM_STATUSES])


def _priority_options(selected):
    return ''.join([f'<option value="{_esc(x)}" {"selected" if x == selected else ""}>{_esc(x)}</option>' for x in CLAIM_PRIORITIES])


def _franchise_options(selected=''):
    names = _available_franchises_for_user()
    if not names:
        return '<option value="">No franchise access assigned</option>'
    return ''.join([f'<option value="{_esc(x)}" {"selected" if str(x) == str(selected) else ""}>{_esc(x)}</option>' for x in names])


@app.route('/claims')
def claims_workflow():
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('dashboard'))
    status_filter = (request.args.get('status') or '').strip()
    allowed = _available_franchises_for_user()
    user = getattr(g, 'user', None) or {}
    params = {}
    where = []
    if status_filter:
        where.append('status = :status')
        params['status'] = status_filter
    if not (user.get('role') == 'admin' or user.get('is_super_admin')):
        if not allowed:
            where.append('1 = 0')
        else:
            where.append('LOWER(TRIM(franchise_name)) = ANY(:allowed)')
            params['allowed'] = [str(x).strip().lower() for x in allowed]
    where_sql = ('WHERE ' + ' AND '.join(where)) if where else ''
    with engine.begin() as conn:
        claims = [dict(r) for r in conn.execute(text(f"""
            SELECT id, claim_ref, franchise_name, claimant_name, policy_number, claim_date,
                   claim_amount, status, priority, assigned_to_email, created_by_email,
                   created_at, updated_at, closed_at
            FROM app_claim_cases
            {where_sql}
            ORDER BY created_at DESC, id DESC
            LIMIT 500
        """), params).mappings().all()]
        summary = [dict(r) for r in conn.execute(text(f"""
            SELECT status, COUNT(*) AS count, COALESCE(SUM(claim_amount),0) AS amount
            FROM app_claim_cases
            {where_sql}
            GROUP BY status
            ORDER BY status
        """), params).mappings().all()]
    summary_rows = ''.join([f"<tr><td>{_esc(r.get('status'))}</td><td>{int(r.get('count') or 0)}</td><td>{money(r.get('amount') or 0)}</td></tr>" for r in summary]) or '<tr><td colspan="3">No claims yet.</td></tr>'
    rows = []
    for c in claims:
        rows.append(f"""
        <tr>
          <td><a class="link" href="/claims/{int(c.get('id'))}">{_esc(c.get('claim_ref') or c.get('id'))}</a></td>
          <td>{_esc(c.get('franchise_name'))}</td>
          <td>{_esc(c.get('claimant_name'))}</td>
          <td>{_esc(c.get('policy_number'))}</td>
          <td>{_esc(c.get('claim_date'))}</td>
          <td>{money(c.get('claim_amount') or 0)}</td>
          <td><strong>{_esc(c.get('status'))}</strong><br><span class="small">{_esc(c.get('priority'))}</span></td>
          <td>{_esc(c.get('created_at'))}</td>
        </tr>
        """)
    status_links = ' | '.join([f'<a class="link" href="/claims?status={_esc(x)}">{_esc(x)}</a>' for x in CLAIM_STATUSES])
    body = f"""<!doctype html><html><head><title>Claims Workflow</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/dashboard">Back to dashboard</a> &nbsp; | &nbsp; <a class="link" href="/claims/new">New Claim</a> &nbsp; | &nbsp; <a class="link" href="/claims">All Claims</a></div>
    <h1>Claims Workflow</h1><p class="muted">Capture, track and manage claim cases. Franchise users only see assigned franchises.</p>
    <p class="small">Filter: {status_links}</p>
    <h2>Summary</h2><table><tr><th>Status</th><th>Count</th><th>Amount</th></tr>{summary_rows}</table>
    <h2>Claim Cases</h2><table><tr><th>Ref</th><th>Franchise</th><th>Claimant</th><th>Policy</th><th>Claim Date</th><th>Amount</th><th>Status</th><th>Created</th></tr>{''.join(rows) or '<tr><td colspan="8">No claims found.</td></tr>'}</table>
    </div></body></html>"""
    return body


@app.route('/claims/new', methods=['GET', 'POST'])
def new_claim_case():
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        franchise_name = (request.form.get('franchise_name') or '').strip()
        claimant_name = (request.form.get('claimant_name') or '').strip()
        policy_number = (request.form.get('policy_number') or '').strip()
        claim_date = (request.form.get('claim_date') or '').strip() or None
        claim_amount = safe_float(request.form.get('claim_amount'), 0)
        priority = request.form.get('priority') if request.form.get('priority') in CLAIM_PRIORITIES else 'Normal'
        description = (request.form.get('description') or '').strip()
        if not franchise_name or not claimant_name:
            flash('Franchise and claimant name are required.', 'danger')
        elif not _can_access_franchise_name(franchise_name):
            flash('You do not have access to that franchise.', 'danger')
        else:
            user = getattr(g, 'user', None) or {}
            with engine.begin() as conn:
                row = conn.execute(text("""
                    INSERT INTO app_claim_cases (
                        claim_ref, franchise_name, claimant_name, policy_number, claim_date,
                        claim_amount, status, priority, description, created_by_id, created_by_email, assigned_to_email
                    ) VALUES (
                        :claim_ref, :franchise_name, :claimant_name, :policy_number, :claim_date,
                        :claim_amount, 'New', :priority, :description, :created_by_id, :created_by_email, :assigned_to_email
                    ) RETURNING id
                """), {
                    'claim_ref': 'CLM-' + datetime.now().strftime('%Y%m%d') + '-' + uuid.uuid4().hex[:6].upper(),
                    'franchise_name': franchise_name,
                    'claimant_name': claimant_name,
                    'policy_number': policy_number,
                    'claim_date': claim_date,
                    'claim_amount': claim_amount,
                    'priority': priority,
                    'description': description,
                    'created_by_id': user.get('id'),
                    'created_by_email': user.get('email'),
                    'assigned_to_email': request.form.get('assigned_to_email') or '',
                }).mappings().first()
                claim_id = int(row['id'])
                conn.execute(text("""
                    INSERT INTO app_claim_notes (claim_id, user_id, user_email, note_text, new_status)
                    VALUES (:claim_id, :user_id, :user_email, :note_text, 'New')
                """), {
                    'claim_id': claim_id,
                    'user_id': user.get('id'),
                    'user_email': user.get('email'),
                    'note_text': 'Claim created.' + (f' Description: {description}' if description else ''),
                })
            claim_for_email = _get_claim_case(claim_id) or {'id': claim_id, 'claim_ref': '', 'franchise_name': franchise_name, 'claimant_name': claimant_name, 'policy_number': policy_number, 'claim_date': claim_date, 'claim_amount': claim_amount, 'status': 'New', 'priority': priority, 'assigned_to_email': request.form.get('assigned_to_email') or '', 'created_by_email': user.get('email')}
            _send_claim_notification(
                f"New claim alert: {claim_for_email.get('claim_ref') or claim_id}",
                _claim_email_body('A new claim has been created.', claim_for_email),
                claim_for_email,
            )
            log_audit('claim_created', f'Claim case {claim_id} created for {franchise_name}')
            flash('Claim case created.', 'success')
            return redirect(url_for('claim_case_detail', claim_id=claim_id))
    body = f"""<!doctype html><html><head><title>New Claim</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/claims">Back to claims</a> &nbsp; | &nbsp; <a class="link" href="/dashboard">Back to dashboard</a></div>
    <h1>New Claim</h1><p class="muted">Create a claim case for tracking and follow-up.</p>
    <form method="post">
      <div class="field"><label>Franchise</label><select name="franchise_name" required>{_franchise_options()}</select></div>
      <div class="field"><label>Claimant name</label><input name="claimant_name" required></div>
      <div class="field"><label>Policy number</label><input name="policy_number"></div>
      <div class="field"><label>Claim date</label><input name="claim_date" type="date"></div>
      <div class="field"><label>Claim amount</label><input name="claim_amount" type="number" step="0.01"></div>
      <div class="field"><label>Priority</label><select name="priority">{_priority_options('Normal')}</select></div>
      <div class="field"><label>Assigned to email</label><input name="assigned_to_email" type="email"></div>
      <div class="field"><label>Description / notes</label><textarea name="description" style="width:100%;min-height:110px;border:1px solid #ccd7e6;border-radius:10px;padding:12px"></textarea></div>
      <button class="btn" type="submit">Create Claim</button>
    </form>
    </div></body></html>"""
    return body


@app.route('/claims/<int:claim_id>', methods=['GET', 'POST'])
def claim_case_detail(claim_id):
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('dashboard'))
    claim = _get_claim_case(claim_id)
    if not claim:
        flash('Claim not found or access denied.', 'danger')
        return redirect(url_for('claims_workflow'))
    user = getattr(g, 'user', None) or {}
    if request.method == 'POST':
        new_status = request.form.get('status') if request.form.get('status') in CLAIM_STATUSES else claim.get('status')
        priority = request.form.get('priority') if request.form.get('priority') in CLAIM_PRIORITIES else claim.get('priority')
        assigned_to_email = (request.form.get('assigned_to_email') or '').strip()
        note_text = (request.form.get('note_text') or '').strip()
        closed_at_sql = ', closed_at = NOW()' if new_status in {'Closed', 'Approved', 'Rejected'} and not claim.get('closed_at') else ''
        with engine.begin() as conn:
            conn.execute(text(f"""
                UPDATE app_claim_cases
                SET status = :status, priority = :priority, assigned_to_email = :assigned_to_email,
                    updated_at = NOW(){closed_at_sql}
                WHERE id = :id
            """), {'id': claim_id, 'status': new_status, 'priority': priority, 'assigned_to_email': assigned_to_email})
            if note_text or new_status != claim.get('status'):
                conn.execute(text("""
                    INSERT INTO app_claim_notes (claim_id, user_id, user_email, note_text, old_status, new_status)
                    VALUES (:claim_id, :user_id, :user_email, :note_text, :old_status, :new_status)
                """), {
                    'claim_id': claim_id,
                    'user_id': user.get('id'),
                    'user_email': user.get('email'),
                    'note_text': note_text or 'Status updated.',
                    'old_status': claim.get('status'),
                    'new_status': new_status,
                })
        if new_status != claim.get('status'):
            updated_claim = _get_claim_case(claim_id) or claim
            _send_claim_notification(
                f"Claim status changed: {updated_claim.get('claim_ref') or claim_id}",
                _claim_email_body(
                    f"Claim status changed from {claim.get('status')} to {new_status}.",
                    updated_claim,
                    [f"Note: {note_text}" if note_text else ''],
                ),
                updated_claim,
            )
        log_audit('claim_updated', f'Claim case {claim_id} updated: {claim.get("status")} -> {new_status}')
        flash('Claim case updated.', 'success')
        return redirect(url_for('claim_case_detail', claim_id=claim_id))
    with engine.begin() as conn:
        notes = [dict(r) for r in conn.execute(text("""
            SELECT user_email, note_text, old_status, new_status, created_at
            FROM app_claim_notes
            WHERE claim_id = :claim_id
            ORDER BY created_at DESC, id DESC
        """), {'claim_id': claim_id}).mappings().all()]
        attachments = [dict(r) for r in conn.execute(text("""
            SELECT id, filename, content_type, file_size, uploaded_by_email, created_at,
                   document_type, verification_status, verification_details, has_id_number, is_certified
            FROM app_claim_attachments
            WHERE claim_id = :claim_id
            ORDER BY created_at DESC, id DESC
        """), {'claim_id': claim_id}).mappings().all()]
    note_rows = ''.join([f"<tr><td>{_esc(n.get('created_at'))}</td><td>{_esc(n.get('user_email'))}</td><td>{_esc(n.get('old_status') or '')} to {_esc(n.get('new_status') or '')}</td><td>{_esc(n.get('note_text'))}</td></tr>" for n in notes]) or '<tr><td colspan="4">No notes yet.</td></tr>'
    doc_status = _claim_document_status(claim_id)
    checklist_rows = ''.join([f'<tr><td>{_esc(i.get("type"))}</td><td><strong>{_esc(i.get("status"))}</strong></td><td>{_esc(i.get("details"))}</td></tr>' for i in doc_status.get('items', [])]) or '<tr><td colspan="3">No checklist items.</td></tr>'
    attachment_rows = ''.join([f'<tr><td><a class="link" href="/claims/{claim_id}/attachments/{int(a.get("id"))}/download">{_esc(a.get("filename"))}</a></td><td>{_esc(a.get("document_type") or "")}</td><td><strong>{_esc(a.get("verification_status") or "Pending")}</strong><br><span class="small">{_esc(a.get("verification_details") or "")}</span></td><td>{"Yes" if a.get("has_id_number") else "No"}</td><td>{"Yes" if a.get("is_certified") else "No"}</td><td>{int(a.get("file_size") or 0)}</td><td>{_esc(a.get("uploaded_by_email"))}</td><td>{_esc(a.get("created_at"))}</td></tr>' for a in attachments]) or '<tr><td colspan="8">No attachments yet.</td></tr>'
    body = f"""<!doctype html><html><head><title>Claim {_esc(claim.get('claim_ref'))}</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/claims">Back to claims</a> &nbsp; | &nbsp; <a class="link" href="/dashboard">Back to dashboard</a></div>
    <h1>Claim {_esc(claim.get('claim_ref') or claim_id)}</h1>
    <table><tr><th>Field</th><th>Value</th></tr>
      <tr><td>Franchise</td><td>{_esc(claim.get('franchise_name'))}</td></tr>
      <tr><td>Claimant</td><td>{_esc(claim.get('claimant_name'))}</td></tr>
      <tr><td>Policy Number</td><td>{_esc(claim.get('policy_number'))}</td></tr>
      <tr><td>Claim Date</td><td>{_esc(claim.get('claim_date'))}</td></tr>
      <tr><td>Claim Amount</td><td>{money(claim.get('claim_amount') or 0)}</td></tr>
      <tr><td>Description</td><td>{_esc(claim.get('description'))}</td></tr>
      <tr><td>Created By</td><td>{_esc(claim.get('created_by_email'))}</td></tr>
    </table>
    <h2>Update Claim</h2>
    <form method="post">
      <div class="field"><label>Status</label><select name="status">{_status_options(claim.get('status'))}</select></div>
      <div class="field"><label>Priority</label><select name="priority">{_priority_options(claim.get('priority'))}</select></div>
      <div class="field"><label>Assigned to email</label><input name="assigned_to_email" type="email" value="{_esc(claim.get('assigned_to_email'))}"></div>
      <div class="field"><label>Add note</label><textarea name="note_text" style="width:100%;min-height:90px;border:1px solid #ccd7e6;border-radius:10px;padding:12px"></textarea></div>
      <button class="btn" type="submit">Save Claim Update</button>
    </form>
    <h2>Document Checklist</h2>
    <div class="card" style="padding:14px;margin:10px 0;border-left:8px solid {'#2f855a' if doc_status.get('approved') else '#a06a00'}">
      <h3 style="margin:0">{_esc(doc_status.get('overall_status'))}</h3>
      <p class="muted">{_esc(doc_status.get('summary'))}</p>
    </div>
    <table><tr><th>Required document</th><th>Status</th><th>Details</th></tr>{checklist_rows}</table>
    <form method="post" action="/claims/{claim_id}/documents/recheck" style="margin:8px 0 16px"><button class="btn btn-muted" type="submit">Recheck documents</button></form>
    <h2>Attachments</h2>
    <form method="post" action="/claims/{claim_id}/attachments" enctype="multipart/form-data" style="margin-bottom:12px">
      <input type="file" name="attachment" multiple required>
      <button class="btn" type="submit">Upload and Verify Document(s)</button>
    </form>
    <table><tr><th>File</th><th>Detected Document</th><th>Verification</th><th>ID No.</th><th>Certified</th><th>Size bytes</th><th>Uploaded By</th><th>Date</th></tr>{attachment_rows}</table>
    <h2>Claim Notes</h2><table><tr><th>Date</th><th>User</th><th>Status Change</th><th>Note</th></tr>{note_rows}</table>
    </div></body></html>"""
    return body



@app.route('/claims/<int:claim_id>/attachments', methods=['POST'])
def upload_claim_attachment(claim_id):
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('dashboard'))
    claim = _get_claim_case(claim_id)
    if not claim:
        flash('Claim not found or access denied.', 'danger')
        return redirect(url_for('claims_workflow'))
    files = request.files.getlist('attachment') or request.files.getlist('attachments')
    files = [f for f in files if f and f.filename]
    if not files:
        flash('Please select one or more files to upload.', 'danger')
        return redirect(url_for('claim_case_detail', claim_id=claim_id))
    user = getattr(g, 'user', None) or {}
    uploaded_names = []
    with engine.begin() as conn:
        for file in files:
            original_name = secure_filename(file.filename) or 'attachment'
            stored_name = f"{uuid.uuid4().hex}_{original_name}"
            claim_dir = os.path.join(CLAIM_ATTACHMENT_DIR, str(claim_id))
            os.makedirs(claim_dir, exist_ok=True)
            path = os.path.join(claim_dir, stored_name)
            file.save(path)
            size = os.path.getsize(path) if os.path.exists(path) else 0
            verification = _verify_claim_attachment(original_name, path)
            conn.execute(text("""
                INSERT INTO app_claim_attachments (
                    claim_id, filename, stored_filename, file_path, content_type, file_size,
                    document_type, verification_status, verification_details, extracted_text,
                    has_id_number, is_certified, uploaded_by_id, uploaded_by_email
                ) VALUES (
                    :claim_id, :filename, :stored_filename, :file_path, :content_type, :file_size,
                    :document_type, :verification_status, :verification_details, :extracted_text,
                    :has_id_number, :is_certified, :uploaded_by_id, :uploaded_by_email
                )
            """), {
                'claim_id': claim_id,
                'filename': original_name,
                'stored_filename': stored_name,
                'file_path': path,
                'content_type': file.mimetype or '',
                'file_size': size,
                'document_type': verification.get('document_type'),
                'verification_status': verification.get('verification_status'),
                'verification_details': verification.get('verification_details'),
                'extracted_text': verification.get('extracted_text'),
                'has_id_number': verification.get('has_id_number'),
                'is_certified': verification.get('is_certified'),
                'uploaded_by_id': user.get('id'),
                'uploaded_by_email': user.get('email'),
            })
            uploaded_names.append(f"{original_name} ({verification.get('document_type')}: {verification.get('verification_status')})")
            conn.execute(text("""
                INSERT INTO app_claim_notes (claim_id, user_id, user_email, note_text)
                VALUES (:claim_id, :user_id, :user_email, :note_text)
            """), {
                'claim_id': claim_id,
                'user_id': user.get('id'),
                'user_email': user.get('email'),
                'note_text': f"Attachment uploaded and checked: {original_name}. {verification.get('verification_details')}",
            })
        doc_status = _update_claim_document_status(claim_id, conn=conn)
    _send_claim_notification(
        f"Claim document(s) uploaded: {claim.get('claim_ref') or claim_id}",
        _claim_email_body('New claim document(s) were uploaded and checked.', claim, uploaded_names + [f"Overall document status: {doc_status.get('overall_status')}", doc_status.get('summary')]),
        claim,
    )
    log_audit('claim_attachment_uploaded', f'Claim case {claim_id}: {", ".join(uploaded_names)}')
    flash(f'{len(uploaded_names)} document(s) uploaded and checked. Status: {doc_status.get("overall_status")}.', 'success')
    return redirect(url_for('claim_case_detail', claim_id=claim_id))


@app.route('/claims/<int:claim_id>/documents/recheck', methods=['POST'])
def recheck_claim_documents(claim_id):
    engine = get_db_engine()
    claim = _get_claim_case(claim_id)
    if engine is None or not claim:
        flash('Claim not found or database not connected.', 'danger')
        return redirect(url_for('claims_workflow'))
    with engine.begin() as conn:
        rows = [dict(r) for r in conn.execute(text("""
            SELECT id, filename, file_path FROM app_claim_attachments WHERE claim_id = :claim_id
        """), {'claim_id': int(claim_id)}).mappings().all()]
        for r in rows:
            verification = _verify_claim_attachment(r.get('filename'), r.get('file_path'))
            conn.execute(text("""
                UPDATE app_claim_attachments
                SET document_type = :document_type, verification_status = :verification_status,
                    verification_details = :verification_details, extracted_text = :extracted_text,
                    has_id_number = :has_id_number, is_certified = :is_certified
                WHERE id = :id
            """), {
                'id': int(r.get('id')),
                'document_type': verification.get('document_type'),
                'verification_status': verification.get('verification_status'),
                'verification_details': verification.get('verification_details'),
                'extracted_text': verification.get('extracted_text'),
                'has_id_number': verification.get('has_id_number'),
                'is_certified': verification.get('is_certified'),
            })
        status = _update_claim_document_status(claim_id, conn=conn)
    flash(f'Document checklist rechecked. Status: {status.get("overall_status")}.', 'success')
    return redirect(url_for('claim_case_detail', claim_id=claim_id))


@app.route('/claims/<int:claim_id>/attachments/<int:attachment_id>/download')
def download_claim_attachment(claim_id, attachment_id):
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('dashboard'))
    claim = _get_claim_case(claim_id)
    if not claim:
        flash('Claim not found or access denied.', 'danger')
        return redirect(url_for('claims_workflow'))
    with engine.begin() as conn:
        row = conn.execute(text("""
            SELECT filename, file_path, content_type FROM app_claim_attachments
            WHERE id = :id AND claim_id = :claim_id
        """), {'id': attachment_id, 'claim_id': claim_id}).mappings().first()
    if not row or not os.path.exists(row.get('file_path')):
        flash('Attachment not found.', 'danger')
        return redirect(url_for('claim_case_detail', claim_id=claim_id))
    return send_file(row.get('file_path'), as_attachment=True, download_name=row.get('filename') or 'attachment', mimetype=row.get('content_type') or None)


@app.route('/admin/claims/send_overdue_reminders')
def send_overdue_claim_reminders():
    engine = get_db_engine()
    if engine is None:
        flash('Database not connected.', 'danger')
        return redirect(url_for('dashboard'))
    cutoff = datetime.now() - timedelta(days=CLAIM_OVERDUE_DAYS)
    with engine.begin() as conn:
        rows = [dict(r) for r in conn.execute(text("""
            SELECT * FROM app_claim_cases
            WHERE COALESCE(archived, false) = false
              AND status NOT IN ('Closed', 'Approved', 'Rejected')
              AND created_at < :cutoff
            ORDER BY created_at ASC, id ASC
        """), {'cutoff': cutoff}).mappings().all()]
    sent = 0
    for claim in rows:
        if _send_claim_notification(
            f"Overdue claim reminder: {claim.get('claim_ref') or claim.get('id')}",
            _claim_email_body(f"This claim is overdue by the {CLAIM_OVERDUE_DAYS}-day reminder rule.", claim),
            claim,
        ):
            sent += 1
    log_audit('claim_overdue_reminders_sent', f'Sent {sent} overdue reminders from {len(rows)} overdue claims')
    flash(f'Overdue reminder emails sent for {sent} claim(s).', 'success')
    return redirect(url_for('claims_workflow'))


@app.route('/admin/database_health')
def database_health():
    report = database_health_report()
    rows = []
    for t in report.get('tables', []):
        rows.append(f"<tr><td>{t['table']}</td><td>{'Yes' if t['exists'] else 'No'}</td><td>{t['rows']}</td></tr>")
    missing_rows = []
    for m in report.get('missing', []):
        missing_rows.append(f"<tr><td>{m['table']}</td><td>{m['column']}</td><td>{m['type']}</td></tr>")
    missing_html = ''.join(missing_rows) or '<tr><td colspan="3">No missing fields</td></tr>'
    return """
    <html><head><title>Database Health</title><style>
        body{{font-family:Arial,sans-serif;margin:24px;background:#f7f7f7;color:#222;}}
        table{{border-collapse:collapse;background:white;margin:12px 0;width:100%;max-width:1000px;}}
        th,td{{border:1px solid #ddd;padding:8px;text-align:left;}}
        th{{background:#222;color:white;}}
        a.button{{display:inline-block;background:#111;color:white;padding:10px 14px;text-decoration:none;border-radius:6px;}}
        </style></head><body>
        <h1>Database Health</h1>
        <p>Status: <strong>{message}</strong></p>
        <p><a class="button" href="/admin/repair_database">Repair Database</a></p>
        <h2>Tables</h2><table><tr><th>Table</th><th>Exists</th><th>Rows</th></tr>{rows}</table>
        <h2>Missing Fields</h2><table><tr><th>Table</th><th>Column</th><th>Type</th></tr>{missing}</table>
        <p><a href="/dashboard">Back to dashboard</a></p>
        </body></html>
        """.format(message=report.get('message',''), rows=''.join(rows), missing=missing_html)

@app.route('/admin/repair_database')
def repair_database():
    ensure_database_schema()
    return redirect(url_for('database_health'))

@app.route('/sample')
def sample():
    single_monthly_premium_total = 0.0
    data = []
    franchises = [('Franchise A', '100% Claim Ratio'), ('Franchise B', '100% Claim Ratio'), ('Franchise C', 'BrightRock')]
    months = pd.date_range('2024-01-01', periods=24, freq='MS')
    for f_idx, (franchise, scenario) in enumerate(franchises):
        for i, month in enumerate(months):
            retail = 100000 + f_idx * 25000 + i * 2500
            original_risk = retail * 0.68
            single_monthly_premium_total = original_risk
            r1_fee = 100 + f_idx * 35 + i * 3
            risk_after_r1 = original_risk - r1_fee
            underwriter_2_1_fee = risk_after_r1 * 0.021
            risk = risk_after_r1 - underwriter_2_1_fee
            claims = risk * (0.42 + f_idx * 0.18 + (i % 6) * 0.035)
            data.append({
                'Franchise': franchise,
                'Month': month,
                'Retail Premium': retail,
                'Original Risk Premium': original_risk,
                'Risk Premium': risk,
                'Risk After R1': risk_after_r1,
                'ADV Fee 2.1%': underwriter_2_1_fee,
                'Underwriter 2.1% Fee': underwriter_2_1_fee,
                'Single Monthly Premium Total': single_monthly_premium_total,
                'Claims': claims,
                'Policy Qty': 100 + f_idx * 35 + i * 3,
                'Current Scenario': scenario,
            })
    path = os.path.join(EXPORT_DIR, 'sample_import_file.xlsx')
    pd.DataFrame(data).to_excel(path, index=False)
    return send_file(path, as_attachment=True)




# Compatibility route for dashboard Phase 8 Claims Reports card.
# The dashboard template links to endpoint `claims_report`; keep this route present
# so the dashboard can always render even if detailed report pages are added later.
@app.route('/claims/report')
def claims_report():
    return redirect(url_for('claims_workflow'))



def load_default_import_file():
    global LAST_RESULT
    if os.path.exists(DEFAULT_IMPORT_FILE):
        try:
            raw = read_excel_file(DEFAULT_IMPORT_FILE)
            monthly, periods, portfolio = analyse(raw, DEFAULT_RATES.copy(), DEFAULT_BOOK_VALUE.copy())
            LAST_RESULT = {
                'raw': raw,
                'monthly': monthly,
                'periods': periods,
                'portfolio': portfolio,
                'rates': DEFAULT_RATES.copy(),
                'book_rates': DEFAULT_BOOK_VALUE.copy(),
            }
        except Exception as exc:
            print(f'Could not load default import file: {exc}')


# Load PostgreSQL data first; fall back to bundled sample only when DB is empty/unavailable.
if not reload_dashboard_from_postgres():
    load_default_import_file()



def admin_required(func):
    """Small route decorator used by late-added admin routes."""
    @wraps(func)
    def wrapper(*args, **kwargs):
        user = getattr(g, 'user', None) or {}
        if not user.get('id'):
            return redirect(url_for('login', next=request.path))
        if not (user.get('role') == 'admin' or user.get('is_super_admin')):
            flash('Admin access required.', 'danger')
            return redirect(url_for('dashboard'))
        return func(*args, **kwargs)
    return wrapper





@app.route('/admin/parlour_fee_settings', methods=['GET', 'POST'])
@admin_required
def admin_parlour_fee_settings():
    ensure_parlour_fee_settings_table()
    monthly = apply_user_franchise_scope(LAST_RESULT.get('monthly', pd.DataFrame()))
    franchises = sorted(monthly['Franchise'].dropna().unique()) if monthly is not None and not monthly.empty and 'Franchise' in monthly.columns else []
    selected = request.values.get('franchise') or (franchises[0] if franchises else 'All')
    if request.method == 'POST':
        try:
            vals = {
                'franchise': selected,
                'commission': float(request.form.get('commission_to_parlour_pct') or 0),
                'admin': float(request.form.get('admin_to_parlour_pct') or 0),
                'facility': float(request.form.get('facility_fee_to_parlour_pct') or 0),
                'notes': request.form.get('notes') or '',
            }
            engine = get_db_engine()
            if engine is not None:
                with engine.begin() as conn:
                    conn.execute(text("""
                        INSERT INTO app_franchise_parlour_fee_settings
                            (franchise_name, commission_to_parlour_pct, admin_to_parlour_pct, facility_fee_to_parlour_pct, notes, updated_at)
                        VALUES (:franchise, :commission, :admin, :facility, :notes, NOW())
                        ON CONFLICT (franchise_name) DO UPDATE SET
                            commission_to_parlour_pct = EXCLUDED.commission_to_parlour_pct,
                            admin_to_parlour_pct = EXCLUDED.admin_to_parlour_pct,
                            facility_fee_to_parlour_pct = EXCLUDED.facility_fee_to_parlour_pct,
                            notes = EXCLUDED.notes,
                            updated_at = NOW()
                    """), vals)
            flash('Parlour fee settings saved.', 'success')
            return redirect(url_for('admin_parlour_fee_settings', franchise=selected))
        except Exception as exc:
            flash(f'Could not save settings: {exc}', 'danger')
    settings = get_parlour_fee_settings(selected)
    rows = []
    try:
        engine = get_db_engine()
        if engine is not None:
            with engine.begin() as conn:
                rows = [dict(r) for r in conn.execute(text("""
                    SELECT franchise_name, commission_to_parlour_pct, admin_to_parlour_pct, facility_fee_to_parlour_pct, updated_at
                    FROM app_franchise_parlour_fee_settings
                    ORDER BY franchise_name
                """)).mappings().all()]
    except Exception:
        rows = []
    opts = ''.join([f'<option value="{html.escape(str(f))}" {"selected" if f==selected else ""}>{html.escape(str(f))}</option>' for f in franchises])
    table_rows = ''.join([f"<tr><td>{html.escape(str(r.get('franchise_name') or ''))}</td><td>{float(r.get('commission_to_parlour_pct') or 0):.2f}%</td><td>{float(r.get('admin_to_parlour_pct') or 0):.2f}%</td><td>{float(r.get('facility_fee_to_parlour_pct') or 0):.2f}%</td><td>{r.get('updated_at') or ''}</td></tr>" for r in rows]) or '<tr><td colspan="5">No configured franchise rates yet. Defaults are used until saved.</td></tr>'
    body = f"""
    <div class="topnav"><a class="link" href="/dashboard">Back to dashboard</a> &nbsp; | &nbsp; <a class="link" href="/admin/system_health">System Health</a></div>
    <h1>Parlour Fee Settings</h1>
    <p class="muted">Enter/change Commission, Admin and Facility Fee percentages per franchise. Dashboard calculations read these settings dynamically.</p>
    <form method="get" class="filter-row"><label>Franchise<br><select name="franchise" onchange="this.form.submit()">{opts}</select></label></form>
    <form method="post" class="admin-form" style="max-width:720px">
      <input type="hidden" name="franchise" value="{html.escape(str(selected))}">
      <label>Commission to Parlour %<br><input name="commission_to_parlour_pct" type="number" step="0.01" value="{settings.get('commission_to_parlour_pct',0):.2f}"></label>
      <label>Admin to Parlour %<br><input name="admin_to_parlour_pct" type="number" step="0.01" value="{settings.get('admin_to_parlour_pct',0):.2f}"></label>
      <label>Facility Fee to Parlour %<br><input name="facility_fee_to_parlour_pct" type="number" step="0.01" value="{settings.get('facility_fee_to_parlour_pct',0):.2f}"></label>
      <label>Notes<br><textarea name="notes" rows="3"></textarea></label>
      <button class="btn" type="submit">Save Settings</button>
    </form>
    <h2>Configured Rates</h2>
    <div class="table-wrap"><table><thead><tr><th>Franchise</th><th>Commission</th><th>Admin</th><th>Facility</th><th>Updated</th></tr></thead><tbody>{table_rows}</tbody></table></div>
    """
    return render_auth_page('Parlour Fee Settings', body)


@app.route('/admin/age_notifications/franchise/<path:franchise_name>')
@admin_required
def admin_age_notifications_franchise(franchise_name):
    selected = unquote_plus(franchise_name or 'All')
    groups = get_policy_age_notification_franchise_age_groups(selected)
    total = sum(int(g.get('count') or 0) for g in groups)
    group_rows = []
    for g in groups:
        band = g.get('age_band') or 'Unknown'
        group_rows.append(f"""
        <tr>
          <td><a class="link" href="/admin/age_notifications/franchise/{quote(selected, safe='')}/band/{quote(str(band), safe='')}">{html.escape(str(band))}</a></td>
          <td>{int(g.get('count') or 0)}</td>
          <td>{g.get('next_due_date') or '-'}</td>
        </tr>
        """)
    return f"""<!doctype html><html><head><title>Age Notifications - {html.escape(selected)}</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/dashboard?franchise={quote_plus(selected)}">Back to dashboard</a> &nbsp; | &nbsp; <a class="link" href="/admin/map_cache_manager">Map Cache Manager</a> &nbsp; | &nbsp; <a class="link" href="/admin/system_health">System Health</a></div>
    <h1>Age Limit Notifications</h1>
    <p class="muted"><strong>{html.escape(selected)}</strong> - {total} active age-limit notices. Click an age band to list affected members.</p>
    <form method="post" action="/admin/policy_age_notifications/refresh" style="margin-bottom:14px">
      <input type="hidden" name="franchise" value="{html.escape(selected)}">
      <input type="hidden" name="limit" value="600000">
      <button class="btn" type="submit">Refresh this franchise</button>
    </form>
    <table><tr><th>Age Band</th><th>Total to Fix</th><th>Next Due</th></tr>{''.join(group_rows) or '<tr><td colspan="3">No active age notifications.</td></tr>'}</table>
    </div></body></html>"""


@app.route('/admin/age_notifications/franchise/<path:franchise_name>/band/<path:age_band>')
@admin_required
def admin_age_notifications_franchise_band(franchise_name, age_band):
    selected = unquote_plus(franchise_name or 'All')
    band = unquote_plus(age_band or 'Unknown')
    rows = get_policy_age_notification_members(selected, limit=int(request.args.get('limit', 10000)), age_band=band)
    body_rows = []
    for r in rows:
        body_rows.append(f"""
        <tr>
          <td>{html.escape(str(r.get('member_name') or '-'))}</td>
          <td>{html.escape(str(r.get('policy_number') or '-'))}</td>
          <td>{html.escape(str(r.get('id_number') or '-'))}</td>
          <td>{html.escape(str(r.get('relation') or '-'))}</td>
          <td>{r.get('age_years') if r.get('age_years') is not None else '-'}</td>
          <td>{html.escape(str(r.get('age_band') or '-'))}</td>
          <td>{r.get('due_date') or '-'}</td>
          <td>{html.escape(str(r.get('message') or '-'))}</td>
        </tr>
        """)
    return f"""<!doctype html><html><head><title>{html.escape(band)} - {html.escape(selected)}</title>{AUTH_PAGE_CSS}</head><body><div class="admin-wrap">
    <div class="topnav"><a class="link" href="/admin/age_notifications/franchise/{quote(selected, safe='')}">Back to age groups</a> &nbsp; | &nbsp; <a class="link" href="/dashboard?franchise={quote_plus(selected)}">Back to dashboard</a></div>
    <h1>{html.escape(band)}</h1>
    <p class="muted"><strong>{html.escape(selected)}</strong> - {len(rows)} affected members in this age band.</p>
    <table><tr><th>Member</th><th>Policy</th><th>ID Number</th><th>Relation</th><th>Age</th><th>Age Band</th><th>Due Date</th><th>Message</th></tr>{''.join(body_rows) or '<tr><td colspan="8">No active age notifications for this age band.</td></tr>'}</table>
    </div></body></html>"""

@app.route('/admin/policy_age_notifications/refresh', methods=['POST'])
@admin_required
def admin_refresh_policy_age_notifications():
    selected = request.form.get('franchise') or request.args.get('franchise') or 'All'
    limit = int(request.form.get('limit') or request.args.get('limit') or os.getenv('AGE_NOTIFICATION_REFRESH_LIMIT', '600000'))
    result = refresh_policy_age_notifications(limit=limit, selected=selected)
    flash(f"Age notification refresh processed {result.get('processed',0)} rows and created/updated {result.get('created',0)} notifications.")
    return redirect(url_for('dashboard', franchise=selected))



def _run_policydata_current_members_rebuild_background():
    """Run rebuild in a real Flask app context. Do not call request/session helpers here."""
    print('Current-member rebuild thread started', flush=True)
    try:
        with app.app_context():
            rebuild_policydata_current_members_job()
        print('Current-member rebuild thread completed', flush=True)
    except Exception as exc:
        print(f'Current-member rebuild thread failed: {exc}', flush=True)
        try:
            _POLICY_CURRENT_REBUILD_STATUS.update({
                'running': False,
                'status': 'error',
                'error': str(exc),
                'updated_at': datetime.now().isoformat()
            })
        except Exception:
            pass

def _start_policydata_current_members_rebuild():
    """Start the current-member rebuild thread exactly once."""
    global _POLICY_CURRENT_REBUILD_STATUS
    if _POLICY_CURRENT_REBUILD_STATUS.get('running'):
        return False
    _POLICY_CURRENT_REBUILD_STATUS.update({
        'running': True,
        'status': 'starting',
        'processed': 0,
        'error': '',
        'updated_at': datetime.now().isoformat()
    })
    # Delay the heavy background thread very slightly so the HTTP response can
    # return immediately. This avoids Edge/Chrome showing the start request as
    # pending while PostgreSQL begins the rebuild.
    def _delayed_start():
        t = threading.Thread(target=_run_policydata_current_members_rebuild_background, daemon=True)
        t.start()
    threading.Timer(0.3, _delayed_start).start()
    return True


@app.after_request
def add_no_cache_for_admin_jobs(response):
    try:
        if request.path.startswith('/admin/policydata_current_members'):
            response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
            response.headers['Pragma'] = 'no-cache'
    except Exception:
        pass
    return response

@app.route('/admin/policydata_current_members/rebuild/start', methods=['GET', 'POST'])
@admin_required
def admin_rebuild_policydata_current_members_start():
    # Direct GET-start endpoint. Return immediately so the browser never waits for the rebuild.
    _start_policydata_current_members_rebuild()
    status_url = url_for('admin_rebuild_policydata_current_members')
    return f"""<!doctype html><html><head><title>Rebuild Started</title>{AUTH_PAGE_CSS}<meta http-equiv='refresh' content='1;url={status_url}'></head><body><div class='admin-wrap'>
    <h1>Rebuild started</h1>
    <p>The current-member rebuild is running in the background.</p>
    <p><a class='btn' href='{status_url}'>View progress</a></p>
    </div></body></html>"""

@app.route('/admin/policydata_current_members/rebuild', methods=['GET', 'POST'])
@admin_required
def admin_rebuild_policydata_current_members():
    if request.method == 'POST':
        # Backward compatible; the visible button now uses the GET /start endpoint.
        _start_policydata_current_members_rebuild()
    st = _POLICY_CURRENT_REBUILD_STATUS
    running = bool(st.get('running'))
    refresh_meta = "<meta http-equiv='refresh' content='3'>" if running else ""
    action_url = url_for('admin_rebuild_policydata_current_members')
    start_url = url_for('admin_rebuild_policydata_current_members_start')
    status_url = url_for('admin_policydata_current_members_status')
    return f"""<!doctype html><html><head><title>Rebuild Current Members</title>{AUTH_PAGE_CSS}{refresh_meta}<script>if('serviceWorker' in navigator){{navigator.serviceWorker.getRegistrations().then(rs=>rs.forEach(r=>r.unregister()));}}</script></head><body><div class='admin-wrap'>
    <div class='topnav'><a class='link' href='/dashboard'>Back to dashboard</a> &nbsp; | &nbsp; <a class='link' href='/admin/system_health'>System Health</a></div>
    <h1>Rebuild Current Policy Members</h1>
    <p class='muted'>This creates one latest/current record per member from PolicyData history, then refreshes Policy Relation Counts. It runs in the background and prevents dashboard duplicates/freezing.</p>
    <p><strong>Status:</strong> {html.escape(str(st.get('status')))}<br>
       <strong>Processed:</strong> {int(st.get('processed') or 0):,}<br>
       <strong>Updated:</strong> {html.escape(str(st.get('updated_at') or '-'))}<br>
       <strong>Error:</strong> {html.escape(str(st.get('error') or '-'))}</p>
    <p><a class='btn' href='{start_url}'>{'Rebuild Running...' if running else 'Start Rebuild'}</a></p>
    <p class='muted'>If the button does not visibly start, open this direct start URL: <a class='link' href='{start_url}'>{start_url}</a></p>
    <p class='muted'>Live status JSON: <a class='link' href='{status_url}'>{status_url}</a></p>
    <p class='muted'>When complete, dashboard Policy Relation Counts will read from app_policy_relation_summary.</p>
    </div></body></html>"""

@app.route('/admin/policydata_current_members/status')
@admin_required
def admin_policydata_current_members_status():
    return jsonify(_POLICY_CURRENT_REBUILD_STATUS)

# Start daily background scheduler only when the Flask server actually starts.
# This avoids silent startup hangs during import/module loading.
if __name__ == '__main__':
    print('Starting Franchise Claims Analytics System...', flush=True)
    print('Loading web server...', flush=True)
    try:
        start_policy_age_scheduler()
    except Exception as exc:
        print(f'Daily policy age scheduler could not start: {exc}', flush=True)
    print('Starting Flask on http://127.0.0.1:5000 ...', flush=True)
    app.run(host='0.0.0.0', port=int(os.getenv('PORT', 5000)), debug=os.getenv('FLASK_DEBUG') == '1', use_reloader=False, threaded=True)
